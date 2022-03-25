#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
"""

__author__    = "Suguru Ishizaki"
__copyright__ = "2017-20 Suguru Ishizaki, Carnegie Mellon University"


import os, sys
import string
import re
import copy
import shutil
import difflib
import itertools
import json
import platform

from anytree import Node, RenderTree
from anytree.render import AsciiStyle

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt

from time import time
from datetime import datetime, date
from operator import itemgetter

from zipfile import ZipFile

CATEGORY    = 1
SUBCATEGORY = 2
CLUSTER     = 3
DIMENSION   = 4
LAT         = 5

import dslib.utils  as utils
import docuscope.dstagger as dstagger

import pprint     
pp = pprint.PrettyPrinter(indent=4)

# spw = "wuDtH3tXopS937hNOvGW"

##################################################
#
# class DS Dict
#
##################################################

untitled_category_count = 0

def adjustCases(patterns):
    adjustedPatterns = list()
    for p in patterns:
        adjusted_p = list()
        for w in p.strip().split():
            if w[0]=='!':
                adjusted_p.append(w.upper())
            else:
                adjusted_p.append(w.lower())
        adjustedPatterns.append(' '.join(adjusted_p))
    return adjustedPatterns
    
ignored_clusters = ['Orphaned', 'SyntacticComplexity', 'Other']

class DSDict():

    def __init__(self, controller, path, compact=False):

        self.controller = controller
        self.is_compact = compact

        # if os.path.isdir(path):
        #     self.directory = path
        #     self.dat_fpath = None

        # elif os.path.isfile(path):
        #     self.directory = None
        #     self.dat_fpath = path

        self.directory = path

        self.tones_edited = False
        self.wordclasses_edited = False

        self.info = None # only used when an encrypted dictionary is used.

        # list of lat names that are no-longer used. They are candidates for removal, 
        # but they may be re-used. So, we can't blindly delete all of them.
        self.deleted_lat_names = list()

        self.tones = list()
        self.help = list()

        self.lats = dict()

        self.tagger = None

        self.zip_basename = ""

        self.is_custom_dict = False

        self.pronouns_visible = False

    def setCustomDict(self, val):
        self.is_custom_dict = val

    def setPronouns(self, val):
        self.pronouns_visible = val

    def pronouns(self):
        return self.pronouns_visible

    def isCustomDict(self):
        return self.is_custom_dict

    def isRuleSet(self):
        if self.directory is not None:
            path = os.path.join(self.directory, "rules.json")
            return os.path.exists(path)
        else:
            return None

    def getCustomDictName(self):
        custom_dict_name = os.path.basename(self.directory)
        return custom_dict_name

    def invalidateTagger(self):
        self.tagger = None

    def loadTones(self):
        """
        Read _tones.txt to create a dictionary hierarchy;
        """
        global untitled_category_count

        clust_count = 0
        dim_count   = 0
        lat_count   = 0

        cluster_d = None
        dim_d     = None
        lat_d     = None

        path = os.path.join(self.directory, "_tones.txt")

        if os.path.exists(path) != True:
            self.controller.showWarningDialog("Warning", "{} does not exist.".format(path))
            return

        with open(path, encoding="utf-8") as tones_file:
            for line in tones_file:
                # cluster
                if(line.startswith('CLUSTER:')):   
                    cluster_name = line[9:].strip()

                    if re.search('untitled[0-9]+$', cluster_name.lower()) is not None:
                    # if cluster_name.lower().startswith('untitled'):
                        untitled_category_count+=1
                        cluster_name = "Untitled{}".format(untitled_category_count)

                    cluster_d = dict()
                    cluster_d['type'] = 'CLUSTER'
                    cluster_d['name'] = cluster_name
                    cluster_d['lc_name'] = cluster_name.lower()
                    cluster_d['dimensions'] = list()
                    cluster_d['default'] = True                       
                    self.tones.append(cluster_d)

                    clust_count += 1

                if cluster_d != None and (line.startswith('DIMENSION:')):
                    dim_name = line[11:].strip()

                    # if dim_name.lower().startswith('untitled'):
                    if re.search('untitled[0-9]+$', dim_name.lower()) is not None:
                        untitled_category_count+=1
                        dim_name = "Untitled{}".format(untitled_category_count)

                    dim_d = dict()
                    dim_d['type'] = 'DIMENSION'
                    dim_d['name'] = dim_name
                    dim_d['lc_name'] = dim_name.lower()
                    dim_d['default'] = True                    
                    dim_d['lats'] = list()
                    cluster_d['dimensions'].append(dim_d)

                    dim_count += 1

                if dim_d != None and (line.startswith('LAT:') or line.startswith('LAT*:')):
                    lat_name = line[5:].strip()

                    if lat_name:
                        # if lat_name.lower().startswith('untitled'):
                        if re.search('untitled[0-9]+$', lat_name.lower()) is not None:
                            untitled_category_count+=1
                            lat_name = "Untitled{}".format(untitled_category_count)
                        
                        lat_d = dict()
                        lat_d['type'] = 'LAT'
                        lat_d['name'] = lat_name
                        lat_d['lc_name'] = lat_name.lower()
                        lat_d['original_patterns'] = None
                        lat_d['revised_patterns']  = None
                        lat_d['original_count'] = 0
                        lat_d['revised_count']  = 0
                        lat_d['added'] = []      # all of the added/new patterns
                        lat_d['deleted'] = []    # deleted patterns
                        lat_d['verified'] = []   # new/added but already verified.
                        lat_d['is_new'] = False
                        lat_d['deleted'] = True
                        lat_d['default'] = True
                        dim_d['lats'].append(lat_d)  

                        lat_count += 1                

                        self.lats[lat_name] = (dim_name, cluster_name)

        # sort the dimension names
        for c in self.tones:
            c['dimensions'] = sorted(c['dimensions'] , key=itemgetter('lc_name'))

        for c in self.tones:
            for d in c['dimensions']:
                d['lats'] = sorted(d['lats'] , key=itemgetter('lc_name'))

        self.clust_count = clust_count
        self.dim_count   = dim_count
        self.lat_count   = lat_count

    def loadInfo(self):
        info_fpath = os.path.join(self.directory, "info.json")
        with open(info_fpath) as fin:
            self.info = json.load(fin)

        self.is_custom_dict = self.info.get('customized', False)

    def loadInfoFromDatFile(self, zipf):
        with zipf.open("info.json") as info_file:
            data = info_file.read()
            data = data.decode('utf-8')
            self.info = json.loads(data)

    def loadTonesFromDatFile(self, zipf):
        """
        Read _tones.txt to create a dictionary hierarchy;
        """
        global untitled_category_count

        cluster_d = None
        dim_d     = None
        lat_d     = None

        clust_count = 0
        dim_count   = 0
        lat_count   = 0

        with zipf.open("_tones.txt") as tones_file:
            for line in tones_file:
                line = line.decode('utf-8')
                # custom cluster
                if(line.startswith('CLUSTER:')):   
                    cluster_name = line[9:].strip()

                    # if cluster_name.lower().startswith('untitled'):
                    if re.search('untitled[0-9]+$', cluster_name.lower()) is not None:
                        untitled_category_count+=1
                        cluster_name = "Untitled{}".format(untitled_category_count)

                    cluster_d = dict()
                    cluster_d['type'] = 'CLUSTER'
                    cluster_d['name'] = cluster_name
                    cluster_d['lc_name'] = cluster_name.lower()
                    cluster_d['dimensions'] = list()
                    cluster_d['default'] = True   
                    self.tones.append(cluster_d)
                    clust_count += 1

                if cluster_d != None and (line.startswith('DIMENSION:')):
                    dim_name = line[11:].strip()

                    # if dim_name.lower().startswith('untitled'):
                    if re.search('untitled[0-9]+$', dim_name.lower()) is not None:
                        untitled_category_count+=1
                        dim_name = "Untitled{}".format(untitled_category_count)

                    dim_d = dict()
                    dim_d['type'] = 'DIMENSION'
                    dim_d['name'] = dim_name
                    dim_d['lc_name'] = dim_name.lower()
                    dim_d['lats'] = list()
                    dim_d['default'] = True      #
                    cluster_d['dimensions'].append(dim_d)
                    dim_count += 1

                if dim_d != None and (line.startswith('LAT:')):
                    lat_name = line[5:].strip()

                    if lat_name.lower().startswith('untitled'):
                        untitled_category_count+=1
                        lat_name = "Untitled{}".format(untitled_category_count)
                    
                    lat_d = dict()
                    lat_d['type'] = 'LAT'
                    lat_d['name'] = lat_name
                    lat_d['lc_name'] = lat_name.lower()
                    lat_d['original_patterns'] = None
                    lat_d['revised_patterns']  = None
                    lat_d['original_count'] = 0
                    lat_d['revised_count']  = 0
                    lat_d['added'] = []      # all of the added/new patterns
                    lat_d['deleted'] = []    # deleted patterns
                    lat_d['verified'] = []   # new/added but already verified.
                    lat_d['default'] = True
                    dim_d['lats'].append(lat_d)  
                    lat_count += 1

        # sort the dimension names
        for c in self.tones:
            c['dimensions'] = sorted(c['dimensions'] , key=itemgetter('lc_name'))

        for c in self.tones:
            for d in c['dimensions']:
                d['lats'] = sorted(d['lats'] , key=itemgetter('lc_name'))

        self.clust_count = clust_count
        self.dim_count   = dim_count
        self.lat_count   = lat_count

    # --------------------------------------------------            
    # Get methods
    # DSDict.getStats()
    # --------------------------------------------------                

    def getStats(self):
        stats = dict()
        stats['cluster_count'] = len(self.tones)
        stats['dimension_count'] = 0
        stats['lat_count'] = 0

        for c in self.tones:
            stats['dimension_count'] += len(c['dimensions'])
            for d in c['dimensions']:
                stats['lat_count'] += len(d['lats'])

        wd = 150
        if platform.system() == 'Windows':
            wd = 150*1.3

        stats_str  = "<p><b>Default Dictionary</b><table>"
        stats_str += "<tr><td width=\"{}\">Total Clusters:  </td><td align=\"right\">{}</td></tr>".format(wd, stats['cluster_count'])
        stats_str += "<tr><td>Total Dimensions:</td><td align=\"right\">{:,}</td></tr>".format(stats['dimension_count'])
#        stats_str += "<tr><td>Total LATs:      </td><td align=\"right\">{:,}</td></tr>".format(stats['lat_count'])
        stats_str += "</table></p>"

        return stats_str

    def getClusterAndDimension(self, cluster_name, dim_name):
        for c in self.tones:
            if c['name'] == cluster_name:
                for d in c['dimensions']:
                    if d['name'] == dim_name:
                        return c, d
        return None, None


    def getCluster(self, cluster_name):
        for c in self.tones:
            if c['name'] == cluster_name:
                return c

        return None

    def getDimension(self, dim_name):
        for c in self.tones:
            for d in c['dimensions']:
                if d['name'] == dim_name:
                    return d
        return None

    def getLAT(self, lat_name):
        for c in self.tones:
            for d in c['dimensions']:
                for lat in d['lats']:
                    if lat['name'] == lat_name:
                        return lat
        return None

    def findCategories(self, pattern, bDims, bLATs):
        matched_dims   = list()
        matched_lats   = list()
        
        for c in self.tones:
            for d in c['dimensions']:

                if bDims:
                    res = re.search(pattern, d['name'])
                    if res:
                        matched_dims.append((c['name'], d['name']))

                if bLATs == False:
                    continue

                for lat in d['lats']:
                    res = re.search(pattern, lat['name'])
                    if res:
                        matched_lats.append((c['name'], d['name'], lat['name']))

        return matched_dims, matched_lats

    def getAddedLATPatterns(self, lat_name):
        lat = self.getLAT(lat_name)
        return lat['added']

    def getDeletedLATPatterns(self, lat_name):
        lat = self.getLAT(lat_name)
        return lat['deleted']

    def getLATNames(self, cluster_name, dimension_name):

        res = list()
        if dimension_name != None:
            for c in self.tones:                   # for each cluster
                if c['name'] == cluster_name:      # if custer_name matches
                    for d in c['dimensions']:      # for each dimension
                        if d['name'] == dimension_name:  # 
                            for lat in d['lats']:
                                res.append(lat['name'])
        else:
            for c in self.tones:                   # for each cluster
                if c['name'] == cluster_name:      # if custer_name matches
                    for d in c['dimensions']:      # for each dimension
                        for lat in d['lats']:
                            res.append(lat['name'])

        return res

    def getTones(self):
        return self.tones

    def getInfo(self):
        return self.info

    def getUnusedDimensions(self, list_of_dims):
        res = list()
        for c in self.tones:
            for d in c['dimensions']:
                if d['name'] not in list_of_dims:
                    res.append(copy.deepcopy(d))
        return res

    def getClusterNames(self):
        res = list()
        for c in self.tones:
            if c['name'] not in ignored_clusters:
                res.append(c['name'])
        return res

    def getDimensionNames(self, cluster_name):
        res = list()
        for c in self.tones:
            if c['name'] == cluster_name:
                for d in c['dimensions']:
                    res.append(d['name'])
                break   
        return res

    def getClusterCount(self):
        return self.clust_count

    def getDimensionCount(self):
        return self.dim_count

    def getLATCount(self):
        return self.lat_count

    # --------------------------------------------------            
    # Untested function
    # --------------------------------------------------            
    def generateHelpDoc(self, savepath):

        dict_name = os.path.basename(self.directory)

        docx = Document()
        styles = docx.styles

        style = styles['Title']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)

        style = styles['Heading 1']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        style.paragraph_format.keep_with_next = True

        style = styles['Normal']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        style.paragraph_format.keep_together = True
 
        p = docx.add_paragraph(dict_name)
        p.style = 'Title'
        for c in self.tones:
            if c['name'] == 'Other':
                continue
            if len(c['dimensions']) == 0:
                self.controller.showWarningDialog("Warning", "{} is not included becasue it has no dimensions.".format(c['name']))
            else:
                p = docx.add_paragraph(c['label'])
                p.style = 'Heading 1'
                p = docx.add_paragraph(c['help'])
                p.style = 'Normal'

        docx.save(savepath)

    def getDirectory(self):
        return self.directory

    def getDATPath(self):
        return self.dat_fpath

    def getVersion(self):
        if self.directory is not None:
            return self.info['version']
        else:
            return "n/a"

    def getBaseDictName(self):
        if self.directory is not None and self.is_custom_dict:
            return self.info.get('base_dict', None)
        else:
            return None

    def getName(self):
        if self.directory is not None:
            return self.info['name']
            # return os.path.basename(self.directory)
        else:
            return os.path.basename(self.dat_fpath)

    def renameDictDirectory(self, new_name):
        if self.directory is not None:
            old_path = self.directory
            new_path = os.path.join(os.path.dirname(self.directory), new_name)

            if os.path.isdir(new_path):
                # The new_path exists. Something sent wrong...
                return False
            else:
                shutil.move(old_path, new_path)
                self.directory = new_path
                return True

    def getPatternsFromTagger(self, lat_name):
        if self.tagger is not None:
            patterns = list(self.tagger.getPatterns(lat_name))
            return patterns

    def getPatterns(self, lat_name):
        """
        Given the name of an LAT, return the latest set of patterns in a single string.
        """

        lat = self.getLAT(lat_name)

        if lat['original_patterns'] != None:
            # This LAT has been opened already, and may have been edited.
            if lat['revised_patterns'] != None:
                # This LAT has been edited.
                return lat['revised_patterns'], lat['original_patterns']
            else:
                return None, lat['original_patterns']
        else:
            # This LAT has never been opened by the user. So, we'll need to 
            # read the patterns from the file.
            path = os.path.join(self.directory, "{}.txt".format(lat_name))
            count = 0

            if os.path.exists(path):
                try:
                    with open(path, encoding='utf-8') as fin:
                        patterns = fin.read()
                except UnicodeDecodeError:
                    patterns = utils.repair_encoding_problems(path)
                    if patterns == None:
                        self.controller.showWarningDialog("Error", "{}.txt include characters that are not encoded in non-UTF-8. Press the OK button to quit the application.")
                        sys.exit(app.exec_())
                    else:
                        msg =  "{}: One or more characters were not properly encoded in UTF-8. "
                        msg += "The system has tried to fix them, but there may still be some errors. "
                        msg += "Review the patterns and make sure all the patterns appear correctly "
                        msg += "and save the dictionary."
                        self.controller.showWarningDialog("Warning", msg.format(lat_name))

                # We assume that the saved LAT file includes no duplicates; so
                # we don't need to clean up the patterns.
                # patterns, count = self.cleanUpLATs(patterns, lat_name)
                # we also assume that saved patterns are sorted.

                tmp = patterns.strip().splitlines()        # strip and split into a list of patterns        
                tmp = list(map(lambda x: x.strip(), tmp))  # strip each pattern, just in case.  
                tmp = [p for p in tmp if p]                # remove empty strings
                count = len(tmp)                           # get the total # of patterns
                patterns = '\r\n'.join(tmp)                # cleaned up patterns
            else:
                # This must be a new LAT that hasn't been saved yet. 
                patterns = ""

            lat['original_patterns'] = patterns.strip()
            lat['original_count'] = count

            return None, lat['original_patterns']

    def getLATStats(self, lat_name):
        """
        Goiven the name of an LAT, return the basic stats:
        (1) coutt, (2) # of added patterns, (3) # of deleted patterns.
        """

        # We assume that 'updateDiff()' has been called before calling this function. i
        # setRevisedPatterns().
        lat_d = self.getLAT(lat_name)
        if lat_d != None:
            if lat_d['revised_count'] > 0:
                # count, added, deleted = diff(lat_d['original_patterns'], lat_d['revised_patterns'])
                count = self.getLATPatternCount(lat_name)
                return count, len(lat_d['added']), len(lat_d['deleted'])
            else:
                count = self.getLATPatternCount(lat_name)
                return count, 0, 0

        return -1, -1, -1

    def getLATPatternCount(self, lat_name):
        lat_d = self.getLAT(lat_name)
        if lat_d != None:
            if lat_d['revised_count'] > 0:
                return lat_d['revised_count']
            else:
                return lat_d['original_count']
        return -1

    def getTagger(self):
        return self.tagger;

    def setTagger(self, tagger):
        self.tagger = tagger
        self.setupTones()
        self.loadInfo()
        self.loadHelp()

    # setup tones from compact dict.
    def setupTones(self):
        """
        Read _tones.txt to create a dictionary hierarchy;
        """

        global untitled_category_count

        cluster_d = None
        dim_d     = None
        lat_d     = None

        clusters = self.tagger.getClustDimTree();

        for clust_name, dims in clusters.items():
            if re.search('untitled[0-9]+$', clust_name.lower()) is not None:
                untitled_category_count+=1
                clust_name = "Untitled{}".format(untitled_category_count)

            cluster_d = dict()
            cluster_d['type'] = 'CLUSTER'
            cluster_d['name'] = clust_name
            cluster_d['lc_name'] = clust_name.lower()
            cluster_d['dimensions'] = list()
            self.tones.append(cluster_d)

            for dim_name in dims:
                if re.search('untitled[0-9]+$', dim_name.lower()) is not None:
                    untitled_category_count+=1
                    dim_name = "Untitled{}".format(untitled_category_count)

                dim_d = dict()
                dim_d['type'] = 'DIMENSION'
                dim_d['name'] = dim_name
                dim_d['lc_name'] = dim_name.lower()
                dim_d['lats'] = list()
                cluster_d['dimensions'].append(dim_d)

                lats = self.tagger.getLats(dim_name)

                for lat_name in lats:
                    if re.search('untitled[0-9]+$', lat_name.lower()) is not None:
                        untitled_category_count+=1
                        lat_name = "Untitled{}".format(untitled_category_count)
                    
                    lat_d = dict()
                    lat_d['type'] = 'LAT'
                    lat_d['name'] = lat_name
                    lat_d['lc_name'] = lat_name.lower()
                    lat_d['original_patterns'] = None
                    lat_d['revised_patterns']  = None
                    lat_d['original_count'] = 0
                    lat_d['revised_count']  = 0
                    lat_d['added'] = []      # all of the added/new patterns
                    lat_d['deleted'] = []    # deleted patterns
                    lat_d['verified'] = []   # new/added but already verified.
                    lat_d['is_new'] = False
                    dim_d['lats'].append(lat_d)  
                
                    self.lats[lat_name] = (dim_name, clust_name)

        # sort the dimension names
        for c in self.tones:
            c['dimensions'] = sorted(c['dimensions'] , key=itemgetter('lc_name'))

        for c in self.tones:
            for d in c['dimensions']:
                d['lats'] = sorted(d['lats'] , key=itemgetter('lc_name'))

    def isLoaded(self):
        if self.tagger is not None:
            return True
        else:
            return False

    def isIncluded(self, search_pattern):
        if self.tagger != None:
            return self.tagger.isIncluded(search_pattern)

    def findAndReplaceLATNames(self, cluster_name, dimension_name, find_pattern, replace_pattern):
        for lat_name in self.getLATNames(cluster_name, dimension_name):
            new_name = lat_name.replace(find_pattern, replace_pattern)
            self.updateLATName(lat_name, new_name)

    def findAndReplaceDimensionNames(self, dimension_names, find_pattern, replace_pattern, bname, blabel):

        num_changes = len(dimension_names)

        if num_changes == 0:
            return
        new_dim_names = []
        for c in self.tones:
            for d in c['dimensions']:
                if d['name'] in dimension_names:

                    if bname:
                        d['name'] = d['name'].replace(find_pattern, replace_pattern)
                        d['lc_name'] = d['name'].lower()
                        new_dim_names.append(d['name'])
                    if blabel:
                        d['label'] = d['label'].replace(find_pattern, replace_pattern)

                    num_changes -= 1

                if num_changes == 0:
                    break;

        for c in self.tones:
            c['dimensions'] = sorted(c['dimensions'] , key=itemgetter('lc_name'))

        self.tones_edited = True
        return new_dim_names

    def findLATs(self, search_pattern, exact_match=False):
        res = list()
        if self.tagger is not None:
            matches = self.tagger.findLats(search_pattern, exact_match)
            for lat_name, patterns in matches.items():
                dim_name, clust_name = self.getDimensionAndCluster(lat_name)                
                for pattern in patterns:
                    res.append( ( (clust_name, dim_name, lat_name), pattern) )
        return res

    def findSuperPatterns(self, search_pattern, lat_name):
        if self.tagger != None:
            return self.tagger.findSuperPatterns(search_pattern, lat_name)

    # --------------------------------------------------            
    # Query functions
    # --------------------------------------------------            
    def existingCategoryName(self, name):
        for c in self.tones:
            if c['name']== name:
                return True

            for d in c['dimensions']:
                if d['name'] == name:
                    return True

                for lat in d['lats']:
                    if lat['name'] == name:
                        return True
        return False

    def isUntitledCluster(self, cluster_name):
        n = 0
        res = re.search('(?<=Untitled)[1-9]+', cluster_name)
        if res != None:
            n = int(res.group())
        return n

    # --------------------------------------------------            
    # Methods for modifying the dictionary
    # --------------------------------------------------            

    def addNewCluster(self, index):
        global untitled_category_count

        untitled_category_count += 1

        cluster_d = dict()
        cluster_d['type']    = 'CLUSTER'
        cluster_d['name']    = 'Untitled{}'.format(untitled_category_count)
        cluster_d['lc_name'] = cluster_d['name'].lower()
        cluster_d['label']   = 'n/a'
        cluster_d['help']    = "n/a"
        cluster_d['dimensions'] = list()
        cluster_d['default'] = False

        if index >= 0:
            self.tones.insert(index, cluster_d)
        else:
            self.tones.append(cluster_d)

        self.tones_edited = True
        return cluster_d['name']

    def addNewDimension(self, cluster_name):
        global untitled_category_count
        untitled_category_count += 1

        # 1. Find the parent cluster
        parent_cluster = self.getCluster(cluster_name)

        if parent_cluster != None:
            # 2. Create a new dimension object

            dim_d = dict()
            dim_d['type']    = 'DIMENSION'
            dim_d['name']    = 'Untitled{}'.format(untitled_category_count)
            dim_d['lc_name'] = dim_d['name'].lower()
            dim_d['label']   = 'n/a'
            dim_d['help']    = "n/a"
            dim_d['default'] = False
            dim_d['lats']    = list()

             # 3. Add the new dimension object to the list of dimensions
            parent_cluster['dimensions'].append(dim_d)

            parent_cluster['dimensions'] = sorted(parent_cluster['dimensions'] , key=itemgetter('lc_name'))

            self.tones_edited = True
            return dim_d['name']
        else:
            # print("ERROR!! dimension = {}".format(dimension))
            return "ERROR"

    def addNewLAT(self, cluster_name, dimension_name):
        global untitled_category_count
        untitled_category_count += 1

        _, dimension = self.getClusterAndDimension(cluster_name, dimension_name)

        lat_d = dict()
        lat_d['type']    = 'LAT'
        lat_d['name']    = 'Untitled{}'.format(untitled_category_count)
        lat_d['lc_name'] = lat_d['name'].lower()
        lat_d['original_patterns'] = None
        lat_d['revised_patterns']  = None
        lat_d['original_count']    = 0
        lat_d['revised_count']     = 0
        lat_d['added']    = []      # all of the added/new patterns
        lat_d['deleted']  = []      # deleted patterns
        lat_d['verified'] = []      # new/added but already verified.
        lat_d['is_new']   = True

        if dimension != None:
            dimension['lats'].append(lat_d)
            self.tones_edited = True
            return lat_d['name']
        else:
            # print("ERROR!! dimension = {}".format(dimension))
            return "ERROR"

    def moveLAT(self, lat_name, dest_clust, dest_dim, src_clust, src_dim):
        """
        Move 'lat_name' from src_dim to dest_dim.
        """
        moved_lat = None

        # Let's find the lat dictionaries and collect them.
        for ci in range(len(self.tones)):
            c = self.tones[ci]
            if c['name'] == src_clust:
                for di in range(len(c['dimensions'])):
                    d = c['dimensions'][di]

                    if d['name'] == src_dim:
                        for lati in range(len(d['lats'])):
                            lat = d['lats'][lati]

                            if lat['name'] == lat_name:
                                moved_lat = copy.deepcopy(lat)
                                d['lats'].remove(lat)
                                break

        if moved_lat is None:
            self.controller.showWarningDialog("ERROR", "{} > {} > {} does not exist.".format(src_clust, src_dim, lat_name))
            return

        for ci in range(len(self.tones)):
            c = self.tones[ci]
            if c['name'] == dest_clust:
                for di in range(len(c['dimensions'])):
                    d = c['dimensions'][di]
                    if d['name'] == dest_dim:
                        d['lats'].append(moved_lat)
                        self.tones_edited = True
                        break

    def deleteLAT(self, lat_name):
        lat = self.getLAT(lat_name)

        for ci in range(len(self.tones)):
            c = self.tones[ci]

            for di in range(len(c['dimensions'])):
                d = c['dimensions'][di]

                for lati in range(len(d['lats'])):
                    lat = d['lats'][lati]

                    if lat['name'] == lat_name:
                        d['lats'].remove(lat)
                        c_name = c['name']
                        d_name = d['name']
                        self.tones_edited = True
                        self.deleted_lat_names.append(lat_name)
                        return c_name, d_name


    def updateLATName(self, curr_lat_name, new_lat_name):

        if curr_lat_name == new_lat_name:
            return

        # find the lat with curr_lat_name
        lat = self.getLAT(curr_lat_name)

        # is the new name taken already by another category?
        if self.existingCategoryName(new_lat_name) == True:
            return False

        # the new name is unique, so let's replace the name and the lower case name
        # of the lat.
        lat['name'] = new_lat_name
        lat['lc_name'] = new_lat_name.lower()

        # print("lat['revised_patterns']  = ", lat['revised_patterns'])
        # print("lat['original_patterns'] = ", lat['original_patterns'])

        if lat['revised_patterns'] is None and lat['original_patterns'] is None:
            # Both revised_patterns and original_patterns are None. This means that the user has
            # not clicked the LAT to look at the patterns yet. The user must have global replaced
            # the file names. We'll just rename the files.
            curr_path = os.path.join(self.directory, "{}.txt".format(curr_lat_name)) 
            new_path  = os.path.join(self.directory, "{}.txt".format(new_lat_name)) 
            if os.path.exists(curr_path):
                os.rename(curr_path, new_path)
                
            self.tones_edited = True
            return            

        elif lat['revised_patterns'] is None:
            # if lat['revised_patterns'] is None, we'll copy lat['original_patterns']
            # to it, so that this lat will be saved when the dictionary is saved.
            lat['revised_patterns'] = lat['original_patterns']
            lat['original_patterns'] = ""
            
        elif lat['revised_patterns'] is not None:
            # if lat['revised_patterns'] is not None, this lat has been modified already.
            # So, we don't change lat['revised_patterns']. We just assign lat['original_patterns']
            # an empyt string, so that all the patterns in lat['revised_patterns'] are 
            # considered new entries.
            lat['original_patterns'] = ""

        self.tones_edited = True

        # Remove new_name from the list of deleted lat names,
        # if it is in the list.
        if new_lat_name in self.deleted_lat_names:
            self.deleted_lat_names.remove(new_lat_name)

        # Add the current (old) lat_name to the list of deleted lat names.
        self.deleted_lat_names.append(curr_lat_name)

        return True

    def deleteCategory(self, cluster_name, dimension_name):
        global untitled_category_count

        is_default = False

        if dimension_name == None:
            help_item = self.getHelpItem('CLUSTER', cluster_name)
        else:
            help_item = self.getHelpItem('DIMENSION', dimension_name)

        if help_item != None and help_item in self.help:
            self.help.remove(help_item)

        removed_category = None
        for c  in self.tones:
            if c['name'] == cluster_name and dimension_name == None:
                self.tones.remove(c)
                n = self.isUntitledCluster(cluster_name)
                if n > 0:
                    untitled_category_count -= 1
                removed_category = c
            else:
                for d in c['dimensions']:
                    if c['name'] == cluster_name and d['name'] == dimension_name:
                       c['dimensions'].remove(d)
                       removed_category = d

        self.tones_edited = True
        return removed_category

    def moveCluster(self, cluster_name, direction):
        for index in range(len(self.tones)):
            c = self.tones[index]
            if c['name'] == cluster_name:
                self.tones.insert(index+direction, self.tones.pop(index))        
                self.tones_edited = True
                break

    def moveDimension(self, src_clust, src_dim, dest_clust):
        # Let's find the dim
        moved_dim = None
        for ci in range(len(self.tones)):
            c = self.tones[ci]

            if c['name'] == src_clust:
                for di in range(len(c['dimensions'])):
                    d = c['dimensions'][di]

                    if d['name'] == src_dim:
                        moved_dim = copy.deepcopy(d)
                        c['dimensions'].remove(moved_dim)
                        break

        if moved_dim is None:
            self.controller.showWarningDialog("ERROR", "{} > {} does not exist.".format(src_clust, src_dim))
            return

        for ci in range(len(self.tones)):
            c = self.tones[ci]
            if c['name'] == dest_clust:
                c['dimensions'].append(moved_dim)
                c['dimensions'] = sorted(c['dimensions'] , key=itemgetter('lc_name'))
                break

        self.tones_edited = True

    def getDimensionAndCluster(self, lat_name):   # bad naming. change it.
        clust = None
        dim = None

        if lat_name == "":
            return None, None

        elif self.lats:
            # return self.lats[lat_name]
            # if the lat_name does not exist, it returns empty strings.
            return self.lats.get(lat_name, ("",""))

        elif self.tagger:
            i = 0
            for cat in self.tagger.getDimAndClust(lat_name):
                if i == 0:
                    dim = cat
                elif i == 1:
                    clust = cat
                i += 1

        if dim == "" or clust == "":
            dim == None
            clust == None

        return dim, clust


    # ----------------------------------------
    # Help related
    # ----------------------------------------

    def updateClusterHelpContent(self, old_name, new_name, label, content):
        cluster = self.getCluster(old_name)

        if cluster == None:
            return False
        else:

            if old_name != new_name and self.existingCategoryName(new_name) == True:
                return False

            if cluster['name'] != new_name:
                cluster['name']   = new_name
                cluster['lc_name'] = new_name.lower()
                self.tones_edited = True

            if cluster['label'] != label:
                cluster['label'] = label
                self.tones_edited = True

            if cluster['help'] != content:
                cluster['help'] = content
                self.tones_edited = True

        return self.tones_edited

    def updateDimensionHelpContent(self, old_cluster, old_dim, new_name, label, content):
        dim = self.getDimension(old_dim)
        if dim == None:
            # this is an error.
            # print("coudn't find dim...")
            return False
        else:
            if old_dim != new_name and self.existingCategoryName(new_name) == True:
                return False

            if dim['name'] != new_name:
                dim['name']  = new_name
                dim['lc_name'] = new_name.lower()
                cluster = self.getCluster(old_cluster)
                cluster['dimensions'] = sorted(cluster['dimensions'] , key=itemgetter('lc_name'))
                self.tones_edited = True

            if dim['label'] != label:
                dim['label'] = label
                self.tones_edited = True

            if dim['help'] != content:
                dim['help']  = content
                self.tones_edited = True

        return self.tones_edited

    def getHelpItem(self, cat_type, name):

        for h in self.help:
            if h['type'] == cat_type and h['name'] == name:
                return h

        return None
 
    def loadHelp(self):

        path = os.path.join(self.directory, "_help.txt")
        if os.path.exists(path) != True:
            self.controller.showWarningDialog("Warning", "{} does not exist.".format(path))
            return

        with open(path) as help_file:
            prev_line = ""
            text = help_file.read()
            lines = text.split("\n")
            lines = [l for l in lines if l != '']
            i = 0

            while i < len(lines):
                line = lines[i]

                if i % 3 == 0:
                    if not (line.startswith("CLUSTER:") or line.startswith("DIMENSION:")):
                        self.controller.showWarningDialog("ERROR!",
                            "Found an extra line break around {} and {} (line {}) in _help.txt".format(prev_line, line, i))
                        return

                if line.startswith("CLUSTER:"):
                    cluster_name = line[9:].strip()
                    cluster_d = dict()
                    cluster_d['type']  = 'CLUSTER'
                    cluster_d['name']  = cluster_name
                    cluster_d['label'] = lines[i+1]
                    cluster_d['help']  = lines[i+2]
                    self.help.append(cluster_d)
                    i+=2
                    
                elif line.startswith("DIMENSION:"):
                    dim_name = line[11:].strip()
                    dim_d = dict()
                    dim_d['type']  = 'DIMENSION'
                    dim_d['name']  = dim_name
                    dim_d['label'] = lines[i+1]
                    dim_d['help']  = lines[i+2]
                    self.help.append(dim_d)
                    i+=2

                prev_line = line
                i+=1

        # add the help content to the tones hierarchy
        for c in self.tones:
            cluster_help = self.getHelpItem('CLUSTER', c['name'])
            if cluster_help is None:
                continue

            c['label'] = cluster_help['label']
            c['help']  = cluster_help['help']
            for d in c['dimensions']:
                dim_help = self.getHelpItem('DIMENSION', d['name'])
                if dim_help == None:
                    d['label'] = "Error"
                    d['help']  = "Error"
                else:
                    d['label'] = dim_help['label']
                    d['help']  = dim_help['help']

    def loadHelpFromDatFile(self, zipf):

        with zipf.open("_help.txt") as help_file:
            text  = help_file.read()
            text  = text.decode('utf-8')
            lines = text.splitlines()
            lines = [l for l in lines if l != '']

            i = 0
            while i < len(lines):
                line = lines[i]
                if i % 3 == 0:
                    if not (line.startswith("CLUSTER:") or line.startswith("DIMENSION:")):
                        self.controller.showWarningDialog("ERROR!", "Found an extra line break around {} (line {}) in _help.txt".format(line, i))
                        return

                if line.startswith("CLUSTER:"):
                    cluster_name = line[9:].strip()
                    cluster_d = dict()
                    cluster_d['type']  = 'CLUSTER'
                    cluster_d['name']  = cluster_name
                    cluster_d['label'] = lines[i+1]
                    cluster_d['help']  = lines[i+2]
                    self.help.append(cluster_d)
                    i+=2
                elif line.startswith("DIMENSION:"):
                    dim_name = line[11:].strip()
                    dim_d = dict()
                    dim_d['type']  = 'DIMENSION'
                    dim_d['name']  = dim_name
                    dim_d['label'] = lines[i+1]
                    dim_d['help']  = lines[i+2]
                    self.help.append(dim_d)
                    i+=2

                i+=1

        # add the help content to the tones hierarchy
        for c in self.tones:
            cluster_help = self.getHelpItem('CLUSTER', c['name'])
            c['label'] = cluster_help['label']
            c['help']  = cluster_help['help']
            for d in c['dimensions']:
                dim_help = self.getHelpItem('DIMENSION', d['name'])
                if dim_help == None:
                    d['label'] = "Error"
                    d['help']  = "Error"
                else:
                    d['label'] = dim_help['label']
                    d['help']  = dim_help['help']

    def removeLiteralWordclassDuplicates(self, patterns, lat):

        def is_literal(p):
            res = True
            for w in p:
                if w[0] == '!':
                    res = False
                    break

            return res

        def str_to_tuple(s):
            s = s.split()
            return tuple(s)

        # Let's create a list of unverified patterns = added patterns
        unverified_patterns = list()
        for p in lat['added']:
            unverified_patterns.append(p)

        if len(unverified_patterns) == 0:
            return patterns

        plural_s = ""  # default 
        tpatterns = list(map(str_to_tuple, patterns))
        tunverified_patterns = list(map(str_to_tuple, unverified_patterns))

        warning_candidates = list()
        removed = list()

        for pair in itertools.filterfalse(lambda p: p[0]==p[1] or len(p[0])!=len(p[1]),
                                          itertools.product(*[tunverified_patterns, tpatterns])):                                          

            p1 = pair[0]
            p2 = pair[1]

            # Let's find out if p1 and p2 match.
            count = len(p1)
            bMatch = True
            for i in range(count):
                if p1[i] == p2[i]:  # match
                    pass
                elif p1[i][0] != '!' and p2[i][0] != '!':   # check if both are literals?
                    bMatch = False                          # since they are different, reject.
                    break;

                elif p1[i][0] == '!' and len(p1[i])>1 and p2[i][0] != '!':
                    # p1 is a classword
                    # p2 is a literal
                    words = self.wordclasses.get(p1[i].upper(), None)
                    if words is None:
                        bMatch = False
                        break
                    elif p2[i] not in words:
                        bMatch = False
                        break
                elif p1[i][0] != '!' and p2[i][0] == '!' and len(p2[i])>1:
                    # p1 is a literal
                    # p2 is a classword
                    words = self.wordclasses.get(p2[i].upper(), None)
                    if words is None:
                        bMatch = False
                        break
                    elif p1[i] not in words:
                        bMatch = False
                        break
                elif p1[i][0] == '!' and len(p1[i])>1 and p2[i][0] == '!' and len(p2[i])>1:
                    # both are classword
                    words1 = self.wordclasses.get(p1[i].upper(), None)
                    words2 = self.wordclasses.get(p2[i].upper(), None)

                    # Slightly optimized...
                    if words1 is None or words2 is None:
                        # one of them is a undefined wordclass. so, no match.
                        bMatch = False                        
                        break
                    elif len(list(set(words1) & set(words2)))==0:  # no overlap
                        bMatch = False
                        break
                    
                else:
                    bMatch = False
                    break

            if bMatch:
                # Add literal patterns to the list of auto erase candidates.
                if is_literal(p1) is True and is_literal(p2) is not True:
                    p = ' '.join(p1)
                    try:
                        patterns.remove(p)
                        removed.append(p)
                        if p in lat['added']:
                            lat['added'].remove(p)                        
                    except:
                        pass

                elif is_literal(p1) is not True and is_literal(p2) is True:
                    p = ' '.join(p2)
                    try:
                        patterns.remove(p)
                        removed.append(p)
                        if p in lat['added']:
                            lat['added'].remove(p)
                    except:
                        pass

                else:
                    warning_candidates.append(pair)

        tmp = list()
        for pair in warning_candidates:
            tmp.append(tuple(sorted(pair)))
        warning_candidates = list(set(tmp))

        if len(removed)>0 or len(warning_candidates)>0:
            warning_msg = ""

            if len(removed)>1:
                warning_msg += "The following patterns have been removed.\n\n"
            elif len(removed)>0:
                warning_msg += "The following pattern has been removed.\n\n"

            for p in removed:
                warning_msg += "{}\n".format(p)

            warning_msg += "\n"

            if len(warning_candidates)>1:
                warning_msg += "\nThe following pairs generate duplicates, which must be fixed manually.\n\n"
            elif len(warning_candidates)>0:
                warning_msg += "\nThe following pair generates one ore more duplicates, and must be fixed manually, or you may decide to keep them.\n\n"

            duplicates = ""
            for wpair in warning_candidates:
                wp1 = ' '.join(wpair[0])
                wp2 = ' '.join(wpair[1])

                # warning_msg += "\"{}\"\n\"{}\"\n\n".format(wp1, wp2)
                duplicates += "\"{}\"\n\"{}\"\n\n".format(wp1, wp2)

            if (len(removed) + len(warning_candidates))>1:
                plural_s = "s"

            self.controller.showBigWarningDialog("Duplicate{} Found in {}".format(plural_s, 
                                                                                    lat['name']),
                                                                                    warning_msg,
                                                                                    duplicates)

        # find undefined wordclasses
        undefined_wordclasses = list()
        res = list()
        for p in patterns:
           l = [w for w in p.split() if len(w)>1 and w[0]=='!']
           if l:
                res.extend(l)

        for wc in list(set(res)):
            if wc not in self.wordclasses:
                undefined_wordclasses.append(wc)

        if len(undefined_wordclasses)>0:
            if len(undefined_wordclasses)>1:
                plural_s = "es"

            undefined_wordclasses = list(set(undefined_wordclasses))
            self.controller.showWarningDialog("Undefined Word Class{} in {}".format(plural_s, lat['name']), 
                                     "{}".format('\n'.join(undefined_wordclasses)))

        return patterns

    def updateDiff(self, lat, original, revised):
        """
        lat:        dict
        original:   list of patterns
        revised:    list of patterns
        """

        if original == "":
            # New LAT or renamed LAT. revised patterns are all new.
            pass
        else:
            #unified_diff
            original.sort()
            revised.sort()
            diff = difflib.unified_diff(original,revised,lineterm="")
            count = 0
            added = list()
            deleted = list()

            # res = ""
            for w in diff:                
                w = w.strip()

                if w.startswith("---") or w.startswith("+++") or w.startswith("@@"):
                    continue                

                 # Skip if there is an unknown line, which usually is an empty line.
                if w:
                    if w[0] == '?' or len(w)==1:
                        continue

                    if w[0]=='+':
                        added.append(w[1:])
                    elif w[0]=='-':
                        deleted.append(w[1:])

                    count+=1

        lat['added']   = added
        lat['deleted'] = deleted
        lat['revised_count'] = count

    def setRevisedPatterns(self, lat_name, revised_patterns):
        """
        Update the lat dictionary for 'lat_name'. If no changes are necessary,
        return None. Otherwise, return the latest count of the patterns in the LAT.

        lat_name:   The name of LAT that ise being revised
        patterns:   The patterns from the patterns editor, including the old and the new patterns,
                    deleted patterns are not inclded.
        """

        def adjust_cases(p):
            """
            p: pattern
            """
            words = p.strip().split()
            res = []
            for w in words:
                if w[0] == "!":
                    res.append(w.upper())
                else:
                    res.append(w.lower())
            return ' '.join(res)

        start = time()

        lat = self.getLAT(lat_name)                    # let's find the lat object (dict)

        if lat['original_patterns'] == None:           # assign "", if 'original_patterns' = None,
            lat['original_patterns'] = ""              # it must be a new LAT, or renamed LAT.

        tmp = revised_patterns.strip().splitlines()
#       tmp = list(map(lambda x: x.strip(), tmp))       # strip each pattern, just in case.  
        tmp = list(map(lambda p: adjust_cases(p), tmp)) # strip each pattern and adjust cases. 
        tmp = [p for p in tmp if p]                     # remove empty strings
        tmp = list(set(tmp))                            # remove literal duplicates

        revised_patterns_list = tmp

        original_patterns_list = lat['original_patterns'].splitlines()

        self.updateDiff(lat, original_patterns_list, revised_patterns_list)

        revised_patterns_list = self.removeLiteralWordclassDuplicates(revised_patterns_list, lat)

        revised_patterns_list.sort()
        sorted_patterns = '\r\n'.join(revised_patterns_list)  # reconstruct the newline separated patterns

        if lat['revised_patterns'] is None and lat['original_patterns'] != sorted_patterns:
            # if the patterns have not been edited (updated) at all yet, AND
            # the new set of patterns minus duplicates is different from the saved patterns,
            # set the revised pattern to the new set of patterns.
            lat['revised_patterns'] = sorted_patterns.strip()
            lat['revised_count']    = len(revised_patterns_list)

            return True
        elif lat['revised_patterns'] is not None and lat['revised_patterns'] != sorted_patterns:
            # if the patterns have already been edited at lines once (i.e., updated), AND
            # if new set of patterns minus duplicates is different from the revised (and updated) patterns,
            # update the revised patterns.
            lat['revised_patterns'] = sorted_patterns.strip()
            lat['revised_count']    = len(revised_patterns_list)
            return True
        else:
            return False

    def removePattern(self, pattern_to_remove, lat_name):

        lat_d = self.getLAT(lat_name)

        if lat_d != None:
            if lat_d['original_patterns'] == None:
                # the LAT hasn't been loaded yet.
                self.getPatterns(lat_name)

            if lat_d['revised_patterns'] != None:
                patterns = lat_d['revised_patterns'].splitlines()
                error = True
                for p in patterns:
                    dsp = re.findall(r"[!?\w'-\u2019]+|[" + string.punctuation + "]", p)
                    dsp = ' '.join(dsp).lower()
                    if pattern_to_remove.lower() == dsp:
                        patterns.remove(p)
                        error = False
                        break

            elif lat_d['original_patterns'] != None:
                patterns = lat_d['original_patterns'].splitlines()
                error = True
                for p in patterns:
                    dsp = re.findall(r"[!?\w'-\u2019]+|[" + string.punctuation + "]", p)
                    dsp = ' '.join(dsp).lower()
                    if pattern_to_remove.lower() == dsp:
                        patterns.remove(p)
                        error = False
                        break

            if error:
                self.controller.showWarningDialog("Warning", "\"{}\" cannot be found in {}. It may be a duplicate.".format(pattern_to_remove, lat_name))
            else:
                lat_d['revised_patterns'] = '\r\n'.join(patterns)

    # ----------
    # Methods for Word Classes
    # ----------
    def loadWordClasses(self):
        path = os.path.join(self.directory, "_wordclasses.txt")

        self.wordclasses = dict()

        try:    
            with open(path) as fin:
                text = fin.read()

        except UnicodeDecodeError:
            text = utils.repair_encoding_problems(path)
            if text == None:
                self.controller.showWarningDialog("Error", "{}.txt include characters that are not encoded in non-UTF-8. Press the OK button to quit the application.")
                sys.exit(app.exec_())
            else:
                msg =  "One or more characters were not properly encoded in UTF-8 in _wordclasses.txt. "
                msg += "The system has tried to fix them, but there may still be some errors. "
                msg += "Review the patterns and make sure all the patterns appear correctly "
                msg += "and save the dictionary."
                self.controller.showWarningDialog("Word Class Error", msg)

        lines = text.splitlines()
        curClass = None
        for r in lines:
            r = r.strip()
            l = r.split()
            if len(l) == 1 and curClass != None:
                self.wordclasses[curClass].append(l[0].lower()) 
            elif len(l) == 2:
                curClass = "!" + l[1].upper()
                self.wordclasses[curClass] = list()

        for key in self.wordclasses:
            self.wordclasses[key].sort()

    def loadWordClassesFromDatFile(self, zipf):

        self.wordclasses = dict()

        with zipf.open("_wordclasses.txt") as fin:
            text = fin.read()
            text = text.decode('utf-8')

        lines = text.splitlines()
        curClass = None
        for r in lines:
            r = r.strip()
            l = r.split()
            if len(l) == 1 and curClass != None:
                self.wordclasses[curClass].append(l[0].lower()) 
            elif len(l) == 2:
                curClass = "!" + l[1].upper()
                self.wordclasses[curClass] = list()

        for key in self.wordclasses:
            self.wordclasses[key].sort()

    def saveWordClasses(self, wc_path):
        with open(wc_path, 'w', encoding='utf-8') as _wc_fout:
            words = list(self.wordclasses.keys())
            words.sort()
            for w in words:
                _wc_fout.write("CLASS: {}\r\n".format(w[1:]))
                for item in self.wordclasses[w]:
                    _wc_fout.write("{}\r\n".format(item))
                _wc_fout.write("\r\n")

    def updateWordClasses(self, wordclasses):
        self.wordclasses = wordclasses
        self.wordclasses_edited = True

    def getWordClasses(self):
        return self.wordclasses

    def deleteContent(self):
        shutil.rmtree(self.directory)

    # ----------
    # DS.Dict.save()
    #
    # ----------
    def save(self):
        timestamp = datetime.now().strftime("%Y-%m-%d-%H%M%S").replace(":", "-")

        # create the backup's root dir if it doesn't exist
        backup_root_dir = "{}_backup".format(self.directory)
        if os.path.exists(backup_root_dir) != True:
            os.makedirs(backup_root_dir)

        # create the backup dir. It shold not exist, so we will just create one.
        backup_dir = os.path.join(backup_root_dir, timestamp)
        os.makedirs(backup_dir)

        # log file
        log_path = os.path.join(backup_dir, "_log.txt")
        log_str  = ""

        if self.tones_edited:
            # setup the paths we need
            tones_path         = os.path.join(self.directory, '_tones.txt')
            help_path          = os.path.join(self.directory, '_help.txt')
            export_filter_path = os.path.join(self.directory, "_export_filter.txt")
            # paths for backup files
            tones_backup_path         = os.path.join(backup_dir, '_tones.txt')
            help_backup_path          = os.path.join(backup_dir, '_help.txt')
            export_filter_backup_path = os.path.join(backup_dir, "_export_filter.txt")

            # if _tones.txt and _help.txt exist (i.e., no errors), archive them.
            if os.path.exists(tones_path):
                shutil.copyfile(tones_path, tones_backup_path)
            if os.path.exists(help_path):
                shutil.copyfile(help_path, help_backup_path)
            if os.path.exists(export_filter_path):
                shutil.copyfile(export_filter_path, export_filter_backup_path)

            # Update _tones.txt with the new content
            with open(tones_path, 'w', encoding="utf-8") as _tones_fout:
                for c in self.tones:
                    _tones_fout.write("CLUSTER: {}\r\n".format(c['name']))
                    _tones_fout.write("\r\n")
                    for d in c['dimensions']:
                        _tones_fout.write("DIMENSION: {}\r\n".format(d['name']))
                        for lat in d['lats']:
                            _tones_fout.write("LAT: {}\r\n".format(lat['name']))
                        _tones_fout.write("\r\n")
                    _tones_fout.write("\r\n")

            with open(export_filter_path, 'w', encoding="utf-8") as _export_filter_fout:
                for c in self.tones:
                    _export_filter_fout.write("CLUSTER: {}\r\n".format(c['name']))

            # Update _help.txt with the new content
            with open(help_path, 'w', encoding="utf-8") as _help_fout:
                for c in self.tones:
                    _help_fout.write("CLUSTER: {}\r\n".format(c['name']))
                    _help_fout.write("{}\r\n".format(c['label']))
                    
                    clust_help_content = c['help'].strip()
                    if clust_help_content:
                        _help_fout.write("{}\r\n".format(clust_help_content))
                    else:
                        _help_fout.write("n/a\r\n")
                    _help_fout.write("\r\n")

                    for d in c['dimensions']:
                        _help_fout.write("DIMENSION: {}\r\n".format(d['name']))
                        _help_fout.write("{}\r\n".format(d['label']))
                        dim_help_content = d['help']
                        if dim_help_content:
                            _help_fout.write("{}\r\n".format(dim_help_content))
                        else:
                            _help_fout.write("n/a\r\n")
                        _help_fout.write("\r\n")                    
                _help_fout.write("\r\n")

            log_str += "_tones.txt and/or _help.txt have been updated.\n"
            self.tones_edited = False
            
        else:
            log_str += "No changes have been made to _tones.txt and _help.txt.\n"

        if self.wordclasses_edited:
            wc_path        = os.path.join(self.directory, '_wordclasses.txt')
            wc_backup_path = os.path.join(backup_dir, '_wordclasses.txt')
            if os.path.exists(wc_path):
                shutil.copyfile(wc_path, wc_backup_path)

            with open(wc_path, 'w', encoding='utf-8') as _wc_fout:
                words = list(self.wordclasses.keys())
                words.sort()
                for w in words:
                    _wc_fout.write("CLASS: {}\r\n".format(w[1:]))
                    for item in self.wordclasses[w]:
                        _wc_fout.write("{}\r\n".format(item))
                    _wc_fout.write("\r\n")

            log_str += "_wordclasses.txt has been updated."
        else:
            log_str += "No changes have been made to _wordclasses.txt."

        log_str += "\n"
        log_str += "Revised LATs:\n"
        count_revised_lats = 0

        for lat_name in self.deleted_lat_names:
            lat_d = self.getLAT(lat_name)
            # Just being conservative, we make sure that the lat file we are
            # deleting is not currently used. It's an expensive operation, so 
            # we may want to skip this...
            if lat_d == None:
                lat_path    = os.path.join(self.directory, '{}.txt'.format(lat_name))
                backup_path = os.path.join(backup_dir, '{}.txt'.format(lat_name))

                if os.path.exists(lat_path):
                    shutil.copyfile(lat_path, backup_path)  # make a backup of the deleted lat file.
                    os.remove(lat_path) 
                    log_str += "{} - Deleted\n".format(lat_name)

        self.deleted_lat_names = list()   # empty the deleted lat names list

        for c in self.tones:
            for d in c['dimensions']:
                for lat in d['lats']:  # for each LAT

                    if lat['revised_patterns'] == '':
                        return (False, lat['name'])

                    if lat['is_new'] and lat['revised_patterns'] == None:
                        # This handles a case where a new LAT is created yet it is empty.
                        return (False, lat['name'])

                    if lat['revised_patterns'] != None:
                        lat_path = os.path.join(self.directory, '{}.txt'.format(lat['name']))
                        backup_path = os.path.join(backup_dir, 
                                                   '{}.txt'.format(lat['name']))

                        if os.path.exists(lat_path):
                            # if the LAT file exist, make a backup
                            shutil.copyfile(lat_path, backup_path)
                            log_str += "{} - Edited\n".format(lat['name'])

                            os.remove(lat_path)
                        else:
                            log_str += "{} - Created\n".format(lat['name'])
                        
                        with open(lat_path, 'w', encoding='utf-8') as fout:
                            fout.write(lat['revised_patterns'])

                        count_revised_lats += 1

                        # Commented because we are not saving unverified patterns. (Jan 29, 2020)
                        # if len(lat['added'])>0:
                        #     added = adjustCases(lat['added'])
                        #     if lat['name'] in unverified_patterns:
                        #         unverified_patterns[lat['name']].extend(added)
                        #     else:
                        #         unverified_patterns[lat['name']] = added

                        # if lat['name'] in unverified_patterns:
                        #     deleted = adjustCases(lat['deleted'])
                        #     unverified_patterns[lat['name']] = list(set(unverified_patterns[lat['name']]) - set(deleted))

                        # reset the lat fields.
                        lat['added']   = []
                        lat['deleted'] = []
                        lat['original_count']    = lat['revised_count']
                        lat['revised_count']     = 0
                        lat['original_patterns'] = lat['revised_patterns']
                        lat['revised_patterns']  = None
                        lat['is_new'] = False

                # lat
            # dim

        if count_revised_lats == 0:
            log_str += "\nNo LAT files have been revised."
        else:
            log_str += "\n{} LAT files have been revised.".format(count_revised_lats)

        with open(log_path, 'w', encoding="utf-8") as fout:
            fout.write(log_str)

        return (True, '')

    def getNumRules(self):
        if self.tagger:
            return self.tagger.getNumRules()
        else:
            return -1

    def getNumLATs(self):
        res = 0
        for c in self.tones:
            for d in c['dimensions']:
                res += len(d['lats'])

        return res

    def cancelExport(self):
        self.bExport = False

    def cleanupExportTempFiles(self):
        parent_dir = os.path.dirname(self.directory)
        temp_dir   = os.path.join(parent_dir, 'temp')

        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)

        zip_fpath = self.zip_basename + ".zip"
        dat_fpath = self.zip_basename + ".dat"

        if os.path.exists(zip_fpath):
            os.remove(zip_fpath)

        if os.path.exists(dat_fpath):
            os.remove(dat_fpath)

    # DSCustomDict.load()
    def load(self):
        self.tagger = dstagger.DsTagger()                 # create a tagger object
        self.tagger.readDictionaryFiles(self.directory)   # read the dictionary files
        self.tagger.setDictionaryInfo("na", "na", "na")

    # DSCDict.export()
    def export(self, dict_info, export_info, backup):

        self.bExport = True

        src_dict_dir    = self.directory                       # the src dictionary folder
        dst_dir         = os.path.join(export_info['dst_dir'],
                                       "{} (v.{})".format(dict_info['name'], dict_info['version']))  # release dict

        dst_dict_fpath  = os.path.join(dst_dir, "dict.ser")                        # serialized obj file

        if export_info.get('is_temp', False) == True:
            # We are creating a serialized object for the find LAT features.
            temp_dir  =  src_dict_dir
            backup = False   # if it's a temporary export, we'll always ignore backup.
        else:
            # We are creating a release package. We will make sure that non-utf8 characters are included.
            dict_name       = os.path.basename(self.directory)     # the name of the src folder            
            src_parent_dir  = os.path.dirname(self.directory)      # parent dir of the src dict folder            
            temp_dir        = os.path.join(src_parent_dir, 'temp') # temp folder
            dst_tones_fpath = os.path.join(dst_dir, "_tones.txt")                      # modified _tones.txt
            dst_wordclasses_fpath = os.path.join(dst_dir, "_wordclasses.txt")                # _wordclasses.txt
            backup_dir            = os.path.join(export_info['dst_dir'], "{}_backup".format(dict_name))

            # Let's make sure that the temp folder exists.
            if os.path.exists(temp_dir) != True:
                os.makedirs(temp_dir)

            if os.path.exists(dst_dir) != True:
                os.makedirs(dst_dir)

            src_wordclasses_fpath  = os.path.join(src_dict_dir, "_wordclasses.txt")
            src_help_fpath         = os.path.join(src_dict_dir, "_help.txt")

            temp_tones_fpath       = os.path.join(temp_dir, "_tones.txt")
            temp_wordclasses_fpath = os.path.join(temp_dir, "_wordclasses.txt")

            dst_help_fpath         = os.path.join(dst_dir, "_help.txt")

            # Save the dictionary info.
            info_fpath = os.path.join(dst_dir, "info.json")
            with open(info_fpath, 'w') as info_fout:
                json.dump(dict_info, info_fout, indent=4)

            shutil.copy(src_help_fpath, dst_help_fpath)
            shutil.copy(src_wordclasses_fpath, temp_wordclasses_fpath)

            num_clusters = len(self.tones)

            # The folloing code merges all the LAT files under a dimension into a single file.
            # We are only doing this when we are releasing the dictionary since the end-users do not
            # need to look at the LAT level information.

            with open(temp_tones_fpath, 'w') as tones_fout:
                for c in self.tones:
                    clust_name = c['name']

                    tones_fout.write("CLUSTER: {}\r\n".format(clust_name))

                    for d in c['dimensions']:
                        dim_name = d['name']
                        tones_fout.write("DIMENSION: {}\r\n".format(dim_name))
                        tones_fout.write("LAT: {}_LAT\r\n".format(dim_name))

                        # make a new lat file
                        new_lat_fpath = os.path.join(temp_dir, "{}_LAT.txt".format(dim_name))  
                        with open(new_lat_fpath, 'w', encoding='utf-8') as fout:
                            for lat in d['lats']:

                                # read each lat file
                                lat_fpath = os.path.join(src_dict_dir, "{}.txt".format(lat['name']))
                                with open(lat_fpath, encoding='utf-8-sig') as fin:                                          # merge them into the new lat file
                                    try:
                                        s = fin.read()
                                        fout.write(s)
                                    except:
                                        patterns = utils.remove_non_utf8_patterns(lat_fpath)
                                        fout.write(patterns)
                                fout.write("\r\n")

            if self.bExport != True:
                self.cleanupExportTempFiles()
                return "", ""

            # create a backup of the original dictionary (not the one in the temp dir)
            if backup:
                if os.path.exists(backup_dir) != True:
                    os.makedirs(backup_dir)
                backup_basename = os.path.join(backup_dir, "{} (v.{}) archive".format(dict_info['name'], 
                                                                                      dict_info['version']))
                shutil.make_archive(backup_basename, 'zip', src_dict_dir)

                if self.bExport != True:
                    self.cleanupExportTempFiles()
                    return "", ""

            # Copy _tones.txt and _wordclasses.txt to dst_dir
            shutil.copy(temp_tones_fpath, dst_tones_fpath)
            shutil.copy(temp_wordclasses_fpath, dst_wordclasses_fpath)

        # Serialize the dictionary
        self.tagger = dstagger.DsTagger()             # create a tagger object
        self.tagger.readDictionaryFiles(temp_dir)     # read the dictionary files
        self.tagger.setDictionaryInfo(dict_info['name'],
                                      dict_info['version'],
                                      dict_info['copyright'])
        self.tagger.exportAsDat(dst_dict_fpath)       # seralize the file and save it as a .ser file.
        
        if self.bExport: 
            # The user didn't cancel the process.
            self.cleanupExportTempFiles()
            # self.controller.updateExportProgress(1.0)
            self.bExport = False

            if backup:
                zip_path = backup_basename+".zip"
            else:
                zip_path = '(no backup)'

            return dst_dir, zip_path
        else:
            self.cleanupExportTempFiles()
            return "", ""

    # DSCDict.exportOriginal()
    def exportOriginal(self, dict_info, export_info, backup):

        self.bExport = True

        src_dict_dir    = self.directory                       # the src dictionary folder
        dst_dir         = os.path.join(export_info['dst_dir'],
                                       "{} (v.{})".format(dict_info['name'], dict_info['version']))  # release dict

        dst_dict_fpath  = os.path.join(dst_dir, "dict.ser")                        # serialized obj file

        # We are creating a release package. We will make sure that non-utf8 characters are included.
        dict_name       = os.path.basename(self.directory)     # the name of the src folder            
        src_parent_dir  = os.path.dirname(self.directory)      # parent dir of the src dict folder            
        temp_dir        = os.path.join(src_parent_dir, 'temp') # temp folder
        dst_tones_fpath = os.path.join(dst_dir, "_tones.txt")                      # modified _tones.txt
        dst_wordclasses_fpath = os.path.join(dst_dir, "_wordclasses.txt")                # _wordclasses.txt
        backup_dir            = os.path.join(export_info['dst_dir'], "{}_backup".format(dict_name))

        # Let's make sure that the temp folder exists.
        if os.path.exists(temp_dir) != True:
            os.makedirs(temp_dir)

        if os.path.exists(dst_dir) != True:
            os.makedirs(dst_dir)

        src_tones_fpath        = os.path.join(src_dict_dir, "_tones.txt")
        src_wordclasses_fpath  = os.path.join(src_dict_dir, "_wordclasses.txt")
        src_help_fpath         = os.path.join(src_dict_dir, "_help.txt")

        temp_tones_fpath       = os.path.join(temp_dir, "_tones.txt")
        temp_wordclasses_fpath = os.path.join(temp_dir, "_wordclasses.txt")

        dst_help_fpath         = os.path.join(dst_dir, "_help.txt")

        # Save the dictionary info.
        info_fpath = os.path.join(dst_dir, "info.json")
        with open(info_fpath, 'w') as info_fout:
            json.dump(dict_info, info_fout, indent=4)

        shutil.copy(src_help_fpath,        dst_help_fpath)
        shutil.copy(src_wordclasses_fpath, temp_wordclasses_fpath)
        shutil.copy(src_tones_fpath,       temp_tones_fpath)

        num_clusters = len(self.tones)

        # Make a copy of the LAT files. Make sure that there are not character errors.
        for c in self.tones:
            for d in c['dimensions']:
                for lat in d['lats']:

                    lat_name = lat['name']

                    new_lat_fpath = os.path.join(temp_dir, "{}.txt".format(lat_name))
                    with open(new_lat_fpath, 'w', encoding='utf-8') as fout:

                        # read each lat file
                        src_lat_fpath = os.path.join(src_dict_dir, "{}.txt".format(lat_name))
                        fin = open(src_lat_fpath, encoding='utf-8-sig')

                        try:
                            s = fin.read()
                            fout.write(s)
                        except:
                            fin.close()
                            patterns = utils.remove_non_utf8_patterns(src_lat_fpath)
                            fout.write(patterns)

                        fout.write("\r\n")

        if self.bExport != True:
            self.cleanupExportTempFiles()
            return "", ""

        # create a backup of the original dictionary (not the one in the temp dir)
        if backup:
            if os.path.exists(backup_dir) != True:
                os.makedirs(backup_dir)
            backup_basename = os.path.join(backup_dir, "{} (v.{}) archive".format(dict_info['name'], 
                                                                                  dict_info['version']))
            shutil.make_archive(backup_basename, 'zip', src_dict_dir)

            if self.bExport != True:
                self.cleanupExportTempFiles()
                return "", ""

        # Copy _tones.txt and _wordclasses.txt to dst_dir
        shutil.copy(temp_tones_fpath, dst_tones_fpath)
        shutil.copy(temp_wordclasses_fpath, dst_wordclasses_fpath)

        # Serialize the dictionary
        self.tagger = dstagger.DsTagger()             # create a tagger object
        self.tagger.readDictionaryFiles(temp_dir)     # read the dictionary files
        self.tagger.setDictionaryInfo(dict_info['name'],
                                      dict_info['version'],
                                      dict_info['copyright'])
        self.tagger.exportAsDat(dst_dict_fpath)       # seralize the file and save it as a .ser file.
        
        if self.bExport: 
            # The user didn't cancel the process.
            self.cleanupExportTempFiles()
            # self.controller.updateExportProgress(1.0)
            self.bExport = False

            if backup:
                zip_path = backup_basename+".zip"
            else:
                zip_path = '(no backup)'

            return dst_dir, zip_path
        else:
            self.cleanupExportTempFiles()
            return "", ""

    def createDictHierarchy(self):

        def add_a_path(parent, new_node):
            if parent.children:
                no_match = True
                for child in parent.children:
                    if child.name == new_node.name:
                        if new_node.children:
                            add_a_path(child, new_node.children[0])
                            no_match = False
                            break

                if no_match:
                    new_node.parent = parent

            else:
                if parent.name == new_node.name:
                    if new_node.children:
                        add_a_path(parent, new_node.children[0])
                else:
                    new_node.parent = parent

        # Remove all the files in the folder first.
        # for f in os.listdir(dst_dict_path):
            # fpath = os.path.join(dst_dict_path, f)
            # os.remove(fpath)

        root = None
        dim_labels = list()
        leaf_count  = 0
        list_of_dims = list() # list of lists

        if self.tones is not None:

            # Create a list of dimention path (a list of nodes)
            for c in self.tones:
                clust_label = c['label']
                clust_name  = c['name']

                for d in c['dimensions']:
                    dim_label = d['label']
                    dim_label = dim_label.strip()  # just in case, trim newlines/spaces
                    # list_of_dims.append(dim_label.split("/"))

                    dim_name  = d['name']
                    list_of_dims.append((dim_label.split("/"), clust_name, dim_name))


            # Review all the list of nodes (dimension path),
            # then, if there is a node without a leaf name, add it.
            new_list_of_dims = list()
            for i in range(len(list_of_dims)):

                dim = list_of_dims[i]                    # dim is a tuple
                if i < len(list_of_dims)-1:
                    next_dim = list_of_dims[i+1]

                leaf = dim[-1]                                         # if the last node is "Dimension General" or "Generic"
                if leaf == "Dimension General" or leaf == "Generic":   # remove the last node
                    dim = dim[:-1]

                if len(dim) == 1:                                      # if there is only one node, append "Generic"
                    dim = dim + ["Generic"]

                num_nodes = len(dim[0])                   # number of nodes
                num_next_nodes = len(next_dim[0])         # number of nodes in the next dimension

                clust_name = dim[1]
                dim_name   = dim[2]

                # If the current dimension does not have a leaf name, add "Generic" as its leaf.
                if num_nodes < num_next_nodes and dim[0] == next_dim[0][:len(dim[0])]:
                    new_list_of_dims.append((dim[0] + ["Generic"], clust_name, dim_name))
                else:
                    new_list_of_dims.append(dim)

            new_list_of_dims.sort()
            
            root = Node("root")        
            for nodes in new_list_of_dims:
                clust_root = Node("dictionary")
                dim_labels.append('/'.join(nodes[0]))

                n = clust_root
                cat = ""
                for i in range(len(nodes[0])):
                    cat = nodes[0][i]
                    clust_name = nodes[1]
                    dim_name   = nodes[2]
                    # n = Node(cat, parent=n)
                    n = Node(cat, parent=n, clust=clust_name, dim=dim_name)

                add_a_path(root, clust_root)
                leaf_count += 1

        return root, dim_labels, leaf_count

    def exportDictHierarchy(self, dst_dict_path):

        # def add_a_path(parent, new_node):
        #     if parent.children:
        #         no_match = True
        #         for child in parent.children:
        #             if child.name == new_node.name:
        #                 if new_node.children:
        #                     add_a_path(child, new_node.children[0])
        #                     no_match = False
        #                     break

        #         if no_match:
        #             new_node.parent = parent

        #     else:
        #         if parent.name == new_node.name:
        #             if new_node.children:
        #                 add_a_path(parent, new_node.children[0])
        #         else:
        #             new_node.parent = parent

 
        # # Remove all the files in the folder first.
        for f in os.listdir(dst_dict_path):
            fpath = os.path.join(dst_dict_path, f)
            os.remove(fpath)

        # dim_labels = list()
        # leaf_count  = 0
        # list_of_dims = list() # list of lists

        # if self.tones is not None:

        #     # Create a list of dimention path (a list of nodes)
        #     for c in self.tones:
        #         clust_label = c['label']

        #         for d in c['dimensions']:
        #             dim_label = d['label']
        #             dim_label = dim_label.strip()  # just in case, trim newlines/spaces
        #             list_of_dims.append(dim_label.split("/"))

        #     # Review all the list of nodes (dimension path),
        #     # then, if there is a node without a leaf name, add it.
        #     new_list_of_dims = list()
        #     for i in range(len(list_of_dims)):

        #         dim = list_of_dims[i]
        #         if i < len(list_of_dims)-1:
        #             next_dim = list_of_dims[i+1]


        #         leaf = dim[-1]                                         # if the last node is "Dimension General" or "Generic"
        #         if leaf == "Dimension General" or leaf == "Generic":   # remove the last node
        #             dim = dim[:-1]

        #         if len(dim) == 1:                                      # if there is only one node, append "Generic"
        #             dim = dim + ["Generic"]

        #         num_nodes = len(dim)                      # number of nodes
        #         num_next_nodes = len(next_dim)            # number of nodes in the next dimension

        #         # If the current dimension does not have a leaf name, add "Generic" as its leaf.
        #         if num_nodes < num_next_nodes and dim == next_dim[:len(dim)]:
        #             new_list_of_dims.append(dim + ["Generic"])
        #         else:
        #             new_list_of_dims.append(dim)

        #     new_list_of_dims.sort()
            
        #     root = Node("root")        
        #     for nodes in new_list_of_dims:
        #         clust_root = Node("dictionary")
        #         dim_labels.append('/'.join(nodes))

        #         n = clust_root
        #         cat = ""
        #         for i in range(len(nodes)):
        #             cat = nodes[i]
        #             if i == 0:
        #                 n = Node(cat, parent=n)
        #             else:
        #                 n = Node(cat, parent=n)

        #         add_a_path(root, clust_root)
        #         leaf_count += 1

        root, dim_labels, leaf_count = self.createDictHierarchy()

        if root is None:
            return

        fpath = "{}/_dim_labels.txt".format(dst_dict_path)
        with open(fpath, 'w', errors="ignore") as fout:
            for dim_label in dim_labels:
                csv_line = "{}\n".format(dim_label)
                fout.write(csv_line)

        for child in root.children[0].children:
            clust_name = child.name
            fpath = "{}/{}.txt".format(dst_dict_path, clust_name)
            dim_count = len(child.children)
            with open(fpath, 'w', errors="ignore") as fout:
                fout.write("Top Level Category:    {}\n".format(clust_name))
                fout.write("Total Subcategories:   {}\n\n".format(dim_count))
                count = 1
                if child.children:
                    for top_cat in child.children:
                        fout.write("{}. {}\n".format(count, top_cat.name))
                        count+=1                

                fout.write("\n---------------------------------------------------------------------\n\n")
                fout.write(RenderTree(child, style=AsciiStyle()).by_attr('name'))

        with open("{}/_all.txt".format(dst_dict_path), 'w', errors="ignore") as fout:
            r = root.children[0]    
            fout.write("Total Top Level Categories:   {}\n".format(len(r.children)))
            fout.write("Total Leaves:                 {}\n\n".format(leaf_count))
            count = 1
            for top_cat in r.children:
                fout.write("{:02d}. {}\n".format(count, top_cat.name))
                count+=1

            fout.write("\n---------------------------------------------------------------------\n\n")
            fout.write(RenderTree(r, style=AsciiStyle()).by_attr('name'))    

        return True

##################################################
#
# DSCustomDict
#
##################################################

class DSCustomDict(DSDict):
    def __init__(self, controller, directory, default_dict):
        super(DSCustomDict, self).__init__(controller, directory)

        self.controller = controller
        self.default_dict = default_dict
        self.pronouns = False
        self.is_common = False
        self.setCustomDict(True)

    # ------------------------------
    # DSCustomDict.loadTones()
    # Read _custom_tones.txt to create a dictionary hierarchy,
    # which conform to the custom _tones.txt file.
    # ------------------------------
    def loadTones(self):
        
        cluster_d = None
        dim_d     = None
        lat_d     = None

        list_of_dims = list()

        with open('{}/_custom_tones.txt'.format(self.directory)) as tones_file:
            for line in tones_file:
                # custom cluster
                if(line.startswith('CLUSTER:')):        
                    cluster_name = line[9:].strip()
                    cluster_d = dict()
                    cluster_d['type'] = 'CLUSTER'
                    cluster_d['name'] = cluster_name
                    cluster_d['dimensions'] = list()
                    cluster_d['default'] = False
                    self.tones.append(cluster_d)

                # from the default dict
                if(line.startswith('CLUSTER*:')):
                    # copy the cluster from the base dictionary.
                    cluster_name = line[9:].strip()

                    default_cluster = self.default_dict.getCluster(cluster_name)
                    if default_cluster != None:
                        cluster_d = copy.deepcopy(default_cluster)
                        self.tones.append(cluster_d)
                    else:
                        self.controller.showWarningDialog("Error", 
                            "There is no cluster labeled \'{}\' in the default dictionary.".format(cluster_name))
                        return

                # from the default dict
                if cluster_d is not None and (line.startswith('DIMENSION:')):                
                    dim_name = line[10:].strip()
                    dim_d = dict()
                    dim_d['type'] = 'DIMENSION'
                    dim_d['name'] = dim_name
                    dim_d['lc_name'] = dim_name.lower()
                    dim_d['lats'] = list()
                    dim_d['default'] = False
                    cluster_d['dimensions'].append(dim_d)
                    list_of_dims.append(dim_name)                    

                if cluster_d is not None and (line.startswith('DIMENSION*:')):
                    dim_name = line[11:].strip()
                    default_dim = self.default_dict.getDimension(dim_name)

                    if default_dim is not None:
                        # Copy a dimension from the base dictionary
                        new_dim_d = copy.deepcopy(default_dim)
                        cluster_d['dimensions'].append(new_dim_d)
                        list_of_dims.append(dim_name)
                    else:
                        # dim_name is not a dimension in the base dictionary. 
                        # Then, it must be a cluster.

                        default_clust = self.default_dict.getCluster(dim_name)
                        if default_clust is not None:
                            # A cluster from the base dict is used as a dimension in this custom dict.
                            # This algorithm assumes that there is only ONE LAT under each dimension. This is
                            # done by the Dictionary Editor by merging all the LATs under each dimension.
                            new_dim_d = copy.deepcopy(default_clust) # make a copy
                            new_dim_d['type'] = 'DIMENSION'
                            new_dim_d['lats'] = list()               # add a new key for listing lats
                            for dim in new_dim_d['dimensions']:      # populate the list with lats.
                                new_dim_d['lats'].append(dim['lats'][0])

                            del new_dim_d['dimensions']

                            cluster_d['dimensions'].append(new_dim_d)
                            list_of_dims.append(dim_name)
                        else:
                            self.controller.showWarningDialog("Error", 
                                "There is no category labeled \'{}\' in the default dictionary.".format(dim_name))
                            return

                if dim_d != None and (line.startswith('LAT*:')):
                    lat_name = line[6:].strip()
                    default_lat = self.default_dict.getLAT(lat_name)
                    if default_lat is not None:
                        new_lat_d = copy.deepcopy(default_lat) # make a copy
                        dim_d['lats'].append(new_lat_d)  
                    else:
                        default_lat = self.default_dict.getDimension(lat_name.replace("_LAT", ""))
                        if default_lat is not None:
                            new_lat_d = copy.deepcopy(default_lat) # make a copy 
                            new_lat_d['type'] = 'LAT'
                            dim_d['lats'].append(new_lat_d)  
                            del new_lat_d['lats']
                        else:
                            self.controller.showWarningDialog("Error", 
                                "There is no category labeled \'{}\' in the default dictionary.".format(lat_name))
                            return

        # sort the dimension names
        for c in self.tones:
            c['dimensions'] = sorted(c['dimensions'] , key=itemgetter('name'))

    # ---------------------------------
    # isLight(self):
    # self.is_light is true IF the custom dictionary consists ONLY of
    # clusters from the default dictionary.
    # ---------------------------------
    # def isLight(self):
        # return self.is_light;

    # ---------------------------------
    # DSCustomDict.getStats(self):
    #
    # ---------------------------------
    def getStats(self):

        list_of_dims = list()

        stats = dict()
        stats['cluster_count']   = len(self.tones)
        stats['dim_count'] = 0
        stats['lat_count'] = 0

        for c in self.tones:
            stats['dim_count'] += len(c['dimensions'])
            for d in c['dimensions']:
                stats['lat_count'] += len(d['lats'])
                list_of_dims.append(d['name'])

        stats['other_dims'] = 0
        stats['other_lats'] = 0

        # OPTIMIZE like this by making the default dictionary remember the counts after the first time.
        default_dim_count = self.default_dict.getDimensionCount()
        default_lat_count = self.default_dict.getLATCount()
        stats['other_dims'] = default_dim_count - stats['dim_count']
        stats['other_lats'] = default_lat_count - stats['lat_count']

        # other_dims = self.default_dict.getUnusedDimensions(list_of_dims) 
        # stats['other_dims'] = len(other_dims)
        # for d in other_dims:
        #     stats['other_lats'] += len(d['lats'])

        wd = 150
        if platform.system() == 'Windows':
            wd = 150*1.3

        stats_str  = "<p><b>Custom Dictionary</b><table>"
        stats_str += "<tr><td width=\"{}\">Custom Clusters:  </td><td align=\"right\">{}</td></tr>".format(wd, stats['cluster_count'])
        stats_str += "<tr><td>Custom Dimensions:</td><td align=\"right\">{:,}</td></tr>".format(stats['dim_count'])
        # stats_str += "<tr><td>Custom LATs:      </td><td align=\"right\">{:,}</td></tr>".format(stats['lat_count'])
        stats_str += "<tr><td>Unused Dimensions:</td><td align=\"right\">{:,}</td></tr>".format(stats['other_dims'])
        # stats_str += "<tr><td>Unused LATs:      </td><td align=\"right\">{:,}</td></tr>".format(stats['other_lats'])
        stats_str += "<tr><td>Total Dimensions: </td><td align=\"right\">{:,}</td></tr>".format(stats['dim_count'] + stats['other_dims'])
        # stats_str += "<tr><td>Total LATs:       </td><td align=\"right\">{:,}</td></tr>".format(stats['lat_count'] + stats['other_lats'])
        stats_str += "</table></p>"

        return stats_str

    # ---------------------------------
    # generateHelpDoc(self, savepath):
    #
    #
    # ---------------------------------
    def generateHelpDoc(self, savepath):

        # self.updateClusterInfo()

        dict_name = os.path.basename(self.directory)

        docx = Document()
        styles = docx.styles

        style = styles['Title']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)

        style = styles['Heading 1']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        style.paragraph_format.keep_with_next = True

        style = styles['Normal']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        style.paragraph_format.keep_together = True
 
        p = docx.add_paragraph(dict_name)
        p.style = 'Title'
        for c in self.tones:
            if c['name'] == 'Other':
                continue
            if len(c['dimensions']) == 0:
                self.controller.showWarningDialog("Warning", "{} is not included becasue it has no dimensions.".format(c['name']))
            else:
                p = docx.add_paragraph(c['label'])
                p.style = 'Heading 1'
                p = docx.add_paragraph(c['help'])
                p.style = 'Normal'

        docx.save(savepath)

    def getDefaultDATPath(self):
        if self.default_dict:
            return self.default_dict.getDATPath()
        else:
            return None


    def isBasedOnBaseDict(self, category_name, category_type):

        if category_type == CLUSTER:
            clust = self.getCluster(category_name)
            return clust['default']
        elif category_type == DIMENSION:
            dim = self.getDimension(category_name)
            return dim['default']            

    # ---------------------------------
    # isDefaultName(self):
    #
    # ---------------------------------
    def isUntitledCluster(self, cluster_name):
        n = 0
        res = re.search('(?<=Untitled_)[1-9]+', cluster_name)
        if res != None:
            n = int(res.group())
        return n

    # ---------------------------------
    # addDimension(self, cluster_name, dimension_name):
    # Add a dimension from the default dictionary
    # ---------------------------------
    def addDefaultDimension(self, dimension_name, cluster_name):
        default_dim = self.default_dict.getDimension(dimension_name)  # get the default dim dict
        dim_copy = copy.deepcopy(default_dim)                         # make a copy of the dict
        cluster = self.getCluster(cluster_name)          # get the custom cluster (which should exist now)
        cluster['dimensions'].append(dim_copy)           # add the copy of the dim to the cluster

    # ---------------------------------
    # addNewCluster(self)
    # Add a new cluster to the default dictionary
    # ---------------------------------    
    def addNewCluster(self, index):

        global untitled_category_count
        untitled_category_count += 1

        cluster_d = dict()
        cluster_d['type']    = 'CLUSTER'
        cluster_d['name']    = 'Untitled{}'.format(untitled_category_count)
        cluster_d['label']   = 'n/a'
        cluster_d['help']    = "n/a"
        cluster_d['default'] = False
        cluster_d['dimensions'] = list()

        if index >= 0:
            self.tones.insert(index, cluster_d)
        else:
            self.tones.append(cluster_d)

        # self.updateClusterInfo()

        return cluster_d['name']

    # ---------------------------------
    # DSCustomDict
    # addDefaultDimensionAsLAT(self, base_cluster_name, custom_cluster_name, index):
    # Add a cluster from the default dictionary to the custom dictony as a dimension.
    # ---------------------------------
    def addDefaultDimensionAsLAT(self, base_dim_name, custom_dim_name):

        base_dim = self.default_dict.getDimension(base_dim_name)

        # create a new lat from the base_dim.
        new_lat_d = dict()
        new_lat_d['type']    = 'LAT'
        new_lat_d['name']    = "{}_LAT".format(base_dim['name'])
        new_lat_d['lc_name'] = new_lat_d['name'].lower()
        new_lat_d['label']   = base_dim['label']
        new_lat_d['help']    = base_dim['help']
        new_lat_d['original_patterns'] = None
        new_lat_d['revised_patterns']  = None
        new_lat_d['original_count']    = 0
        new_lat_d['revised_count']     = 0
        new_lat_d['added']    = []      # all of the added/new patterns
        new_lat_d['deleted']  = []      # deleted patterns
        new_lat_d['verified'] = []      # new/added but already verified.
        new_lat_d['is_new']   = False
        new_lat_d['default']  = True

        dim = self.getDimension(custom_dim_name)   # get the cluster
        dim['lats'].append(new_lat_d)              # add the copy of the dim to the cluster

    # ---------------------------------
    # DSCustomDict
    # addDefaultClusterAsDimension(self, base_cluster_name, custom_cluster_name, index):
    # Add a cluster from the default dictionary to the custom dictony as a dimension.
    # ---------------------------------
    def addDefaultClusterAsDimension(self, base_cluster_name, custom_cluster_name, index):

        base_cluster = self.default_dict.getCluster(base_cluster_name)

        new_dim_d = dict()
        new_dim_d['type']    = 'DIMENSION'
        new_dim_d['name']    = base_cluster['name']
        new_dim_d['lc_name'] = new_dim_d['name'].lower()
        new_dim_d['label']   = base_cluster['label']
        new_dim_d['help']    = base_cluster['help']
        new_dim_d['lats']    = list()
        new_dim_d['default'] = True

        for dim in base_cluster['dimensions']:
            # for each dimension under a default cluster.
            lat_d = dict()
            lat_d['type']    = 'LAT'
            lat_d['name']    = "{}_LAT".format(dim['name'])
            lat_d['lc_name'] = lat_d['name'].lower()
            lat_d['label']   = dim['label']
            lat_d['help']    = dim['help']
            # lat_d['original_patterns'] = None
            # lat_d['revised_patterns']  = None
            # lat_d['original_count']    = 0
            # lat_d['revised_count']     = 0
            # lat_d['added']    = []      # all of the added/new patterns
            # lat_d['deleted']  = []      # deleted patterns
            # lat_d['verified'] = []      # new/added but already verified.
            # lat_d['is_new']   = False
            lat_d['default']  = True
            new_dim_d['lats'].append(lat_d)

        cluster = self.getCluster(custom_cluster_name)       # get the cluster
        cluster['dimensions'].append(new_dim_d)              # add the copy of the dim to the cluster

    # ---------------------------------
    # DSCustomDict
    # addDefaultCluster(self, cluster_name):
    # Add a cluster from the default dictionary
    # ---------------------------------
    def addDefaultCluster(self, cluster_name, index):
        default_cluster = self.default_dict.getCluster(cluster_name)
        cluster_copy = copy.deepcopy(default_cluster)

        if index >=0:
            self.tones.insert(index, cluster_copy)
        else:
            self.tones.append(cluster_copy)

    # ----------------------------------------
    # deleteCategory(self, cluster, dimension):
    # Delete a category that matches 'cluster' and 'dimension'.
    # ----------------------------------------
    def deleteCategory(self, cluster_name, dimension_name, lat_name):
        global untitled_category_count

        is_default = False

        if dimension_name == None:
            help_item = self.getHelpItem('CLUSTER', cluster_name)
        else:
            help_item = self.getHelpItem('DIMENSION', dimension_name)

        if help_item != None and help_item in self.help:
            self.help.remove(help_item)

        removed_category = None

        if cluster_name is not None and dimension_name is None and lat_name is None:
            # cluster
            for c  in self.tones:
                if c['name'] == cluster_name:
                    self.tones.remove(c)
                    n = self.isUntitledCluster(cluster_name)
                    if n > 0:
                        untitled_category_count -= 1
                    removed_category = c
                    break

        elif dimension_name is not None and lat_name is None:
            # dimension
            for c  in self.tones:
                for d in c['dimensions']:
                    if d['name'] == dimension_name:
                       c['dimensions'].remove(d)
                       removed_category = d
                       break

        elif lat_name is not None:
            # lat
            for c in self.tones:
                for d in c['dimensions']:            
                    for lat in d['lats']:
                        if lat['name'] == lat_name:
                            d['lats'].remove(lat)
                            removed_category = lat
                            break
                                    
        return removed_category

    def moveCluster(self, cluster_name, direction):
        for index in range(len(self.tones)):
            c = self.tones[index]
            if c['name'] == cluster_name:
                self.tones.insert(index+direction, self.tones.pop(index))        
                break

    # ----------------------------------------
    # updateClusterName(self, old_name, new_name):
    #
    # ----------------------------------------
    def updateDimensionName(self, old_name, new_name):
        dim = self.getDimension(old_name)
        dim['name'] = new_name

        return dim

    # ----------------------------------------
    # updateClusterName(self, old_name, new_name):
    #
    # ----------------------------------------
    def updateClusterName(self, old_name, new_name):
        cluster = self.getCluster(old_name)
        cluster['name'] = new_name

        return cluster
    # ----------------------------------------
    # setDefaultDict(self, default_dict)
    #
    # ----------------------------------------
    def setDefaultDict(self, default_dict):
        self.default_dict = default_dict

    # ----------------------------------------
    # deleteCompactDic(self)
    #
    # ----------------------------------------
    def deleteCompactDic(self):
        filepath = '{}/_dict.dat'.format(self.directory)
        if os.path.exists(filepath):
            os.remove(filepath)

    # ----------------------------------------
    # saveExportFilter(dict_name, custom_clusters):
    # 
    # ----------------------------------------
    def saveExportFilter(self):
        with open('{}/_export_filter.txt'.format(self.directory), 'w') as export_filter:
            for c in self.tones:
                if c['name'] != 'Other':
                    export_filter.write("CLUSTER: {}\n".format(c['name']))


    ##########################################
    # Help related methods
    ##########################################

    def isEmpty(self):
        if len(self.tones) == 0:
            return True

    # ----------------------------------------
    # DSCustomDict.loadHelp(self):
    # This function reads _custom_help.txt and add the help content
    # to the dictionary hierarchy (i.e., self.tones)
    # ----------------------------------------
    # def loadHelp(self):
    #     with open('{}/_custom_help.txt'.format(self.directory)) as help_file:
    #         try:
    #             text = help_file.read()
    #         except:
    #             print("Non UTF-8 characters found in _help.txt")
    #             text = repair_encoding_problems('{}/_custom_help.txt'.format(self.directory))

    #         lines = text.split("\n")
    #         lines = [l for l in lines if l != '']
    #         i = 0
    #         while i < len(lines):
    #             line = lines[i]
    #             if i % 3 == 0:
    #                 if not (line.startswith("CLUSTER:") or line.startswith("DIMENSION:")):
    #                     dialogs.WarningDialog("ERROR!", "Found an extra line break around {} (line {}) in _help.txt".format(line, i))
    #                     return

    #             if line.startswith("CLUSTER:"):
    #                 cluster_name = line[9:].strip()
    #                 cluster_d = dict()
    #                 cluster_d['type']  = 'CLUSTER'
    #                 cluster_d['name']  = cluster_name
    #                 cluster_d['label'] = lines[i+1]
    #                 cluster_d['help']  = lines[i+2]
    #                 self.help.append(cluster_d)
    #                 i+=2
    #             elif line.startswith("DIMENSION:"):
    #                 dim_name = line[11:].strip()
    #                 dim_d = dict()
    #                 dim_d['type']  = 'DIMENSION'
    #                 dim_d['name']  = dim_name
    #                 dim_d['label'] = lines[i+1]
    #                 dim_d['help']  = lines[i+2]
    #                 self.help.append(dim_d)
    #                 i+=2
    #             i+=1

    #     # add the help content to the tones hierarchy
    #     for c in self.tones:
    #         cluster_help = self.getHelpItem('CLUSTER', c['name'])
    #         c['label'] = cluster_help['label']
    #         c['help']  = cluster_help['help']

    #         for d in c['dimensions']:
    #             dim_help = self.getHelpItem('DIMENSION', d['name'])
    #             d['label'] = dim_help['label']
    #             d['help']  = dim_help['help']

    # def loadHelp_old(self):

    #     with open('{}/_custom_help.txt'.format(self.directory)) as help_file:
    #         try:
    #             text = help_file.read()
    #         except:
    #             print("Non UTF-8 characters found in _help.txt")
    #             text = repair_encoding_problems('{}/_custom_help.txt'.format(self.directory))

    #         lines = text.split("\n")
    #         lines = [l for l in lines if l != '']
    #         i = 0
    #         while i < len(lines):
    #             line = lines[i]
    #             if i % 3 == 0:
    #                 if not (line.startswith("CLUSTER:") or line.startswith("DIMENSION:")):
    #                     dialogs.WarningDialog("ERROR!","Found an extra line break around {}".format(line))
    #                     return

    #             if(line.startswith("CLUSTER:")):
    #                 cluster_name = line[9:].strip()
    #                 cluster_d = dict()
    #                 cluster_d['type']  = 'CLUSTER'
    #                 cluster_d['name']  = cluster_name
    #                 cluster_d['label'] = lines[i+1]
    #                 cluster_d['help']  = lines[i+2]
    #                 cluster_d['dimensions'] = list()
    #                 self.help.append(cluster_d)
    #                 i+=2
    #             i+=1

    #     # add the help content to the tones hierarchy
    #     for c in self.tones:
    #         cluster_help = self.getHelpItem('CLUSTER', c['name'])
    #         if c['default'] == True:
    #             if cluster_help == None:
    #                 cluster_help = self.default_dict.getHelpItem('CLUSTER', c['name'])

    #             if cluster_help == None:
    #                 dialogs.WarningDialog("Error","{} does not exist. _cusom_help may be damaged.".format(c['name']))
    #             else:
    #                 c['label'] = cluster_help['label']
    #                 c['help']  = cluster_help['help']
    #         else:
    #             if cluster_help == None:
    #                 dialogs.WarningDialog("Error","{} does not exist. _cusom_help may be damaged.".format(c['name']))
    #             else:                    
    #                 c['label'] = cluster_help['label']
    #                 c['help']  = cluster_help['help']

    #         for d in c['dimensions']:
    #             dim_help = self.default_dict.getHelpItem('DIMENSION', d['name'])
    #             if dim_help != None:
    #                 d['label'] = dim_help['label']
    #                 d['help']  = dim_help['help']
    #             else:
    #                 dialogs.WarningDialog("loadHelp() Error", "{} doesn't exist".format(d['name']))

    # ----------------------------------------
    # DSCustomDict.save()
    #
    # ----------------------------------------
    def save(self):

        info_doc_path = '{}/info.json'.format(self.directory)
        dict_fpath      = os.path.join(self.directory, "dict.ser")

        # If there is already a .ser file (serialized object), delete it since
        # it is not in sync with the newly saved dictionary.
        if os.path.exists(dict_fpath):
            os.remove(dict_fpath)        

        # delete the existing file
        if os.path.exists(info_doc_path):
            os.remove(info_doc_path)        

        info = dict()
        info['name']      = os.path.basename(self.directory)
        info['base_dict_path'] = self.default_dict.getDirectory()
        info['pronouns']  = self.pronouns
        info['is_common'] = self.is_common

        if self.isEmpty():
            info['customized'] = True
            with open(info_doc_path, 'w') as fout:
                json.dump(info, fout, indent=4)
            # return []
        else:
            info['customized'] = True
            with open(info_doc_path, 'w') as fout:
                json.dump(info, fout, indent=4)

        list_of_dims = list()
        list_of_lats = list()

        custom_wrodclasses_path = "{}/_wordclasses.txt".format(self.directory)            
        custom_tones_path       = '{}/_custom_tones.txt'.format(self.directory)
        custom_help_path        = '{}/_custom_help.txt'.format(self.directory)
        tones_path              = '{}/_tones.txt'.format(self.directory)
        help_path               = '{}/_help.txt'.format(self.directory)
        collisions_path         = '{}/_Dict_Collisions.txt'.format(self.directory)
        dicterrors_path         = '{}/_dicterrors_log.csv'.format(self.directory)
        badwc_path              = '{}/_Bad_Word_Classes.txt'.format(self.directory)

        # delete the existing files
        if os.path.exists(custom_wrodclasses_path):
            os.remove(custom_wrodclasses_path)

        if os.path.exists(custom_tones_path):
            os.remove(custom_tones_path)
        if os.path.exists(custom_help_path):
            os.remove(custom_help_path)        

        if os.path.exists(tones_path):
            os.remove(tones_path)
        if os.path.exists(help_path):
            os.remove(help_path)        

        if os.path.exists(collisions_path):
            os.remove(collisions_path)
        if os.path.exists(dicterrors_path):
            os.remove(dicterrors_path)        
        if os.path.exists(badwc_path):
            os.remove(badwc_path)        

        skipped_clusters = []
        # _custom_tones.txt
        with open(custom_tones_path, 'w') as _custom_tones:
            for c in self.tones:
                if c['name'] == 'Other':
                    continue

                if len(c['dimensions']) == 0:
                    skipped_clusters.append(c['name'])
                else:
                    if c['default'] == True:
                        _custom_tones.write("CLUSTER*: {}\n".format(c['name']))
                    else:
                        _custom_tones.write("CLUSTER: {}\n".format(c['name']))

                        for d in c['dimensions']:
                            if d['default'] == True:
                                _custom_tones.write("DIMENSION*: {}\n".format(d['name']))
                            else:
                                _custom_tones.write("DIMENSION: {}\n".format(d['name']))

                                for lat in d['lats']:
                                    _custom_tones.write("LAT*: {}\n".format(lat['name']))

                _custom_tones.write("\n")


        # _tones.txt
        with open(tones_path, 'w') as _tones:

            for c in self.tones:
                _tones.write("CLUSTER: {}\n".format(c['name']))

                for d in c['dimensions']:
                    dim_name = d['name']
                    _tones.write("DIMENSION: {}\n".format(dim_name))
                    list_of_dims.append(dim_name)

                    for lat in d['lats']:
                        lat_name = lat['name']
                        _tones.write("LAT*: {}\n".format(lat_name))
                        list_of_lats.append(lat_name)
                    _tones.write("\n")
                _tones.write("\n")

            for lat_name in list_of_lats:
                list_of_dims.append(lat_name.replace("_LAT", ""))

            other_dims = self.default_dict.getUnusedDimensions(list_of_dims)
            _tones.write("CLUSTER: Other\n")
            for d in other_dims:
                dim_name = d['name']
                _tones.write("DIMENSION: {}\n".format(dim_name))
                for lat in d['lats']:
                    lat_name = lat['name']
                    _tones.write("LAT*: {}\n".format(lat_name))
                _tones.write("\n")

        self.default_dict.saveWordClasses(custom_wrodclasses_path)

        # _custom_help.txt
        with open(custom_help_path, 'w') as _custom_help:
            for c in self.tones:
                if c['default']:
                    continue                
                _custom_help.write("CLUSTER: {}\n".format(c['name']))
                _custom_help.write("{}\n".format(c['label']))
                _custom_help.write("{}\n".format(c['help']))
                _custom_help.write("\n")

                for d in c['dimensions']:
                    if d['default']:
                        continue
                    _custom_help.write("DIMENSION: {}\n".format(d['name']))
                    _custom_help.write("{}\n".format(d['label']))
                    _custom_help.write("{}\n".format(d['help']))
                    _custom_help.write("\n")

        # _help.txt
        with open(help_path, 'w') as _help:
            for c in self.tones:
                _help.write("CLUSTER: {}\n".format(c['name']))
                _help.write("{}\n".format(utils.convert_invalid_xml_chars(c['label'])))
                _help.write("{}\n".format(utils.convert_invalid_xml_chars(c['help'])))
                _help.write("\n")

                for d in c['dimensions']:
                    _help.write("DIMENSION: {}\n".format(d['name']))
                    _help.write("{}\n".format(utils.convert_invalid_xml_chars(d['label'])))
                    _help.write("{}\n".format(utils.convert_invalid_xml_chars(d['help'])))
                    _help.write("\n")

                    # for lat in d['lats']:
                    #     _help.write("LAT: {}\n".format(lat['name']))
                    #     _help.write("{}\n".format(utils.convert_invalid_xml_chars(lat['label'])))
                    #     _help.write("{}\n".format(utils.convert_invalid_xml_chars(lat['help'])))
                    #     _help.write("\n")

                _help.write("\n")

        return skipped_clusters

    def getDirectory(self):
        return self.directory


    def getItem(self, cluster_name, dim_name):
        if self.isEmpty():
            return self.default_dict.getItem(cluster_name, dim_name)
        else:
            if cluster_name is not None:
                for cluster in self.tones:
                    if cluster['name'] == cluster_name and dim_name is None:
                        return cluster
                    elif dim_name is not None:
                        for dim in cluster['dimensions']:
                            if dim['name'] == dim_name:
                                return dim
        return None                                        

    def setPronouns(self, val):
        self.pronouns = val

    def setCommon(self, val):
        self.is_common = val

    def isCommon(self):
        return self.is_common

    def setDefaultTagger(self, tagger):
        self.default_dict.setTagger(tagger)

    # DSCustomDict.serialize
    def export(self, dict_info, export_info, backup):

        self.bExport = True

        base_dict_fpath = os.path.join(self.default_dict.getDirectory(), "dict.ser")

        src_dict_dir    = self.directory                     # the src dictionary folder
        dict_name       = os.path.basename(src_dict_dir)     # the name of the src folder

        dst_dir         = os.path.join(export_info['dst_dir'], 
                                       "{} (v.{})".format(dict_info['name'], dict_info['version']))
        dst_dict_fpath  = os.path.join(dst_dir, "dict.ser")

        if export_info.get('is_temp', False) == True:
            # We are creating a serialized object for the find LAT features.
            temp_dir  =  src_dict_dir
            backup = False   # if it's a temporary export, we'll always ignore backup.
            self.bExport = False
            
        else:
            src_parent_dir     = os.path.dirname(src_dict_dir)      # parent dir of the src dict folder            
            src_help_fpath     = os.path.join(src_dict_dir, "_help.txt")
            src_tones_fpath    = os.path.join(src_dict_dir, "_tones.txt")            
            src_rules_fpath    = os.path.join(src_dict_dir, "rules.json")
            src_patterns_fpath = os.path.join(src_dict_dir, "patterns.json")
            src_cv_fpath       = os.path.join(src_dict_dir, "values.json")            

            dst_help_fpath     = os.path.join(dst_dir, "_help.txt")
            dst_tones_fpath    = os.path.join(dst_dir, "_tones.txt")
            dst_rules_fpath    = os.path.join(dst_dir, "rules.json")
            dst_patterns_fpath = os.path.join(dst_dir, "patterns.json")
            dst_cv_fpath       = os.path.join(dst_dir, "values.json")
            backup_dir         = os.path.join(export_info['dst_dir'], "{}_backup".format(dict_name))

            if os.path.exists(dst_dir) != True:
                os.makedirs(dst_dir)

            # Copy the help doc
            shutil.copy(src_help_fpath, dst_help_fpath)
            shutil.copy(src_tones_fpath, dst_tones_fpath)

            if os.path.exists(src_cv_fpath) == True:
                shutil.copy(src_cv_fpath, dst_cv_fpath)

            # Copy the rules files if it exists.
            if os.path.exists(src_rules_fpath) == True:
                shutil.copy(src_rules_fpath, dst_rules_fpath)

            # Copy the pattens file if it exits.
            if os.path.exists(src_patterns_fpath) == True:
                shutil.copy(src_patterns_fpath, dst_patterns_fpath)

            # Save the dictionary info.
            info_fpath = os.path.join(dst_dir, "info.json")
            if os.path.exists(info_fpath):
                os.remove(info_fpath)

            dict_info['base_dict'] = os.path.basename(self.default_dict.getDirectory())

            with open(info_fpath, 'w') as info_fout:
                dict_info['pronouns'] = self.pronouns
                dict_info['is_common'] = self.is_common
                dict_info['customized'] = True
                json.dump(dict_info, info_fout, indent=4)


            # Generate a serialized dictionary object
            if os.path.exists(dst_dict_fpath):
                os.remove(dst_dict_fpath)

            # create a backup of the original dictionary where the custom dictionary defintions are stored.
            if backup:
                if os.path.exists(backup_dir) != True:
                    os.makedirs(backup_dir)
                backup_basename = os.path.join(backup_dir, 
                                               "{} (v.{}) archive".format(dict_info['name'], dict_info['version']))
                shutil.make_archive(backup_basename, 'zip', src_dict_dir)

                if self.bExport != True:
                    self.cleanupExportTempFiles()
                    return "", ""

        self.tagger = dstagger.DsTagger()               # create a tagger object
        self.tagger.loadDat(base_dict_fpath)            # load the base dictionary
        self.tagger.customizeDictionary(src_dict_dir)   # customize the tagger.
        self.tagger.exportAsDat(dst_dict_fpath)         # seralize the file, save it as a .ser file.

        self.cleanupExportTempFiles()
        
        if self.bExport: 
            self.bExport = False
            if backup:
                zip_path = backup_basename+".zip"
            else:
                zip_path = '(no backup)'

            return dst_dir, zip_path
        else:
            return "", ""

    # DSCustomDict.load()
    def load(self):
        base_dict_fpath = os.path.join(self.default_dict.getDirectory(), "dict.ser")
        self.tagger = dstagger.DsTagger()                 # create a tagger object
        self.tagger.loadDat(base_dict_fpath)              # load the base dictionary
        self.tagger.customizeDictionary(self.directory)   # customize the tagger.
        
    def generateCommonDict(self, dict_info, export_info):

        dst_dir_path = export_info['dst_dir']

        common_dict_json_path  = os.path.join(dst_dir_path, "common_dict.json")
        if os.path.exists(common_dict_json_path):
            retval = self.controller.showYesNoDialog("File Exists", 
                "Do you want to overwrite {}? \n\nThe existing file will be permanently deleted.\n".format(common_dict_json_path))

            if retval == False:
                return

        base_dict_name, ext = os.path.splitext(self.default_dict.getName())
        base_dict_version   = self.default_dict.getVersion()

        common_dict_data = dict()
        common_dict_data['base_dict'] = "{} ({})".format(base_dict_name, base_dict_version)

        custom_dict_name = os.path.basename(self.getDirectory())
        common_dict_data['custom_dict']  = custom_dict_name

        common_dict_data['copyright'] = dict_info['copyright']

        common_dict_data['use_base_dict']  = True

        # let's remember when it wa saved
        now = datetime.now() # current date and time
        common_dict_data['timestamp'] = now.strftime("%m/%d/%Y, %H:%M:%S")

        # WRite the custom dictionary data.
        common_dict_data['categories'] = list()
        for clust in self.tones:
            new_cat = dict()
            new_cat['label'] = clust['label']
            new_cat['help']  = clust['help']
            new_cat['subcategories'] = list()
            common_dict_data['categories'].append(new_cat)

            for dim in clust['dimensions']:
                new_subcat = dict()
                new_subcat['label'] = dim['label']
                new_subcat['help']  = dim['help']
                new_subcat['clusters'] = list()
                new_cat['subcategories'].append(new_subcat)

                temp_clusters = list()
                for lat in dim['lats']:
                    dim_name, clust_name = self.default_dict.getDimensionAndCluster(lat['name'])
                    clust = self.default_dict.getCluster(clust_name)
                    if clust_name not in temp_clusters:
                        new_clust = dict()
                        new_clust['name']  = clust_name
                        new_clust['label'] = clust['label']
                        new_clust['help']  = clust['help']
                        new_subcat['clusters'].append(new_clust)
                        temp_clusters.append(clust_name)


        with open(common_dict_json_path, 'w') as fout:
            json.dump(common_dict_data, fout, indent=4)


##################################################
#
# DSCommonDict
#
##################################################

class DSCommonDict():
    def __init__(self, controller, custom_dict, default_dict):

        self.controller = controller

        self.cat_count     = 0
        self.subcat_count  = 0
        self.cluster_count = 0

        self.default_dict  = default_dict
        self.custom_dict   = custom_dict
        self.directory     = self.custom_dict.getDirectory()

        self.categories_edited = False

        self.categories = list()
        json_path = os.path.join(custom_dict.getDirectory(), "common_dict.json")
        if os.path.isfile(json_path):
            with open(json_path) as fin:
                d = json.load(fin)
                self.categories = d['categories']

            # expand the clusters
            for cat in self.categories:
                for subcat in cat['subcategories']:
                    for cluster in subcat['clusters']:
                        if self.custom_dict.isEmpty():
                            clust = self.default_dict.getCluster(cluster['name'])
                        else:
                            clust = self.custom_dict.getCluster(cluster['name'])
                        cluster['label'] = clust['label']
                        cluster['help']  = clust['help']
                        cluster['dimensions'] = copy.deepcopy(clust['dimensions'])
        else:
            self.categories = list()

    def print(self):
        print("-----------------------------")
        for cat in self.categories:
            print(cat['label'])
            print(cat['help'])
            for subcat in cat['subcategories']:
                print(subcat['label'])
                print(subcat['help'])

                for cluster in subcat['clusters']:
                    print(cluster['name'])
                    print(cluster['label'])
                    print(cluster['help'])
        print("-----------------------------")

    def getDirectory(self):
        return self.directory

    def save(self):

        res = dict()

        # ds_dict_<lang>_yyyy.mm.dd    -- e.g., d_dict_en_2020.09.28
        defatul_dict_name, ext = os.path.splitext(self.default_dict.getName())
        res['default_dict'] = defatul_dict_name

        custom_dict_name = os.path.basename(self.custom_dict.getDirectory())
        res['custom_dict']  = custom_dict_name

        if self.custom_dict.isEmpty():
            res['use_default_dict']  = True
        else:
            res['use_default_dict']  = False

        # let's remember when it wa saved
        now = datetime.now() # current date and time
        res['timestamp'] = now.strftime("%m/%d/%Y, %H:%M:%S")

        # dictionary categories
        res['categories'] = list()
        for cat in self.categories:
            new_cat = dict()
            new_cat['label'] = cat['label']
            new_cat['help']  = cat['help']
            new_cat['subcategories'] = list()
            res['categories'].append(new_cat)

            for subcat in cat['subcategories']:
                new_subcat = dict()
                new_subcat['label'] = subcat['label']
                new_subcat['help']  = subcat['help']
                new_subcat['clusters'] = list()
                new_cat['subcategories'].append(new_subcat)

                for cluster in subcat['clusters']:
                    new_clust = dict()
                    new_clust['name']  = cluster['name']
                    new_clust['label'] = cluster['label']
                    new_clust['help']  = cluster['help']
                    new_subcat['clusters'].append(new_clust)

        path = self.custom_dict.getDirectory() 
        with open(os.path.join(path, "common_dict.json"), 'w') as fout:
            json.dump(res, fout, indent=4)

    def existingCategoryLabel(self, label):

        for cat in self.categories:
            if cat['label']== label:
                return True

            for subcat in cat['subcategories']:
                if subcat['label'] == label:
                    return True

                # for clust in subcat['cluster']:
                    # if clust['label'] == label:
                        # return True
        return False

    def getItem(self, cat_label, subcat_label, cluster_label, dim_label):
        for cat in self.categories:
            if cat['label'] == cat_label and subcat_label is None:
                return cat
            elif subcat_label is not None:
                for subcat in cat['subcategories']:
                    if subcat['label'] == subcat_label and cluster_label is None:
                        return subcat
                    elif cluster_label is not None:
                        for cluster in subcat['clusters']:
                            if cluster['name'] == cluster_label and dim_label is None:
                                return cluster
                            elif dim_label is not None:
                                for dim in cluster['dimensions']:
                                    if dim['name'] == dim_label:
                                        return dim
        return None                                        

    def getCategory(self, cat_label):
        for cat in self.categories:
            if cat['label'] == cat_label:
                return cat
        return None

    def getSubcategory(self, subcat_label):
        for cat in self.categories:
            for subcat in cat['subcategories']:
                if subcat['label'] == subcat_label:
                    return subcat
        return None

    def getCategories(self):
        return self.categories

    def getCategoryCount(self):
        return self.cat_count

    def getSubcategoryCount(self):
        return self.subcat_count

    def getClusterCount(self):
        return self.cluster_count

    #
    # Add a NEW category
    #
    def addNewCategory(self, index):

        global untitled_category_count
        untitled_category_count += 1

        cat_d = dict()
        cat_d['type']    = 'CATEGORY'
        cat_d['label']   = 'Untitled{}'.format(untitled_category_count)
        cat_d['help']    = "n/a"
        cat_d['subcategories'] = list()

        if index >= 0:
            self.categories.insert(index, cat_d)
        else:
            self.categories.append(cat_d)

        # self.updateClusterInfo()

        return cat_d['label']

    def addNewSubcategory(self, cat_label):
        global untitled_category_count
        untitled_category_count += 1

        # 1. Find the parent cluster
        parent_cat = self.getCategory(cat_label)

        if parent_cat != None:
            # 2. Create a new dimension object

            subcat_d = dict()
            subcat_d['type']     = 'SUBCATEGORY'
            subcat_d['label']    = 'Untitled{}'.format(untitled_category_count)
            subcat_d['help']     = "n/a"
            subcat_d['clusters'] = list()

             # 3. Add the new dimension object to the list of dimensions
            parent_cat['subcategories'].append(subcat_d)
            parent_cat['subcategories'] = sorted(parent_cat['subcategories'] , key=itemgetter('label'))

            self.tones_edited = True
            return subcat_d['label']
        else:
            print("ERROR!! category {} not found.".format(cat_label))
            return "ERROR"

    def isEmpty(self):
        if self.categories == []:
            return True
        else:
            return False

    def isUntitledCategory(self, name):
        n = 0
        res = re.search('(?<=Untitled_)[1-9]+', name)
        if res != None:
            n = int(res.group())
        return n

    #
    # DELETE a category
    #
    def deleteCategory(self, deleted):
        global untitled_category_count

        dcat      = deleted['cat']
        dsubcat   = deleted['subcat']
        dclusters = deleted['clusters']
        level     = deleted['level']

        is_default = False

        removed_category = None
        for cat  in self.categories:
            # a category is deleted
            if cat['label'] == dcat:
                if level == CATEGORY:
                    self.categories.remove(cat)
                    n = self.isUntitledCategory(dcat)
                    if n > 0:
                        untitled_category_count -= 1
                    removed_category = cat
                    break

                elif level == SUBCATEGORY:
                    for subcat in cat['subcategories']:
                        if subcat['label'] == dsubcat:
                           cat['subcategories'].remove(subcat)
                           n = self.isUntitledCategory(dsubcat)
                           if n > 0:
                               untitled_category_count -= 1
                           removed_category = subcat
                           break

                elif level == CLUSTER:
                    is_done = False
                    for subcat in cat['subcategories']:
                        if subcat['label'] == dsubcat:
                            for cluster in subcat['clusters']:
                                if cluster['name'] in dclusters:
                                   subcat['clusters'].remove(cluster)
                                   removed_category = cluster
                                   is_done = True
                                   break
                        if is_done:
                            break

        return removed_category

    # MOVE a category UP/DOWN
    #
    def moveCategory(self, cat_name, direction):

        for index in range(len(self.categories)):
            c = self.categories[index]
            if c['label'] == cat_name:
                self.categories.insert(index+direction, self.categories.pop(index))        
                break


    def addCustomCluster(self, subcat_name, cluster_name, index, cat_type=SUBCATEGORY):

        if self.custom_dict.isEmpty():
            custom_cluster = self.default_dict.getCluster(cluster_name)
        else:
            custom_cluster = self.custom_dict.getCluster(cluster_name)
        cluster_copy = copy.deepcopy(custom_cluster)

        subcat = self.getSubcategory(subcat_name)
        if subcat is not None:
            if index >=0:
                subcat['clusters'].insert(index, cluster_copy)
            else:
                subcat['clusters'].append(cluster_copy)

        # self.updateClusterInfo()            


    # ----------------------------------------
    # Help related
    # ----------------------------------------

    def updateCategoryHelpContent(self, old_label, new_label, content):
        cat = self.getCategory(old_label)

        if cat == None:
            return False
        else:

            if old_label != new_label and self.existingCategoryLabel(new_label) == True:
                return False

            if cat['label'] != new_label:
                cat['label'] = new_label
                self.categories_edited = True

            if cat['help'] != content:
                cat['help'] = content
                self.categories_edited = True

        return self.categories_edited

    def updateSubcategoryHelpContent(self, old_cat, old_subcat, new_label, content):

        subcat = self.getSubcategory(old_subcat)
        if subcat == None:
            # this is an error.

            return False
        else:
            if old_subcat != new_label and self.existingCategoryLabel(new_label) == True:
                return False

            if subcat['label'] != new_label:
                subcat['label'] = new_label
                self.categories_edited = True

            if subcat['help'] != content:
                subcat['help']  = content
                self.categories_edited = True

        return self.categories_edited

    # ---------------------------------
    # generateHelpDoc(self, savepath):
    #
    #
    # ---------------------------------
    def generateHelpDoc(self, savepath):

        docx = Document()
        styles = docx.styles

        style = styles['Title']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        font.size = Pt(28)

        style = styles['Heading 1']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        font.size = Pt(18)

        style.paragraph_format.keep_with_next = True

        style = styles['Heading 2']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        font.size = Pt(13)
        style.paragraph_format.keep_with_next = True

        style = styles['Heading 3']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        font.size = Pt(11)
        font.italic = True
        style.paragraph_format.keep_with_next = True

        style = styles['Normal']   
        font = style.font
        font.color.rgb = RGBColor(0,0,0)
        font.size = Pt(11)
        style.paragraph_format.keep_together = True
 
        #
        # content
        #
        p = docx.add_paragraph("Common Dictionary")
        p.style = 'Title'

        p = docx.add_paragraph("Default Dictionary: {}".format(self.default_dict.getName()))
        p.style = 'Normal'

        if self.custom_dict.isEmpty():
            custom_dict_type = "No Customization"
        else:
            custom_dict_type = "Customized"

        custom_dict_name = os.path.basename(self.custom_dict.getDirectory())
        p = docx.add_paragraph("Custom Dictionary: {} ({})".format(custom_dict_name, custom_dict_type))
        p.style = 'Normal'

        now = datetime.now() # current date and time
        p = docx.add_paragraph(now.strftime("%m/%d/%Y, %H:%M:%S"))
        p.style = 'Normal'

        for cat in self.categories:
            p = docx.add_paragraph(cat['label'])
            p.style = 'Heading 1'
            p = docx.add_paragraph(cat['help'])
            p.style = 'Normal'

            for subcat in cat['subcategories']:
                p = docx.add_paragraph(subcat['label'])
                p.style = 'Heading 2'
                p = docx.add_paragraph(subcat['help'])
                p.style = 'Normal'

                for cluster in subcat['clusters']:
                    p = docx.add_paragraph(cluster['label'])
                    p.style = 'Heading 3'
                    p = docx.add_paragraph(cluster['help'])
                    p.style = 'Normal'


        docx.save(savepath)

    def generateHierarchyDoc(self, savepath):
        with open(savepath, 'w') as fout:
            for cat in self.categories:
                fout.write("\"{}\"".format(cat['label']))

                subcat_count = 0
                for subcat in cat['subcategories']:
                    if subcat_count == 0:
                        fout.write(",\"{}\"".format(subcat['label']))
                    else:
                        fout.write(",\"{}\"".format(subcat['label']))
                    subcat_count += 1

                    clust_count = 0
                    for cluster in subcat['clusters']:
                        if clust_count == 0:
                            fout.write(",\"{}\"".format(cluster['label']))
                        else:
                            fout.write(",,\"{}\"".format(cluster['label']))
                        fout.write("\n")
                        clust_count += 1




