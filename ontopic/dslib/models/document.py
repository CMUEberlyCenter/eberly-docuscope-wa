#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
"""

__author__    = "Suguru Ishizaki"
__copyright__ = "2017-21 Suguru Ishizaki, Carnegie Mellon University"

# from time import time
import logging
import json
import re
import string
import os
import pprint                                  # pretty prnting for debugging
# import sys
from collections import Counter

import regex

# from urllib.parse import urlparse
# import base64 # unused-import
# from PIL import Image # unused-import
# from io import BytesIO # unused-import

from bs4 import BeautifulSoup as bs
from bs4 import Comment
import bs4

# from pydocx import PyDocX # unused-import
import unidecode

from docx import Document
from docx.shared import RGBColor
# from docx.enum.style import WD_BUILTIN_STYLE, WD_STYLE_TYPE # unused-import
from docx.shared import Inches # unused-import Pt


from spacy.lang.en.stop_words import STOP_WORDS
import spacy                                   # SpaCy NLP library
# from spacy.attrs import ORTH, LEMMA, POS

import dslib.models.stat as ds_stat


##################################################
# English ONLY
##################################################
stop_words = list(STOP_WORDS)
stop_words += ['therefore', 'however', 'thus', 'hence', '\'s', 'mr.', 'mrs.', 'ms.', 'dr.', 'prof.']
pronouns = ['it',  'i',    'we',   'you',   'he',  'she',  'they',
            'its'  'my',   'our',  'your',  'his', 'her',  'their',
                   'me',   'us',            'him',         'them',
                   'mine', 'ours', 'yours',        'hers', 'theirs']

for p in pronouns:
    if p in stop_words:
        stop_words.remove(p)

no_space_patterns = ["\u2019ve", "n\u2019t", ".\u201D", ".\u2019",
                     "\u2019s", "\u2019t", "\u2019m", "\u2019ll", "\u2019re", "\u2019d",
                     "\'ve" ,"n\'t", ".\"", ".\'",
                     "\'s", "\'t", "\'m", "\'ll", "\'re", "\'d",
                     "%"]
                    # note: use a single quote for apostorophes

left_quotes  = ['\u201C', '\u2018', '(', '[', '$', '\"', '\'', '\u00BF', '<']  # todo: should be 'left_punct'
right_quotes = ['\u201D', '\u2019', ')', ']', '>']                             # todo: should be 'right_punct'

dashes       = ['\u2014', '\u2013', '\u2015']  # em dash, en dash, horizontal bar
hyphen_slash = ['-', '/',]  # hyphen, slash,

ellipsis     = '\u2026'

end_puncts   = ['.', ',', ';', ':', '?', '!']

be_verbs = ['be', 'been', 'am', 'is', 'are', 'was', 'were', '\'m', '\'re', '\'s', "\u2019s", "\u2019m", "\u2019re"]

##################################################

# POS dependency labels (see: https://spacy.io/api/annotation)

NLP_MODEL_DEFAULT = 0
NLP_MODEL_LARGE   = 1

TOPIC_FILTER_LEFT       = 0     # LEFT
TOPIC_FILTER_LEFT_RIGHT = 1     # LEFT + RIGHT
TOPIC_FILTER_ALL        = 2     # ALL

TOPIC_SORT_APPEARANCE   = 0
TOPIC_SORT_LEFT_COUNT   = 1

# Default font info used to create HTML strings.
default_font_info = {'Title':          24,
                     'Normal':         14,
                     'List Bullet':    14,
                     'List Bullet 2':  14,
                     'List Bullet 3':  14,
                     'List Number':    14,
                     'List Number 2':  14,
                     'List Number 3':  14,
                     'List Paragraph': 14,
                     'Heading 1':      22,
                     'Heading 2':      20,
                     'Heading 3':      18,
                     'Heading 4':      16}

pp = pprint.PrettyPrinter(indent=4)            # create a pretty printing object used for debugging.

TEXT_COLOR      = RGBColor( 64,  64,  64)
TEXT_VERB_COLOR = RGBColor(  0, 188, 242)
TEXT_NP_VCOLOR  = RGBColor(125,   0, 125)

#
# Constants for sentence['text_w_info']. Each word is represented using a tuple with
# the following elements.
# TODO: Create a enumeration class to represent this. The indexes have grown too long.
#
POS      = 0          # Part of Speech
WORD     = 1          # Original Word
LEMMA    = 2          # Lemmatized word
ISLEFT   = 3          # True if the word is on the left side of the main verb
SUBJ     = 3          # True if the word is on the left side of the main verb (SUBJ = ISLEFT) same thing!!
DEP      = 4          # Dependency Type (e.g., 'pobj', 'nsubj', etc.)
STEM     = 5          # Stem of the word
LINKS    = 6          # If it is a verb, it is the total # of links from the verb. Otherwise, None.
QUOTE    = 7          # QUOTED or not
WORD_POS = 8          # Indexes between 8 and 13 have been changed on July 17.
DS_DATA  = 9          # DocuScope Data (i.e., models.DSWord)
PARA_POS = 10 #  9
SENT_POS = 11 # 10
NEW      = 12 # 11
GIVEN    = 13 # 12
IS_SKIP  = 14 # 13
IS_TOPIC = 15 # 14


IGNORE_ADVCL_FIRST_WORDS = ['after', 'before', 'until', 'soon', 'once', 'now',
                            'during', 'while', 'when', 'whenever',
                            'if', 'whether', 'provided', 'in', 'unless', 'even',
                            'because', 'as', 'since', 'so', 'isasmuch',
                            'where', 'although', 'though',
                            'thanks', 'based', 'to', 'based']

nlp = None
LOAD_TRAINED_MODEL = False  # If True, will check for the existence of models in the filesystem.

##
# https://spacy.io/api/top-level#spacy.info
##
def setLanguageModel(lang, model=NLP_MODEL_DEFAULT):

    global nlp

    try:
        if lang == 'en':
            default_model = "data/default_model" # resource_path("data/default_model")
            large_model = "data/large_model" # resource_path("data/large_model")

            if LOAD_TRAINED_MODEL and model == NLP_MODEL_DEFAULT and os.path.exists(default_model):
                logging.info("Loading Spacy default model ...")
                nlp = spacy.load(default_model)
            elif LOAD_TRAINED_MODEL and os.path.exists(large_model):
                logging.info("Loading Spacy large model ...")
                nlp = spacy.load(large_model)
            else:
                nlp = spacy.load("en_core_web_sm")

    except Exception as e:
        logging.error(e)
    else:
        logging.info("Spacy language model loaded successfully")
        logging.info(spacy.info())

##
#
##
def isModelLoaded():
    return nlp is not None

def removeQuotedSlashes(s):
    if s.find('“') > 0 and s.find('”') > 0:
        pattern = r'\s\/\s'
        res = regex.sub(pattern, ' ', s)
        return res
    elif s.count('\"') == 2:
        pattern = r'\s\/\s'
        res = regex.sub(pattern, ' ', s)
        return res
    else:
        return s

def remove_hrefs(text):
    """
    This function removes a URL from a string.
    """
    utlpattern = r"(https|http):[A-z0-9/.-]+"  # UTL regex pattern
    for m in re.finditer(utlpattern, text):    # find one or more URL patterns,
        text = text.replace(m.group(), "")     # remove the URL from the message.
    return text

def inchesToDoubleQuotes(p):
    """
    Helper function: Given a paragraph, replace all the inche
    marks (i.e., '\"') to actual smart/curly quotes.
    """

    buf = ""

    bQuote = False
    for c in p:
        if bQuote == False and c == "\"": # open
            buf += "\u201C"
            bQuote = True
        # elif bQuote == False and c == "\u201C":
            # buf += c
            # bQuote = True
        elif bQuote == True and c == "\"": # close
            buf += "\u201D"
            bQuote = False
        # elif bQuote == False and c == "\u201D":
            # buf += c
            # bQuote = False
        else:
            buf += c

    return buf

##
#
##
def adjustSpaces(text):
    """
    This function takes a string, often directly read from a file, and cleans up the string for parsing.
    It's only been tested with English texts.
    """

    text = text.strip()                         # trim

    text = remove_hrefs(text)

    # Just in case a period is found in front of a word (e.g, ".And")  This handles
    # a case where 4 dot ellipses is used (3 dot ellipses followed by a period
    # followed by no space.)
    text = re.sub(r'\.([a-zA-Z]+)(?!\.)', r'. \1', text)

    # If there is a comma followed by a character (i.e., missing a sapce), add a space.
    text = re.sub(',(=?[a-zA-Z])',  r', \1', text)

    # Remove the image HTLM elements.
    text = re.sub(r'\<img.+>', '', text)

    # Replace the space character(s) in multiword patterns with underscor character "_"
    for mw_topic in DSDocument.multiword_topics:
        connected_mw_topic = mw_topic.replace(' ', '_')
        text = text.replace(mw_topic, connected_mw_topic)
        text = text.replace(mw_topic.lower(), connected_mw_topic.lower())
        text = text.replace(mw_topic.title(), connected_mw_topic.title())
        text = text.replace(mw_topic.capitalize(), connected_mw_topic.capitalize())

    text = text.replace(u'\u200b', '')               # remove the zero width space
    text = text.replace(u'\t', u' ')                 # sorry. we don't do tabs. replace it with a space.
    text = text.replace(u'\xa0', u' ')               # replace non-breaking spaces with a regular space.
    text = text.replace(u'\u2013', u' \u2014 ')      # replace en-dashes with em-dashes.
                                                     # not the best way to deal with en-dashes, though.

    text = text.replace(u' cannot ', u' can not ')   # replace cannot with can not.
    text = text.replace(u' cannot, ', u' can not, ') # replace cannot with can not.

    text = text.replace(u'\u2014', u' \u2014 ')      # always add a space before/after an em-dash
    text = text.replace(u'--', u' \u2014 ')          # replace double dashes with an em-dash.

    text = text.replace(u'/ ', u' ')                 # a slash followed by a space is replaced by a space

    text = text.replace(u' )', u')')                 # remove a space followed by a close parentesis
    text = text.replace(u'( ', u'(')                 # remove a space followed by a open parentesis

    text = text.replace(u'. . .', u' \u2026')        # convert the MLA style ellipsis to an ellipsis character
    text = text.replace(u'...', u' \u2026')          # convert 3 dots to an ellipsis character

    text = text.replace(u'??', u'?').replace(u'??', u'?')    # double/triple punctuations are not allowed
    text = text.replace(u'!!', u'!').replace(u'!!', u'!')    # double/triple punctuations are not allowed
    text = text.replace(u',,', u',').replace(u',,', u',')    # double/triple punctuations are not allowed

    text = inchesToDoubleQuotes(text)          # convert inch characters (") to curly double-quotes.
    text = removeQuotedSlashes(text)

    text = text.replace(u'  ', u' ').replace(u'  ', u' ').replace(u'  ', u' ') # remove extra spaces

    return text

def is_skip(elem, left_count, topic_filter):

    # if theme_only == True:
    if topic_filter == TOPIC_FILTER_LEFT:

        # if the left only mode is on
        if not elem[IS_TOPIC]:
            # skip, if it is not a topic word
            return True # skip

        elif left_count < 1 and not elem[ISLEFT]:
            # skip if it's a right-side word and left_count < 1 (==0)
            return True # skip

        else:
            return False            # otherwise, it's not a skip word.
    else:
        return False

def contains_html_tags(text):
    # Regular expression to match HTML tags
    pattern = r'<[^>]+>'

    # Search for the pattern in the text
    match = re.search(pattern, text)

    # Return True if a match is found, False otherwise
    return match is not None

class DSWord():
    def __init__(self, w, lw, end, lat, pos):

        # strip double quoted numbers
        x = re.search('(?<=\")[0-9,.]+(?=\")', w)
        if x:
            self.word  = x.group()   # original word
        else:
            self.word = w

        x = re.search('(?<=\")[0-9,.]+(?=\")', lw)
        if x:
            self.lword = x.group()   # lower-case word
        else:
            self.lword = lw

        self.end   = end
        self.lat   = lat
        self.pos   = pos

    def getWord(self):
        return self.word

    def getLower(self):
        return self.lword

    def getEnd(self):
        return self.end

    def getLAT(self):
        return self.lat

    def getPos(self):
        return self.pos

class AdjacencyStats():
    """
    The AdjacencyStats object is used to reprsent the stats for
    for each paragraph (or any unit).
    """

    def __init__(self, topic="", controller=None):
        self.topic           = topic
        self.controller      = controller
        self.topic_paras     = list()
        self.topic_sents     = list()
        self.topic_positions = list()

        self.clust_counter   = Counter()
        self.dim_counter     = Counter()
        self.lat_counter     = Counter()

        self.para_stats      = dict()
        self.sent_stats      = dict()

    def getTopic(self):
        return self.topic

    def addParagraphID(self, p_id):
        if p_id not in self.topic_paras:
            self.topic_paras.append(p_id)

    def addSentenceID(self, p_id, s_id):
        if (p_id, s_id) not in self.topic_sents:
            self.topic_sents.append((p_id, s_id))

    def addTopicPosition(self, p_id, s_id, t_pos):
        if (p_id, s_id, t_pos) not in self.topic_positions:
            self.topic_positions.append((p_id, s_id, t_pos))

    def addLAT(self, lat, p_id, s_id):

        if lat in ('UNRECOGNIZED', ''):
            return
        if self.controller is None:
            return
        dim, clust = self.controller.getDimensionAndCluster(lat)

        if clust is None or dim is None:
            return

        self.clust_counter[clust] += 1
        self.dim_counter[dim]     += 1
        self.lat_counter[lat]     += 1

        stat = self.para_stats.get(p_id, None)
        if stat is None:
            stat = ds_stat.DSStat()
            self.para_stats[p_id] = stat

        stat.addCategory(clust, dim, lat)

        # Keep track of the sentence level statistics
        sent_stat = self.sent_stats.get((p_id, s_id), None)
        if sent_stat is None:
            sent_stat = ds_stat.DSStat()
            self.sent_stats[(p_id, s_id)] = sent_stat
        sent_stat.addCategory(clust, dim, lat)

    def setText(self, text, p_id):
        stats = self.para_stats.get(p_id, None)
        if stats is not None:
            stats.setText(text)

    def getParaClusterCount(self, clust, p_id):
        stats = self.para_stats.get(p_id, None)
        if stats is None:
            return None
        else:
            res = stats.getClusterCount(clust)
            if res is None:
                return 0
            else:
                return res

    def getParaClusterFreq(self, clust, p_id, method=ds_stat.DSStat.DOC):
        stats = self.para_stats.get(p_id, None)
        if stats is None:
            return None
        else:
            res = stats.getClusterFreq(clust, method=method)
            if res is None:
                return 0
            else:
                return res

    def getParaDimensionCount(self, dimension, p_id):
        stats = self.para_stats.get(p_id, None)
        if stats is None:
            return None
        else:
            res = stats.getDimensionCount(dimension)
            if res is None:
                return 0
            else:
                return res

    def getParaDimensionFreq(self, dimension, p_id, method=ds_stat.DSStat.DOC):
        stats = self.para_stats.get(p_id, None)

        if stats is None:
            return None
        else:
            res = stats.getDimensionFreq(dimension, method=method)
            if res is None:
                return 0
            return res

    def getParaStats(self):
        return self.para_stats

    def getSentStats(self):
        return self.sent_stats

    def getStats(self, p_id):

        if p_id == -1:
            return self.clust_counter, self.dim_counter, self.lat_counter
        else:
            clust_counter = Counter()
            for clust in self.clust_counter.keys():
                clust_c = self.getParaClusterCount(clust, p_id)
                if clust_c is not None:
                    clust_counter[clust] = clust_c

            dim_counter = Counter()
            for dim in self.dim_counter.keys():
                dim_c = self.getParaClusterCount(dim, p_id)
                if dim_c is not None:
                    dim_counter[dim] = dim_c

            lat_counter = Counter()
            for lat in self.lat_counter.keys():
                lat_c = self.getParaClusterCount(lat, p_id)
                if lat_c is not None:
                    lat_counter[lat] = lat_c

            return clust_counter, dim_counter, lat_counter

    # def getParaSentIDs(self):
        # return self.topic_paras, self.topic_sents

    def getTopicPositions(self):
        return self.topic_positions

    def getParaSentIDs(self, clust_filter=None):
        if clust_filter is None:
            return self.topic_paras, self.topic_sents
        else:
            new_topic_paras = list()
            new_topic_sents = list()
            for key in self.topic_sents:   # key = tuple like (para_id, sent_id)

                stat = self.sent_stats.get(key, None)
                if stat is not None:
                    cluster_count = stat.getClusterCount(clust_filter)
                    if cluster_count is not None and cluster_count > 0:
                        new_topic_sents.append(key)
                        new_topic_paras.append(key[0])

            new_topic_paras = list(set(new_topic_paras))

            return new_topic_paras, new_topic_sents

    def isTopicSent(self, sent_id, para_id):
        return (para_id, sent_id) in self.topic_sents

    def isTopicPara(self, para_id):
        return para_id in self.topic_paras

    def print(self):
        print("AdjacencyStats:")
        print("topic =", self.topic)
        pp.pprint(self.clust_counter)
        pp.pprint(self.dim_counter)
        pp.pprint(self.lat_counter)
        print(self.topic_paras)
        print(self.topic_sents)
        print("-------------------------------")


SUBJECTS = ["nsubj", "nsubjpass", "csubj", "csubjpass", "agent", "expl"]

class DSDocument():
    """
    This class is used to holds the textual data.
    """

    ##########################################
    # Class variables/methods
    ##########################################

    noun_subject_options     = ["nsubj", "nsubjpass", "csubj", "csubjpass", "expl", "agent"]

    multiword_topics         = []
    deleted_multiword_topics = []

    user_defined_synonyms      = None
    user_defined_synonym_names = []
    user_defined_topics        = []

    @classmethod
    def setNounSubjectOptions(cls, options):
        DSDocument.noun_subject_options = options

    @classmethod
    def setUserDefinedSynonyms(cls, synsets):
        # convert a list of synonyms into a dictionary for faster look up
        if synsets == [] or synsets == None:
            DSDocument.user_defined_synonyms = None
            return

        DSDocument.user_defined_synonyms = dict()
        DSDocument.user_defined_synonym_names = list()
        undefined_count = 0

        for synset in synsets:
            lemma = synset.getLemma_()
            synonyms = synset.getSynonyms()

            for synonym in synonyms:
                s = synonym.replace(' ', '_')

                DSDocument.user_defined_synonyms[s] = lemma
                DSDocument.user_defined_synonyms[s.lower()] = lemma
                DSDocument.user_defined_synonyms[s.capitalize()] = lemma
                DSDocument.user_defined_synonyms[s.title()] = lemma

            if synonyms == []:
                key = "undefined_{}"
                DSDocument.user_defined_synonyms[key] = lemma   # undefined synonym
                undefined_count+=1

            DSDocument.user_defined_synonym_names.append(lemma)

    @classmethod
    def isUserDefinedSynonym(cls, lemma):
        if DSDocument.user_defined_synonyms is not None and lemma in DSDocument.user_defined_synonym_names:
            return True
        else:
            return False

    @classmethod
    def isUserDefinedSynonymDefined(cls, lemma):
        if DSDocument.user_defined_synonyms is not None and lemma in DSDocument.user_defined_synonym_names:
            for synonym, name in DSDocument.user_defined_synonyms.items():
                if name == lemma:
                    if synonym.startswith('undefined_'):
                        return False
                    else:
                        return True
            return False
        else:
            return False

    @classmethod
    def setMultiwordTopics(cls, topics):
        if topics is None:
            DSDocument.multiword_topics = []
            DSDocument.deleted_multiword_topics = []
        else:
            topics.sort(reverse=True, key=lambda x: len(x.split()))
            DSDocument.deleted_multiword_topics = list(set(DSDocument.multiword_topics) - set(topics))
            DSDocument.multiword_topics = topics

    @classmethod
    def setUserDefinedTopics(cls, topics):
        DSDocument.user_defined_topics = [t.lower().replace(' ', '_') for t in topics]

    @classmethod
    def isUserDefinedTopic(cls, lemma):
        if lemma in DSDocument.user_defined_topics:
            return True
        else:
            return False

    @classmethod
    def addUserDefinedTopic(cls, topic):
        if topic is not None and topic != '':
            lc_topic = topic.lower()
            if lc_topic not in DSDocument.user_defined_topics:
                DSDocument.user_defined_topics.append(lc_topic)

    @classmethod
    def deleteUserDefinedTopic(cls, topic_to_delete):
        if topic_to_delete in DSDocument.user_defined_topics:
            DSDocument.user_defined_topics.remove(topic_to_delete)

    @classmethod
    def clearUserDefinedTopics(cls):
        DSDocument.user_defined_topics = []

    @classmethod
    def getUserDefinedTopics(cls):
        return DSDocument.user_defined_topics

    @classmethod
    def getUserDefinedSynonyms(cls):
        return DSDocument.user_defined_synonym_names

    @classmethod
    def clearUserDefinedSynonyms(cls):
        DSDocument.user_defined_synonym_names = []

    ##########################################
    # Instance methods
    ##########################################

    def __init__(self):

        self.controller = None
        self.dictionary = None

        # Data & Stats
        self.stats    = None         # Stats (dictionary)
        self.sections = None         # Document Data
        self.current_section = 0

        self.filename = None
        self.header_labels = []

        # Visualizer options
        self.noun = True
        self.verb = False
        self.adj  = False
        self.adv  = False
        self.prp  = False

        self.max_paras        = 500
        self.skip_paras       = 0

        self.is_collapsed         = True

        self.total_words          = 0

        self.para_accum_mode      = True
        self.sent_accum_mode      = True

        self.selectedTopic        = (None, None)
        self.selectedTopicCluster = []
        self.keyTopics            = []

        self.keyParaTopics        = []
        self.keySentTopics        = []

        self.selectedSent         = -1
        self.selectedWord         = None
        self.showKeyTopics        = False
        self.showKeySentTopics    = False

        self.selected_sent_info   = None

        self.bQuote = False        # temp variable

        self.progress_callback = None

        self.img_count = 0

        self.global_topical_prog_data = None

        # variables used by the methods for the online version of write & audit
        self.local_topics_dict = None
        self.global_header           = []
        self.local_header            = []
        self.global_topics           = []
        self.local_topics            = []
        self.topic_location_dict     = None

        self.lexical_overlaps_analyzed = False

    def setController(self, c):
        self.controller = c

    def setDictionary(self, d):
        self.dictionary = d

    def setProgressCallback(self, cb_func):
        self.progress_callback = cb_func

    def clearData(self):
        self.sections = None

    def setFilters(self, vis_mode, max_paras, skip_paras):
        """
        Update the visualizer options.
        """
        self.verb             = False
        self.noun             = True
        self.adj              = False
        self.adv              = False

        self.para_accum_mode  = True
        self.sent_accum_mode  = True

        self.max_paras        = max_paras
        self.skip_paras       = skip_paras

    ########################################
    #
    # Setting Methods
    #
    ########################################

    def setHeaderLabels(self, val):
        self.header_labels = val

    def getHeaderLabels(self, val):
        return self.header_labels

    def setSection(self, val):
        self.current_section = val

    def setPronounsVisible(self, val):
        self.prp = val

    def setUserTopics(self, user_topics):
        res = []
        if nlp is not None:
            doc = nlp(user_topics)
            for t in doc:
                res.append(t.lemma_)
        self.userTopics = res

    def setKeySentTopics(self, keytopics):
        self.keySentTopics = keytopics

    def setKeyParaTopics(self, keytopics):
        self.keyParaTopics = keytopics
        self.keyTopics     = keytopics

    def setSelectedTopic(self, topic):
        if topic is not None:
            self.selectedTopic = topic
        else:
            self.selectedTopic = (None,None)

    def setSelectedTopicCluster(self, topic_cluster):
        if topic_cluster is not None:
            self.selectedTopicCluster = topic_cluster
        else:
            self.selectedTopicCluster = []

    def setTextWrap(self, val):
        self.text_wrap_mode = val

    def setSelectedSentence(self, sent_id):
        if sent_id < 0:
            self.selectedSent = None
        else:
            self.selectedSent = sent_id

    def setSelectedWord(self, selected_word):
        self.selectedWord = selected_word    # word is a tuple = elem

    # def setSynonymThreshold(self, val):
        # self.synonym_threshold = val

    def setShowKeyTopics(self, val):
        self.showKeyTopics = val

    def setShowKeySentTopics(self, val):
        self.showKeySentTopics = val

    def setCollapsed(self, val):
        self.is_collapsed = val


    ########################################
    #
    # Query Mthods
    #
    ########################################

    def userTopicsDefined(self):
        if len(self.userTopics) > 0:
            return True
        else:
            return False

    def isUserTopic(self, topic):
        if len(topic) == 0:
            return True
        if topic in self.userTopics:
            return True
        else:
            return False

    def isSelectedTopic(self, lemma):
        for w in self.selectedTopic[1:]:
            if w == lemma:
                return True
        return False

    def isPronounsVisible(self):

        return self.prp

    def isPOSVisible(self, pos):
        """
        Given a POS tag, returns a WordNet POS tag (one of 4 tags), or None.
        """
        if pos.startswith("NOUN") or pos.startswith("PROPN"):
            # res = self.noun
            return self.noun
        elif pos.startswith("VERB"):
            # res = self.verb
            return self.verb
        elif pos.startswith("ADJ"):
            # res = self.adj
            return self.adj
        elif pos.startswith("ADV"):
            # res = self.adv
            return self.adv
        elif pos.startswith("PRP"):
            # res = self.prp      # pronoun, personal
            return self.prp
        elif pos.startswith("PRON"):
            res = True

        return False

    ########################################
    #
    # Get Methods
    #
    ########################################

    def getCurrentSection(self):
        if self.sections is not None and self.current_section < len(self.sections):
            return self.sections[self.current_section]
        else:
            return None

    # def getSections(self):
    #     if self.sections is not None:
    #         return self.sections
    #     else:
    #         return []

    # def getSectionCount(self):
    #     if self.sections is not None:
    #         return len(self.sections)
    #     else:
    #         return 0

    # def getSectionAt(self, i):
    #     if self.sections is not None and i < len(self.sections):
    #         return self.sections[i]
    #     else:
    #         return None

    # def getCurrentTagDicts(self):
    #     section = self.getCurrentSection()
    #     if section is not None:
    #         return section.get('tag_dicts', None)
    #     else:
    #         return None

    # def getWordCount(self):
    #     section = self.getCurrentSection()
    #     if section is not None:
    #         return section.get('num_tokens', None)
    #     else:
    #         return None

    # def getParagraphs(self):
    #     data = self.sections[self.current_section]['data']
    #     return data['paragraphs']

    # def getRealParaCount(self):
    #     data = self.sections[self.current_section]['data']
    #     return data['real_para_count']

    # def getParaCount(self):
    #     data = self.sections[self.current_section]['data']
    #     return len(data['paragraphs'])

    # def getSentCount(self):
    #     data = self.sections[self.current_section]['data']
    #     sent_count = 0
    #     for p in data['paragraphs']:
    #         sent_count += len(p['sentences'])
    #     return sent_count

    # def getSentCounts(self):
    #     res = list()
    #     data = self.sections[self.current_section]['data']
    #     num_paras = len(data['paragraphs'])
    #     for i in range(num_paras):
    #         p = data['paragraphs'][i]
    #         num_sents = len(p['sentences'])
    #         res.append(num_sents)
    #     return res

    # def getParaMenuItems(self, num_words=5):
    #     """
    #     Create a list of menu labels for the paragraphs in the text.
    #     The first 'num_words' words from each paragraph are used for the label.
    #     """
    #     res = list()
    #     pcount = 1
    #     data = self.sections[self.current_section]['data']
    #     for p in data['paragraphs']:        # for each paragraph
    #         res.append("\u00B6{} {} ...".format(pcount, ' '.join(p['text'].split()[0:num_words])))
    #         pcount += 1
    #     return res

    def getNumParagraphs(self):
        if self.sections is None:
            return 0
        data = self.sections[self.current_section]['data']
        if 'paragraphs' in data:
            return len(data['paragraphs'])
        return 0

    # def getNumSections(self):
    #     if self.sections:
    #         return len(self.sections)
    #     else:
    #         return 0

    # def getKeyTopics(self):
    #     return self.keyTopics

    # def getKeyTopicsFirstSent(self):
    #     return self.keyTopicsFirstSent

    def getKeyParaTopics(self):
        return self.keyParaTopics

    # def setHighlightQuotes(self, val):
    #     self.show_quotes = val

    # def getNumQuotedWords(self):
    #     return self.num_quoted_words

    # def getTotalWords(self):
    #     return self.total_words

    # def getFilename(self):
    #     if self.filename is not None:
    #         return self.filename
    #     else:
    #         return 'Undefined'

    # def getSelectedSentInfo(self):
    #     return self.selected_sent_info

    def findSentences(self, ruleset):

        if self.sections is None:
            return "Error in findSentences()"
        else:
            try:
                data = self.sections[self.current_section]['data']
                if data is None:
                    raise ValueError
            except:
                logging.warning(self.sections[self.current_section])
                return []

        res = []
        pcount = 1
        for para in data['paragraphs']:

            scount = 1
            for sent in para['sentences']:
                b_topic = False
                b_experience = False

                # get the list of lemmas (i.e., topics)
                lemmas   = sent['lemmas']

                # get the list of clusters (i.e., experiences)
                clusters = list()
                for w in sent['text_w_info']:
                    dsw   = w[DS_DATA]      # docuscope word
                    if dsw is not None:
                        lat   = dsw.getLAT()
                        dim, clust = self.dictionary.getDimensionAndCluster(lat)
                        if clust is not None and clust not in clusters:
                            clusters.append(clust)
                    else:
                        logging.warn("ERROR: dsw is none.")
                        continue
                res = dict
                for rule in ruleset.getRules():

                    b_match = rule.isMatching(lemmas, clusters)

                    if res.get(rule, False) and sent not in res[rule]:
                        res[rule].append(sent)

                scount += 1

            pcount += 1

        """
        Note. We are returning a python dictionary, where keys are rules, and the values are sentences.
        """
        return res

    ########################################
    #
    # Processing Methods
    #
    ########################################

    def processSent(self, sent, start=0):
        """
        Given a sentence 'sent' in a string format, return a list [] of analyzed words.
        Each entry in the list is a tuple (POS, Word, Lemma, bLeft_of_the_Main_Verb).
        e.g., ('NOUN', 'dogs', 'dog', False)
        """

        word_pos = start

        def processHeading(heading):
            nonlocal word_pos
            res = list()
            doc = nlp(heading)

            # t = 0.POS, 1.WORD, 2.LEMMA, 3.ISLEFT,
            #     4.DEP, 5.STEM, 6.LINKS, 7.QUOTE, 8.WORD_POS, 9.DS_DATA
            # Note: the nouns in headings/titels are always considered 'pre-verb' ISLEFT + TRUE
            is_left = False
            if DSDocument.user_defined_synonyms is not None:
                for token in doc:

                    if token.pos_ == 'NOUN' or token.pos_ == 'PROPN':     # if the token is a noun
                        is_left = True
                        t = token.text
                        lemma = DSDocument.user_defined_synonyms.get(t, None)

                        if lemma is None:
                            lemma = token.lemma_.lower()

                    elif token.tag_ == 'PRP' or token.tag_ == 'PRP$':   # if the token is a pronoun
                        is_left = True
                        t = token.text
                        lemma = DSDocument.user_defined_synonyms.get(t, None)

                        if lemma is None:
                            lemma = token.lemma_.lower()

                    else:
                        t = token.text.lower()
                        lemma = getPronounLemma(t)
                        if lemma is None:
                            lemma = token.lemma_.lower()

                    if token.text in end_puncts:
                        pos = 'PUNCT'
                    elif token.pos_ == 'PROPN':                            # Proper Nouns
                        pos = 'NOUN'
                    elif token.tag_ == 'PRP' or token.tag_ == 'PRP$':      # Pronouns / Pronoun Possessive
                        pos = 'PRP'
                    elif token.is_punct:                                   # Punctuations
                        pos = 'PUNCT'
                    elif token.pos_ == 'PRON' and token.tag_ == 'NN':
                        pos = 'NOUN'
                    else:                                                  # Everything else
                        pos = token.pos_

                    res.append((pos, token.text, lemma, is_left,
                               token.dep_, '', None, False, word_pos, None))
                    word_pos += 1
            else:
                for token in doc:
                    # t = 0.POS, 1.WORD, 2.LEMMA, 3.ISLEFT,
                    #     4.DEP, 5.STEM, 6.LINKS, 7.QUOTE, 8.WORD_POS, 9.DS_DATA
                    # Note: the nouns are always considered 'pre-verb' ISLEFT + TRUE
                    is_left = False
                    if token.pos_ in ['NOUN', 'PROPN', 'PRON']:
                        is_left = True

                    if token.text in end_puncts:
                        pos = 'PUNCT'
                    elif token.pos_ == 'PROPN':                            # Proper Nouns
                        pos = 'NOUN'
                    elif token.tag_ == 'PRP' or token.tag_ == 'PRP$':      # Pronouns / Pronoun Possessive
                        pos = 'PRP'
                    elif token.is_punct:                                   # Punctuations
                        pos = 'PUNCT'
                    elif token.pos_ == 'PRON' and token.tag_ == 'NN':
                        pos = 'NOUN'
                    else:                                                  # Everything else
                        pos = token.pos_

                    res.append((pos, token.text, token.lemma_.lower(), is_left,
                               token.dep_, '', None, False, word_pos, None))
                    word_pos += 1

            temp = []
            for n in doc.noun_chunks:
                n = str(n)
                n = n.translate(n.maketrans('', '', string.punctuation))     # removes punctuation
                temp.append(n)

            noun_phrases = tuple(temp)

            return res, noun_phrases


        def getPronounLemma(t):
            lemma = None
            if t in ['he', 'his', 'him', 'himself']:
                lemma = 'he'
            elif t in ['she', 'her', 'hers', 'herself']:
                lemma = 'she'
            elif t in ['i', 'my', 'me', 'mine', 'myself']:
                lemma = 'I'
            elif t in ['you', 'your', 'yours', 'yourself', 'yourselves']:
                lemma = 'you'
            elif t in ['we', 'our', 'us', 'ours', 'ourselves']:
                lemma = 'we'
            elif t in ['they', 'their', 'them', 'theirs', 'themselves']:
                lemma = 'they'
            elif t in ['it', 'its', 'itself']:
                lemma = 'it'

            return lemma

        # TODO: We should not do this here... This shold be passed to this method as an argument...
        if self.controller:
            incl_nsubj = self.controller.postMainVerbTopics() # true or false
        else:
            incl_nsubj = True # default

        res = list()
        is_left = True
        if type(sent) != str:
            doc = sent
            root = None

            for token in doc:
                lemma = ""

                if DSDocument.user_defined_synonyms is not None:
                    if token.pos_ == 'NOUN' or token.pos_ == 'PROPN':     # if the token is a noun

                        t = token.text.lower()                                 # original spelling
                        lemma = DSDocument.user_defined_synonyms.get(t, None)  #

                        if lemma is None:
                            t = token.lemma_.lower()                           # SpaCy's lemma
                            lemma = DSDocument.user_defined_synonyms.get(t, None)

                        if lemma is None:
                            lemma = token.lemma_.lower()

                    elif token.tag_ == 'PRP' or token.tag_ == 'PRP$':
                        t = token.text.lower()
                        lemma = DSDocument.user_defined_synonyms.get(t, None)   # e.g., He

                        if lemma is None:
                            lemma = getPronounLemma(t)  # e.g., his
                            if lemma is None:
                                lemma = token.lemma_.lower()
                        else:
                            token.pos_ = 'NOUN'
                            token.tag_ = 'NN'       # if it is a synonym, treat them like a noun

                    else:
                        t = token.text.lower()
                        lemma = DSDocument.user_defined_synonyms.get(t, None)   # non-noun phrases included in a topic cluster

                        if lemma is None:
                            lemma = getPronounLemma(t)
                            if lemma is None:
                                lemma = token.lemma_.lower()
                        else:
                            token.pos_ = 'NOUN'
                            token.tag_ = 'NN'

                elif token.tag_ == 'PRP' or token.tag_ == 'PRP$':
                    t = token.text.lower()
                    lemma = getPronounLemma(t)
                    if lemma is None:
                        lemma = token.lemma_.lower()
                else:
                    t = token.text.lower()
                    lemma = getPronounLemma(t)
                    if lemma is None:
                        lemma = token.lemma_.lower()

                # temporary fix about 'data'
                if lemma == "datum":
                    lemma = "data"

                if token.text in end_puncts:
                    pos = 'PUNCT'
                elif token.pos_ == 'PROPN':                            # Proper Nouns
                    pos = 'NOUN'
                elif token.tag_ == 'PRP' or token.tag_ == 'PRP$':      # Pronouns / Pronoun Possessive
                    pos = 'PRP'
                elif token.is_punct:                                   # Punctuations
                    pos = 'PUNCT'
                elif token.pos_ == 'PRON' and token.tag_ == 'NN':
                    pos = 'NOUN'
                else:                                                  # Everything else
                    pos = token.pos_

                # t = 0.POS, 1.WORD, 2.LEMMA, 3.ISLEFT, 4.DEP, 5.STEM,
                # 6.LINKS (for ROOT), 7.QUOTE, 8.WORD_POS, 9.DS_DATA,
                # 10.PARA_POS, 11.SENT_POS, 12.NEW, 13.GIVEN, 14.IS_SKIP, 15.IS_TOPIC
                # 10 thr 15 are added later.

                if incl_nsubj and token.dep_ in DSDocument.noun_subject_options:
                    # if incl_nsubj is True (user option), and the POS tag is in
                    # the options selected by the user
                    left = True
                elif lemma.lower() in DSDocument.user_defined_topics:
                    left = True
                else:
                    left = is_left

                res.append((pos, token.text, lemma, left, token.dep_, '',
                            None, False, word_pos, None))

                if token.dep_ == "ROOT":
                    is_left = False

                word_pos += 1

            temp = []
            for n in doc.noun_chunks:
                n = str(n)
                n = n.translate(n.maketrans('', '', string.punctuation))     # removes punctuation
                temp.append(n)
            noun_phrases = tuple(temp)

        else:
            res, noun_phrases = processHeading(sent)

        return res, noun_phrases, word_pos

    def analyzeSent(self, sent_dict, sent):
        """
        Perform a basic analysis of a given sentence in'sent_data'.
        """
        def traverse(token, is_left):
            l = None
            r = None
            res = None

            l = []
            for w in token.lefts:
                l += traverse(w, is_left)

            m = [token.text]

            r = []
            for w in token.rights:
                r += traverse(w, is_left)

            res = l + m + r

            return res   # end of traverse

        def isNP(d, i):
            if d is not None:
                for np in d:
                    if np[0] <= i and i < np[1]:
                        return True
            return False

        sent_data = sent_dict['text_w_info']

        res = {}         # create a dictionary

        res['NOUNS']     = 0   # total # of nouns
        res['HNOUNS']    = 0   # total # of head nouns
        res['L_HNOUNS']  = 0   # head nouns on the left
        res['R_HNOUNS']  = 0   # head nouns on the right

        res['L_NOUNS']   = []  # nouns on the left
        res['R_NOUNS']   = []  # nouns on the right

        res['MV_LINKS']  = 0   # total # of links fron the root verb
        res['MV_L_LINKS']= 0   # total # of left links from the root veb
        res['MV_R_LINKS']= 0   # total # of right links from the root verb

        res['V_LINKS']   = 0   # total # of links from all the non-root verbs
        res['V_L_LINKS'] = 0   # total # of left links from all the non-root verbs
        res['V_R_LINKS'] = 0   # total # of right links fromn all the non-root verbs

        res['NR_VERBS']  = 0   # total # of non-root verbs
        res['NPS']       = []
        res['NUM_NPS']   = 0   # total # of NPs
        res['L_NPS']     = 0
        res['R_NPS']     = 0
        res['BE_VERB']   = False
        res['HEADING']   = False

        word_count = 1
        root_pos = -1
        for data in sent_data:
            if data[POS] == 'NOUN' or data[POS] == 'PRP':
                res['NOUNS'] += 1

                if data[ISLEFT] == True:
                    res['L_NOUNS'].append(data)
                else:
                    res['R_NOUNS'].append(data)

            if data[DEP] in SUBJECTS:  # we may want to change SUBJECTS to DSDocument.noun_subject_options:
                res['HNOUNS'] += 1

                if data[ISLEFT] == True:
                    res['L_HNOUNS'] += 1
                else:
                    res['R_HNOUNS'] += 1

            if data[DEP] == 'ROOT':
                links = data[LINKS]
                if links is not None:
                    res['MV_LINKS']   = links['CCOUNT']
                    res['MV_L_LINKS'] = links['LCOUNT']
                    res['MV_R_LINKS'] = links['RCOUNT']

                if data[WORD] in be_verbs:
                    res['BE_VERB'] = True

                root_pos = word_count

            elif data[POS] == 'VERB':
                links = data[LINKS]
                if links is not None:
                    res['V_LINKS'] += links['CCOUNT']
                    res['V_L_LINKS'] += links['LCOUNT']
                    res['V_R_LINKS'] += links['RCOUNT']
                    res['NR_VERBS']  += 1

            word_count+=1


        text = sent_dict['text']
        if contains_html_tags(text):
            # Strip HTML tags, if any.
            soup = bs(text, "html.parser")
            text = soup.text

        text = text.strip().replace('\n', ' ').replace('   ', ' ').replace('  ', ' ')
        doc = nlp(text)

        if type(sent) == str:
            res['HEADING'] = True

        for token in doc:
            if token.dep_ == 'ROOT':
                root_pos = token.i
                break

        left_nps = list()
        for np in doc.noun_chunks:
            if (np.end-1) < root_pos: # LEFT
                res['L_NPS'] += 1
                left_nps.append(np)
            elif (np.end-1) > root_pos: # RIGHT
                res['R_NPS'] += 1
            res['NPS'].append(np.text)
        res['NUM_NPS'] = len(res['NPS'])

        # res['NOUN_CHUNKS'] = list(doc.noun_chunks)
        # 2022 May 9. Use a python dictionary instead of SpaCy's Span object. So that we can convert
        # the 'sent_analysis' property to JSON easily later.
        res['NOUN_CHUNKS'] = [ {'text': np.text, 'start': np.start, 'end': np.end} for np in doc.noun_chunks]
        res['TOKENS'] = [ {'text': token.text, 'is_root': token.dep_ == 'ROOT'} for token in doc]

        advcl_root = None
        for token in doc:
            if token.dep_ == 'ROOT':     # we are only intersted in the advcl before the main verb
                break
            elif token.dep_ == 'advcl':
                advcl_root = token
                break

        bIgnore = False
        if advcl_root is not None:
            left_span  = []
            right_span = []
            count = 0
            start = 0
            for w in advcl_root.lefts:
                if count==0:
                    start = w.i
                    if w.text.lower() in IGNORE_ADVCL_FIRST_WORDS or start != 0:
                        bIgnore = True
                        break

                left_span += traverse(w, True)
                count+=1
            if bIgnore == False:
                for w in advcl_root.rights:
                    right_span += traverse(w, False)

                tmp = left_span + [advcl_root.text] + right_span
                end = start + len(tmp)

                mod_cl = ' '.join(tmp)
                # Let's count how many NPs are in the modifier clause
                last_np = 0
                for np in left_nps:
                    if np.end <= end:
                        last_np = np.end
                    else:
                        break

                res['MOD_CL'] = (start, end, mod_cl, last_np)
            else:
                res['MOD_CL'] = None

        else:
            res['MOD_CL'] = None

        return res

    def clearGlobalTopicalProgDataCache(self):
        self.global_topical_prog_data = None # clear the cache

    def processDoc(self, section_data):
        """
        This function iterates through all the paragraphs in the given docx document.
        and find all the unique lemmas in each sentence, and in each paragraph.
        """

        if (isModelLoaded () == False):
            logging.warning("Warning: language model not loaded yet, loading ...")
            setLanguageModel ("en", NLP_MODEL_DEFAULT)

        if (isModelLoaded () == False):
            logging.error("Error: unable to load language model!")
            return list()

        logging.info("Language model appears to be loaded, processing text ...")

        self.global_topical_prog_data = None # clear the cache

        doc = section_data['doc']

        if self.progress_callback:
            self.progress_callback(max_val=20, msg="Preprocessing...")
        self.num_quoted_words = 0

        word_pos = section_data['start']

        # Optimized by only adding NOUNs. We ignore all the other POSs.
        def listLemmas(sent):
            global stop_words

            lemmas = list()
            temp = list()

            for w in sent:
                if w[POS] is not None and  w[LEMMA] not in stop_words:
                    if w[POS] == 'NOUN' or w[POS] == 'PRP':
                        if (w[POS], w[LEMMA]) not in temp:
                            temp.append( (w[WORD], w[POS], w[LEMMA]) )   # 'they' + 'NOUN' + 'man'
                            lemmas.append(w)
            return lemmas

        def listStems(sent):
            stems = list()
            temp = list()
            for w in sent:
                if w[POS] is not None and w[POS] != 'PUNCT' and w[POS] != 'SYM' and w[LEMMA] not in stop_words:
                    if w[STEM] not in temp:
                        temp.append(w[STEM])
                        stems.append(w)
            return stems

        # Optimization. We are only interested in nouns now.
        def accumulateParaLemmas(paragraphs):
            temp = list()
            accum = list()
            for para in paragraphs:
                lemmas = para['lemmas']
                for l in lemmas:
                    if l[POS] == 'NOUN':
                        if l[LEMMA] not in temp:
                            temp.append(l[LEMMA])
                            accum.append(l)
            return accum


        self.bQuote = False
        num_paras = len(doc.paragraphs)
        count_para  = 0

        data = dict()
        data['paragraphs'] = list()
        data['text'] = ""
        data['real_para_count'] = num_paras

        skip_paras = self.skip_paras

        for para in doc.paragraphs:                             # for each paragraph in the file.
            count_para += 1

            if count_para <= skip_paras:
                continue
            elif count_para > (skip_paras + self.max_paras):
                break

            if self.progress_callback:
                self.progress_callback(new_val=round(100 * count_para/num_paras))

            p = para.text

            if p.startswith("<img "): # skip if p is an image element.
                para_dict = dict()
                para_dict['text'] = p
                para_dict['sentences']          = list()

                para_dict['lemmas']             = list()
                para_dict['accum_lemmas']       = list()
                para_dict['given_lemmas']       = list()
                para_dict['new_lemmas']         = list()
                para_dict['given_accum_lemmas'] = list()
                para_dict['new_accum_lemmas']   = list()

                para_dict['style']              = 'Normal'
                data['paragraphs'].append(para_dict)
                continue

            if p == "":                       # skip if it is an empty line.
                continue

            if para.style.name not in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title',
                                       'List Bullet', 'List Bullet 2', 'List Bullet 3',
                                       'List Number', 'List Number 2', 'List Number 3']:
                style = 'Normal'
            else:
                style = para.style.name

            para_dict = dict()                                     # initalizes a dict for the paragraph
            para_dict['text']               = p                    # assign the paragraph obj to para_dict['text']
            para_dict['sentences']          = list()               # and initialize other fields with an empty list.

            para_dict['lemmas']             = list()
            para_dict['accum_lemmas']       = list()
            para_dict['given_lemmas']       = list()
            para_dict['new_lemmas']         = list()
            para_dict['given_accum_lemmas'] = list()
            para_dict['new_accum_lemmas']   = list()

            para_dict['style']              = style                 # set the paragraph style (from the docx file).

            if style in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6', 'Title']:
                slist = [para.text]
            else:
                parsed_para = nlp(p)
                slist = [sent for sent in parsed_para.sents]

            for s in slist:                                         # for each sentence in the paragraph 'p'
                sent_dict = dict()

                if type(s) == str:
                    sent_dict['text'] = s.strip()
                    sent_dict['text_w_info'], NPs, word_pos = self.processSent(s.strip(), start=word_pos+1)
                else:
                    sent_dict['text'] = s.text.strip()
                    sent_dict['text_w_info'], NPs, word_pos = self.processSent(s, start=word_pos+1)

                # sent_dict['sent_analysis'] = self.analyzeSent(sent_dict['text_w_info'], s) # analyze sentences
                sent_dict['sent_analysis'] = self.analyzeSent(sent_dict, s) # analyze sentences

                sent_dict['lemmas'] = listLemmas(sent_dict['text_w_info'])    # list lemmas in the paragraph
                sent_dict['accum_lemmas'] = []
                sent_dict['given_lemmas'] = []                  # initializes the 'given_lemmas' field.
                sent_dict['new_lemmas'] = []                    # initializes the 'new_lemmas' field.
                sent_dict['given_accum_lemmas'] = []            # initializes the 'given_lemmas' field.
                sent_dict['new_accum_lemmas'] = []              # initializes the 'new_lemmas' field.

                para_dict['sentences'].append(sent_dict)        # add sent_dict to para_dict.

                # update the paragraph's 'lemmans' field by adding new lemmas from the new sentence 's'
                for sl in sent_dict['lemmas']:                  # for each lemma in the new sentence,
                    match = False
                    for pl in para_dict['lemmas']:              # check against each lemma in its paragraph.
                        if sl[POS] == pl[POS] and sl[LEMMA] == pl[LEMMA]:  # if there is a match
                            match = True                                   # mark 'match' as True, and beak
                            break
                    if match != True:
                        para_dict['lemmas'].append(sl)

                sent_dict['accum_lemmas'] = list(para_dict['lemmas'])

            data['paragraphs'].append(para_dict)                # add para_dict to doc_dict

            # here, we need to create para_dict['accmum_lemmas'] by going through all the previous
            # paragraphs and accumulating lemmas.

            para_dict['accum_lemmas'] = accumulateParaLemmas(data['paragraphs'])

        section_data['data']    = data
        section_data['end_pos'] = word_pos

        # self.recalculateGivenWords(data=data) # 2025.07.03

    def recalculateGivenWords(self, data=None, section=-1):
        """
        (re)Calculate given words at the sentence and the paragraph level.
        """
        if data is None:
            if section < 0 or self.sections is None:
                # print("Error in recalculateGivenWords()")
                return
            else:
                data = self.sections[section]['data']

        if self.lexical_overlaps_analyzed:
            return

        if self.progress_callback:
            self.progress_callback(max_val=50, msg="Finding lexical overlaps between paragraphs...")

        self.findGivenWordsPara(data)               # find given words between paragraphs

        if self.progress_callback:
            self.progress_callback(max_val=80, msg="Finding lexical overlaps between sentences...")

        self.findGivenWordsSent(data)               # find given words between sentences in each paragraph

        self.clearGlobalTopicalProgDataCache()
        self.lexical_overlaps_analyzed = True

    def isGiven(self, l1, l2, pos1, pos2, pronoun=False):

        if pos1 == pos2:
            # these 2 lemmas share the same POS tag.
            if pronoun == False and pos1 == 'PRP':
                # we'll ignore this since they are pronouns and the pronoun display is OFF.
                return False
            elif l1 == l2:  # lemma 1 == lemma 2 and POS1 == POS1
                return True
            else:
                return False

        elif pronoun == True and pos1 == 'PRP':
            # The pronoun display is ON, and pos1 is a pronoun.
            # Pronouns are ALWAYS given.
            return True

        else:
            return False

    def findGivenWordsSent(self, data):
        """
        This method iterates through the sentences in each paragraph
        of the document, and find all the given and 'new' lemmas. Notice
        that this function only finds given words/lemmas between sentences
        within a given paragraph. (i.e., No given words are discovered
        between the last sentence of a paragraph and the first sentence
        of the next paragraph.)
        """
        num_paras = len(data['paragraphs'])
        count_para = 0
        for para in data['paragraphs']:                    # for each paragraph
            count_para += 1

            if self.progress_callback:
                self.progress_callback(new_val=round(100 * count_para/num_paras))

            prev_s = None
            for sent in para['sentences']:                      # for each sentence in the paragraph
                sent['given_accum_lemmas'] = list()
                sent['new_accum_lemmas']   = list()

                if prev_s is not None:                          # start with the 2nd sentence.
                    for gl in sent['lemmas']:                   # for each given lemmas in the sentence

                        for cl in prev_s['accum_lemmas']:
                            # for each given accum_lemmas in the prev sentence
                            # if a lemma ('gl') is in the previous paragraph AND their POSs maatch

                            # Note: POS is no longer checked for optmization.
                            # if self.isPOSVisible(gl[POS]) and self.isGiven(gl[LEMMA], cl[LEMMA], gl[POS], cl[POS], pronoun=True):
                            if self.isGiven(gl[LEMMA], cl[LEMMA], gl[POS], cl[POS], pronoun=self.prp):

                                # and if 'gl' is not already in the sentence's given lemma's list
                                if gl[LEMMA] not in [x[LEMMA] for x in sent['given_accum_lemmas']]:
                                    sent['given_accum_lemmas'].append(gl)            # add 'gl' to the list

                        # commented out 9/16/2021
                        for temp_sent in para['sentences']:
                            # for each sentence in the paragraph
                            if temp_sent == sent:  # break if temp_sent is the currente sent.
                                break
                            for cl in temp_sent['accum_lemmas']:  # for each accum_lemma in the sentence
                                # if a lemma ('gl') is in the sentence AND their POSs maatch
                                # Note: POS is no longer checked for optmization.
                                if self.isGiven(gl[LEMMA], cl[LEMMA], gl[POS], cl[POS], pronoun=self.prp):
                                    # if 'cl' is not already in the sentence's new lemmas list
                                    if gl[POS] != 'PRP' and cl[LEMMA] not in [x[LEMMA] for x in temp_sent['new_accum_lemmas']]:
                                        temp_sent['new_accum_lemmas'].append(gl)     # add 'cl' to the list ol

                        # if 'lemma' a user defined topic, force it to be 'given'
                        if gl[LEMMA].lower() in DSDocument.user_defined_topics:
                            if gl[LEMMA].lower() not in sent['given_accum_lemmas']:
                                sent['given_accum_lemmas'].append(gl)

                    prev_s = sent

                else:
                    prev_s = sent
                    for gl in sent['lemmas']:
                        if gl[LEMMA].lower() in DSDocument.user_defined_topics:
                            if gl[LEMMA].lower() not in sent['given_accum_lemmas']:
                                sent['given_accum_lemmas'].append(gl)

    def findGivenWordsPara(self, data):
        """
        This method iterates through the paragraphs in the data (Document),
        and find all the given and 'new' lemmas. (connection lemmmas are those words that
        appear in the following paragraph.
        """
        num_paras = len(data['paragraphs'])
        count_para = 0
        prev_p = None
        for para in data['paragraphs']:                         # for each paragraph
            count_para += 1

            if self.progress_callback:
                self.progress_callback(new_val=round(100 * count_para/num_paras))

            para['given_accum_lemmas'] = list()
            para['new_accum_lemmas']   = list()
            foo = list()
            if prev_p is not None:                      # start with the 2nd paragraph.
                for gl in para['lemmas']:               # for each lemma in the paragraph

                    for cl in prev_p['accum_lemmas']:   # for each given lemmas in the prev. paragraph
                        # if a lemma ('gl') is in the previous paragraphs AND their POSs maatch
                        #if self.isPOSVisible(gl[POS]) and self.isGiven(gl[LEMMA], cl[LEMMA], gl[POS], cl[POS]):
                        if self.isGiven(gl[LEMMA], cl[LEMMA], gl[POS], cl[POS], pronoun=self.prp):
                            para['given_accum_lemmas'].append(gl)         # add 'gl' to the list
                            break # 2024.12.02

                    for temp_para in data['paragraphs']:
                        if temp_para == para:
                            break

                        for cl in temp_para['accum_lemmas']:     # for each paragraph's accum_lemmas
                            # if a lemma ('gl') is in the previous paragraph AND their POSs maatch
                            if self.isGiven(gl[LEMMA], cl[LEMMA], gl[POS], cl[POS], pronoun=self.prp):
                                # if 'cl' is not already in the previous paragraph's new lemmas list
                                if gl[LEMMA] not in [x[LEMMA] for x in temp_para['new_accum_lemmas']]:
                                    temp_para['new_accum_lemmas'].append(gl)     # add 'gl' to the list

                    # if 'lemma' is a user defined topic, force it to be 'given' (not used by myProse...)
                    if DSDocument.isUserDefinedTopic(gl[LEMMA].lower()) or \
                       DSDocument.isUserDefinedSynonym(gl[LEMMA]):
                        if gl[LEMMA].lower() not in para['given_accum_lemmas']:
                            para['given_accum_lemmas'].append(gl)

                # remove the duplicates. It may be a bit faster to do it here than checking duplicates in the loop above...
                para['given_accum_lemmas'] = list(set(para['given_accum_lemmas']))

                prev_p = para

            else:
                prev_p = para
                for gl in para['lemmas']:
                    #### Adding new (not-given) lemmas to the first paragraph (2024.12.02)
                    if gl[LEMMA] not in [x[LEMMA] for x in para['new_accum_lemmas']]:
                        para['new_accum_lemmas'].append(gl)

                    if DSDocument.isUserDefinedTopic(gl[LEMMA].lower()) or \
                       DSDocument.isUserDefinedSynonym(gl[LEMMA]):
                        if gl[LEMMA].lower() not in para['given_accum_lemmas']:
                            para['given_accum_lemmas'].append(gl)

    ########################################
    #
    # Loading Methods
    #
    ########################################
    def loadFromTxt(self, aText):
        #print ("loadFromTxt ()")
        logging.info(aText)

        self.current_section = 0

        paragraphs = aText.splitlines()

        doc = Document()

        for para in paragraphs:
            if para.strip() != "":
                para = adjustSpaces(para)
                doc.add_paragraph(para)

        section_data = dict()
        section_data['doc']     = doc
        section_data['data']    = dict()
        section_data['heading'] = "n/a"
        section_data['start']   = 0
        section_data['pos']     = 0
        section_data['para_data'] = []

        self.sections = [section_data]

        self.processDoc(section_data)

        section_data['start'] = 0

        self.processDoc(section_data)

    def loadFromTxtFile(self, src_dir, file):
        """
        Load a text from a plain text file.
        It then process the document by calling the processDoc method.
        It calls back the application object to update the headings.
        """

        if self.progress_callback:
            self.progress_callback(new_val=10, msg="Opening file...")

        self.current_section = 0

        with open(os.path.join(src_dir,file), errors="ignore") as fin:
            text = fin.read()
            paragraphs = text.splitlines()

        doc = Document()
        for para in paragraphs:
            if para:
                para = adjustSpaces(para)
                doc.add_paragraph(para)

        if self.progress_callback:
            self.progress_callback(new_val=10)

        section_data = dict()
        section_data['doc']     = doc
        section_data['data']    = dict()
        section_data['heading'] = "n/a"
        section_data['start']   = 0
        section_data['pos']     = 0
        section_data['para_data'] = []
        self.sections = [section_data]

        self.filename = file

        self.processDoc(section_data)
        section_data['start'] = 0
        self.processDoc(section_data)

    def loadFromMSWordFile(self, src_dir, file):
        """
        Load a text from a MS Word file. If the file contains headings using
        the "Title" or th "Heading 1" style, the document is split into multiple sections.
        It then process the document by calling the processDoc method.
        """
        def get_list_type_and_level(paragraph):
            """
            Determines if a paragraph is part of a bulleted or numbered list.
            Returns:
            - "bullet" for bulleted lists
            - The numFmt value (e.g., "decimal", "lowerLetter") for numbered lists
            - None if not a list paragraph or can't determine
            """
            # Access the XML element
            p_element = paragraph._element

            # The namespace prefix is already registered in python-docx
            # Check for numbering properties
            num_pr = p_element.xpath('./w:pPr/w:numPr')
            if not num_pr:
                return None, None

            # Get numbering ID and level
            num_id_elements = num_pr[0].xpath('./w:numId')
            ilvl_elements = num_pr[0].xpath('./w:ilvl')

            if not num_id_elements or not ilvl_elements:
                return None, None

            # Get the values
            num_id = num_id_elements[0].attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            ilvl = ilvl_elements[0].attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')

            # Get the document's numbering part
            try:
                numbering_part = paragraph._parent._parent.part.numbering_part
                if not numbering_part:
                    return None, None
            except AttributeError:
                return None, None

            # Get the XML from numbering definitions
            numbering_xml = numbering_part._element

            # Find the abstract numbering ID
            # We need to use full XPath with namespace prefixes already registered
            abstract_num_xpath = f'.//w:num[@w:numId="{num_id}"]/w:abstractNumId'
            abstract_num_id_elements = numbering_xml.xpath(abstract_num_xpath)

            if not abstract_num_id_elements:
                return None, None

            abstract_num_id = abstract_num_id_elements[0].attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')

            # Find the level format
            level_xpath = f'.//w:abstractNum[@w:abstractNumId="{abstract_num_id}"]/w:lvl[@w:ilvl="{ilvl}"]/w:numFmt'
            num_fmt_elements = numbering_xml.xpath(level_xpath)

            if not num_fmt_elements:
                return None, None

            num_fmt = num_fmt_elements[0].attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            return num_fmt, ilvl

        self.current_section = 0

        if self.progress_callback:
            self.progress_callback(max_val=10, msg="Opening file...")

        self.sections = list()
        fpath = os.path.join(src_dir,file)
        doc = Document(fpath)     # create a Document object from a file.

        if self.progress_callback:
            self.progress_callback(new_val=12)

        para_count = 0
        sect_count = 0
        headings = list()
        p = None
        for para in doc.paragraphs:

            ptext = para.text.strip()

            if ptext == '':
                continue

            ptext = adjustSpaces(ptext)

            if para.style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title',
                                   'List Bullet', 'List Bullet 2', 'List Bullet 3',
                                   'List Number', 'List Number 2', 'List Number 3']:
                style = para.style.name

            elif para.style.name == 'List Paragraph':

                list_type, list_level = get_list_type_and_level(para) # check the list type and its level

                if list_type == 'bullet':
                    if list_level == '0':
                        style = "List Bullet"
                    else:
                        style = f"List Bullet {int(list_level)+1}"

                elif list_type == 'decimal':
                    if list_level == '0':
                        style = "List Number"
                    else:
                        style = f"List Number {int(list_level)+1}"

                else:
                    style = 'Normal'
            else:
                style = 'Normal'

            if para_count == 0 and sect_count == 0:
                # if the first paragraph is not a heading. We need to crreate a doc w/o a heading
                section_data = dict()
                section_data['doc'] = Document()
                section_data['heading'] = ptext
                section_data['pos'] = sect_count
                section_data['para_data'] = []
                self.sections.append(section_data)
                headings.append("")
                sect_count += 1
                p = None

            p = section_data['doc'].add_paragraph(ptext, style=style)

            para_count += 1

        offset = 0
        for section_data in self.sections:
            section_data['start'] = offset
            self.processDoc(section_data)
            offset = section_data['end_pos'] + 1

        # if there are more than 1 sections, AND the first section does not have a heading,
        # use '(No Sections Header/Title)'
        if sect_count > 1 and headings[0] == "":
            headings[0] = "(No Sections Header/Title)"
            text = self.sections[0]['heading']
            if self.controller:
                self.controller.showWarningDialog("Warning", "The first section does not include a title or heading. " +
                                                             "Consider assigning the Title or Heading style to the following line " +
                                                             "in Microsoft Word. \n\n\"{}\"".format(text))
        self.filename = file
        return headings

    def loadFromListOfParagraphs(self, doc, section):
        """ UNUSED """
        self.current_section = section

        if self.filename is not None:
            file = self.filename
        else:
            file = "Undefined.docx"
            self.sections = list()

        if self.progress_callback:
            self.progress_callback(new_val=25)

        para_count = 0
        # sect_count = 0
        headings = list()
        new_sections = list()

        for para in doc.paragraphs:

            ptext = para.text

            if ptext.startswith("<img "):
                section_data['doc'].add_paragraph(ptext)
                para_count += 1
                continue

            for mw_topic in DSDocument.multiword_topics + DSDocument.deleted_multiword_topics:

                connected_mw_topic = mw_topic.replace(' ', '_')
                ptext = ptext.replace(connected_mw_topic, mw_topic)
                ptext = ptext.replace(connected_mw_topic.lower(), mw_topic.lower())
                ptext = ptext.replace(connected_mw_topic.title(), mw_topic.title())
                ptext = ptext.replace(connected_mw_topic.capitalize(), mw_topic.capitalize())

            ptext = adjustSpaces(ptext)

            if para_count == 0:

                if para.style.name in ('Heading 1', 'Title'):
                    section_data = dict()
                    section_data['doc'] = Document()
                    section_data['heading'] = ptext
                    section_data['pos'] = section
                    section_data['para_data'] = []
                    new_sections.append(section_data)
                    headings.append(ptext)

                else:
                    section_data = dict()
                    section_data['doc'] = Document()
                    section_data['heading'] = ptext
                    section_data['pos'] = section
                    section_data['para_data'] = []
                    headings.append("(No Heading)")
                    new_sections.append(section_data)

            # if para.style.name in ['Normal', 'Heading 1', 'Heading 2', 'Title']:
            if para.style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6', 'Title', 'List Bullet']:
                style = para.style.name
            else:
                style = 'Normal'

            section_data['doc'].add_paragraph(ptext, style=style)
            para_count += 1

        if self.sections:
            if len(self.sections) <= section:   # if len (sections) = 0 > 0
                return

            offset = self.sections[section]['start']  # get the offset of the current section
            section_data['start'] = offset

            self.processDoc(section_data)

            self.sections.pop(self.current_section)
            self.sections[self.current_section:self.current_section] = new_sections

        else:  # We get here only if the user copy & pasted text into the editor
            self.sections = new_sections
            offset = 0
            for section_data in self.sections:
                section_data['start'] = offset
                self.processDoc(section_data)
                offset = section_data['end_pos'] + 1

        return headings

    def loadFromHtmlFile(self, src_dir, html_file):
        with open(os.path.join(src_dir,html_file), errors="ignore") as fin:
            html_str = fin.read()
        self.loadFromHtmlString(html_str)

    def loadFromHtmlString(self, html_str):

        def add_list_recursively(tag, level, list_type):

            p = None
            for child in tag.children:

                if type(child) == bs4.element.NavigableString:
                    text = child.text
                    text = text.strip()
                    if text:
                        if level == 1:
                            style = f"List {list_type}"
                        else:
                            style = f"List {list_type} {level}"

                        if p is None:
                            p = doc.add_paragraph("", style=style)

                        p.add_run(text + " ")

                elif child.name == 'li':
                    if child.children:
                        add_list_recursively(child, level, list_type)

                    else: # no children
                        text = child.decode_contents()

                        if level == 1:
                            style = f"List {list_type}"
                        else:
                            style = f"List {list_type} {level}"

                        p = doc.add_paragraph(text, style=style)

                elif child.name == 'ul':
                    add_list_recursively(child, level+1, "Bullet")

                elif child.name == 'ol':
                    add_list_recursively(child, level+1, "Number")

                else:
                    text = child.decode_contents()
                    text = text.strip()

                    if level == 1:
                        style = f"List {list_type}"
                    else:
                        style = f"List {list_type} {level}"

                    if p is None:
                        p = doc.add_paragraph("", style=style)

                    if child.name == 'b':
                        text = f"<b>{text}</b> "
                    elif child.name == 'i':
                        text = f"<i>{text}</i> "
                    elif child.name == 'u':
                        text = f"<u>{text}</u> "
                    r = p.add_run(text)

        soup = bs(html_str, "html.parser")
        body = soup.find("body")
        doc = Document()

        for tag in body.children:

            if isinstance(tag, bs4.element.NavigableString):
                # skip floating strings w/ no tags.
                continue

            if isinstance(tag, Comment) or not hasattr(tag, 'decode_contents'):
                # skip comments and other elements that do not contain content.
                continue

            tag_name = tag.name

            if tag_name == 'p':
                text = tag.decode_contents()
                text = text.replace('\n', ' ')
                doc.add_paragraph(text)

            elif tag_name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = tag_name[-1]
                style_str = f"Heading {level}"

                text = tag.decode_contents()
                text = text.replace('\n', ' ')
                doc.add_paragraph(text, style_str)

            elif tag_name == 'ul':
                add_list_recursively(tag, 1, "Bullet")

            elif tag_name == 'ol':
                add_list_recursively(tag, 1, "Number")
            else:
                text = tag.decode_contents()

        section_data = dict()
        section_data['doc']       = doc
        section_data['data']      = dict()
        section_data['heading']   = "n/a"
        section_data['start']     = 0
        section_data['pos']       = 0
        section_data['para_data'] = []

        self.sections = [section_data]
        self.processDoc(section_data)


    ########################################
    #
    # Methods for creating HTML/XML strings from the text data
    # Doesn't rely on DocuScope tagging, does rely on PythonDocX
    #
    ########################################

    def toHtml_OTOnly(self, topics=[], para_pos=-1, font_info=default_font_info):
        """
        This method is used if there is no docusope dictionary. In Write & Audit, this means
        that it is either used without the Impressions panel, or it is used from the Online version
        where docuscope is not integrated into the Coherence visualization.
        """
        data = self.sections[self.current_section]['data']

        if data is None:
            return ""

        total_paras = len(data['paragraphs'])
        html_str = ""
        pcount  = 1
        li_tag = 0
        ol_tag = 0
        ul_tag = 0
        for para in data['paragraphs']: # for each paragraph

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1
                continue

            total_sents = len(para['sentences'])
            para_style = para['style']

            if li_tag > 0: # previous li_tag
                # If the prev tag was li, and the current tag is also li, do nothing.
                # Otherwise, we need to close the ul tag.
                if para_style.startswith('List Bullet'):
                    html_str += '</li>'   # Close the list tag

                    level_str = para_style[-1]
                    if level_str.isnumeric():
                        level = int(level_str)
                    else:
                        level = 1

                    if li_tag < level:                # The current level is greater
                        html_str += "<ul>"
                        ul_tag += 1
                    elif li_tag > level:
                        html_str += "</ul></li>"         # Th current level is lesser
                        ul_tag -= 1

                elif para_style.startswith('List Number'):
                    html_str += '</li>'   # Close the list tag

                    level_str = para_style[-1]
                    if level_str.isnumeric():
                        level = int(level_str)
                    else:
                        level = 1

                    if li_tag < level:
                        html_str += "<ol>"
                        ol_tag += 1
                    elif li_tag > level:
                        html_str += "</ol></li>"
                        ol_tag -= 1

                elif ul_tag > 0:                  # prev is an li tag, but the current is not.
                    for n in range(ul_tag):
                        if n == (ul_tag-1):
                            html_str += '</ul>'
                        else:
                            html_str += '</ul></li>'
                        ul_tag -= 1
                    pcount += 1

                elif ol_tag > 0:
                    for n in range(ol_tag):
                        if n == (ol_tag-1):
                            html_str += '</ol>'
                        else:
                            html_str += '</ol></li>'
                        ol_tag -= 1
                    pcount += 1

            elif li_tag == 0:
                # If the prev tag was not li, and the current tag is li, we will fist add the ul tag.
                if para_style.startswith('List Bullet'):
                    if ul_tag > 0:
                        html_str += "<ul>"
                    else:
                        html_str += '<ul data-ds-paragraph="{0}" id="p{0}" class="paragraph">'.format(pcount)
                    level_str = para_style[-1]
                    if level_str.isnumeric():
                        level = int(level_str)
                    else:
                        level = 1
                    ul_tag = level
                    scount = 1  # reset the sentence count
                elif para_style.startswith('List Number'):
                    if ol_tag > 0:
                        html_str += "<ol>"
                    else:
                        html_str += '<ol data-ds-paragraph="{0}" id="p{0}" class="paragraph">'.format(pcount)
                    level_str = para_style[-1]
                    if level_str.isnumeric():
                        level = int(level_str)
                    else:
                        level = 1
                    ol_tag = level
                    scount = 1  # reset the sentence count
            li_tag = 0

            if para_style.startswith('List Bullet') or para_style.startswith('List Number'):
                html_str += '<li>'
                level_str = para_style[-1]
                if level_str.isnumeric():
                    li_tag = int(level_str)
                else:
                    li_tag = 1

                # we don't reset 'scount' since we are in <ul> or <ol>, which is treated like a paragraph
            else:
                if para_style == 'Heading 1':
                    html_str += '<h1 data-ds-paragraph="{0}" id="p{0}" class="paragraph">'.format(pcount)
                elif para_style == 'Heading 2':
                    html_str += '<h2 data-ds-paragraph="{0}" id="p{0}" class="paragraph">'.format(pcount)
                elif para_style == 'Heading 3':
                    html_str += '<h3 data-ds-paragraph="{0}" id="p{0}" class="paragraph">'.format(pcount)
                elif para_style == 'Heading 4':
                    html_str += '<h4 data-ds-paragraph="{0}" id="p{0}" class="paragraph">'.format(pcount)
                else:
                    html_str += '<p data-ds-paragraph="{0}" id="p{0}" class="paragraph">'.format(pcount)
                scount = 1

            for sent in para['sentences']: # for each sentence

                total_words = len(sent['text_w_info'])
                is_combo = False
                word_count = 0
                html_str += '<span id="p{0}s{1}" class="sentence" data-ds-paragraph="{0}" data-ds-sentence="{1}"> '.format(pcount, scount)

                while word_count < total_words:

                    w = sent['text_w_info'][word_count]
                    word  = w[WORD]         # ontopic word
                    lemma = w[LEMMA]        # lemma/topic

                    # if w[LEMMA] in topics:
                    if lemma in topics and (w[POS] == 'NOUN' or w[POS] == 'PRP'):
                        # we need to remove periods.
                        topic = unidecode.unidecode(w[LEMMA].replace('.',''))
                        html_word = '<span class="word" data-ds-paragraph="{}" data-ds-sentence="{}" data-topic="{}">{}</span> '.format(pcount, scount, topic, word)
                    else:
                        html_word = word

                    is_space = True
                    is_combo = False
                    temp_word_count = 0

                    if word_count < total_words-1:

                        # Let's deal with hyphenations and words followed by a no-space units.
                        next_w = sent['text_w_info'][word_count+1]
                        next_next_w = None

                        if word_count+1 < total_words-1:
                            next_next_w = sent['text_w_info'][word_count+2]

                        if word in right_quotes and html_str[-1] == ' ':
                            # if w[WORD] is a right side quotes, and the current html_str string ends
                            # with a space, remove the space.
                            html_str = html_str[:-1]

                        if word in left_quotes or word in hyphen_slash:
                            # w is a left side punct or a dash (not em-dash or en-dash)
                            is_space = False

                        elif word == '.' and next_w[WORD].isdigit():
                            # if word is a period and the next word is a digit + %, it's a decimal point.
                            is_space = False

                        elif word == '.' and len(next_w[WORD]) == 1 and next_w[WORD].isalpha():
                            # if word is a period and the next word is single alphabet, it's a decimal point
                            # or multi-level figure numbers.
                            is_space = False

                        elif word == 'can' and next_w[WORD] == 'not' and next_next_w is not None and next_next_w[WORD] != 'only':
                            # we always spell 'can not' without a space, except when it is followd by 'only'.
                            is_space = False

                        elif next_w[WORD] not in left_quotes:

                            # next_w is a punct, which is not on one of the left quotes
                            temp_word_count = 0
                            if next_w[WORD] in no_space_patterns:

                                combo_word = html_word + next_w[WORD]
                                temp_word_count = word_count + 2

                                if temp_word_count < total_words:    # is the folloiwng word punct?
                                    temp_w = sent['text_w_info'][temp_word_count]

                                    if temp_w[WORD] in no_space_patterns or temp_w[WORD] in right_quotes:
                                        combo_word = combo_word + temp_w[WORD]
                                        word_count = temp_word_count
                                        is_space = True
                                        is_combo = True

                                    elif temp_w[POS] == 'PUNCT' and temp_w[WORD] not in dashes and temp_w[WORD] != ellipsis:
                                        combo_word = combo_word + temp_w[WORD]
                                        word_count = temp_word_count
                                        is_space = True
                                        is_combo = True

                                    else:
                                        word_count += 1
                                        is_space = True
                                        is_combo = True

                                else:
                                    word_count += 1  # skip the next word
                                    is_space = True
                                    is_combo = True

                            elif next_w[WORD] in hyphen_slash:
                                combo_word = html_word + next_w[WORD] # word + '-' or '/'
                                temp_word_count = word_count + 2

                                if temp_word_count >= total_words:
                                    pass

                                temp_next_w = None
                                while True:

                                    if temp_word_count >= (total_words):
                                        break

                                    # get the next pair
                                    temp_w = sent['text_w_info'][temp_word_count]
                                    if temp_w[LEMMA] in topics and (temp_w[POS] == 'NOUN' or temp_w[POS] == 'PRP'):
                                        topic = unidecode.unidecode(w[LEMMA].replace('.',''))
                                        html_temp_word = '<span class="word" data-ds-paragraph="{}" data-ds-sentence="{}" data-topic="{}">{}</span> '.format(pcount, scount, topic, temp_w[WORD])
                                    else:
                                        html_temp_word = temp_w[WORD]

                                    if temp_word_count+1 < (total_words):
                                        temp_next_w = sent['text_w_info'][temp_word_count+1]

                                        if temp_next_w[WORD] in hyphen_slash:
                                            combo_word += html_temp_word
                                            combo_word += temp_next_w[WORD]  # word + '-' or '/'
                                            temp_word_count += 2
                                        else:
                                            combo_word += html_temp_word
                                            break
                                    else:
                                        combo_word += html_temp_word
                                        break

                                if temp_next_w is not None:

                                    if temp_next_w[WORD] in no_space_patterns or \
                                       temp_next_w[WORD] in right_quotes:
                                        is_space = False

                                    elif temp_next_w[WORD] in end_puncts and temp_next_w[WORD] not in dashes \
                                                                         and temp_next_w[WORD] != ellipsis:
                                        is_space = False

                                    else:
                                        is_space = True
                                else:
                                    is_space = True

                                word_count = temp_word_count # skip all the words in the hyphenated word
                                # is_space = True
                                is_combo = True

                            elif next_w[WORD] in right_quotes:
                                is_space = False

                            elif next_w[POS] == 'PUNCT' and next_w[WORD] not in dashes and next_w[WORD] != ellipsis:
                                is_space = False

                            elif next_w[POS] == 'PUNCT' and next_w[WORD] in end_puncts:
                                is_space = False
                            else:
                                is_space = True

                        if is_space:
                            next_char = ' '
                        else:
                            next_char = ''

                    else:  # last word/char in 'sent'
                        if word in ['\u201C', '\u2018']:
                            # If the last word is an opening/left single/double quote, no space should be added.
                            # This case should only happen if the NLP parser's result is incorrect.
                            next_char = ''

                        elif word in right_quotes:
                            if html_str[-1] == ' ':
                                html_str = html_str.rstrip()
                            next_char = ' '
                        else:
                            next_char = ''       # end of 'sent' Jan 14

                    if is_combo:
                        html_str += combo_word
                    else:
                        html_str += html_word

                    html_str += next_char
                    word_count += 1

                # if bTagOpen == True:
                    # html_str += '</ds>'
                    # bTagOpen = False

                #  word
                html_str += '</span> '
                html_str += next_char
                scount += 1

            # sent
            if ol_tag > 0 or ul_tag > 0:
                pass
            elif li_tag > 0:
                pass
            else:
                if para_style == 'Heading 1':
                    html_str += '</h1>'
                elif para_style == 'Heading 2':
                    html_str += '</h2>'
                elif para_style == 'Heading 3':
                    html_str += '</h3>'
                elif para_style == 'Heading 4':
                    html_str += '</h4>'
                else:
                    html_str += '</p>'
                pcount += 1

        if li_tag > 0:
            if ul_tag > 0:
                for n in range(ul_tag):
                    html_str += '</ul>'
            elif ol_tag > 0:
                for n in range(ol_tag):
                    html_str += '</ol>'

        return html_str

    ########################################
    #
    # Methods for creating HTML/XML strings from the text data
    # Doesn't rely on DocuScope tagging, does rely on PythonDocX
    #
    # DSWA specific version that makes it easier to display and
    # manager in the React displace
    #
    ########################################

    def toHtml_OTOnly_DSWA(self, topics=[], para_pos=-1, font_info=default_font_info):
        """
        This method is used if there is no docusope dictionary. In Write & Audit, this means
        that it is either used without the Impressions panel, or it is used from the Online version
        where docuscope is not integrated into the Coherence visualization.
        """

        data = self.sections[self.current_section]['data']

        if data is None or 'paragraphs' not in data:
            return ""

        # total_paras = len(data['paragraphs'])
        html_str = ""
        pcount  = 1

        for para in data['paragraphs']: # for each paragraph

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1
                continue

            html_str += '<p data-ds-paragraph="{0}" id="p{0}" class="paragraph">'.format(pcount)

            scount = 1

            for sent in para['sentences']: # for each sentence

                total_words = len(sent['text_w_info'])
                is_combo = False
                word_count = 0
                html_str += '<span id="p{0}s{1}" class="sentence" data-ds-paragraph="{0}" data-ds-sentence="{1}">'.format(pcount, scount)

                while word_count < total_words:

                    w = sent['text_w_info'][word_count]
                    word  = w[WORD]         # ontopic word
                    lemma = w[LEMMA]        # lemma/topic

                    # if w[LEMMA] in topics:
                    if lemma in topics and (w[POS] == 'NOUN' or w[POS] == 'PRP'):
                        # we need to remove periods.
                        topic = unidecode.unidecode(w[LEMMA].replace('.',''))
                        html_word = '<span class="word" data-ds-paragraph="{}" data-ds-sentence="{}" data-topic="{}">{}</span>'.format(pcount, scount, topic, word)
                    else:
                        html_word = word

                    is_space = True
                    is_combo = False
                    temp_word_count = 0

                    if word_count < total_words-1:

                        # Let's deal with hyphenations and words followed by a no-space units.
                        next_w = sent['text_w_info'][word_count+1]
                        next_next_w = None

                        if word_count+1 < total_words-1:
                            next_next_w = sent['text_w_info'][word_count+2]

                        if word in right_quotes and html_str[-1] == ' ':
                            # if w[WORD] is a right side quotes, and the current html_str string ends
                            # with a space, remove the space.
                            html_str = html_str[:-1]

                        if word in left_quotes or word in hyphen_slash:
                            # w is a left side punct or a dash (not em-dash or en-dash)
                            is_space = False

                        elif word == '.' and next_w[WORD].isdigit():
                            # if word is a period and the next word is a digit + %, it's a decimal point.
                            is_space = False

                        elif word == '.' and len(next_w[WORD]) == 1 and next_w[WORD].isalpha():
                            # if word is a period and the next word is single alphabet, it's a decimal point
                            # or multi-level figure numbers.
                            is_space = False

                        elif word == 'can' and next_w[WORD] == 'not' and next_next_w is not None and next_next_w[WORD] != 'only':
                            # we always spell 'can not' without a space, except when it is followd by 'only'.
                            is_space = False

                        elif next_w[WORD] not in left_quotes:

                            # next_w is a punct, which is not on one of the left quotes
                            temp_word_count = 0
                            if next_w[WORD] in no_space_patterns:

                                combo_word = html_word + next_w[WORD]
                                temp_word_count = word_count + 2

                                if temp_word_count < total_words:    # is the folloiwng word punct?
                                    temp_w = sent['text_w_info'][temp_word_count]

                                    if temp_w[WORD] in no_space_patterns or temp_w[WORD] in right_quotes:
                                        combo_word = combo_word + temp_w[WORD]
                                        word_count = temp_word_count
                                        is_space = True
                                        is_combo = True

                                    elif temp_w[POS] == 'PUNCT' and temp_w[WORD] not in dashes and temp_w[WORD] != ellipsis:
                                        combo_word = combo_word + temp_w[WORD]
                                        word_count = temp_word_count
                                        is_space = True
                                        is_combo = True

                                    else:
                                        word_count += 1
                                        is_space = True
                                        is_combo = True

                                else:
                                    word_count += 1  # skip the next word
                                    is_space = True
                                    is_combo = True

                            elif next_w[WORD] in hyphen_slash:
                                combo_word = html_word + next_w[WORD] # word + '-' or '/'
                                temp_word_count = word_count + 2

                                if temp_word_count >= total_words:
                                    pass

                                temp_next_w = None
                                while True:

                                    if temp_word_count >= (total_words):
                                        break

                                    # get the next pair
                                    temp_w = sent['text_w_info'][temp_word_count]
                                    if temp_w[LEMMA] in topics and (temp_w[POS] == 'NOUN' or temp_w[POS] == 'PRP'):
                                        topic = unidecode.unidecode(w[LEMMA].replace('.',''))
                                        html_temp_word = '<span data-ds-paragraph="{}" data-ds-sentence="{}" class="word" data-topic="{}">{}</t>'.format(pcount, scount, topic, temp_w[WORD])
                                    else:
                                        html_temp_word = temp_w[WORD]

                                    if temp_word_count+1 < (total_words):
                                        temp_next_w = sent['text_w_info'][temp_word_count+1]

                                        if temp_next_w[WORD] in hyphen_slash:
                                            combo_word += html_temp_word
                                            combo_word += temp_next_w[WORD]  # word + '-' or '/'
                                            temp_word_count += 2
                                        else:
                                            combo_word += html_temp_word
                                            break
                                    else:
                                        combo_word += html_temp_word
                                        break

                                if temp_next_w is not None:

                                    if temp_next_w[WORD] in no_space_patterns or \
                                       temp_next_w[WORD] in right_quotes:
                                        is_space = False

                                    elif temp_next_w[WORD] in end_puncts and temp_next_w[WORD] not in dashes \
                                                                         and temp_next_w[WORD] != ellipsis:
                                        is_space = False

                                    else:
                                        is_space = True
                                else:
                                    is_space = True

                                word_count = temp_word_count # skip all the words in the hyphenated word
                                is_space = True
                                is_combo = True

                            elif next_w[WORD] in right_quotes:
                                is_space = False

                            elif next_w[POS] == 'PUNCT' and next_w[WORD] not in dashes and next_w[WORD] != ellipsis:
                                is_space = False

                            else:
                                is_space = True

                        if is_space:
                            next_char = ' '
                        else:
                            next_char = ''

                    else:  # last word/char in 'sent'
                        if word in ['\u201C', '\u2018']:
                            # If the last word is an opening/left single/double quote, no space should be added.
                            # This case should only happen if the NLP parser's result is incorrect.
                            next_char = ''

                        elif word in right_quotes:
                            if html_str[-1] == ' ':
                                html_str = html_str.rstrip()
                            next_char = ' '
                        else:
                            next_char = ' '       # end of 'sent'


                    if is_combo:
                        html_str += combo_word
                    else:
                        html_str += html_word

                    html_str += next_char
                    word_count += 1

                # if bTagOpen == True:
                    # html_str += '</ds>'
                    # bTagOpen = False

                #  word
                html_str += '</span>'
                html_str += next_char
                scount += 1

            # sent
            pcount += 1
            html_str += '</p>'   # Close the paragraph tag

        return html_str

    def toHtml(self, topics=[], para_pos=-1, font_info=default_font_info):
        """
        This method is used only by the desktop version of Write & Audit. It genreates an HTML string
        for the edtiror from the parsed data.
        """

        if self.controller is None:
            return "DSDocument.toHtml(). No controller is available."

        if self.controller.isDocTagged() is False:
            return self.toHtml_OTOnly(topics=topics, para_pos=para_pos, font_info=font_info)
            # return self.toHtml_OTOnly_DSWA(topics=topics, para_pos=para_pos, font_info=font_info)

        if self.sections is None:
            return "Error in toHtml()"
        else:
            try:
                data = self.sections[self.current_section]['data']
                doc  = self.sections[self.current_section]['doc']
                if doc is None:
                    raise ValueError
                docx_paras = doc.paragraphs

                if data is None:
                    raise ValueError
            except:
                logging.error(self.sections[self.current_section])
                return

        bTagOpen = False
        bParaCrossingPattern = False

        dsw   = None
        clust = ''
        dim   = ''
        lat   = ''
        end   = ''
        pos   = ''
        next_char = ''

        html_str = ""
        pcount = 1

        for para in data['paragraphs']:    # for each paragraph

            scount = 1

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1
                bParaCrossingPattern = False
                bTagOpen = False
                continue

            # If the paragraph is an URL, we assume that that's an image.
            docx_p = docx_paras[pcount-1]
            if docx_p.text.startswith("<img "):
                html_str += ("<p>" + docx_p.text + "</p>\n")
                pcount += 1
                continue

            para_style = para['style']
            font_size = font_info[para_style]
            if para_style  in ('Heading 1', 'Heading 2'):
                html_str += '<p class=\'p{}\' style=\'font-weight: bold; font-size: {}pt;\'>'.format(pcount, font_size)
            elif para_style in ('Heading 3', 'Heading 4'):
                html_str += '<p class=\'p{}\' style=\'font-weight: bold; font-size: {}pt;\'>'.format(pcount, font_size)
            else:
                html_str += '<p class=\'p{}\' style=\'font-size: {}pt;\'>'.format(pcount, font_size)

            for sent in para['sentences']: # for each sentence
                total_words = len(sent['text_w_info'])
                is_combo = False
                word_count = 0

                # For some reason (usualy it is a NLP parser problem), if a sentence is starting
                # with a punctionation that's supposed to be part of the previous sentence, remove
                # a space from html_str.
                first_w = sent['text_w_info'][0]
                if word_count == 0:
                    if first_w[LEMMA] in no_space_patterns:
                        html_str = html_str.rstrip()
                        next_char = ''

                    elif first_w[WORD] in end_puncts:
                        html_str = html_str.rstrip()
                        next_char = ''

                html_str += '<sent class=\'p{} s{}\'>'.format(pcount, scount)  # Sentence

                while word_count < total_words:
                    w = sent['text_w_info'][word_count]     # for each word.

                    word  = w[WORD]         # ontopic word
                    lemma = w[LEMMA]        # lemma/topic

                    if lemma in topics and (w[POS] == 'NOUN' or w[POS] == 'PRP'):
                        topic = unidecode.unidecode(lemma.replace('.',''))
                        html_word = '<t class=\'{} p{} s{}\'>{}</t>'.format(topic, pcount, scount, word)
                    else:
                        html_word = word

                    dsw   = w[DS_DATA]      # docuscope word
                    if dsw is not None:
                        end   = dsw.getEnd()
                        lat   = dsw.getLAT()
                        pos   = dsw.getPos()
                        dim, clust = self.controller.getDimensionAndCluster(lat)
                    else:
                        if bTagOpen:
                            bTagOpen = False
                            html_str += '</ds>' + ' '
                        else:
                            html_str += ' '
                        html_str += html_word
                        word_count += 1
                        continue

                    if pos != '':
                        pos = int(float(pos))
                    else:
                        pos = -1

                    if dsw is not None:
                        # Let's put an open tags around DS patterns.
                        if pos == 0:                     # The first word of a DS TAG

                            if bTagOpen is True:         # if a tag is open, close it.
                                html_str += '</ds>'
                                html_str += next_char    # add a space (or nothing)
                                bTagOpen = False
                            else:
                                html_str += next_char

                            if bTagOpen is False:        # if a tag is not open, start a new one.

                                if word_count == (total_words-1) and \
                                    word == '.' and pos == 0 and lat != '':
                                    # If the period at the end of a pragraph or a sentence
                                    # is part of a pattern, make it as such.
                                    bParaCrossingPattern = True
                                else:
                                    if word in end_puncts or word in right_quotes:
                                        html_str = html_str.rstrip()

                                    if lat == '' or lat == 'UNRECOGNIZED':
                                        html_str += '<ds class=\'na\'>'
                                    else:
                                        html_str += '<ds class=\'{} {} {}\'>'.format(clust, dim, lat)
                                    bTagOpen = True

                        elif bParaCrossingPattern:

                            html_str += next_char
                            if lat in ('', 'UNRECOGNIZED'):
                                html_str += '<ds class=\'na\'>'
                            else:
                                html_str += '<ds class=\'{} {} {}\'>'.format(clust, dim, lat)

                            bTagOpen = True

                            bParaCrossingPattern = False
                        else:
                            html_str += next_char

                    is_space = True
                    is_combo = False
                    temp_word_count = 0

                    # --------------------------------------------------------
                    # Now, we will deal with OnTOpic Tags.
                    # --------------------------------------------------------

                    if word_count < total_words-1:
                        # Let's deal with hyphenations and words followed by a no-space units.
                        next_w = sent['text_w_info'][word_count+1]
                        next_next_w = None

                        if word_count+1 < total_words-1:
                            next_next_w = sent['text_w_info'][word_count+2]

                        if word in right_quotes and html_str[-1] == ' ':
                            # if w[WORD] is a right side quotes, and the current html_str string ends
                            # with a space, remove the space.
                            html_str = html_str[:-1]

                        if word in left_quotes or word in hyphen_slash:
                            # w is a left side punct or a dash (not em-dash or en-dash)
                            is_space = False

                        elif word == '.' and next_w[WORD].isdigit():
                            # if word is a period and the next word is a digit + %, it's a decimal point.
                            is_space = False

                        elif word == '.' and len(next_w[WORD]) == 1 and next_w[WORD].isalpha():
                            # if word is a period and the next word is single alphabet, it's a decimal point
                            # or multi-level figure numbers.
                            is_space = False

                        elif word == 'can' and next_w[WORD] == 'not' and next_next_w is not None and next_next_w[WORD] != 'only':
                            # we always spell 'can not' without a space, except when it is followd by 'only'.
                            is_space = False

                        elif next_w[WORD] not in left_quotes:

                            # next_w is a punct, which is not on one of the left quotes
                            temp_word_count = 0
                            if next_w[WORD] in no_space_patterns:

                                combo_word = html_word + next_w[WORD]
                                temp_word_count = word_count + 2

                                if temp_word_count < total_words:    # is the folloiwng word punct?
                                    temp_w = sent['text_w_info'][temp_word_count]

                                    if temp_w[WORD] in no_space_patterns or temp_w[WORD] in right_quotes:
                                        combo_word = combo_word + temp_w[WORD]
                                        word_count = temp_word_count
                                        is_space = True
                                        is_combo = True

                                    elif temp_w[POS] == 'PUNCT' and temp_w[WORD] not in dashes and temp_w[WORD] != ellipsis:
                                        combo_word = combo_word + temp_w[WORD]
                                        word_count = temp_word_count
                                        is_space = True
                                        is_combo = True

                                    else:
                                        word_count += 1
                                        is_space = True
                                        is_combo = True

                                else:
                                    word_count += 1  # skip the next word
                                    is_space = True
                                    is_combo = True

                            elif next_w[WORD] in hyphen_slash:
                                combo_word = html_word + next_w[WORD] # word + '-' or '/'
                                temp_word_count = word_count + 2

                                if temp_word_count >= total_words:
                                    pass

                                temp_next_w = None
                                while True:

                                    if temp_word_count >= (total_words):
                                        break

                                    # get the next pair
                                    temp_w = sent['text_w_info'][temp_word_count]
                                    if temp_w[LEMMA] in topics and (temp_w[POS] == 'NOUN' or temp_w[POS] == 'PRP'):
                                        topic = unidecode.unidecode(temp_w[LEMMA].replace('.',''))
                                        html_temp_word = "<t class=\'{} p{} s{}\'>{}</t>".format(topic,
                                                                                  pcount, scount,
                                                                                  temp_w[WORD])
                                    else:
                                        html_temp_word = temp_w[WORD]

                                    if temp_word_count+1 < (total_words):
                                        temp_next_w = sent['text_w_info'][temp_word_count+1]

                                        if temp_next_w[WORD] in hyphen_slash:
                                            combo_word += html_temp_word
                                            combo_word += temp_next_w[WORD]  # word + '-' or '/'
                                            temp_word_count += 2
                                        else:
                                            combo_word += html_temp_word
                                            break
                                    else:
                                        combo_word += html_temp_word
                                        break

                                if temp_next_w is not None:

                                    if temp_next_w[WORD] in no_space_patterns or \
                                       temp_next_w[WORD] in right_quotes:
                                        is_space = False

                                    elif temp_next_w[WORD] in end_puncts and temp_next_w[WORD] not in dashes \
                                                                         and temp_next_w[WORD] != ellipsis:
                                        is_space = False

                                    else:
                                        is_space = True
                                else:
                                    is_space = True

                                word_count = temp_word_count # skip all the words in the hyphenated word
                                is_combo = True

                            elif next_w[WORD] in right_quotes:
                                is_space = False

                            elif next_w[POS] == 'PUNCT' and next_w[WORD] not in dashes and next_w[WORD] != ellipsis:
                                is_space = False

                            else:
                                is_space = True

                        if is_space:
                            next_char = ' '
                        else:
                            next_char = ''

                    else:  # last word/char in 'sent'
                        if word in ['\u201C', '\u2018']:
                            # If the last word is an opening/left single/double quote, no space should be added.
                            # This case should only happen if the NLP parser's result is incorrect.
                            next_char = ''

                        elif word in right_quotes:
                            if html_str[-1] == ' ':
                                html_str = html_str.rstrip()
                            next_char = ' '
                        else:
                            next_char = ' '       # end of 'sent'

                    if is_combo:
                        html_str += combo_word
                    else:
                        html_str += html_word

                    word_count += 1

                if bTagOpen == True:
                    html_str += '</ds>'
                    bTagOpen = False

                #  word
                html_str += '</sent>'
                html_str += next_char
                scount += 1

            # sent
            pcount += 1
            html_str += '</p>'   # Close the paragraph tag

        html_str = "<html><body>\n" + html_str + "\n</body></html>"

        return html_str

    def saveAsDocx(self, filepath):
        """
        This method saves the currently analyzed document as a MS docx file.
        """
        docx = self.toPlainDoc()
        docx.save(filepath)

    def toPlainDoc(self):
        """
        Convert the current document into a Word document w/ no highlights.
        """
        res_docx = Document()                        # create a new Word Document
        for section in self.sections:
            data = section['data']
            paragraphs = data['paragraphs']              # paragraphs is a list of paragraph dictionaries.
            num_paragraphs = len(paragraphs)             # count the total # of paragraphs

            for i in range(num_paragraphs):              # for each paragraph in the current document.

                style = paragraphs[i]['style']
                if style not in ['Normal', 'Heading 1', 'Heading 2', 'Heasing 3', 'Title',
                                'List Bullet',   'List Number'
                                'List Bullet 2', 'List Number 2'
                                'List Bullet 3', 'List Number 3'
                                ]:
                    style = 'Normal'

                text = paragraphs[i]['text']

                if text.startswith("<img"):
                    soup = bs(text, "html.parser")
                    tag = soup.find("img")
                    src = tag.get('src', None)
                    src = src.replace("file:///", "")
                    w   = tag.get('width', None)
                    h   = tag.get('height', None)
                    w = int(float(w)/96.0)
                    h = int(float(h)/96.0)
                    if src is not None:
                        p = res_docx.add_paragraph()           # add a paragraph object
                        r = p.add_run()
                        r.add_picture(src, width=Inches(w), height=Inches(h))
                    continue

                p = res_docx.add_paragraph('')           # add a paragraph object
                p.style = style

                # replace underscores with spaces if there are multi-word topics
                for mw_topic in DSDocument.multiword_topics:
                    connected_mw_topic = mw_topic.replace(' ', '_')
                    text = text.replace(connected_mw_topic, mw_topic)
                    text = text.replace(connected_mw_topic.lower(), mw_topic.lower())
                    text = text.replace(connected_mw_topic.title(), mw_topic.title())
                    text = text.replace(connected_mw_topic.capitalize(), mw_topic.capitalize())

                # if style.startswith('Heading') or style == 'Title':
                #     r = p.add_run(text)
                #     r.font.color.rgb = RGBColor(0,0,0)
                #     if style == 'Title':
                #         r.font.size = Pt(24)
                #     else:
                #         r.font.size = Pt(18)
                #     r.underline = False
                #     r.bold = False
                # else:
                #     r = p.add_run(text)

                r = p.add_run(text)

        return res_docx

    def toXml(self):
        """
        This method is used to generate a XML tagged representation of the document. It's primary purpose
        is to create a structured content that can be submitted to LLM for analysis.
        """
        data = self.sections[self.current_section]['data']

        if data is None:
            return ""

        total_paras = len(data['paragraphs'])
        xml_str = ""
        pcount = 1
        scount = 1
        ul_tag = False
        ol_tag = False
        li_tag = False
        title_tag = False
        htag_level = 0

        for para in data['paragraphs']: # for each paragraph

            para_style = para['style']

            if li_tag: # previous li_tag
                # If the prev tag was li, and the current tag is also li, do nothing.
                # Otherwise, we need to close the ul tag.
                if para_style.startswith('List Bullet') or para_style.startswith('List Number'):
                    pass
                elif ul_tag:
                    xml_str += "</ul>\n\n"
                    ul_tag = False
                    pcount += 1
                elif ol_tag:
                    xml_str += "</ol>\n\n"
                    ol_tag = False
                    pcount += 1
            else:
                # If the prev tag was not li, and the current tag is li, we will fist add the ul tag.
                if para_style.startswith('List Bullet'):
                    xml_str += f"<ul id=\"p{pcount}\">"
                    ul_tag = True
                    scount = 1
                elif para_style.startswith('List Number'):
                    xml_str += f"<ol id=\"p{pcount}\">"
                    ol_tag = True
                    scount = 1

            li_tag = False
            title_tag = False
            htag_level = 0

            if para_style == 'Title':
                xml_str += f"<title id=\"p{pcount}\">"
                title_tag = True
            elif para_style == 'Heading 1':
                xml_str += f"<h1 id=\"p{pcount}\">"
                htag_level = 1
            elif para_style == 'Heading 2':
                xml_str += f"<h2 id=\"p{pcount}\">"
                htag_level = 2
            elif para_style == 'Heading 3':
                xml_str += f"<h3 id=\"p{pcount}\">"
                htag_level = 3
            elif para_style == 'Heading 4':
                xml_str += f"<h4 id=\"p{pcount}\">"
                htag_level = 4
            elif para_style == 'Heading 5':
                xml_str += f"<h5 id=\"p{pcount}\">"
                htag_level = 5
            elif para_style == 'Heading 6':
                xml_str += f"<h6 id=\"p{pcount}\">"
                htag_level = 6
            elif para_style.startswith('List Bullet'):
                xml_str += '<li>'
                li_tag = True
            elif para_style.startswith('List Number'):
                xml_str += '<li>'
                li_tag = True
            elif ol_tag is not True and ul_tag is not True:
                xml_str += f'<p id=\"p{pcount}\">'

            if ol_tag is not True and ul_tag is not True:
                scount = 1

            for sent in para['sentences']: # for each sentence

                sent_text = sent['text']

                if htag_level == 0:
                    xml_str += f"<span id=\"p{pcount}s{scount}\">{sent_text}</span>"
                else:
                    # this text is inisde a heading tag. We do not mark it as a sentence.
                    sent_text = sent['text']
                    xml_str += sent_text

                scount += 1
            # sent

            if htag_level > 0:
                xml_str += f'</h{htag_level}>\n\n'
                pcount += 1
            elif title_tag:
                xml_str += '</title>\n'
                pcount += 1
            elif li_tag:
                xml_str += '</li>\n'
            elif ol_tag is not True and ul_tag is not True:
                xml_str += '</p>\n\n'
                pcount += 1

        return xml_str

    ##

    def getLocalTopicalProgData(self, selected_paragraphs):

        selected_paragraphs.sort()

        all_lemmas = list()
        temp = list()

        curr_section = self.sections[self.current_section]

        data = curr_section.get('data', None)

        if data is None:
            return []

        selected_para_list = list()
        for pos in selected_paragraphs:
            if pos > len(data['paragraphs']):
                break
            selected_para_list.append(data['paragraphs'][pos-1])

        for p in selected_para_list:
            for s in p['sentences']:                 # for each sentence
                for l in s['new_accum_lemmas']:      # for each new lemma
                    t = (l[POS], l[LEMMA])           # tuple = (POS, LEMMA)
                    if t not in temp:                # if the lemma 'l' is not already in the list,
                        all_lemmas.append((l[POS], l[LEMMA], l[8]))  # append a tuple (POS, LEMMA, first word_pos)
                        temp.append(t)

                for l in s['given_accum_lemmas']:    # for each given lemma
                    t = (l[POS], l[LEMMA])           # tuple = (POS, LEMMA)
                    if t not in temp:                # if the lemma 'l' is not already in the list,
                        all_lemmas.append((l[POS], l[LEMMA], l[8]))   # append a tuple (POS, LEMMA, first_word_pos)
                        temp.append(t)

        # Adding local given/new lemmas from adjacent paragraphs.
        if len(selected_paragraphs) > 1:        # applicable if 2 or more paragrpahs are selected.
            temp_count = 0
            p_all_lemmas = list() # all the global given & new lemmas.
            p_temp = temp
            for p in selected_para_list:
                for l in p['new_accum_lemmas']:
                    t = (l[POS], l[LEMMA])
                    if t not in p_temp:
                        p_all_lemmas.append((l[POS], l[LEMMA], l[8]))
                        p_temp.append(t)

                for l in p['given_accum_lemmas']:
                    t = (l[POS], l[LEMMA])
                    if t not in p_temp:
                        p_all_lemmas.append((l[POS], l[LEMMA], l[8]))
                        p_temp.append(t)

                temp_count += 1

            # Exclude the lemmas that do not appear in more than 2 paragraphs
            for l in p_all_lemmas:
                count = 0
                for p in selected_para_list:
                    if (l[0], l[1]) in [(t[POS], t[LEMMA]) for t in p['lemmas']]:
                        count += 1
                if count >= 2:
                    all_lemmas.append(l)

        all_lemmas.sort(key=lambda tup: tup[2])  # sort by the order of appearance.

        res = list()
        res.append(all_lemmas)
        p_count = 1

        bSkipPunct = False
        data = self.sections[self.current_section]['data']

        new_lemmas   = []
        given_lemmas = []
        prev_given_lemmas = []
        prev_para = None
        for p in selected_para_list:
            s_count = 1

            # we'll need to find a sentence with a single punctuation character.
            for s in p['sentences']:
                bSkipPunct = len(s['text']) == 1 and \
                       (s['text'] in dashes or \
                        s['text'] in hyphen_slash or \
                        s['text'] in left_quotes or \
                        s['text'] in right_quotes or \
                        s['text'] == ellipsis)

                # The following changes support the local cohesion visualization with 2 ore more paragraphs
                # remove any lemmas that have appeared in the previous paragraphs from new_lemmas list
                # new_lemmas   = list(set(s['new_accum_lemmas']) - set(prev_given_lemmas)) # OLD

                # add the given_accum_lemmas from the last sentence of the previous paragraphs
                # and remove duplicates

                s_lemmas = [(t[POS],t[LEMMA]) for t in s['lemmas']]
                s_new_accum_lemmas   = [(t[POS],t[LEMMA]) for t in s['new_accum_lemmas']]
                s_given_accum_lemmas = [(t[POS],t[LEMMA]) for t in s['given_accum_lemmas']]
                new_lemmas = list(set(s_lemmas) & (set(s_new_accum_lemmas) - set(prev_given_lemmas)))

                if prev_para is not None:
                    prev_para_given_accum_lemmas = [(t[POS],t[LEMMA]) for t in prev_para['given_accum_lemmas'] ]
                    prev_para_new_accum_lemmas   = [(t[POS],t[LEMMA]) for t in prev_para['new_accum_lemmas'] ]

                    given_lemmas = list(set(s_lemmas) &
                                        set(s_given_accum_lemmas + prev_given_lemmas + \
                                                                   prev_para_given_accum_lemmas + \
                                                                   prev_para_new_accum_lemmas))
                else:
                    given_lemmas = list(set(s_lemmas) &
                                        set(s_given_accum_lemmas + prev_given_lemmas))

                sres = list() # get an empty list
                for l in all_lemmas:

                    is_new   = False
                    is_given = False

                    # if (l[0], l[1]) in [(t[POS], t[LEMMA]) for t in new_lemmas]:
                    if (l[0], l[1]) in new_lemmas:
                        is_new = True

                    # if (l[0], l[1]) in [(t[POS], t[LEMMA]) for t in given_lemmas]:
                    if (l[0], l[1]) in given_lemmas:
                        is_given = True

                    is_match = False
                    if is_new or is_given:
                        is_left = False
                        left_t = None
                        for sl in s['lemmas']:

                            if sl[LEMMA] == l[1]:

                                # LOCAL
                                # t = 0.POS, 1.WORD, 2.LEMMA, 3.ISLEFT, 4.DEP, 5.STEM,
                                # 6.LINKS, 7.QUOTE, 8.WORD_POS, 9.DS_DATA,
                                # 10.PARA_POS, 11.SENT_POS, 12.NEW, 13.GIVEN, 14.IS_SKIP, 15.IS_TOPIC

                                if sl[ISLEFT]:
                                    is_left = True
                                    left_t = sl + tuple([p_count, s_count, is_new, is_given, bSkipPunct, True])
                                elif is_left is False:
                                    t = sl + tuple([p_count, s_count, is_new, is_given, bSkipPunct, False])

                                is_match = True
                                break   # 2021.11.21

                        if is_match:
                            if is_left and left_t is not None:
                                sres.append(left_t)
                            else:
                                sres.append(t)

                    if is_match is False:
                        sres.append(tuple([None, bSkipPunct]))

                # for each lemmas in 'all_lemmas'

                res.append(sres)
                s_count += 1

                prev_given_lemmas = given_lemmas
                prev_para = p

            res.append([-1]*(len(all_lemmas)+2))
            p_count += 1

        res = res[:-1]

        return res

    def getGlobalTopicalProgData(self, sort_by=TOPIC_SORT_APPEARANCE):

        if self.lexical_overlaps_analyzed:                # if lexical overlaps have been analyzed already
            if self.global_topical_prog_data is not None: # if global_topical_prog_data are already genreated, return themn.
                return self.global_topical_prog_data      # otherwise, we'll generate new global_topical_prog_data

        # first we should make a list of given lemmas as they appear in the text
        all_lemmas = list()
        temp = list()

        if self.sections is None or len(self.sections) <= self.current_section:
            return

        data      = self.sections[self.current_section]['data']
        #para_data = self.sections[self.current_section]['para_data']

        if data is None or 'paragraphs' not in data:
            return

        for p in data['paragraphs']:                 # for each paragraph
            for l in p['new_accum_lemmas']:      # for each new lemma
                t = (l[POS], l[LEMMA])           # tuple = (POS, LEMMA)
                if t not in temp:                # if the lemma 'l' is not already in the list,
                    all_lemmas.append((l[POS], l[LEMMA],l[8]))  # append a tuple (POS, LEMMA, word_pos)
                    temp.append(t)
            for l in p['given_accum_lemmas']:    # for each new lemma
                t = (l[POS], l[LEMMA])           # tuple = (POS, LEMMA)
                if t not in temp:                # if the lemma 'l' is not already in the list,
                    all_lemmas.append((l[POS], l[LEMMA],l[8]))  # append a tuple (POS, LEMMA, word_pos)
                    temp.append(t)

        # Count how many times each topic appears on the left side of a sentence.
        sent_topic_counter = Counter()
        para_topic_set = set()
        p_count = 1
        for p in data['paragraphs']:        # for each paragraph
            s_count = 1
            for s in p['sentences']:             # for each sentence
                for l in all_lemmas:             # for each topic candidate
                    is_new   = False
                    is_given = False

                    if (l[0], l[1]) in [(t[POS], t[LEMMA]) for t in p['new_accum_lemmas']]:
                        is_new = True
                    if (l[0], l[1]) in [(t[POS], t[LEMMA]) for t in p['given_accum_lemmas']]:
                        is_given = True

                    if is_new or is_given:            # if 'l' is given or new
                        for sl in s['lemmas']:
                            if sl[LEMMA] == l[1] and sl[POS] == l[0]:
                                if sl[ISLEFT]:
                                    sent_topic_counter[sl[LEMMA]] += 1
                                    para_topic_set.add( (p_count, sl[LEMMA]) )
            p_count += 1

        para_topic_counter = Counter()
        for t in para_topic_set:
            para_topic_counter[t[1]] += 1

        all_lemmas = list()
        temp = list()
        for p in data['paragraphs']:                 # for each paragraph
            for l in p['new_accum_lemmas']:      # for each new lemma
                t = (l[POS], l[LEMMA])           # tuple = (POS, LEMMA)
                if t not in temp:                # if the lemma 'l' is not already in the list,
                    s_topic_count = sent_topic_counter[l[LEMMA]]
                    p_topic_count = para_topic_counter[l[LEMMA]]
                    all_lemmas.append((l[POS], l[LEMMA],l[8], s_topic_count, p_topic_count))  # append a tuple (POS, LEMMA, word_pos)
                    temp.append(t)
            for l in p['given_accum_lemmas']:    # for each new lemma
                t = (l[POS], l[LEMMA])           # tuple = (POS, LEMMA)
                if t not in temp:                # if the lemma 'l' is not already in the list,
                    s_topic_count = sent_topic_counter[l[LEMMA]]
                    p_topic_count = para_topic_counter[l[LEMMA]]
                    all_lemmas.append((l[POS], l[LEMMA],l[8], s_topic_count, p_topic_count))  # append a tuple (POS, LEMMA, word_pos)
                    temp.append(t)

        ######
        all_lemmas.sort(key=lambda tup: tup[2])      # sort by the order of appearance.
        if sort_by == TOPIC_SORT_LEFT_COUNT:
            all_lemmas.sort(key=lambda tup: tup[3], reverse=True)      # sort by the total left count

        res = list()
        res.append(all_lemmas)
        collapsed_res = list()
        res.append([-1]*(len(all_lemmas)+2))

        p_count = 1
        pres = [None] * (len(data['paragraphs'])+2)
        pres[0] = all_lemmas
        pres[1] = ([-1]*(len(all_lemmas)+2))

        bSkipPunct = False

        # NOTE: Create a dictionary that keeps track of given/new counts
        # given_new_counts = dict()

        for p in data['paragraphs']:        # for each paragraph
            s_count = 1

            for s in p['sentences']:             # for each sentence
                # if s['text'] is a single character, and one of these punct chracters
                # we'll make the row as bSkipPunct == True.
                bSkipPunct = len(s['text']) == 1 and \
                       (s['text'] in dashes or \
                        s['text'] in hyphen_slash or \
                        s['text'] in left_quotes or \
                        s['text'] in right_quotes or \
                        s['text'] == ellipsis)
                sres = list()                    # get an empty list
                for l in all_lemmas:             # for each topic candidate
                    is_new   = False
                    is_given = False

                    if (l[0], l[1]) in [(t[POS], t[LEMMA]) for t in p['new_accum_lemmas']]:
                        is_new = True
                    if (l[0], l[1]) in [(t[POS], t[LEMMA]) for t in p['given_accum_lemmas']]:
                        is_given = True

                    is_match = False
                    if is_new or is_given:            # if 'l' is given or new
                        is_left = False
                        left_t = None
                        for sl in s['lemmas']:
                            # and the lemma is also in the sentence.
                            if sl[LEMMA] == l[1] and sl[POS] == l[0]:

                                # Global
                                # t = 0.POS, 1.WORD, 2.LEMMA, 3.ISLEFT, 4.DEP, 5.STEM,
                                # 6.LINKS, 7.QUOTE, 8.WORD_POS, 9.DS_DATA, 10.PARA_POS, 11.SENT_POS
                                # 12.NEW, 13.GIVEN, 14.IS_SKIP, 15.IS_TOPIC

                                if sl[ISLEFT]:
                                    is_left = True
                                    left_t = sl + tuple([p_count, s_count, is_new, is_given, bSkipPunct, is_left])
                                elif is_left is False:
                                    t = sl + tuple([p_count, s_count, is_new, is_given, bSkipPunct, is_left])

                                is_match = True
                                break   # 2021.11.21

                        if is_match:
                            if is_left and left_t is not None:
                                pres[p_count+1] = left_t
                                sres.append(left_t)
                            else:
                                pres[p_count+1] = t
                                sres.append(t)

                    if is_match is False:
                        sres.append(tuple([None, bSkipPunct]))

                    # end of 'for l in all_lemmas:'

                # append a sentencce
                collapsed_res.append(pres)
                res.append(sres)

                s_count += 1
                # end of 'for s in p['sentences']:'

            res.append([-1]*(len(all_lemmas)+2))
            p_count += 1

            # end of 'for p in data['paragraphs']:''

        # print("para_data: ----------------")  # for debugging...
        # for pd in para_data:
            # pd.printData()

        self.global_topical_prog_data = {'data': res, 'para_data': []}

        return self.global_topical_prog_data

    def getSentStructureData(self):
        p_count = 1
        res = list()
        res.append("\n") # paragraph break
        bSkipPunct = False

        if self.sections is None:
            return res

        if len(self.sections) <= self.current_section:
            return res

        data = self.sections[self.current_section]['data']
        if data is None or 'paragraphs' not in data:
            return res

        num_paras = len(data['paragraphs'])
        i = 0

        while i < num_paras:

            p = data['paragraphs'][i]
            style = p['style']

            if style.startswith("List"):
                s_count = 1
                while style.startswith("List"):
                    for s in p['sentences']:             # for each sentence
                        if len(s['text']) == 1 and \
                               (s['text'] in dashes or \
                                s['text'] in hyphen_slash or \
                                s['text'] in left_quotes or \
                                s['text'] in right_quotes or \
                                s['text'] == ellipsis):
                            bSkipPunct = True
                        else:
                            bSkipPunct = False
                        res.append(tuple([p_count, s_count, s, bSkipPunct]))
                        s_count += 1

                    i += 1
                    p = data['paragraphs'][i]
                    style = p['style']

            else:
                s_count = 1
                for s in p['sentences']:             # for each sentence
                    bSkipPunct = len(s['text']) == 1 and \
                           (s['text'] in dashes or \
                            s['text'] in hyphen_slash or \
                            s['text'] in left_quotes or \
                            s['text'] in right_quotes or \
                            s['text'] == ellipsis)
                    res.append(tuple([p_count, s_count, s, bSkipPunct]))
                    s_count += 1

                i+= 1

            res.append("\n")                     # paragraph break
            p_count += 1

        return res

    ########################################
    #
    # Debugging tools
    #
    ########################################
    def saveJSON(self, filepath):
        """ Write json to <filepath>.json file."""
        filename = "{}.json".format(filepath)
        with open(filename, "w") as fout:
            fout.write(json.dumps(self.sections, indent=4))

    def exportToCSV(self, filepath):
        """
        e.g., ["Dog", 'dog']
        POS      = 0          # Part of Speech
        WORD     = 1          # Original Word
        LEMMA    = 2          # Lemmatized word
        ISLEFT   = 3          # True if the word is on the left side of the main verb
        SUBJ     = 3          # True if the word is on the left side of the main verb (SUBJ = ISLEFT) same thing!!
        DEP      = 4          # Dependency Type (e.g., 'pobj', 'nsubj', etc.)
        STEM     = 5          # Stem of the word
        LINKS    = 6          # If it is a verb, it is the total # of links from the verb. Otherwise, None.
        QUOTE    = 7          # QUOTED or not
        WORD_POS = 8          # Indexes between 8 and 13 have been changed on July 17.
        DS_DATA  = 9          # DocuScope Data (i.e., models.DSWord)

        """

        output = []
        if self.sections:
            data = self.sections[self.current_section]['data']
            for para in data['paragraphs']:    # for each paragraph
                for sent in para['sentences']: # for each sentence
                    for w in sent['text_w_info']:
                        line = list(w[:DS_DATA])
                        dsw = w[DS_DATA]
                        if dsw is not None:
                            line = line + [dsw.getEnd(), dsw.getLAT(), dsw.getPos()]
                            output.append(line)

            with open(filepath, 'w') as fout:
                for line in output:
                    fout.write("{}\n".format(str(line)[1:-1]))

    def getSentencesWithDimension(self, cluster_name, dimension_name):

        if cluster_name is None or dimension_name is None:
            return []

        if self.sections:
            data = self.sections[self.current_section]['data']

            sent_list = list()
            pcount = 1
            for para in data['paragraphs']:    # for each paragraph

                scount = 1
                for sent in para['sentences']: # for each sentence

                    wcount = 1
                    for w in sent['text_w_info']:

                        dsw = w[DS_DATA]
                        if dsw is not None:
                            lat = dsw.getLAT()
                            dim, clust = self.dictionary.getDimensionAndCluster(lat)

                            if clust == cluster_name and \
                               dim == dimension_name and \
                               (pcount, scount) not in sent_list:

                                sent_list.append((pcount, scount, wcount))

                        wcount += 1

                    scount += 1

                pcount += 1

            if sent_list:
                return sent_list
        return []

    def getSentencesWithCluster(self, cluster_name):

        if cluster_name is None:
            return []

        if self.sections:
            data = self.sections[self.current_section]['data']

            sent_list = list()
            pcount = 1
            for para in data['paragraphs']:    # for each paragraph

                scount = 1
                for sent in para['sentences']: # for each sentence

                    wcount = 1
                    for w in sent['text_w_info']:

                        dsw = w[DS_DATA]
                        if dsw is not None:
                            lat = dsw.getLAT()

                            dim, clust = self.dictionary.getDimensionAndCluster(lat)
                            if clust == cluster_name and (pcount, scount) not in sent_list:
                                sent_list.append((pcount, scount, wcount))

                        wcount += 1

                    scount += 1

                pcount += 1

            if sent_list:
                return sent_list
        return []

    def getSentencesWithImpression(self, impression):

        if impression is None:
            return []

        if self.sections:
            data = self.sections[self.current_section]['data']

            sent_list = list()
            pcount = 1
            for para in data['paragraphs']:    # for each paragraph

                scount = 1
                for sent in para['sentences']: # for each sentence

                    wcount = 1
                    for w in sent['text_w_info']:
                        lemma = w[LEMMA]
                        if lemma == impression:
                            sent_list.append((pcount, scount, wcount))
                        wcount += 1

                    scount += 1

                pcount += 1

            if sent_list:
                return sent_list
            return []

    ### UNUSED NOW ####
    # def getAdjacentDSCategories(self, topic):

    #     def isTopicInSent(sent, topic):
    #         for w in sent['text_w_info']:
    #             if w[lemma] == topic:
    #                 return True
    #         return False

    #     if self.sections:
    #         data = self.sections[self.current_section]['data']

    #         patterns = []

    #         for para in data['paragraphs']:    # for each paragraph
    #             for sent in para['sentences']: # for each sentence

    #                 lats_sentence  = []
    #                 lats_paragraph = []
    #                 lats_adjacent  = []

    #                 if isTopicInSent(sent):
    #                     # if 'topic' is in this sentence, add all the docuscope categories
    #                     # in this sentence to the result.
    #                     for w in sent['text_w_info']:
    #                         dsw = w[DS_DATA]
    #                         end = dsw.getEnd()
    #                         lat = getLAT()
    #                         pos = getPos()

    #                     line = line + [dsw.getEnd(), dsw.getLAT(), dsw.getPos()]

    #         return adj_categories

    #     return None

    def calculateAdjacencyStats(self, topic=None, para_pos=-1):
        """
        This method is used for calculating collocations between Topics and DocuScope categories.

        TODO: para_pos starts with 0, but pcount starts with 1. :-) FIXIT.
        """

        if self.controller is None:
            return

        if self.sections is None:
            return "Error in calculateAdjacencyStats()"
        else:
            try:
                data = self.sections[self.current_section]['data']
                if data is None:
                    raise ValueError
            except:
                return

        adj_stats = AdjacencyStats(topic=topic, controller=self.controller)

        topic_filter = self.controller.getTopicFilter()

        # Let's find which paragraphs/sentences the selected 'topic' is included.
        pcount = 1
        for para in data['paragraphs']:
            scount = 1

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1
                continue

            for sent in para['sentences']:

                wcount = 1
                for w in sent['text_w_info']:

                    if w[LEMMA] == topic and (w[POS] == 'NOUN' or w[POS] == 'PRP'):
                        if topic_filter == TOPIC_FILTER_LEFT and w[ISLEFT] == False:  ## NEW 3/2/21
                            pass
                        else:
                            adj_stats.addParagraphID(pcount)
                            adj_stats.addSentenceID(pcount, scount)
                            adj_stats.addTopicPosition(pcount, scount, wcount)

                    wcount += 1

                scount+=1

            pcount+=1

        bTagOpen = False
        bParaCrossingPattern = False

        dsw   = None
        lat   = ''
        end   = ''
        pos   = ''
        clust_list = list()
        dim_list   = list()
        lat_list   = list()
        pcount = 1

        for para in data['paragraphs']:    # for each paragraph
            scount = 1

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1
                continue

            if not adj_stats.isTopicPara(pcount):
                # skip the paragraphs that do not include the topic word.
                pcount += 1
                continue

            for sent in para['sentences']:

                if not adj_stats.isTopicSent(scount, pcount):
                    # skip the sentences that do not include the topic word.
                    bParaCrossingPattern = False
                    scount += 1
                    continue

                total_words = len(sent['text_w_info'])
                word_count = 0

                for w in sent['text_w_info']:
                    dsw   = w[DS_DATA]      # docuscope word

                    if dsw is not None:
                        lat   = dsw.getLAT()

                        if lat is not None:
                            # pos = int(round(float(pos)))
                            pos   = dsw.getPos()
                            if pos == 0:                # The first word of a DS TAG
                                if bTagOpen is True:    # if a tag is open, close it.
                                    # New tag
                                    # adj_stats.addLAT(lat) no no... we are closing the tag here.
                                    bTagOpen = False

                                if bTagOpen is False:
                                    if word_count == (total_words-1) and w[WORD] == '.' and pos == 0 and lat != '':
                                        # new tag starting with a period at the end of a sentence.
                                        bParaCrossingPattern = True

                                    else:
                                        # New tag
                                        adj_stats.addLAT(lat, pcount, scount)
                                        bTagOpen = True

                            elif bParaCrossingPattern is True:
                                bTagOpen = True
                                # New tag
                                adj_stats.addLAT(lat, pcount, scount)
                                bParaCrossingPattern = False

                    word_count += 1

                scount += 1
            # sent
            # adj_stats.setText(para['text'], pcount)

            pcount += 1

        # adj_stats.print()
        return adj_stats

    def calculateAggregatedAdjacencyStats(self, topics=None, para_pos=-1):
        """
        This method is used for calculating collocations between Topics and DocuScope categories.

        para_pos starts with 0, but pcount starts with 1. :-) FIXIT.
        """

        if self.controller is None:
            return

        if self.sections is None:
            return "Error in calculateAggregatedAdjacencyStats()"
        else:
            try:
                data = self.sections[self.current_section]['data']
                if data is None:
                    raise ValueError
            except:
                return

        topic_filter = self.controller.getTopicFilter()
        adj_stats = AdjacencyStats(controller=self.controller)

        # Let's find which paragraphs/sentences the selected 'topic' is included.
        pcount = 1
        for para in data['paragraphs']:
            scount = 1

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1
                continue

            for sent in para['sentences']:
                for w in sent['text_w_info']:
                    if w[LEMMA] in topics:
                        if topic_filter == TOPIC_FILTER_LEFT and w[ISLEFT] is False:  # NEW 3/2/21
                            pass
                        adj_stats.addParagraphID(pcount)
                        adj_stats.addSentenceID(pcount, scount)
                scount+=1
            pcount+=1

        bTagOpen = False
        bParaCrossingPattern = False

        dsw   = None
        lat   = ''
        end   = ''
        pos   = ''
        clust_list = list()
        dim_list   = list()
        lat_list   = list()
        pcount = 1

        for para in data['paragraphs']:    # for each paragraph
            scount = 1

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1
                continue

            if not adj_stats.isTopicPara(pcount):
                # skip the paragraphs that do not include the topic word.
                pcount += 1
                continue

            for sent in para['sentences']:

                if not adj_stats.isTopicSent(scount, pcount):
                    # skip the sentences that do not include the topic word.
                    bParaCrossingPattern = False
                    scount += 1
                    continue

                total_words = len(sent['text_w_info'])
                word_count = 0

                for w in sent['text_w_info']:
                    dsw = w[DS_DATA]      # docuscope word

                    if dsw is not None:
                        lat = dsw.getLAT()
                        pos = dsw.getPos()
                        pos = int(round(float(pos)))

                        if lat is not None:
                            if pos == 0:                # The first word of a DS TAG
                                if bTagOpen is True:    # if a tag is open, close it.
                                    # New tag
                                    # adj_stats.addLAT(lat) no no... we are closing the tag here.
                                    bTagOpen = False

                                if bTagOpen is False:
                                    if word_count == (total_words-1) and w[WORD] == '.' and pos == 0 and lat != '':
                                        # new tag starting with a period at the end of a sentence.
                                        bParaCrossingPattern = True

                                    else:
                                        # New tag
                                        adj_stats.addLAT(lat, pcount, scount)
                                        bTagOpen = True

                            elif bParaCrossingPattern is True:
                                bTagOpen = True
                                # New tag
                                adj_stats.addLAT(lat, pcount, scount)
                                bParaCrossingPattern = False

                            elif word_count == 0 and pos == 1:
                                # A pattern started in the previous sentence (e.g,. ". He"), but
                                # the previous sentenc is not a topic sentence.
                                adj_stats.addLAT(lat, pcount, scount)
                                bTagOpen = True

                    word_count += 1

                scount += 1
            # sent

            adj_stats.setText(para['text'], pcount)

            pcount += 1

        return adj_stats


    def calculateParaStats(self, para_pos=-1):
        """
        This function is used to get the stats spefific to the given paragraph.
        It is used by the controller to update the dictionary tree.
        """

        if self.controller is None:
            return

        if self.sections is None:
            return "Error in calculateAdjacencyStats()"
        else:
            try:
                data = self.sections[self.current_section]['data']
                if data is None:
                    raise ValueError
            except:
                return

        adj_stats = AdjacencyStats(topic=None, controller=self.controller)

        # Let's find which paragraphs/sentences the selected 'topic' is included.
        pcount = 1
        for para in data['paragraphs']:
            scount = 1

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1

                continue

            for sent in para['sentences']:
                for w in sent['text_w_info']:
                    adj_stats.addParagraphID(pcount)
                    adj_stats.addSentenceID(pcount, scount)
                scount+=1
            pcount+=1

        bTagOpen = False
        bParaCrossingPattern = False

        dsw   = None
        lat   = ''
        end   = ''
        pos   = ''
        clust_list = list()
        dim_list   = list()
        lat_list   = list()
        pcount = 1

        for para in data['paragraphs']:    # for each paragraph
            scount = 1

            if para_pos != -1 and (para_pos+1) != pcount:
                pcount += 1
                continue

            for sent in para['sentences']:

                total_words = len(sent['text_w_info'])
                word_count = 0

                for w in sent['text_w_info']:
                    dsw   = w[DS_DATA]      # docuscope word

                    if dsw is not None:
                        lat   = dsw.getLAT()
                        pos   = dsw.getPos()
                        pos = int(round(float(pos)))

                        if lat is not None:
                            if pos == 0:                # The first word of a DS TAG
                                if bTagOpen is True:    # if a tag is open, close it.
                                    # New tag
                                    # adj_stats.addLAT(lat) no no... we are closing the tag here.
                                    bTagOpen = False

                                if bTagOpen is False:
                                    if word_count == (total_words-1) and w[WORD] == '.' and pos == 0 and lat != '':
                                        # new tag starting with a period at the end of a sentence.
                                        bParaCrossingPattern = True

                                    else:
                                        # New tag
                                        adj_stats.addLAT(lat, pcount, scount)
                                        bTagOpen = True

                            elif bParaCrossingPattern is True:
                                bTagOpen = True
                                # New tag
                                adj_stats.addLAT(lat, pcount, scount)
                                bParaCrossingPattern = False

                    word_count += 1

                scount += 1
            # sent
            pcount += 1

        return adj_stats

    ########################################
    #
    # DocuScope Data Processing
    #
    ########################################

    def setDSCategories(self, section_id=None):
        list_of_errors = list()
        if self.sections is not None:
            if section_id is not None:
                errors = self.setDSCategoriesOne(section_id)
                if errors['positions']:
                    list_of_errors = [errors]
                self.calculateDSTemporalStats(section_id)
            else:
                for section_id in range(len(self.sections)):
                    errors = self.setDSCategoriesOne(section_id)
                    if errors['positions']:
                        list_of_errors.append(errors)
                    self.calculateDSTemporalStats(section_id)

        return list_of_errors

    # per section
    def setDSCategoriesOne(self, section_id):

        errors = dict()
        errors['positions'] = list()

        def makeError(section_id, pcount, scount, patterns):

            pkey = 'sect{}_p{}'.format(section_id, pcount)
            if pkey not in errors:
                errors[pkey] = para['text']

            skey = 'sect{}_p{}s{}'.format(section_id, pcount, scount)
            if skey not in errors:
                errors[skey] = sent['text']

                d = dict()
                d['sect_pos'] = section_id
                d['ppos'] = pcount
                d['spos'] = scount
                d['patterns'] = (dsw.getWord(), otw[WORD])
                errors['positions'].append(d)

        section_data = self.sections[section_id]
        ot_data = section_data['data']
        ds_data = self.generateOTCompatibleData(section_id)

        ds_count = 0
        if ot_data is not None:
            is_ot_data = True
            pcount = 0

            for para in ot_data['paragraphs']:    # for each paragraph
                scount = 0
                for sent in para['sentences']:     # for each sentence
                    for ot_count in range(len(sent['text_w_info'])):

                        if ds_count >= len(ds_data):
                            break

                        dsw = ds_data[ds_count]                  # docuscope data
                        otw = sent['text_w_info'][ot_count]      # ontopic data

                        if dsw.getWord() == otw[WORD]: # match
                            sent['text_w_info'][ot_count] = otw[:-1] + (dsw,)

                        elif (ds_count+1) < len(ds_data): # no match (misalignement)

                            dsw_word = dsw.getWord()
                            next_dsw = ds_data[ds_count+1]
                            next_dsw_word = next_dsw.getWord()
                            otw_word = otw[WORD]

                            if next_dsw_word == '.':
                                # We have a case where a period is used in abbreviations etc.
                                # (not at the end of a sentence.)
                                sent['text_w_info'][ot_count] = otw[:-1] + (dsw,)
                                ds_count += 1

                            elif otw_word == '\'s' and dsw_word == '\'' and next_dsw_word == 's':
                                # Apostorophe + s.
                                sent['text_w_info'][ot_count] = otw[:-1] + (dsw,)
                                ds_count += 1

                            elif otw_word.startswith(dsw_word):
                                pos = len(dsw_word)
                                temp_word = dsw_word
                                while pos < len(otw_word):
                                    next_dsw = ds_data[ds_count+1]
                                    next_dsw_word = next_dsw.getWord()
                                    temp_word += next_dsw_word
                                    ds_count+=1
                                    pos += len(next_dsw_word)

                                if temp_word == otw_word:
                                    sent['text_w_info'][ot_count] = otw[:-1] + (dsw,)

                                else:
                                    # we couldn't find a way to resolve the error. It can be
                                    # caused by generateOTCompatibleData() or OT parser messed up the
                                    # order of words.
                                    makeError(section_id, pcount, scount, (dsw.getWord(), otw[WORD]))

                            else:
                                makeError(section_id, pcount, scount, (dsw.getWord(), otw[WORD]))

                        ds_count += 1
                    scount += 1
                pcount += 1

        return errors

    def generateOTCompatibleData(self, section_id, debug=False):

        def endswith_no_space_pattern(w):
            for p in no_space_patterns:
                if w.endswith(p):
                    return len(w) - len(p)
            return 0

        def slashed_word(w):

            l = w.split('/')

            if len(l) == 1:
                return False

            temp = list()
            for i in range(len(l)):
                temp.append(l[i])
                if i < len(l)-1:
                    temp.append('/')

            res = list()
            for i in range(len(temp)):
                if temp[i] != '':
                    res.append(temp[i])

            return res

        def hyphenated_word(w):
            l = w.split('-')

            if len(l) == 1:
                return False

            temp = list()
            for i in range(len(l)):
                temp.append(l[i])
                if i < len(l)-1:
                    temp.append('-')

            res = list()
            for i in range(len(temp)):
                if temp[i] != '':
                    res.append(temp[i])

            return res

        if self.sections is None or self.current_section >= len(self.sections):
            return

        token_data = self.sections[section_id].get('token_data','')

        output = list()
        word_count = 0

        # b_in_html_element = False
        word = ''
        prev_word = ''

        for line in token_data.split('\n'):   # each line is a word
            if line:
                # replace commas with '\u0000'
                line = re.sub('\",\"(?=,)', '\u0000', line)
                token_info = line.split(',')

                # get a list of numbers with thousand separator(s)
                # num_list = re.findall("\"[0-9,.]+\"(?=,)", line)
                num_list = re.findall(r"\"[\sA-Za-z=0-9,.]+\"(?=,)", line)   # revised Apr 20, 2021
                if num_list:
                    token_info = num_list + token_info[-3:]
            else:
                continue

            if len(token_info) != 5:
                continue

            word_count += 1
            prev_word = word
            word  = token_info[0]    # original word
            lword = token_info[1]    # lowercase word
            end   = token_info[2]    # end reason
            lat   = token_info[3]    # category
            pos   = token_info[4]    #

            if pos != '':
                pos = int(token_info[4])
            else:
                pos = -1

            if word in right_quotes:   # New (May 4, 2020)
                end = 'c'

            words = None

            # if word ends with a no_space_pattern:
            split_at = endswith_no_space_pattern(word)
            if split_at > 0:
                w1 = word[:split_at]
                w2 = word[split_at:]

                res = hyphenated_word(w1)
                if res is not False:
                    res.append(w2)
                    words = res
                else:
                    words = [w1, w2]
            else:
                res = hyphenated_word(word)
                if res is not False:
                    words = res
                else:
                    res = slashed_word(word)
                    if res is not False:
                        words = res

            if words is not None:
                # 'word' is split now so len(words) should be > 1.
                num_words = len(words)  # 2
                if num_words > 1:
                    # just in case, check if num_words is > 1.
                    for i in range(num_words):  # 0, 1
                        # add each word to 'output'
                        w = words[i]                    # get the i-th word.   # i=0, 1
                        if i < num_words-1:             # num_words-1 = 0
                            line = [w, w.lower(), 'c', lat, pos+i]
                        else:
                            # use 'end' at the end.
                            line = [w, w.lower(), end, lat, pos+i]
                        output.append(line)
                else:
                    line = [word, lword, end, lat, pos]
                    output.append(line)
            else:
                if word == '\u0000':
                    word  = ','
                    lword = ','

                line = [word, lword, end, lat, pos]
                output.append(line)

        # [word, lword, end, lat, pos]
        res = [DSWord(line[0], line[1], line[2], line[3], line[4]) for line in output]

        # # Debugging
        if debug:
            error_table = ""
            ot_words = list()
            for section in self.sections:
                data = section['data']
                for p in data['paragraphs']:
                    for s in p['sentences']:
                        for w in s['text_w_info']:
                            ot_words.append(w[WORD])


            for i in range(len(output)):
                line = output[i]
                if i >= len(ot_words):
                    break
                w = ot_words[i]
                error_table += "{},{},{},{},{},{}\n".format(line[0], line[1], line[2], line[3], line[4], w)

            # csv_file = "debugging_{}.csv".format(section_id)
            # with open(csv_file, 'w') as fout:
                # fout.write(error_table)
            return error_table
        return res

    def processTokenData(self, section_id=None, section=None):
        """ UNUSED!!!
        This function uses the token dataset (via csv file) created by DocuScope for a section,
        and generate a set of tags (with positions etc.). Tags data are used to identify
        the category of DS patterns in the editor.
        """
        token_data = ''
        if section is None and section_id is not None:
            section = self.sections[section_id]
            token_data = section.get('token_data', '')
        if section is not None:
            token_data = section.get('token_data', '')

        bTagOpen = False
        bParaCrossingPattern = False

        end_str = ''

        pcount = 0

        pattern = ""
        tags = dict()
        tags[0] = list()

        word = ''
        pat_start = 0
        pat_end   = 0
        pattern_t = None

        lines = token_data.split('\n')

        for i in range(len(lines)):
            line = lines[i]

            if line:
                # replace commas with '\u0000'
                line = re.sub('\",\"(?=,)', '\u0000', line)
                token_info = line.split(',')

                # Get a list of numbers with thousand separator(s), followed by a comma separator
                # This also captures patterns like N=2,000 etc.
                # num_list = re.findall("\"[0-9,.]+\"(?=,)", line)  # old
                num_list = re.findall(r"\"[\sA-Za-z=0-9,.]+\"(?=,)", line)   # revised Apr 20, 2021
                if num_list:
                    token_info = num_list + token_info[-3:]

            else:
                continue

            prev_word = word
            word = token_info[0]    # original word
            end  = token_info[2]    # end reason
            lat  = token_info[3]    # category
            pos  = token_info[4]

            quoted_num = re.search('\"[0-9,.]+\"', word)
            if quoted_num is not None:
                word = word[1:-1]

            if word in right_quotes and prev_word == ellipsis:          # an ellipsis followed by a right quote
                end_str = end_str[:-1]

            elif word in end_puncts and prev_word in right_quotes:      # a right quote followed by an end punct.
                end_str = end_str[:-1]

            elif word in end_puncts and prev_word == ellipsis:          # an ellipsis followed by an end punct
                end_str = end_str[:-1]

            elif word == ellipsis and end_str and end_str[-1] != ' ':   # no space before an ellipsis. add a space.
                end_str += ' '

            elif word in hyphen_slash and end_str and end_str[-1] == ' ':  # a space before a hyphen/slash, there
                end_str = end_str[:-1]                                     # must be a space after a hyphen/slash

            elif word == 'not' and prev_word == 'can' and i+1 < len(lines):
                # a special case. We need to remove the space between "can" and "not",
                # unless it is followed by 'only'.
                next_line = lines[i+1]                    # get the next line
                next_token_info = next_line.split(',')    # we don't worry about it being a number.
                next_word = next_token_info[0]            # next token
                if next_word != 'only':                   # if the next word is not "only",
                    end_str = end_str[:-1]                # remove the space after "can"

            elif word == "%" and end_str and end_str[-1] == ' ':
                end_str = end_str[:-1]


            dim, clust = self.dictionary.getDimensionAndCluster(lat)

            if pos != '':
                pos = int(token_info[4])
            else:
                pos = -1

            if pos == 0:
                if bTagOpen == True:         # if a tag is open, close it.
                    bTagOpen = False

                    pat_end = pat_start + len(pattern)
                    tags[pcount].append( ((pat_start, pat_end), pattern, (prev_clust, prev_dim, prev_lat)) )

                    if len(end_str) <= 1:
                        pat_start = pat_end + len(end_str)
                    else:
                        pat_start = pat_end

                    pattern = ''

                if bTagOpen == False:        # if a tag is not open, start a new one.
                    if end == 'n' and word == '.' and pos == 0 and lat != '':
                        # If the period at the end of a pragraph is part of a pattern.
                        bParaCrossingPattern = True
                    else:
                        if lat == '' or lat == 'UNRECOGNIZED':
                            lat = 'na'
                        else:
                            pass
                        bTagOpen = True

            elif bParaCrossingPattern:
                if len(end_str) <= 1:
                    pattern += end_str

                bTagOpen = True
                bParaCrossingPattern = False
            else:
                if len(end_str) <= 1:
                    pattern += end_str

            if word == '\"\"\"\"':
                pattern += '\"'
            elif word == '\'\'\'\'':
                pattern += '\"'
            elif word == '\\':
                pattern += '\\'
            elif word == '\u0000':
                pattern += ','
            else:
                pattern += word

            end_str = ''
            if end in ['n', 'nn', 'nnn', 'nnnn', 'nnnnn']:
                pcount+=1
                end_str += ''
                tags[pcount] = list()
                pat_start = 0
                pattern = ''

            elif word == ellipsis:
                end_str += ' '

            elif word in left_quotes or word in hyphen_slash:
                end_str += ''

            elif end == 's':           # space
                end_str += ' '

            elif end == 'c':           # no space
                end_str += ''

            prev_clust = clust
            prev_dim   = dim
            prev_lat   = lat

        if bTagOpen is True:
            pat_end = pat_start + len(pattern)
            tags[pcount].append( ((pat_start, pat_end), pattern, (clust, dim, lat)) )

            if len(end_str) <= 1:
                pat_start = pat_end + len(end_str)
            else:
                pat_start = pat_end
            # pat_start = pat_end + 1
            pattern = ''

        section['tags'] = tags

    def calculateDSTemporalStats(self, section_id):

        section_data = self.sections[section_id]
        token_data = section_data.get('token_data', None)

        if token_data is None or self.dictionary is None:
            return

        bTagOpen = False
        bParaCrossingPattern = False

        para_count = 1

        lats_by_para    = list()        # list of lists
        para_data    = ds_stat.DSStat()
        para = ""

        word_count = 0
        for line in token_data.split('\n'):

            line = re.sub('\",\"(?=,)', '\u0000', line)
            token_info = line.split(',')

            if len(token_info) != 5:
                continue

            word_count += 1

            try:
                word = token_info[0]    # original word
                end  = token_info[2]    # end reason
                lat  = token_info[3]    # category
                pos  = token_info[4]

                if pos != '':
                    pos = int(float(token_info[4]))
                else:
                    pos = -1

            except:
                pass

            if pos == 0:                    # first token
                if bTagOpen == True:        # if a tag is open, close it.
                    bTagOpen = False
                    para    += end_str

                if bTagOpen == False:       # if a tag is not open, start a new one.
                    if end == 'n' and word == '.' and pos == 0 and lat != '':
                        # If the period at the end of a pragraph is part of a pattern.
                        bParaCrossingPattern = True
                    else:
                        bTagOpen = True
                        dim, clust = self.dictionary.getDimensionAndCluster(lat)

                        if clust is not None:
                            para_data.addCluster(clust)
                            para_data.addDimension(dim)
                            para_data.addLAT(lat)

            elif bParaCrossingPattern:
                dim, clust = self.dictionary.getDimensionAndCluster(lat)
                if clust is not None:
                    para_data.addCluster(clust)
                    para_data.addDimension(dim)
                    para_data.addLAT(lat)

                bTagOpen = True
                bParaCrossingPattern = False

            else:
                para    += end_str

            if word == '\"\"\"\"':
                para += '\"'

            elif word == '\'\'\'\'':
                para += '\"'

            elif word == '\u0000':
                para += ','

            else:
                para += word

            end_str = ''

            if end in ['n', 'nn', 'nnn', 'nnnn', 'nnnnn']:
                para_data.setText(para.strip())
                lats_by_para.append(para_data)
                para = ""
                para_data = ds_stat.DSStat()
                end_str += ''

                if para_count > (self.skip_paras + self.max_paras - 1):
                    break
                para_count += 1

            elif end == 's':           # no space
                end_str += ' '
            elif end == 'c':
                end_str += ''

        para    = para.strip()

        if para:
            para_data.setText(para.strip())
            lats_by_para.append(para_data)

        self.sections[section_id]['para_data'] = lats_by_para

###############################################################################
#
# The Methods below here are used only by the online version of DocuScope Write & Audit
#
# The following are the top level functions that should be used to retrieve the data.
#     DSDocument.generateLocalVisData(self, selected_paragraphs, max_topic_sents=1, min_topics=2)
#     DSDocument.generateGlobalVisData(self, min_topics=2, max_topic_sents=1, sort_by=TOPIC_SORT_APPEARANCE)
#
# You can use the following methods to test the data genereated by the above methods.
# These functions print the visualization using ASCII characters.
#
#    testPrintLocalVisData(data)
#    testPrintGlobalVisdata(data)
#
###############################################################################

    def generateLocalVisData(self, selected_paragraphs, max_topic_sents=1, min_topics=2):
        """
        This method returns a python dictionary that contains the data that are needed for
        the visualization of coherence across sentences within 1 or more pragraphs.

        selected_paragraphs:    A list of paragraph IDs.
        max_topic_sents:        The total number of sentences at the beginning of paragraphs that
                                should be considered a topic sentence.
        min_topics:             The minimum number of lexical overlaps between sentences.

        """

        self.local_topics = list()

        data = self.getLocalTopicalProgData(selected_paragraphs)
        if data is None:
            return ({'error': 'data is None'})

        topic_filter = TOPIC_FILTER_LEFT_RIGHT
        key_para_topics = self.getKeyParaTopics()
        key_sent_topics = []

        header = data[0]   # list of tuples (POS, LEMMA)
        self.local_header = header
        nrows  = len(data)
        ncols  = len(header)

        if ncols == 0:
            return ({'error': 'ncols is 0'})

        vis_data = dict()
        vis_data['num_topics'] = ncols
        vis_data['data'] = list()

        if selected_paragraphs:
            sent_filter     = self.filterLocalTopics(data, nrows, ncols)
            selected_paragraphs.sort()
            b_para_break = False
            true_left_count = 0
            l_count = 0
            num_topics = 0
            num_sents  = 0
            for ci in range(ncols):

                topic = header[ci][1]
                topic_data = [None] * (nrows-1)
                b_skip     = False
                sent_pos   = 0
                sent_id    = 0
                count = 0

                if topic is not None:
                    true_left_count = sent_filter[topic]['left_count']

                    if topic_filter == TOPIC_FILTER_ALL:
                        count = sent_filter[topic]['count']
                        l_count = sent_filter[topic]['left_count']
                    else:
                        l_count = sent_filter[topic]['left_count']
                        if l_count == 1:
                            count = 2
                        else:
                            count = l_count

                if count < min_topics:
                    continue

                is_global = False
                if topic in self.global_topics:
                    is_global = True

                is_tc = DSDocument.isUserDefinedSynonym(topic)

                self.local_topics.append((topic, is_global))

                topic_info = None
                para_count = 0
                para_id = selected_paragraphs[para_count]

                for ri in range(1,nrows):     # for each row
                    elem= data[ri][ci]      # get the elem

                    if isinstance(elem, int) and elem < 0:
                        b_para_break = True
                    elif isinstance(elem, tuple) and elem[0] is not None and \
                                not is_skip(elem, true_left_count, topic_filter):
                        if elem[IS_SKIP] is False:
                            d = dict()
                            # d['topic'] = elem
                            d['sent_pos'] = sent_id
                            d['para_pos'] = para_id
                            d['is_left'] = elem[ISLEFT]
                            d['is_topic_sent'] = sent_id < max_topic_sents

                            topic_data[sent_pos] = d
                            topic_info = elem

                    elif isinstance(elem, tuple) and elem[0] is not None \
                                and is_skip(elem, true_left_count, topic_filter) is True:
                        b_skip = True

                    elif isinstance(elem, tuple) and elem[0] is None:
                        b_skip = True

                    if b_para_break:
                        para_count += 1
                        para_id = selected_paragraphs[para_count]
                        sent_id = 0
                        b_para_break = False
                    else:
                        sent_pos += 1
                        sent_id  += 1

                vis_data['data'].append({'sentences':        topic_data,
                                         'is_topic_cluster': is_tc,
                                         'is_global':        is_global,
                                         'num_sents':        len(topic_data),
                                         'topic':            list(topic_info[0:3])})
                num_topics += 1

            vis_data['num_topics'] = num_topics

            # debug
            # with open("sample_local_coherence_data.json", 'w') as fout:
                # json.dump(vis_data, fout, indent=4)

            return vis_data

    def generateGlobalVisData(self, min_topics=2, max_topic_sents=1, sort_by=TOPIC_SORT_APPEARANCE):
        """
        This method returns a python dictionary that contains the data that are needed for
        the visualization of coherence across paragraphs.
        min_topics:             The minimum number of lexical overlaps between sentences.
        max_topic_sents:        The total number of sentences at the beginning of paragraphs that
                                should be considered a topic sentence.
        sort_by:                The sorting method option.
        """

        self.recalculateGivenWords(section=0)

        global_data = self.getGlobalTopicalProgData(sort_by=sort_by)
        self.updateLocalTopics()
        self.updateGlobalTopics(global_data)

        topics = self.getCurrentTopics()         # These 2 lines create a data structure that allows us
        self.locateTopics(topics)                # to count the # of sentences each topic appears later. 8/23/2022.

        if global_data is None:
            return {'error': 'ncols is 0'}

        data = global_data['data']
        # para_data = global_data['para_data']

        if data is None:
            return {'error': 'ncols is 0'}

        header = data[0]   # list of tuples (POS, LEMMA, POS, COUNT)

        nrows  = len(data)              # Initialize the number of rows and the number of columns.
        ncols  = len(header)

        self.global_topics = list()

        if ncols == 0:
            return ({'error': 'ncols is 0'})

        vis_data = dict()
        vis_data['num_topics'] = ncols
        vis_data['data'] = list()

        # Filters
        para_filter     = self.filterParaTopics(data, nrows, ncols)
        sent_filter     = self.filterTopics(data, nrows, ncols)
        topic_filter    = TOPIC_FILTER_LEFT_RIGHT

        num_sents_per_topic = 0
        b_break = False
        true_left_count = 0
        l_count = 0
        count = 0

        for ci in range(ncols):                             # for each topic entry,

            topic = header[ci][1]                           # find a topic from the header list.
            topic_data = [None] * (nrows-1)                 # initialize the topic data
            p_ri = 0                                        # initialize the row index w/in paragraph

            if topic is not None:                           # topic exists
                if  self.isLocalTopic(topic) is False and \
                    DSDocument.isUserDefinedSynonym(topic) is False:         # topic cluster = user defined synonym
                    # if the topic is NOT a local topic AND it is NOT a topic cluster,
                    # we should skip this topic.
                    continue

                # Count how manu times the topic appears on the left side of the main verb.
                if DSDocument.isUserDefinedSynonym(topic):                # topic is a topic cluster.
                    true_left_count = sent_filter[topic]['left_count']
                    count   = max(sent_filter[topic]['count'], 2)
                    l_count = sent_filter[topic]['left_count']
                else:                                                            # topic is not a topic sluster
                    true_left_count = sent_filter[topic]['left_count']
                    if topic_filter == TOPIC_FILTER_ALL:                  # rarely used.
                        count   = sent_filter[topic]['count']
                        l_count = sent_filter[topic]['left_count']
                    else:                                                        # default
                        l_count = sent_filter[topic]['left_count']
                        if l_count == 1:                                         # one left + one right case.
                            count = 2
                        else:
                            count = l_count

            if count < min_topics:                    # min_topics == 2 by default. Skip topics that do not apper in more than 2 paragraphs.
                continue

            self.global_topics.append(topic)                    # if we get here, the topic is a global topic.

            is_tc = DSDocument.isUserDefinedSynonym(topic)     # check if topic is a topic cluster.

            topic_info = None
            sent_count = 0

            # we will start with the index == 2 because the first index is the header row,
            # and the second index is the pragraph break indicator.

            for ri in range(2,nrows):     # for each column (r & c are flipped!)

                elem = data[ri][ci]       # get the elem

                if isinstance(elem, int) and elem < 0:       # paragraph brek
                    b_break = True

                # Not the first column (not the paragraph ID/Number)
                elif isinstance(elem, tuple) and elem[0] is not None and \
                            is_skip(elem, true_left_count, topic_filter) is False:

                    curr_elem = topic_data[p_ri]

                    # d['sent_id'] captures the sent id of the first occurence of the topic on the left side.
                    if curr_elem is not None and \
                       elem[ISLEFT] is True and \
                       curr_elem['topic'][ISLEFT] is False:
                        # 'elem' not the first instance for this paragraph
                        # the existing element 'curr_elem' is on the right side.
                        d = dict()
                        d['topic']   = list(elem)
                        d['first_left_sent_id'] = sent_count
                        d['para_pos'] = p_ri
                        d['is_left'] = True
                        # d['is_topic_cluster'] = is_tc

                        if sent_count < max_topic_sents:
                            d['is_topic_sent'] = True
                        else:
                            d['is_topic_sent'] = False

                        topic_data[p_ri] = d

                    elif curr_elem is None:
                        d = dict()
                        d['topic'] = list(elem)
                        d['para_pos'] = p_ri

                        if elem[ISLEFT] is True:
                            d['first_left_sent_id'] = sent_count
                            d['is_left'] = True
                        else:
                            d['first_left_sent_id'] = -1
                            d['is_left'] = False

                        if sent_count < max_topic_sents:
                            d['is_topic_sent'] = True
                        else:
                            d['is_topic_sent'] = False

                        topic_data[p_ri] = d
                        topic_info = elem

                elif isinstance(elem, tuple) and elem[0] is not None \
                            and is_skip(elem, true_left_count, topic_filter) is True:
                    pass

                elif isinstance(elem, tuple) and elem[0] is None:                     # if empty slot
                    pass

                if b_break:
                    p_ri += 1
                    b_break = False
                    sent_count = 0
                else:
                    sent_count += 1

            topic_data = topic_data[0:p_ri]

            for i in range(len(topic_data)):          # delete the topic data.
                d = topic_data[i]
                if d is not None:
                    del d['topic']

            is_non_local = not self.isLocalTopic(topic)

            num_sents_per_topic = self.countSentencesWithTopic(topic)

            vis_data['data'].append({'paragraphs':       topic_data,
                                     'is_topic_cluster': is_tc,
                                     'is_non_local':     is_non_local,
                                     'topic':            list(topic_info[0:3]),
                                     'sent_count':       num_sents_per_topic})

            vis_data['num_paras'] = p_ri

        # Add missing topic clusters, if any
        tcs = DSDocument.getUserDefinedSynonyms()
        if tcs is not None:
            missing_tcs = list(set(tcs) - set(self.global_topics))
            for tc  in missing_tcs:

                topic_info = ['NOUN', '', tc]

                vis_data['data'].append({'paragraphs': [],
                                         'is_topic_cluster': True,
                                         'is_non_local':     False,
                                         'topic':            topic_info,
                                         'sent_count':       0})

        # debug
        # with open("sample_global_coherence_data.json", 'w') as fout:
            # json.dump(vis_data, fout, indent=4)

        return vis_data

    def filterTopics(self, data, nrows, ncols):
        res = dict()
        for c in range(1, ncols+1):  # for each topic

            l_start = -1 # left only
            l_end   = -1
            start = -1  # left or right
            end   = -1
            given_count = 0
            given_left_count  = 0
            given_right_count = 0
            para_count  = 0
            num_paras   = 0
            header = data[0][c-1][1]
            given_left_paras  = list()
            given_right_paras = list()
            first_new_paras   = list()

            for r in range(1, nrows):
                elem = data[r][c-1]

                if isinstance(elem, tuple) and elem[0] is not None:
                    if start < 0:
                        start = r
                    elif start >= 0:
                        end = r

                    if elem[ISLEFT]:
                        if l_start < 0:
                            l_start = r
                        elif l_start >= 0:
                            l_end = r

                        given_left_count += 1
                        given_left_paras.append(para_count)
                    else:
                        given_right_count += 1
                        given_right_paras.append(para_count)

                    given_count += 1

                elif isinstance(elem, int) and elem < 0:
                    para_count += 1

            l_skip_lines = 0
            skip_lines = 0
            sent_count = 0
            for r in range(1, nrows):
                elem= data[r][c-1]

                if r > l_start and r < l_end:
                    if isinstance(elem, str) and elem in ['heading', 'title']:
                        l_skip_lines += 1
                    elif isinstance(elem, int) and elem < 0:
                        l_skip_lines += 1

                if r > start and r < end:
                    if isinstance(elem, str) and elem in ['heading', 'title']:
                        skip_lines += 1
                    elif isinstance(elem, int) and elem < 0:
                        skip_lines += 1

                if not isinstance(elem, (str, int)):
                    sent_count += 1

            given_left_paras  = list(set(given_left_paras))
            given_right_paras = list(set(given_right_paras))

            is_topic = False
            if len(given_left_paras) > 1:
                is_topic = True
            elif len(given_left_paras) == 1:
                # this topic only appears in one paragraph. Let's see if it appears in
                # another pragraph on the right side...
                if len(set(given_right_paras)-set(given_left_paras)) > 0:
                    # This topic satisfies the min requirement to be a topic.
                    is_topic = True

            # left only
            l_span = max((l_end - l_start + 1) - l_skip_lines, 0)
            norm_l_span = (l_span / sent_count) * 100
            norm_l_coverage = (given_left_count / sent_count) * 100

            # left or right
            span = max((end - start + 1) - skip_lines, 0)
            norm_span = (span / sent_count) * 100
            norm_coverage = (given_count / sent_count) * 100

            res[header] = {'type':            0,
                           'is_topic':        is_topic,
                           'left_span':       norm_l_span,
                           'left_coverage':   norm_l_coverage,
                           'span':            norm_span,         # left+right
                           'coverage':        norm_coverage,     # left+right
                           'count':           given_count,
                           'left_count':      given_left_count,
                           'right_count':     given_right_count,
                           'sent_count':      sent_count,
                           'para_count':      None}

        return res

    def filterParaTopics(self, data, nrows, ncols):
        res = dict()
        for c in range(1, ncols+1):  # for each topic

            p_l_start = -1
            p_l_end   = -1
            p_start = -1
            p_end   = -1
            given_count = 0
            given_left_count = 0
            given_right_count = 0

            para_count = 0
            given_paras = list()
            given_left_paras  = list()
            given_right_paras = list()

            header = data[0][c-1][1]

            for r in range(1, nrows):

                elem= data[r][c-1]

                if isinstance(elem, tuple) and elem[0] is not None:
                    if p_start < 0:
                        p_start = para_count
                    elif p_start >= 0:
                        p_end = para_count

                    if elem[ISLEFT]:
                        if p_l_start < 0:
                            p_l_start = para_count
                        elif p_l_start >= 0:
                            p_l_end = para_count

                        given_left_paras.append(para_count)
                        given_left_count += 1
                    else:
                        given_right_paras.append(para_count)
                        given_right_count += 1

                    given_paras.append(para_count)
                    given_count += 1

                elif isinstance(elem, int) and elem < 0:
                    para_count += 1

            para_count -= 1

            p_l_span = (p_l_end - p_l_start + 1)
            norm_l_span = (p_l_span / para_count) * 100

            p_span = (p_end - p_start + 1)
            norm_span = (p_span / para_count) * 100

            given_paras       = list(set(given_paras))
            given_left_paras  = list(set(given_left_paras))
            given_right_paras = list(set(given_right_paras))

            norm_l_coverage   = (len(given_left_paras) / para_count) * 100
            norm_coverage     = (len(given_paras) / para_count) * 100

            is_topic = False
            if len(given_left_paras) > 1:
                is_topic = True
            elif len(given_left_paras) == 1:
                # this topic only appears in one paragraph. Let's see if it appears in
                # another pragraph on the right side...
                if len(set(given_right_paras)-set(given_left_paras)) > 0:
                    # This topic satisfies the min requirement to be a topic.
                    is_topic = True

            res[header] = {'type':           4,
                           'is_topic':       is_topic,
                           'left_span':      norm_l_span,
                           'left_coverage':  norm_l_coverage,
                           'span':           norm_span,
                           'coverage':       norm_coverage,
                           'count':          len(given_paras),
                           'left_count':     len(given_left_paras),
                           'right_count':    len(given_right_paras),
                           'sent_count':     None,
                           'para_count':     para_count}

        return res

    def filterLocalTopics(self, data, nrows, ncols):

        res = dict()
        for c in range(1, ncols+1):  # for each topic

            start = -1
            end   = -1
            given_count = 0
            given_left_count = 0
            # para_count = 0
            first_new_count = 0
            header = data[0][c-1][1]

            for r in range(1, nrows):
                elem = data[r][c-1]

                if isinstance(elem, tuple) and elem[0] is not None:
                    # if para_count == para_pos:
                    bSkip = False

                    if elem[IS_TOPIC] is False: # it's not a topic word
                        bSkip = True

                    if bSkip is False: # it's a topic word
                        if start < 0:
                            start = r
                        elif start >= 0:
                            end = r

                        given_left_count += 1

                        if elem[ISLEFT] is False:
                            first_new_count += 1

                    given_count += 1

            if (given_left_count - first_new_count) == 0:
                start = -1
                end = -1

            skip_lines = 0
            sent_count = 0
            # para_count = 0
            for r in range(1, nrows):
                elem= data[r][c-1]

                if start < r < end:
                    if isinstance(elem, str) and elem == 'heading':
                        skip_lines += 1
                    elif isinstance(elem, int) and elem < 0:
                        skip_lines += 1

                if not isinstance(elem, (str, int)):
                    sent_count += 1

            if start >= 0 and end >= 0:
                span = (end - start + 1) - skip_lines
                norm_span = (span / sent_count) * 100
            else:
                span = 0
                norm_span = 0.0

            if sent_count > 0:
                norm_coverage = (given_left_count / sent_count) * 100
            else:
                norm_coverage = 0

            res[header] = {'type':            1,
                           'span':            norm_span,
                           'coverage':        norm_coverage,
                           'count':           given_count,
                           'left_count':      given_left_count,
                           'first_new_count': first_new_count,
                           'sent_count':      sent_count,
                           'para_count':      None,
                           'top':             False}

        top_left_count = 0
        for stats in res.values():  # find the top left count
            if stats['left_count'] > top_left_count:
                top_left_count = stats['left_count']

        for stats in res.values():  # update the 'top' status of each topic.
            if stats['left_count'] == top_left_count and \
                stats['left_count'] - stats['first_new_count'] != 0:
                stats['top'] = True

        return res

    def updateGlobalTopics(self, global_data, min_topics=2):
        if global_data is None:
            return []

        data = global_data['data']

        header = data[0]   # list of tuples (POS, LEMMA, POS, SENT_COUNT, PARA_COUNT)
        self.global_header = header.copy()   # Let's make a copy so that we don't actually change the global data.

        nrows  = len(data)
        ncols  = len(header)

        self.global_topics = list()

        if ncols == 0:
            return []

        topic_filter    = TOPIC_FILTER_ALL
        sent_filter     = self.filterTopics(data, nrows, ncols)

        true_left_count = 0
        l_count = 0

        for ci in range(ncols):    # for each colum (word) in a row

            topic = header[ci][1]
            p_ri = 0

            if topic is not None:
                if not self.isLocalTopic(topic) and \
                   not DSDocument.isUserDefinedSynonym(topic):
                    continue

                if DSDocument.isUserDefinedSynonym(topic):
                    true_left_count = sent_filter[topic]['left_count']
                    count = sent_filter[topic]['count']
                    l_count = sent_filter[topic]['left_count']
                    count = max(count, 2)
                else:
                    true_left_count = sent_filter[topic]['left_count']

                    if topic_filter == TOPIC_FILTER_ALL:
                        count = sent_filter[topic]['count']
                        l_count = sent_filter[topic]['left_count']
                    else:
                        l_count = sent_filter[topic]['left_count']
                        if l_count == 1:
                            count = 2
                        else:
                            count = l_count

            if count < min_topics:
                continue

            self.global_topics.append(topic)

        # find the topic clusters that are not included in self.global_topics.
        # or simply add all the topic clusters and call list(set(...)) so that there
        # no duplicates!!

        tcs = DSDocument.getUserDefinedSynonyms()
        if tcs is not None:
            missing_tcs = list(set(tcs) - set(self.global_topics))
            self.global_topics = self.global_topics + missing_tcs
            count = 0
            for tc in missing_tcs:
                 # list of tuples (POS, LEMMA, POS, SENT_COUNT, PARA_COUNT)
                if tc not in [h[1] for h in self.global_header]:
                    self.global_header.append(('NOUN', tc, 10000000+count, -1, -1))
                    count += 1

        return self.global_topics

    def updateLocalTopics(self):

        def filter_local_topics(data, nrows, ncols):

            res = dict()
            for c in range(1, ncols+1):  # for each topic

                start = -1
                end   = -1
                given_count = 0
                given_left_count = 0
                # para_count = 0
                first_new_count = 0
                header = data[0][c-1][1]

                for r in range(1, nrows):
                    elem = data[r][c-1]

                    if isinstance(elem, tuple) and elem[0] is not None:
                        # if para_count == para_pos:
                        bSkip = False

                        if not elem[IS_TOPIC]: # it's not a topic word
                            bSkip = True

                        if not bSkip: # it's a topic word
                            if start < 0:
                                start = r
                            elif start >= 0:
                                end = r

                            given_left_count += 1

                            if not elem[ISLEFT]:
                                first_new_count += 1

                        given_count += 1

                if (given_left_count - first_new_count) == 0:
                    start = -1
                    end = -1

                skip_lines = 0
                sent_count = 0
                # para_count = 0
                for r in range(1, nrows):
                    elem= data[r][c-1]

                    # if type(elem) == int and elem < 0:
                        # para_count += 1
                    # elif para_count == para_pos:

                    if start < r < end:
                        if isinstance(elem, str) and elem == 'heading':
                            skip_lines += 1
                        elif isinstance(elem, int) and elem < 0:
                            skip_lines += 1

                    if not isinstance(elem, (str, int)):
                        sent_count += 1

                if start >= 0 and end >= 0:
                    span = (end - start + 1) - skip_lines
                    norm_span = (span / sent_count) * 100
                else:
                    span = 0
                    norm_span = 0.0

                if sent_count > 0:
                    norm_coverage = (given_left_count / sent_count) * 100
                else:
                    norm_coverage = 0

                res[header] = {'type':            1,
                               'span':            norm_span,
                               'coverage':        norm_coverage,
                               'count':           given_count,
                               'left_count':      given_left_count,
                               'first_new_count': first_new_count,
                               'sent_count':      sent_count,
                               'para_count':      None,
                               'top':             False}

            top_left_count = 0
            for stats in res.values():  # find the top left count
                if stats['left_count'] > top_left_count:
                    top_left_count = stats['left_count']

            for stats in res.values():  # update the 'top' status of each topic.
                if stats['left_count'] == top_left_count and \
                    stats['left_count'] - stats['first_new_count'] != 0:
                    stats['top'] = True

            return res

        def get_local_topics(local_data, min_topics=2):
            data = local_data

            topic_filter = TOPIC_FILTER_LEFT_RIGHT
            header = data[0]
            nrows  = len(data)
            ncols  = len(header)

            if ncols == 0:
                return []

            sent_filter = filter_local_topics(data, nrows, ncols)

            # true_left_count = 0 # unused-variable
            l_count = 0
            self.local_topics = list()

            for ci in range(ncols):

                topic = header[ci][1]
                count = 0

                if topic is not None:
                    true_left_count = sent_filter[topic]['left_count']

                    if topic_filter == TOPIC_FILTER_ALL:
                        count = sent_filter[topic]['count']
                        l_count = sent_filter[topic]['left_count']
                    else:
                        l_count = sent_filter[topic]['left_count']
                        if l_count == 1:
                            count = 2
                        else:
                            count = l_count

                if count < min_topics:
                    continue

                is_global = False
                if topic in self.global_topics:
                    is_global = True

                self.local_topics.append((topic, is_global))

            return self.local_topics

        num_paragraphs = self.getNumParagraphs()
        all_local_topics = list()
        self.local_topics_dict = dict()

        for i in range(num_paragraphs):
            local_data = self.getLocalTopicalProgData([i])
            local_topics = get_local_topics(local_data)
            all_local_topics += local_topics

            self.local_topics_dict[i] = local_topics

        return list(set(all_local_topics))

    def isLocalTopic(self, topic):
        """
        Returns True if <topic> is a local topic.
        """
        if self.local_topics_dict is not None:
            for key, value in self.local_topics_dict.items():    # for each paragraph
                local_topics = [t[0] for t in value]             # make a list of lemma/topic strings
                if topic in local_topics:                        # if <topic> is in it, return True
                    return True
        return False


    def testPrintGlobalVisData(self, vis_data):

        if vis_data['data'] == []:
            logging.warning("No global topics")
            return

        data       = vis_data['data']
        num_paras  = vis_data['num_paras']
        # num_topics = vis_data['num_topics'] # unused-variable

        # Print paragraph IDs
        line = " " * 30
        for i in range(num_paras):
            line += "{} ".format(i+1)

        logging.debug(line)

        # For each data for a specific topic,
        for topic_data in data:
            topic = topic_data['topic']
            lemma = topic[LEMMA]
            is_topic_cluster = topic_data['is_topic_cluster']
            is_non_local = topic_data['is_non_local']

            # if the paragraph data is not empty OR if the topic is a topic cluster
            if topic_data['paragraphs'] != [] or is_topic_cluster:

                # header section
                line = lemma
                line += " " * (30 - len(lemma))

                # Write a symbol for each paragraph.
                # L = 'topic' appears on the left side of the main verb at least once.
                # l = a topic cluster that is not a global topic.
                # R = 'topic' appears on the right side of the main verb.
                # r = a topic cluster that is not a global topic.
                # * = topic sentence.
                for para_data in topic_data['paragraphs']:
                    if para_data is not None:
                        # para_pos = para_data['para_pos'] # unused-variable
                        is_left  = para_data['is_left']
                        is_topic_sent = para_data['is_topic_sent']

                        if is_left:                     # is 'topic' on the left side?
                            c = 'l' if is_non_local else 'L'  # is it a non-local global topic?

                            if is_topic_sent:           # is 'topic' appear in a topic sentence?
                                line += (c + "*")
                            else:
                                line += (c + " ")
                        else:
                            c = 'r' if is_non_local else 'R'

                            if is_topic_sent:
                                line += (c + "*")
                            else:
                                line += (c + " ")
                    else:
                        line += "  "

                logging.debug(line)

    def testPrintLocalVisData(self, vis_data):

        if vis_data['data'] == []:
            logging.warning("No local topics")
            return

        data       = vis_data['data']
        num_topics = vis_data['num_topics']
        num_sents  = len(data[0]['sentences'])

        # Print paragraph IDs
        line = " " * 30
        for i in range(num_sents):
            line += "{} ".format(i+1)

        logging.debug(line)

        # For each topic data,
        for topic_data in data:
            topic = topic_data['topic']
            is_topic_cluster = topic_data['is_topic_cluster']
            lemma = topic[LEMMA]
            is_global = topic_data['is_global']

            # If 'topic' appears in one or more sentences OR it is a topic cluster
            if topic_data['sentences'] != [] or is_topic_cluster:
                line = lemma
                line += " " * (30 - len(lemma))

                # Write a symbol for each sentence sentence.
                for sent_data in topic_data['sentences']:
                    if sent_data is not None:
                        # sent_pos = sent_data['para_pos']
                        is_left = sent_data['is_left']
                        is_topic_sent = sent_data['is_topic_sent']

                        if is_left:
                            if is_global:
                                c = 'L'
                            else:
                                c = 'l'

                            if is_topic_sent:
                                line += (c + '*')
                            else:
                                line += (c + ' ')
                        else:
                            if is_global:
                                c = 'R'
                            else:
                                c = 'r'

                            if is_topic_sent:
                                line += (c + '*')
                            else:
                                line += (c + ' ')
                    else:
                        line += "  "

                logging.debug(line)

    def getCurrentTopics(self):

        topics = self.global_topics

        if self.local_topics_dict is not None:
            for _, value in self.local_topics_dict.items():
                local_topics = [t[0] for t in value]             # make a list of lemma/topic strings
                topics += local_topics

        return list(set(topics))

    def countSentencesWithTopic(self, topic):
        """
        Given a topic <string>, this method returns the toal number of sentences
        that includes the topic.
        """
        if self.topic_location_dict is None:
            return 0
        locations = self.topic_location_dict.get(topic, None)
        if locations is None:
            return 0
        temp = set([pos[:2] for pos in locations.getTopicPositions()])
        # topic_positions = locations.getTopicPositions()
        # for t in topic_positions:
        #     if t[:2] not in temp:
        #         temp.append(t[:2])
        return len(temp)

    def locateTopics(self, topics):
        """
        Given a set of topics, this method locates the topics in a current document,
        and update self.topic_location_dict.
        """

        def locateATopic(topic):
            """
            This method is called by locateTopics(), and should not
            """

            if self.sections is None:
                return "Error in locateATopic()"
            else:
                try:
                    data = self.sections[self.current_section]['data']
                    if data is None:
                        raise ValueError
                except:
                    return

            adj_stats = AdjacencyStats(topic=topic, controller=self.controller)
            topic_filter = TOPIC_FILTER_LEFT_RIGHT

            # Let's find which paragraphs/sentences the selected 'topic' is included.
            pcount = 1
            for para in data['paragraphs']:
                scount = 1

                # if para_pos != -1 and (para_pos+1) != pcount:
                    # pcount += 1
                    # continue

                for sent in para['sentences']:

                    wcount = 1
                    for w in sent['text_w_info']:
                        if w[LEMMA] == topic and (w[POS] == 'NOUN' or w[POS] == 'PRP'):
                            if topic_filter == TOPIC_FILTER_LEFT and w[ISLEFT] is False:  ## NEW 3/2/21
                                pass
                            else:
                                adj_stats.addParagraphID(pcount)
                                adj_stats.addSentenceID(pcount, scount)
                                adj_stats.addTopicPosition(pcount, scount, wcount)

                        wcount += 1

                    scount+=1

                pcount+=1

            return adj_stats

        if topics is None or topics == []:
            self.topic_location_dict = None
            return None

        self.topic_location_dict = dict()
        for topic in topics:
            self.topic_location_dict[topic] = locateATopic(topic)

    def getHtmlSents(self, data):
        """
        This method returns a list that contains a set of lists. Each contains a set of sentences
        in HTML formatted for the Clarity panel.
        """

        def generate_html(sent_data, sent_pos, word_pos_offset):

            if isinstance(sent_data, str) and sent_data == "\n":
                return ""

            np_positions = list()
            analysis = sent_data[2]['sent_analysis']
            is_be_verb = analysis['BE_VERB']

            def getNPTag(wpos):
                for pos in np_positions:
                    if wpos == pos[0]:
                        # if is_be_verb:
                            # return "<b class=\"topic-text\">"
                        # else:
                        return "<b class=\"topic-text\">"

                    elif wpos == pos[1]:
                        return "</b>"
                return ""

            for np in analysis['NOUN_CHUNKS']:
                np_positions.append((np['start'], np['end']))

            html_str = '<p class="non-topic-text">'

            if is_be_verb:
                verb_class = "be-verb"
            else:
                verb_class = "active-verb"

            wpos = 0
            tokens = analysis['TOKENS']

            while wpos < len(tokens):
                token = tokens[wpos]
                w = token['text']
                tag = getNPTag(wpos)

                if tag.startswith("</b>") and html_str[-1] == ' ':
                    # if tag is a closing tag, no space before it.
                    html_str = html_str[:-1]

                if w in right_quotes or w in end_puncts or w == "%" or w in no_space_patterns:
                    if html_str[-1] == ' ':
                        # w shouldn't have a space before it.
                        html_str = html_str[:-1]

                html_str += tag

                if tag.startswith("</b>"):
                    if w not in right_quotes and w not in end_puncts and w != "%" and w not in no_space_patterns:
                        html_str += " "

                if token['is_root']:
                    html_str += "<span class=\"{}\">{}</span>".format(verb_class, w)
                else:
                    html_str += f"{w}"

                if w not in left_quotes and w not in hyphen_slash:
                    html_str += " "

                wpos += 1

            html_str += "</p>"

            return html_str

        #
        res = list()
        first_word_pos = 0
        first_sent = True
        sent_pos = 0
        para_count = 1
        html_strs = list()

        for sent_data in data:

            if isinstance(sent_data[0], str) and sent_data[0] == '\n':
                if para_count != 1:
                    res.append(html_strs)
                para_count += 1
                html_strs = list()

            elif not isinstance(sent_data[0], str) and first_sent:
                w = sent_data[2]['text_w_info'][0] # get the first word
                first_word_pos = w[WORD_POS]

            s = generate_html(sent_data, sent_pos, first_word_pos)

            if s != "<p></p>":
                html_strs.append(s)

            if isinstance(sent_data[0], str) and sent_data[0] == '\n':    # paragraph break
                first_sent = True
                sent_pos = 0
            else:
                first_sent = False
                sent_pos += 1

        return res

    def generateCSVString(self, vis_data):

        if vis_data['data'] == []:
            logging.error("No global topics")
            return

        data       = vis_data['data']
        num_paras  = vis_data['num_paras']
        num_topics = vis_data['num_topics']

        csv_str = ""

        # Print paragraph IDs
        line = "Topics,"
        for i in range(1, num_paras+1):
            line += f"{i},"

        # print(line)
        line += "\n"

        csv_str += line

        # For each data for a specific topic,
        for topic_data in data:
            topic = topic_data['topic']
            lemma = topic[LEMMA]
            is_topic_cluster = topic_data['is_topic_cluster']
            is_non_local = topic_data['is_non_local']

            # if the paragraph data is not empty OR if the topic is a topic cluster
            if topic_data['paragraphs'] != [] or is_topic_cluster:

                # header section
                line = f"{lemma},"

                # Write a symbol for each paragraph.
                # L = 'topic' appears on the left side of the main verb at least once.
                # l = a topic cluster that is not a global topic.
                # R = 'topic' appears on the right side of the main verb.
                # r = a topic cluster that is not a global topic.
                # * = topic sentence.
                for para_data in topic_data['paragraphs']:
                    if para_data is not None:
                        para_pos = para_data['para_pos']
                        is_left  = para_data['is_left']
                        is_topic_sent = para_data['is_topic_sent']

                        if is_left:                     # is 'topic' on the left side?
                            c = 'L'
                            if is_topic_sent:           # is 'topic' appear in a topic sentence?
                                line += (c + "*")
                            else:
                                line += c
                        else:
                            c = 'R'

                            if is_topic_sent:
                                line += f"{c}*"
                            else:
                                line += c

                        line += ","

                    else:
                        line += ","

                # print(line)
                line += "\n"

                csv_str += line

        return csv_str

    def getSentenceIDs(self, sentences, fuzzy=True):
        """ UNUSED!!! """
        def is_overlapping_fuzzy(_text, _sentences):
            # threshold = 80  # FIXME define fuzzy_cmpstr
            # for s in _sentences:
            #     if fuzzy_cmpstr(s, text, threshold):    # approximate string comparison
            #         return True
            return False

        def is_overlapping(text, _sentences):

            clean_text = text.replace('\u2019', "'").replace('\u2018', "'")
            clean_text = clean_text.replace('\u201C', "'").replace('\u201D', "'").lower()

            for s in _sentences:
                clean_s = s.replace('\u2019', "'").replace('\u2018', "'")
                clean_s = clean_s.replace('\u201C', "'").replace('\u201D', "'").lower()
                if clean_s in clean_text:
                    return True
                if len(clean_text.split())/len(clean_s.split()) > 0.8 and clean_text in clean_s:
                    return True

            return False

        if self.sections:
            data = self.sections[self.current_section]['data']

            sent_list = list()
            pcount = 1
            for para in data['paragraphs']:    # for each paragraph

                # we will skip headings
                if para.get('style', '') in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title']:
                    pcount += 1
                    continue

                scount = 1
                for sent in para['sentences']: # for each sentence

                    if fuzzy:
                        if is_overlapping_fuzzy(sent['text'], sentences):
                            sent_list.append((pcount, scount, -1))
                        else:
                            if is_overlapping(sent['text'], sentences):
                                sent_list.append((pcount, scount, -1))
                    else:
                        if is_overlapping(sent['text'], sentences):
                            sent_list.append((pcount, scount, -1))

                    scount += 1
                pcount += 1

            if sent_list:
                return sent_list
            return []
        else:
            logging.error("    self.sections is None!!!!...")