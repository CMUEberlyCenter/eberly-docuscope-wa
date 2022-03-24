import os, sys
import platform
from time import time
from datetime import datetime, date
import math
import json

import threading

from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *

import dslib.views as views
import dslib.models.dict as ds_dict
import dslib.utils as utils
import dslib.views.dict_hierarchy as dict_hierarchy

from dslib.utils import resource_path, remove_punct_and_space
from dslib.views.dialogs import WarningDialog, AboutDialog, YesNoDialog, ConfirmationDialog, SaveWarningDialog, SaveWarningDialog2

from pathlib import Path
from docx import Document
from docx.shared import RGBColor

import difflib

import pprint     
pp = pprint.PrettyPrinter(indent=4)

UP   = -1
DOWN = 1

CLUSTER   = 1
DIMENSION = 2


class DictEditor(QMainWindow):
    """
    The main dialog (window) for editing DocuScope's DEFAULT dictionaries.
    """

    def __init__(self, editor_win, controller, parent=None):

        super(DictEditor, self).__init__(parent)

        self.editor_win = editor_win
        self.controller = controller

        self.dictionary = None
        self.dict_hierarchy_viwer = None

        self.html_str         = ""
        self.dicterrors       = None

        self.bShowDeletedLATs = True
        self.bForceLATUpdate  = False
        self.dict_changed     = False
        self.dict_not_saved   = False

        self.ignorePatternsChanges = False

        self.curr_num_patterns = -1

        self.selectedLAT           = None
        self.selectedLATIndexModel = None
        self.selectedLATIndex      = -1
        self.selected_lat_validated  = True

        self.initMenus()
        self.initUI()

    def initUI(self):

        self.setWindowTitle('DocuScope Dictionary Editor')

        font_size = views.default_ui_font_size

        vbox = QVBoxLayout()

        # Top Control UI
        hbox = QHBoxLayout()
        self.nameHeader = QLabel("Dictionary Path: ")
        hbox.addWidget(self.nameHeader)

        self.dictPathLabel = QLabel("")
        self.dictPathLabel.setStyleSheet("font-size: {}pt".format(font_size))
        hbox.addWidget(self.dictPathLabel)

        hbox.addStretch()

        vbox.addLayout(hbox)

        line = QFrame()
        line.setFrameShape(QFrame.HLine);
        line.setFrameShadow(QFrame.Sunken);
        vbox.addWidget(line)

        # hbox will hold 3 equally spaced vertical layout: left/center/right
        hbox = QHBoxLayout()

        # Cluster/Dimension Browser
        vbox_left = QVBoxLayout()

        heading = QLabel("Clusters/Dimensions")
        heading.setStyleSheet("font-size: {}pt; font-weight: bold;".format(font_size));
        vbox_left.addWidget(heading)

        hcontrol_area = QHBoxLayout()

        self.add_button = QComboBox()
        self.add_button.addItem("   Add ...")
        self.add_button.addItem("Add Cluster")
        self.add_button.addItem("Add Dimension")
        self.add_button.setCurrentIndex(0)
        self.add_button.currentIndexChanged.connect(self.addNewCategory)
        self.add_button.setStyleSheet("font-size: {}pt; selection-background-color: #0080ff;".format(font_size))
        hcontrol_area.addWidget(self.add_button)

        self.delete_buttom = QPushButton("Delete")
        self.delete_buttom.clicked.connect(self.deleteSelectedCategory)
        self.delete_buttom.setStyleSheet("font-size: {}pt;".format(font_size))
        self.delete_buttom.setAutoDefault(False)
        hcontrol_area.addWidget(self.delete_buttom)

        self.move_dim_buttom = QPushButton("Move")
        self.move_dim_buttom.clicked.connect(self.moveSelectedDimensions)
        self.move_dim_buttom.setStyleSheet("font-size: {}pt;".format(font_size))
        self.move_dim_buttom.setAutoDefault(False)
        hcontrol_area.addWidget(self.move_dim_buttom)

        self.move_up_button = QPushButton()
        self.move_up_button.setIcon(QIcon(QPixmap(resource_path("data/icons/up_icon.png"))))
        self.move_up_button.clicked.connect(self.moveSelectedCategoryUp)
        self.move_up_button.setAutoDefault(False)
        hcontrol_area.addWidget(self.move_up_button)

        self.move_down_button = QPushButton()
        self.move_down_button.setIcon(QIcon(QPixmap(resource_path("data/icons/down_icon.png"))))
        self.move_down_button.clicked.connect(self.moveSelectedCategoryDown)
        self.move_down_button.setAutoDefault(False)
        hcontrol_area.addWidget(self.move_down_button)

        self.find_button = QPushButton()
        self.find_button.setIcon(QIcon(QPixmap(resource_path("data/icons/find_icon.png"))))
        self.find_button.clicked.connect(self.findCategory)
        self.find_button.setAutoDefault(False)
        hcontrol_area.addWidget(self.find_button)
        
        self.global_replace_dim_names_button = QPushButton()
        self.global_replace_dim_names_button.setIcon(QIcon(QPixmap(resource_path("data/icons/find_and_replace_icon.png"))))
        self.global_replace_dim_names_button.clicked.connect(self.openGlobalReplaceDimNamesDialog)
        self.global_replace_dim_names_button.setAutoDefault(False)
        self.global_replace_dim_names_button.setEnabled(False)
        self.global_replace_dim_names_button.setToolTip("Find & Replace...")
        hcontrol_area.addWidget(self.global_replace_dim_names_button)

        vbox_left.addLayout(hcontrol_area)

        self.dictTree = EditableDictTree(editor_win=self)
        vbox_left.addWidget(self.dictTree)

        self.helpForm = HelpForm(editor_win=self)
        self.dictTree.setHelpForm(self.helpForm)
        vbox_left.addWidget(self.helpForm)
        hbox.addLayout(vbox_left)

        # LAT Browser
        vbox_center = QVBoxLayout()

        heading = QLabel("Language Action Types")
        heading.setStyleSheet("font-size: {}pt; font-weight: bold;".format(font_size));
        vbox_center.addWidget(heading)

        hbox_lat_ui = QHBoxLayout()

        self.lat_add_button  = QPushButton("Add")                        # add button
        self.lat_add_button.clicked.connect(self.addNewLAT)
        self.lat_add_button.setStyleSheet("font-size: {}pt".format(font_size));
        self.lat_add_button.setAutoDefault(False)
        hbox_lat_ui.addWidget(self.lat_add_button)

        self.lat_delete_button  = QPushButton("Delete")                  # delete button
        self.lat_delete_button.clicked.connect(self.deleteSelectedLATs)
        self.lat_delete_button.setStyleSheet("font-size: {}pt".format(font_size));
        self.lat_delete_button.setAutoDefault(False)
        hbox_lat_ui.addWidget(self.lat_delete_button)

        self.lat_move_button  = QPushButton("Move")                  # delete button
        self.lat_move_button.clicked.connect(self.moveSelectedLATs)
        self.lat_move_button.setStyleSheet("font-size: {}pt".format(font_size));
        self.lat_move_button.setAutoDefault(False)
        hbox_lat_ui.addWidget(self.lat_move_button)


        self.latFilterInput = QLineEdit()                                # filter input
        self.latFilterInput.textChanged.connect(self.filterLATs)
        self.latFilterInput.setStyleSheet("QLineEdit {background-color: #fff; font-size: " + str(font_size)+ "pt;}")
        self.latFilterInput.setDragEnabled(True)
        hbox_lat_ui.addWidget(self.latFilterInput)

        self.global_replace_lat_names_button = QPushButton()
        self.global_replace_lat_names_button.setIcon(QIcon(QPixmap(resource_path("data/icons/find_and_replace_icon.png"))))
        self.global_replace_lat_names_button.clicked.connect(self.openGlobalReplaceLATNamesDialog)
        self.global_replace_lat_names_button.setAutoDefault(False)
        self.global_replace_lat_names_button.setEnabled(False)
        self.global_replace_lat_names_button.setToolTip("Find & Replace...")
        hbox_lat_ui.addWidget(self.global_replace_lat_names_button)

        vbox_center.addLayout(hbox_lat_ui)

        self.latListView = QListView(self)
        # self.latListView.setStyleSheet("font-size: {}pt;".format(font_size))
        self.latListView.clicked.connect(self.latSelectionChanged)
        self.latListView.setSelectionMode(QAbstractItemView.ContiguousSelection)

        self.latListView.setStyleSheet("QListView { font-size: " + str(font_size) + " pt;" + \
                                        "background-color: #fff;" + \
                                        "selection-color: #fff;" + \
                                        "selection-background-color: #0080ff;};")


        vbox_center.addWidget(self.latListView)

        hbox_lat_properties = QHBoxLayout()
        header = QLabel("LAT Properties")
        header.setStyleSheet("QLabel {font-weight: bold; font-size: " + str(font_size) + "pt;}")
        hbox_lat_properties.addWidget(header)

        hbox_lat_properties.addStretch()

        self.lat_name_update_button = QPushButton("Update")                           # LAT name
        self.lat_name_update_button.clicked.connect(self.latNameChanged)
        self.lat_name_update_button.setEnabled(False)
        self.lat_name_update_button.setAutoDefault(False)
        hbox_lat_properties.addWidget(self.lat_name_update_button)
        vbox_center.addLayout(hbox_lat_properties)

        # LAT details
        form_box = QFormLayout()
        form_box.setContentsMargins(0,0,0,0)

        self.lat_name_field = QLineEdit()
        self.lat_name_field.setFocusPolicy(Qt.StrongFocus)
        self.lat_name_field.setStyleSheet("QLineEdit { background-color: #fff; font-size: " + \
                                          str(font_size)+ "pt;}")
        #self.lat_name_field.editingFinished.connect(self.latNameChanged)
        self.lat_name_field.textEdited.connect(self.latNameEdited)
        form_box.addRow(QLabel("Name: "), self.lat_name_field)

        self.lat_count = QLineEdit()
        self.lat_count.setReadOnly(True)
        self.lat_count.setStyleSheet("QLineEdit { background-color: #eee; font-size: " + 
                                      str(font_size) + "pt;}")
        form_box.addRow(QLabel("Count: "), self.lat_count)

        ahbox = QHBoxLayout()
        self.lat_num_added = QLineEdit()
        self.lat_num_added.setReadOnly(True)
        self.lat_num_added.setStyleSheet("QLineEdit { background-color: #eee; font-size: " + \
                                         str(font_size) + "pt;}")
        ahbox.addWidget(self.lat_num_added)
        show_added_button = QPushButton()
        show_added_button.setIcon(QIcon(QPixmap(resource_path("data/icons/view_icon.png"))))
        show_added_button.clicked.connect(self.showAddedLATPatterns)
        show_added_button.setAutoDefault(False)
        ahbox.addWidget(show_added_button)
        form_box.addRow(QLabel("Added: "), ahbox)

        dhbox = QHBoxLayout()
        self.lat_num_deleted = QLineEdit()
        self.lat_num_deleted.setReadOnly(True)
        self.lat_num_deleted.setStyleSheet("QLineEdit { background-color: #eee; font-size: " + \
                                           str(font_size) + "pt;}")
        dhbox.addWidget(self.lat_num_deleted)
        show_deleted_button = QPushButton()
        show_deleted_button.setIcon(QIcon(QPixmap(resource_path("data/icons/view_icon.png"))))
        show_deleted_button.clicked.connect(self.showDeletedLATPatterns)
        show_deleted_button.setAutoDefault(False)
        dhbox.addWidget(show_deleted_button)
        form_box.addRow(QLabel("Deleted: "), dhbox)

        vbox_center.addLayout(form_box)

        hbox.addLayout(vbox_center)

        # Pattern Editor
        vbox_right = QVBoxLayout()

        hbox_ui = QHBoxLayout()

        self.patterns_heading = QLabel("Patterns")
        self.patterns_heading.setStyleSheet("font-size: {}pt; font-weight: bold;".format(font_size));
        hbox_ui.addWidget(self.patterns_heading)

        hbox_ui.addStretch()
        vbox_right.addLayout(hbox_ui)

        hbox_ui = QHBoxLayout()        
        self.patterns_update_button = QPushButton("Update")                             # Pattern update
        self.patterns_update_button.clicked.connect(self.updatePatterns)
        self.patterns_update_button.setStyleSheet("font-size: {}pt".format(font_size));
        self.patterns_update_button.setAutoDefault(False)
        hbox_ui.addWidget(self.patterns_update_button)

        hbox_ui.addStretch()

        self.find_pattern_buttom = QPushButton()
        self.find_pattern_buttom.setIcon(QIcon(QPixmap(resource_path("data/icons/find_icon.png"))))
        self.find_pattern_buttom.clicked.connect(self.showPatternDialog)
        self.find_pattern_buttom.setAutoDefault(False)
        hbox_ui.addWidget(self.find_pattern_buttom)

        vbox_right.addLayout(hbox_ui)
        self.patternEditor = QPlainTextEdit()
        self.patternEditor.setStyleSheet("QPlainTextEdit {background-color: #fff; font-size: " + \
                           str(font_size) + "pt;}" + \
                           "QMenu {selection-background-color: #0080ff;}")
        self.patternEditor.modificationChanged.connect(self.latPatternsChanged)
        vbox_right.addWidget(self.patternEditor)

        self.patterns_update_button.setEnabled(False)

        hbox.addLayout(vbox_right)

        self.repairDialog = RepairDialog(editor_win=self)
        self.wordClassDialog = WordClassDialog(editor_win=self)

        vbox.addLayout(hbox)

        container = QWidget()

        container.setLayout(vbox)
        self.setCentralWidget(container) 

    def initMenus(self):

        menubar = self.menuBar()
        menubar.setNativeMenuBar(True)

        fileMenu = menubar.addMenu('&File')

        self.newAction = QAction('&New Dictionary...', self)
        self.newAction.triggered.connect(self.newDict)
        fileMenu.addAction(self.newAction)

        self.openAction = QAction('&Open Dictionary...', self)
        self.openAction.triggered.connect(self.openDict)
        fileMenu.addAction(self.openAction)

        self.saveAction = QAction('&Save Dictionary...', self)
        self.saveAction.triggered.connect(self.saveDict)
        self.saveAction.setEnabled(False)
        fileMenu.addAction(self.saveAction)

        fileMenu.addSeparator()

        self.exportAction = QAction('Export Dictionary...', self)
        self.exportAction.triggered.connect(self.export)
        self.exportAction.setEnabled(False)
        fileMenu.addAction(self.exportAction)        

        self.generateHelpDocAction = QAction('Export Dictionary Help Content (.docx)...', self)
        self.generateHelpDocAction.triggered.connect(self.generateHelpDoc)
        self.generateHelpDocAction.setEnabled(False)
        fileMenu.addAction(self.generateHelpDocAction)

        fileMenu.addSeparator()

        self.closeAction = QAction('Close', self)        
        self.closeAction.triggered.connect(self.close)
        fileMenu.addAction(self.closeAction)


        toolsMenu = menubar.addMenu("&Tools")

        self.editWCAction = QAction('Edit Word Classes...', self)
        self.editWCAction.triggered.connect(self.editWordClasses)
        self.editWCAction.setDisabled(True)
        toolsMenu.addAction(self.editWCAction)

        # self.loadDictAction = QAction('Load Dictionary', self)
        # self.loadDictAction.triggered.connect(self.loadDictionary)
        # self.loadDictAction.setDisabled(True)
        # toolsMenu.addAction(self.loadDictAction)

        self.findLATsAction = QAction('Find in LATs...', self)
        self.findLATsAction.triggered.connect(self.openFindLATsDialog)        
        self.findLATsAction.setDisabled(True)
        toolsMenu.addAction(self.findLATsAction)

        self.openThesaurusAction = QAction('Open Thesaurus...', self)
        self.openThesaurusAction.triggered.connect(self.openThesaurusDialog)
        self.openThesaurusAction.setDisabled(True)
        toolsMenu.addAction(self.openThesaurusAction)

        self.showDictHierarchyAction = QAction('Show Dictionary Hierarchy', self)
        self.showDictHierarchyAction.triggered.connect(self.showDictHierarchy)
        self.showDictHierarchyAction.setDisabled(True)
        toolsMenu.addAction(self.showDictHierarchyAction)

    # End of UI related methods
    # ^^^^^^^^^^^^^^^^^^^^^^^^^^^^^

    ##################################################
    #
    # File Menu Actions
    #
    ##################################################

    def newDict(self):
        new_dialog = NewDictDialog(editor_win=self)

    def openDict(self):
        """
        Open in Editor
        This function is used to open an existing dictionary in the DictEditorDialog.
        """
        def isValidDictionaryFolder(folder):
            tones_path = os.path.join(folder, "_tones.txt")
            help_path  = os.path.join(folder, "_help.txt")
            wordclasses_path = os.path.join(folder, "_wordclasses.txt")

            if os.path.exists(tones_path) and \
               os.path.exists(help_path) and \
               os.path.exists(wordclasses_path):
                return True
            else:
                return False

        folder = QFileDialog.getExistingDirectory(None, 'Select a dictionary folder::', '', QFileDialog.ShowDirsOnly)

        if folder == "":
            return

        # make the dictionary folder backword compatible
        info_path  = os.path.join(folder, "info.json")
        if os.path.exists(info_path) != True:
            info = dict()
            info['name']       = os.path.basename(folder)
            info['version']    = "n/a"
            info['copyright']  = "n/a"            
            info['customized'] = False
            with open(info_path, 'w') as fout:
                json.dump(info, fout, indent=4)

        if isValidDictionaryFolder(folder):
            self.unloadDictionary()
            self.setupDictionary(folder)
            self.clearHelpContent()
            self.saveAction.setEnabled(False)
            self.exportAction.setEnabled(True)
            self.generateHelpDocAction.setEnabled(True)
            # self.verifyAllAction.setEnabled(True)
            self.findLATsAction.setEnabled(True)
            self.editWCAction.setEnabled(True)
            self.openThesaurusAction.setEnabled(True)
            self.showDictHierarchyAction.setEnabled(True)

        else:
            WarningDialog("Error", "{} is not a valid dictionary folder.".format(folder))

    def loadDictionary(self):
        if self.isDictNotSaved():
            WarningDialog("Dictionary needs to be saved first...", 
                          "The dictionary has been edited. Save it first.")
            return

        self.controller.loadDictionary(self.dictionary, self)

    def dictionaryLoaded(self):
        pass

    def openFindLATsDialog(self):
        if self.controller.isDictionaryLoaded():
            if self.isDictNotSaved():
                d = YesNoDialog("Dictionary has been edited", 
                                "Press OK to save and reload the dictionary now. \n\nIt may take a few minutes.")
                if d.retval == QMessageBox.Ok:
                    self.updateLATPatterns()  # update the patterns
                    self.saveDict()           # save the dictionary
                    self.loadDictionary()
            else:
                FindLATsDialog(self.dictionary, editor_win=self)
        else:
            if self.isDictNotSaved():
                YesNoDialog("Dictionary has not been saved and loaded.", 
                            "Press OK to save and load the dictionary now. \n\nIt may take a few minutes.")
            else:
                d = YesNoDialog("Dictionary has not been loaded.", 
                                "Press OK to load the dictionary now. \n\nIt may take a few minutes.")
                
                if d.retval == QMessageBox.Ok:
                    self.updateLATPatterns()  # update the patterns
                    self.saveDict()           # save the dictionary
                    self.loadDictionary()


    def showPatternDialog(self):
        dialog = FindPatternDialog(editor_win=self)

    def findPattern(self, pattern, reset=False):
        if reset:
            self.patternEditor.moveCursor(QTextCursor.Start)

        return self.patternEditor.find(pattern)

    def openGlobalReplaceLATNamesDialog(self):
        if self.isDictNotSaved():
            WarningDialog("Dictionary is not saved", 
                          "The dictionary has been edited. Save it before using this feature.")
            return

        c_name, d_name = self.dictTree.getSelectedItem()
        if d_name is None:
            WarningDialog("Select a Dimension", "Select a dimension first.")

        d = GlobalReplaceLATNamesDialog(d_name)
        if d.retval == QMessageBox.Cancel:
            return
        elif d.retval == QMessageBox.Apply:
            find, replace = d.getFindAndReplace()
            replace = remove_punct_and_space(replace)  # make sure there are no spaces and puncts.
            self.dictionary.findAndReplaceLATNames(c_name, d_name, find, replace)
            self.latNamesChanged()
            self.saveDict()

    def openGlobalReplaceDimNamesDialog(self):
        if self.isDictNotSaved():
            WarningDialog("Dictionary is not saved", 
                          "The dictionary has been edited. Save it before using this feature.")
            return

        indexes = self.dictTree.selectedIndexes() 

        # c_name, d_name = self.dictTree.getSelectedItem()
        if indexes == []:
            WarningDialog("Select Dimensions", "Select dimensions first.")
            return

        dimension_names = []
        clust_names = []
        parent = None
        for index in indexes:
            parent = index.parent()
            if parent.isValid():
                item = self.dictTree.itemFromIndex(index)            
                dimension_names.append(item.text(0))
            else:
                item = self.dictTree.itemFromIndex(index)            
                clust_names.append(item.text(0))

        if clust_names != [] or parent == None:   # cluster names should be empty, and there should be a parent.
            WarningDialog("Select Dimensions Only", "You cannot rename clusters using this tool.")
            return

        parent_item = self.dictTree.itemFromIndex(parent)
        clust_name = parent_item.text(0)

        d = GlobalReplaceDimNamesDialog()

        if d.retval == QMessageBox.Cancel:
            return
        elif d.retval == QMessageBox.Apply:
            find, replace, bname, blabel = d.getFindAndReplace()
            replace = remove_punct_and_space(replace)  # make sure there are no spaces and puncts.
            new_dimension_names = self.dictionary.findAndReplaceDimensionNames(dimension_names, find, replace, bname, blabel)
            self.dictTree.resetHierarchy(self.dictionary.getTones(), clust_name, new_dimension_names)
            self.saveDict()

    def enableGlobalReplaceDimNames(self):
        self.global_replace_dim_names_button.setEnabled(True)

    def disableGlobalReplaceDimNames(self):
        self.global_replace_dim_names_button.setEnabled(False)

    def export(self):
        """
        export() opens a dialog that will be used for collection information about 
        the current dictionary before it is exported.
        """

        if self.dictionary != None:
            self.controller.openExportDialog(self.dictionary)

    def latNameEdited(self, bModified):
        if bModified:
            self.lat_name_update_button.setEnabled(True)
        else:
            self.lat_name_update_button.setEnabled(False)        

    def latPatternsChanged(self, bModified):
        if self.ignorePatternsChanges == False:
            if bModified:
                self.patterns_update_button.setEnabled(True)
                self.selected_lat_validated = False
            else:
                self.patterns_update_button.setEnabled(False)
                self.selected_lat_validated = True

    def showAddedLATPatterns(self):
        if self.selectedLAT:
            patterns = self.dictionary.getAddedLATPatterns(self.selectedLAT)
            ListPatternsDialog("Added Patterns in {}".format(self.selectedLAT), patterns, self)
                

    def showDeletedLATPatterns(self):
        if self.selectedLAT:
            patterns = self.dictionary.getDeletedLATPatterns(self.selectedLAT)
            ListPatternsDialog("Deleted Patterns in {}".format(self.selectedLAT), patterns, self)

    def editWordClasses(self):
        wordclasses = self.dictionary.getWordClasses()
        self.wordClassDialog.setContent(wordclasses)
        self.wordClassDialog.exec_()

    def setDictionary(self, dictionary):
        # this method sets the dictionary to the editor, but it does not load it.
        if dictionary != None and self.dictTree != None:
            self.dictionary = dictionary
            self.dictTree.setHierarchy(dictionary.getTones())
            self.dictPathLabel.setText(dictionary.getDirectory())
            self.setWindowTitle(dictionary.getName())
            self.dict_not_saved = False
            self.dict_changed = False
            self.updateLATlist(None, None)
            self.updateLATlist(None, None)

    def setTaggerInstance(self, tagger_instance):
        self.tagger_instance = tagger_instance

    def saveDict(self):

        # Let's save the help elements that hasn't been saved.
        res = self.dictTree.saveLatestChanges()
        if res == False:
            return

        # Let's save the lat name
        res = self.latNameChanged()
        if res == False:
            return

        if self.isSelectedLATInvalid():
            # if the currently selected LAT has been changed, but not "updated" (i.e., invalid),
            # validate it (i.e., clean up/remove duplicates etc.)
            self.validateSelectedLAT()

        res = self.dictionary.save()

        if res[0] == True:
            self.dictSaved()

            # Update the lat properties in the UI
            if self.selectedLAT:
                count, count_added, count_deleted = self.dictionary.getLATStats(self.selectedLAT)
                self.lat_count.setText("{}".format(count))
                self.lat_num_added.setText("{}".format(count_added))
                self.lat_num_deleted.setText("{}".format(count_deleted))

            ConfirmationDialog("Dictionary Saved", "{} has been saved.".format(self.dictionary.getName()))
            
        else:
            WarningDialog("No Patterns", "There are no patterns in {}. ".format(res[1]) + \
                           "Enter some patterns and try again.")

    def closeEditor(self):

        if self.isDictNotSaved():

            wd = SaveWarningDialog2("\nThe dictionary has not been saved.", 
                                   "Do you want to save the changes before closing the editor?")

            if wd.retval == QMessageBox.Discard:
                self.controller.invalidateDictEditor()

            elif wd.retval == QMessageBox.Save:
                self.saveDict()

            self.close()

    def showDictHierarchy(self):
        if self.dictionary is not None:
            root_node, dict_labels, leaf_count = self.dictionary.createDictHierarchy()
            if self.dict_hierarchy_viwer is None:
                self.dict_hierarchy_viwer = dict_hierarchy.DictHierarchyViewer(self.dictionary, root_node)
            self.dict_hierarchy_viwer.updateHierarchy(root_node)
            self.dict_hierarchy_viwer.show()
            self.dict_hierarchy_viwer.raise_()

    def exportDictHierarchy(self):
        if self.dictionary is not None:
            dst_dir_path = QFileDialog.getExistingDirectory(None, 'Select a folder:', '', QFileDialog.ShowDirsOnly)

            if os.path.exists(os.path.join(dst_dir_path, "_tones.txt")):
                WarningDialog("Bad Fodler", "{} is a dictionary folder.\n\n".format(dst_dir_path) + \
                                            " You cannot export a hierarchy to a dictionary folder.")
                return
            elif dst_dir_path == "": # Cancelled
                return
            else:
                ret = self.dictionary.exportDictHierarchy(dst_dir_path)

    def openThesaurusDialog(self):

        if self.controller.isDictionaryLoaded():
            if self.isDictNotSaved():
                d = YesNoDialog("Dictionary has been edited", 
                                "Press OK to save and reload the dictionary now. \n\nIt may take a few minutes.")
                if d.retval == QMessageBox.Ok:
                    self.updateLATPatterns()  # update the patterns
                    self.saveDict()           # save the dictionary
                    self.loadDictionary()
            else:
                ThesaurusDialog(self.dictionary, editor_win=self)                    
                # FindLATsDialog(self.dictionary, editor_win=self)
        else:
            if self.isDictNotSaved():
                YesNoDialog("Dictionary has not been saved and loaded.", 
                            "Press OK to save and load the dictionary now. \n\nIt may take a few minutes.")
            else:
                d = YesNoDialog("Dictionary has not been loaded.", 
                                "Press OK to load the dictionary now. \n\nIt may take a few minutes.")
                
                if d.retval == QMessageBox.Ok:
                    self.updateLATPatterns()  # update the patterns
                    self.saveDict()           # save the dictionary
                    self.loadDictionary()                                

    def findCategory(self):
        new_dialog = FindCategoryDialog(self.dictionary, editor_win=self)

    def selectACategory(self, categories):
        if len(categories) >= 2:
            cluster_name   = categories[0]
            dimension_name = categories[1]

        if len(categories) == 3:
            lat_name = categories[2]
        else:
            lat_name = None

        self.updateLATPatterns()
        self.patternEditor.clear()

        res = self.dictTree.selectADimension(dimension_name)

        if res and lat_name != None:
            qModelIndex = self.latProxyModel.match(self.latProxyModel.index(0, 0), Qt.DisplayRole, lat_name)[0]
            self.latListView.setCurrentIndex(qModelIndex)
            self.latListView.setFocus()

            # self.selectedLAT = lat_name
            self.selectedLATIndexModel = qModelIndex

            self.updateLATPatterns()

            # Update the lat details in the UI
            self.lat_name_field.setText(lat_name)
            count, count_added, count_deleted = self.dictionary.getLATStats(lat_name)
            self.lat_count.setText("{}".format(count))
            self.lat_num_added.setText("{}".format(count_added))
            self.lat_num_deleted.setText("{}".format(count_deleted))

    # ----------------------------------------
    #
    # Methods for editing and working with LATs
    #
    # ----------------------------------------
    def toggleDeletedLATs(self, val):
        if val == 0:
            self.bShowDeletedLATs = False
        else:
            self.bShowDeletedLATs = True

        self.updateLATPatterns()

    def addNewLAT(self):
        """
        If no dimension is selected, return (or the add button should be disabled?)
        If a dimension is selected, 
        1. Add a new LAT with a default name (e.g., Untitled-1) to the list view
        2. Add a new LAT object to the DSDict object (w/ no patterns, i.e., "")
        """

        # Save the currently selected LAT first.
        self.updateLATPatterns() 

        # Add a new LAT to the dictionary
        c_name, d_name = self.dictTree.getSelectedItem()

        if d_name == None:
            WarningDialog("Warning", "Select a Dimension first.")
            return

        lat_name = self.dictionary.addNewLAT(c_name, d_name)
        self.updateLATlist(c_name, d_name)
        self.patternEditor.clear()

        qModelIndex = self.latProxyModel.match(self.latProxyModel.index(0, 0), Qt.DisplayRole, lat_name)[0]
        self.latListView.setCurrentIndex(qModelIndex)
        self.latListView.setFocus()

        self.selectedLAT = lat_name
        self.selectedLATIndexModel = qModelIndex

        # Update the lat details in the UI
        self.lat_name_field.setText(lat_name)
        count = self.dictionary.getLATPatternCount(lat_name)
        self.lat_count.setText("0")
        self.lat_num_added.setText("0")
        self.lat_num_deleted.setText("0")
        self.dictChanged(1)

    def helpContentChanged(self):
        self.dictChanged(2)

    def deleteSelectedLATs(self):
        """
        If no LAT is selected, do nothing (or the delete button should be disabled?)
        If an LAT is selected, ask the user if they are SURE they want to  delete this LAT.
        if the user presses the OK button, delete the LAT from DSDict.
        !! NOTE that wa don't acutally delete the file until the user hits the "save" button.
        But, it IS deleted from the DSDict object. So, the user can't UNDO it while the 
        application is running.
        """

        indexes = self.latListView.selectedIndexes() 

        if len(indexes)>1:
            msg = "\n"
            for index in indexes:
                msg += "{}\n".format(index.data())
            msg += "\n"
            res  = YesNoDialog("Warning", "Are you sure you want to delete the following LATs?".format(msg))
            if res.retval == QMessageBox.Ok:
                for index in indexes:
                   c_name, dim_name = self.dictionary.deleteLAT(index.data())

            # Update the list view
            self.updateLATlist(c_name, dim_name)
            self.clearLATDetails()
            self.dictChanged(3)

        elif self.selectedLAT:
            res  = YesNoDialog("Warning", "Are you sure you want to delete {}?".format(self.selectedLAT))
            if res.retval == QMessageBox.Ok:
               # Update the dictionary
               c_name, dim_name = self.dictionary.deleteLAT(self.selectedLAT)
                       # Update the list view
               self.updateLATlist(c_name, dim_name)
               self.clearLATDetails()
               self.dictChanged(4)

    def moveSelectedLATs(self):
        # Lets create a list of selected LAT names.
        indexes = self.latListView.selectedIndexes() 

        if len(indexes)==0:
            WarningDialog("Warning", "No LAT is selected. Select an LAT first; then click the Move button.")
            return

        # ask the user to pick a dimension
        dim_picker = DimensionPickerDialog(self.dictionary.getTones(), editor_win=self)

        retval = dim_picker.exec_()
        # 0 == cancel
        # 1 == OK
        if  retval == 0:                
            # Do nothing and return if the user cancels. 
            return

        dest_clust, dest_dim = dim_picker.getSelectedDimension()

        if dest_dim == None and dest_clust != None:
            WarningDialog("Warning", "You must select a dimension. Only a cluster is selected.")
            return
            
        elif dest_dim == None and dest_clust == None:
            WarningDialog("Warning", "You must select a dimension. Nothing was selected.")
            return

        # If we get here, we have one or more LATs selected, the user has selected 
        # a destination dimension.
        for index in indexes:
            lat = index.data()
            curr_clust, curr_dim = self.dictTree.getSelectedItem()
            self.dictionary.moveLAT(lat, dest_clust, dest_dim, curr_clust, curr_dim)

        self.updateLATlist(curr_clust, curr_dim)
        self.clearLATDetails()
        self.dictChanged(5)

    def filterLATs(self):
        filter_str = self.latFilterInput.text()
        self.latProxyModel.setFilterWildcard(filter_str)

    def resetLATFilter(self):
        self.latFilterInput.clear()

    def updateLATlist(self, cluster_name, dimension_name):

        if dimension_name is None:
            self.global_replace_lat_names_button.setEnabled(False)
        else:
            self.global_replace_lat_names_button.setEnabled(True)

        self.latModel = QStringListModel(self)

        if dimension_name != None:
            lats = self.dictionary.getLATNames(cluster_name, dimension_name)
            sorted_lats = sorted(lats, key=lambda s: s.casefold())
            self.latModel.setStringList(sorted_lats)

        self.latProxyModel = QSortFilterProxyModel  ()
        self.latProxyModel.setSourceModel(self.latModel)
        self.latListView.setModel(self.latProxyModel)
        self.latProxyModel.setFilterCaseSensitivity(Qt.CaseInsensitive)

        self.selectedLAT = None
        self.selectedLATIndexModel = None
#        self.clearLATDetails()

    def clearLATDetails(self):
        self.ignorePatternsChanges = True
        self.lat_name_field.setText("")
        self.lat_count.setText("")
        self.lat_num_added.setText("")
        self.lat_num_deleted.setText("")
        self.patternEditor.clear()
        self.lat_name_update_button.setEnabled(False)
        self.patterns_update_button.setEnabled(False)
        self.selected_lat_validated = True
        self.ignorePatternsChanges = False        

    def removePattern(self, pattern, lat_name):
        self.dictionary.removePattern(pattern, lat_name)
        self.dictChanged(6)

    def updatePatterns(self):
        self.updateLATPatterns()  # old

    # ----------------------------------------
    #
    # Methods for dealing with dict erros.
    #
    # ----------------------------------------
    def verifyDictAll(self):
        if self.isDictNotSaved():
            WarningDialog("Dictionary must be Saved", 
                          "The dictionary has not been saved yet. Save the dictionary and try again.")
            return

        num_lats = self.dictionary.getNumLATs()

        if num_lats > 100:
            WarningDialog("Dictionary is too large.", "Sorry. This feature only works with a small dictionary.")        
            return
            
        estimated_sec = (num_lats * 60)
        estimated_min = round(estimated_sec / 60.0)

        if num_lats > 0 and estimated_min > 1:
            msg = "There are {:,} LATs. This may take about {} minutes or more to complete.".format(num_lats, estimated_min)
        elif num_lats > 0 and estimated_min <= 1 and estimated_sec >= 1.0:
            msg = "There are {:,} LATs. This may take about {} seconds or more to complete.".format(num_lats, round(estimated_sec))
        else:
            msg = "There are {:,} LATs.".format(num_lats)

        res  = YesNoDialog("Verify the Entire Dictionary", msg)

        if res.retval == QMessageBox.Cancel:
            pass
        elif res.retval == QMessageBox.Ok:
            self.clearLATDetails()
            self.controller.verifyDictionary(self.dictionary, self, None, 0)

    def verifyDict(self):
        if self.isDictNotSaved():
            WarningDialog("Dictionary must be Saved", "The dictionary has not been saved yet. Save the dictionary and try again.")
            return

        unverified_rules, num_nvps = self.dictionary.getUnverifiedPatterns()

        # print("-------------------------")
        # print("unverified_rules")
        # pp.pprint(unverified_rules)
        # print("-------------------------")

        if not unverified_rules:
            WarningDialog("Nothing to verify", "All the patterns have been verified.")
            return     

        estimated_sec = (num_nvps * 0.20)  + 3
        estimated_min = round(estimated_sec / 60.0)

        if estimated_min > 1:
            msg = "There are {:,} unverified patterns. This may take about {} minutes or more to complete.".format(num_nvps, estimated_min)
        elif estimated_min <= 1 and estimated_sec >= 1.0:
            msg = "There are {:,} unverified patterns. This may take about {} seconds or more to complete.".format(num_nvps, round(estimated_sec))
        else:
            if num_nvps > 0:
                if num_nvps == 1:
                    msg = "There is one unverified pattern.".format(num_nvps)
                else:
                    msg = "There are {:,} unverified patterns.".format(num_nvps)
            else:
                 msg = "There is no pattern to verify. Press OK to find unddefined word classes.".format(num_rules)            

        res  = YesNoDialog("Verify Dictionary", msg)

        if res.retval == QMessageBox.No:
            pass
        elif res.retval == QMessageBox.Ok:
            self.clearLATDetails()
            self.verifyDictionary(unverified_rules, num_nvps)

    def displayDictErrors(self, dicterrors):

        self.dicterrors = dicterrors
        self.bFoundCollisions = False 

        collisions = self.makeListOfCollisions()
        html_str =  ""
        html_str += "<h2>Collisions</h2>"
        html_str += "<table cellpadding=\"0\" cellspacing=\"0\">"
        html_str += "<tr><td width=\"50%\">&nbsp;</td><td width=\"50%\">&nbsp;</td></tr>"
        count = 0
        for collision in collisions:

            for item in collision:
                p = ' '.join(item[1])
                # label = "{} ({})".format(p, item[0])
                if count % 2 == 0:
                    html_str += "<tr bgcolor=\"{}\">".format("#e9f3ff")
                else:
                    html_str += "<tr>"
                html_str += "<td><b>{}&nbsp;&nbsp;</b></td><td>{}</td>".format(p, item[0])
                html_str += "</tr>"
            #html_str += "<tr><td>&nbsp;</td><td>&nbsp;</td></tr>"
            count+=1
        html_str += "</table>"

        if len(dicterrors['undefined_wcs'])>0: 
            html_str += "<h2>Undefined Word Classes</h2>"
            html_str += "<ul>"

            wcs = [[wc_err[0],', '.join(wc_err[1])] for wc_err in dicterrors['undefined_wcs'].items()]

            for wc in wcs:
                self.bFoundCollisions = True
                html_str += "<li><b>{}</b> ({})</li>".format(wc[0], wc[1])
            html_str += "</ul>"

        if html_str != "":
            self.dictErrorDialog = DictErrorDialog(html_str, editor_win=self)
            self.dictErrorDialog.exec_()
        else:
            ConfirmationDialog("No Collisions.", 
                               "No Collisions or missing wordclasses are found.")

    def makeListOfCollisions(self):
        collisions = list()

        def addCollisions(cpair):

            cp0 = cpair[0]
            cp1 = cpair[1]
            bMatch = False
            
            for i in range(len(collisions)):
                for p in collisions[i]:
                    if cp0[1] == p[1] or cp1[1] == p[1]:
                        collisions[i] = list(set(cpair + collisions[i]))
                        bMatch = True
                        break

            if bMatch == False:
                collisions.append(cpair)


        # single collisions are detected by pairs, but it is poissble that some of 
        # the single collisions involve 3 or more LATs. So, we are going to merge
        # those collisions here.
        d = dict()
        for c in self.dicterrors['single_collisions']:
            if c[0] in d:
                d[c[0]].extend(c[1])
                d[c[0]] = list(set(d[c[0]]))
            else:
                d[c[0]] = c[1]

        tmp = list()
        for c in d.items():
            tmp.append((c[0], c[1]))

        # for c in self.dicterrors['single_collisions']:
        #     print("c", c)
        #     pair = []
        #     for cat in c[1]:
        #         pair.append( (cat, (c[0],) ))
        #     print("pair: ", pair)
        #     collisions.append(pair)

        for c in tmp:
            pair = []
            for cat in c[1]:
                pair.append( (cat, (c[0],) ))
            collisions.append(pair)

        for pair in self.dicterrors['collisions']:
            addCollisions(pair)

        return collisions

    def repairCollisions(self):
        if self.dicterrors == None:
            WarningDialog("Error", "Something went wrong. (self.dicterrors == None)")
            return

        collisions = self.makeListOfCollisions()

        if len(collisions)>0:
            self.repairDialog.setCollisions(collisions, self.dicterrors.get('wordclasses', []))
            self.repairDialog.exec_()
        else:
            ConfirmationDialog("No Collisions.", "No Collisions or missing wordclasses are found.")

    # ----------------------------------------
    #
    # Methods for editing the dictionary structure
    #
    # ----------------------------------------
    def deleteSelectedCategory(self):
        # get the category to be deleted
        c_name, d_name = self.dictTree.getSelectedItem()
        if d_name == None:        
            res  = YesNoDialog("Warning", "Are you sure you want to delete {}?".format(c_name))
        else:
            res  = YesNoDialog("Warning", "Are you sure you want to delete {}?".format(d_name))

        if res.retval == QMessageBox.Ok:
            # delete the currentlyselected category
            c_name, d_name = self.dictTree.deleteSelectedItem()
            if c_name == None and d_name == None:
                return
            dict_item = self.dictionary.deleteCategory(c_name, d_name)

            self.helpForm.clearAll()

            c_name, d_name = self.dictTree.getSelectedItem()
            self.updateHelpForm(c_name, d_name)
            self.updateLATlist(c_name, d_name)
            self.dictChanged(7)

    def moveSelectedDimension(self):
        src_clust, src_dim = self.dictTree.getSelectedItem()

        if src_dim == None:
            WarningDialog("Warning", "Select a dimension first.")
            return

        # ask the user to pick a dimension
        clust_picker = ClusterPickerDialog(self.dictionary.getTones(), editor_win=self)
        retval = clust_picker.exec_()
        # 0 == cancel
        # 1 == OK
        if  retval == 0:                
            # Do nothing and return if the user cancels. 
            return

        dest_clust = clust_picker.getSelectedCluster()

        if dest_clust == None:
            WarningDialog("Warning", "You must select a cluster.")
            return

        # If we get here, the user has selected a dimension, then a cluster in the cluster picker
        # dialog
        self.dictionary.moveDimension(src_clust, src_dim, dest_clust)

        # update the tree
        # self.dictTree.resetHierarchy(self.dictionary.getTones(), dest_clust, src_dim)

        self.updateLATlist(dest_clust, src_dim)
        self.clearLATDetails()
        self.dictChanged(8)

    def moveSelectedDimensions(self):
        indexes = self.dictTree.selectedIndexes() 

        b_dim = False
        b_clust = False
        for index in indexes:
            parent = index.parent()        
            if parent.isValid():
                b_dim = True
            else:
                b_clust = True

        if b_dim == True and b_clust == True:
            WarningDialog("Warning", "Dimension and clusters cannot be selected together.")
            return

        selected_items = []
        for index in indexes:
            parent = index.parent()
            if parent.isValid():
                clust_item = self.dictTree.itemFromIndex(parent)
                dim_item   = self.dictTree.itemFromIndex(index)
                selected_items.append((dim_item, clust_item))
            else: # it's a cluster
                clust_item = self.dictTree.itemFromIndex(index)
                for di in range(clust_item.childCount()):
                    dim_item = clust_item.child(di)                
                    selected_items.append((dim_item, clust_item))

        # debug -----
        # for item in selected_items:
            # print("{} {}".format(item[0].text(0), item[1].text(0)))
        # -----------

        if selected_items == []:
            WarningDialog("Warning", "Select one or more dimensions first. The selected cluster may be empty.")
            return

        # ask the user to pick a cluster
        clust_picker = ClusterPickerDialog(self.dictionary.getTones(), editor_win=self)
        retval = clust_picker.exec_()
        # 0 == cancel
        # 1 == OK
        if  retval == 0:                
            # Do nothing and return if the user cancels. 
            return

        dest_clust = clust_picker.getSelectedCluster()

        if dest_clust == None:
            WarningDialog("Warning", "You must select a cluster.")
            return

        # If we get here, the user has selected one or more dimensions, then a cluster in the cluster picker
        # dialog
        dims = []
        for item in selected_items:
            src_dim   = item[0].text(0)   # dim
            src_clust = item[1].text(0)   # parent
            self.dictionary.moveDimension(src_clust, src_dim, dest_clust)
            dims.append(src_dim)

        # update the tree
        self.dictTree.resetHierarchy(self.dictionary.getTones(), dest_clust, dims)

        self.updateLATlist(dest_clust, src_dim)
        self.clearLATDetails()
        self.dictChanged(8)

    def moveSelectedCategoryUp(self):
        if self.dictTree != None:
            cluster_name = self.dictTree.moveSelectedItem(UP)
            if cluster_name != None:
                self.dictionary.moveCluster(cluster_name, UP)
                self.dictChanged(9)

    def moveSelectedCategoryDown(self):
        if self.dictTree != None:
            cluster_name = self.dictTree.moveSelectedItem(DOWN)
            if cluster_name != None:
                self.dictionary.moveCluster(cluster_name, DOWN)
                self.dictChanged(10)

    def addNewCategory(self, cat_type):
        if self.dictionary == None or cat_type == 0:
            return

        dimension_name = None
        cluster_index = self.dictTree.getSelectedClusterIndex()

        if cluster_index < 0 and cat_type == DIMENSION:
            WarningDialog("Warning", "Select a cluster first before adding a dimension")

        # print("cat_type:      ", cat_type)
        # print("cluster_index: ", cluster_index)
        # print("cluster_name:  ", cluster_name)
        # print("dimension_name:", dimension_name)

        cluster_name, dimension_name = self.dictTree.getSelectedItem()

        if cat_type == CLUSTER:
            if cluster_index < 0:
                cluster_name = self.dictionary.addNewCluster(cluster_index)
                self.dictTree.addNewCluster(cluster_index, cluster_name)        
            else:
                cluster_name = self.dictionary.addNewCluster(cluster_index+1)
                self.dictTree.addNewCluster(cluster_index+1, cluster_name)

        elif cat_type == DIMENSION and cluster_name != None:
            dimension_name = self.dictionary.addNewDimension(cluster_name)
            self.dictTree.addNewDimension(cluster_index, dimension_name)

#        self.updateHelpForm(cluster_name, dimension_name)

        self.add_button.setCurrentIndex(0)
        self.dictChanged(11)

    # ----------------------------------------
    #
    # Methods for editing category names
    #
    # ----------------------------------------
    def latNamesChanged(self):

        # If we get this far, we have successfuly replaced the LAT's name.
        # Let's get the selected cluster/dimenion in the left-side pane,
        # and update the list of LATs.
        c_name, d_name = self.dictTree.getSelectedItem()
        self.updateLATlist(c_name, d_name)
        self.selectACategory([c_name, d_name])
        self.dictChanged(12)
        self.lat_name_update_button.setEnabled(False)            

        return True

    def latNameChanged(self):
        lat_name = self.lat_name_field.text()

        # if self.selectedLAT is not None, remove spaces from lat_name
        # if lat_name is not the current lat's name (self.selectedLAT),
        # replace the name of the currenty selected LAT with the new name (=lat_name)
        # if lat_name is already taken, do nothing and display a warning.
        if self.selectedLAT != None:
            lat_name = remove_punct_and_space(lat_name)
            if self.selectedLAT != lat_name:
                res = self.dictionary.updateLATName(self.selectedLAT, lat_name)
                if res == False:  # lat_name exists.
                    WarningDialog("Warning", "{} is already used. Enter a different name.".format(lat_name))
                    return False

            # If we get this far, we have successfuly replaced the LAT's name.
            # Let's get the selected cluster/dimenion in the left-side pane,
            # and update the list of LATs.
            c_name, d_name = self.dictTree.getSelectedItem()
            self.updateLATlist(c_name, d_name)

            # Re-select the selected LAT, Since the list of LAT has been re-drawn, 
            # we need to select the one that was selected before the name change.
            qModelIndex = self.latProxyModel.match(self.latProxyModel.index(0, 0), Qt.DisplayRole, lat_name)[0]
            self.latListView.setCurrentIndex(qModelIndex)
            self.latListView.setFocus()

            self.selectedLAT = lat_name
            self.selectedLATIndexModel = qModelIndex

            # if self.selectedLATIndexModel != None:
            #     self.latSelected(self.selectedLATIndexModel)

            # Update the lat properties in the UI
            self.lat_name_field.setText(lat_name)
            count, count_added, count_deleted = self.dictionary.getLATStats(lat_name)
            self.lat_count.setText("{}".format(count))
            self.lat_num_added.setText("{}".format(count_added))
            self.lat_num_deleted.setText("{}".format(count_deleted))
            self.dictChanged(13)

            # we are no longer updating pattersn when the LAT name is chaned.
            # self.updateLATPatterns()

            self.lat_name_update_button.setEnabled(False)            

            return True

        return True

    def updateClusterName(self, old_name, new_name):
        d = self.dictionary.updateClusterName(old_name, new_name)
        self.helpForm.setHelpContent(d)
        self.dictChanged(14)

    def getClusterName(self):
        d = self.helpForm.getHelpContent()
        if d != None:
            return d['name']
        else:
            return None

    def updateWordClasses(self, wordclasses):
        self.dictionary.updateWordClasses(wordclasses)
        self.dictChanged(15)



    # ----------------------------------------
    #
    # Utilities
    #
    # ----------------------------------------       

    def isSelectedLATValidated(self):
        return self.selected_lat_validated

    def isSelectedLATInvalid(self):
        if self.selected_lat_validated:
            return False
        else:
            return True

    def dictChanged(self, val=-1):
        """
        This function is called when the dictionary is altered.
        """

        title = self.windowTitle()
        if title.endswith("*") != True:
            self.setWindowTitle("{}*".format(title))

        self.invalidateTagger()

        self.saveAction.setEnabled(True)
        self.exportAction.setDisabled(True)
        self.generateHelpDocAction.setDisabled(True)

        self.dict_not_saved = True
        self.dict_changed = True


    def dictSaved(self):
        """
        This function is called when the custom dictionary is saved to files.
        """
        title = self.windowTitle()
        if title.endswith("*"):
            self.setWindowTitle(title[:-1])
        self.dict_not_saved = False

        self.saveAction.setEnabled(False)
        self.exportAction.setDisabled(False)
        self.generateHelpDocAction.setDisabled(False)

    def isDictNotSaved(self):
        """
        this function returns True if the custom dictionary has been changes, yet
        not yet saved.
        """
        if self.dict_not_saved == True:
            return True
        else:   # None or False
            return False

    def isDictChanged(self):
        """
        this function returns True if the custom dictionary has been changes, yet
        not yet saved.
        """
        if self.dict_changed == True:
            return True
        else:   # None or False
            return False


    # ----------------------------------------
    # 
    # Saving selected LAT patterns
    #
    # ----------------------------------------
    def validateSelectedLAT(self):

        bNameChanged = False
        bRevised = None
        if self.selectedLAT != None:
            res = False
            new_lat_name = self.lat_name_field.text()
            new_lat_name = remove_punct_and_space(new_lat_name)

            if new_lat_name:
                if new_lat_name != self.selectedLAT:
                    res = self.dictionary.updateLATName(self.selectedLAT, new_lat_name)
                    if res == False:
                        WarningDialog("Warning", "{} is already used. Enter a different name.".format(new_lat_name))
                        return False
                    else:
                        c_name, d_name = self.dictTree.getSelectedItem() # get the currently selected c & d
                        self.updateLATlist(c_name, d_name)               # update the LAT list
                        self.selectedLAT = new_lat_name
                        bNameChanged = True
            else:
                WarningDialog("Missing Name", "LAT must have a name. Keeping \"{}\"".format(self.selectedLAT))

            patterns = self.patternEditor.toPlainText()
            bRevised = self.dictionary.setRevisedPatterns(self.selectedLAT, patterns)

            if bRevised:
                self.dictChanged(16)

        self.selected_lat_validated = True

        return bNameChanged

    def latSelectionChanged(self, val):

        # indexes = self.latListView.selectedIndexes() 

        # if len(indexes)==1:
        #     # The user selected a single LAT, and no other LATs are selected.   
        #     self.selectedLATs = None
        #     self.latSelected(val)
        # else:
        #     # The use selected 2 or more LATs.
        #     self.selectedLATs = list()
        #     for index in indexes:
        #         self.selectedLATs.append(index.data())


        # if self.selectedLAT != None:
        #     new_lat_name = self.lat_name_field.text()
        #     new_lat_name = remove_punct_and_space(new_lat_name)
        #     if new_lat_name != self.selectedLAT:
        #         res = self.dictionary.updateLATName(self.selectedLAT, new_lat_name)
        #         if res == False:  # lat_name exists.
        #             WarningDialog("Warning", "{} is already used. Enter a different name.".format(lat_name))
        #             return False
        #         else:

        #     else:
        #         self.latSelected(val)
        # else:
        self.latSelected(val)

    def updateLATPatterns(self):

        self.bForceLATUpdate = True
        if self.selectedLATIndexModel != None:
            self.latSelected(self.selectedLATIndexModel)
        self.bForceLATUpdate = False
        self.dictChanged(17)

    def latSelected(self, index):
        """
        If there is already an open LAT, AND if the LAT file has been edited,
        SAVE (1) the original as an archive, and (2) the revised LAT.
        The revised LAT patterns MUST BE sorted, and all the duplicates are removed.
        """
        count = 0
        count_added = 0
        count_deleted = 0

        lat_name = index.data()

        start = time()

        if self.bForceLATUpdate or self.selectedLAT != lat_name:

            # if self.selectedLAT != None:    # Jan 22 2022 - optimization.
            if self.selectedLAT != None and self.selected_lat_validated == False:
                bLATNameChanged = self.validateSelectedLAT()
                if bLATNameChanged:
                    lat_name = self.selectedLAT

            self.selectedLAT = lat_name
            self.selectedLATIndexModel = index
            
            revised_patterns, original_patterns = self.dictionary.getPatterns(lat_name)

            count, count_added, count_deleted = self.dictionary.getLATStats(lat_name)

            self.curr_num_patterns = count
            if revised_patterns == None:
                self.ignorePatternsChanges = True
                self.patternEditor.clear()
                self.patternEditor.setPlainText(original_patterns)
                self.ignorePatternsChanges = False
            else:
                self.patternEditor.clear()
                self.patternEditor.setPlainText(revised_patterns)

        self.patterns_update_button.setEnabled(False) ##

        # Update the lat details in the UI
        self.lat_name_field.setText(lat_name)
        self.lat_count.setText("{}".format(count))
        self.lat_num_added.setText("{}".format(count_added))
        self.lat_num_deleted.setText("{}".format(count_deleted))

        return True

    # ----------------------------------------
    #
    # Help Form related methods
    #
    # ----------------------------------------
    def updateHelpForm(self, cluster_name, dimension_name):
        """
        updateHelpForm() is called when the user clicks an item in the custom
        dictionary tree.
        """
        if dimension_name == None:  # cluster
            c = self.dictionary.getCluster(cluster_name)
            self.helpForm.setHelpContent(c)
        else:  # dimension
            d = self.dictionary.getDimension(dimension_name)
            self.helpForm.setHelpContent(d)

    def categoryHelpContentChanged(self, new_name, new_label, new_content):
        """
        categoryHelpContentChanged() is called when the user select a cluster or a
        dimension in the tree view. updateDimensionHelpContent() checks if any of the
        help content items have been changed, and returns True if the dictionary
        is in fact edited.
        """

        new_content = new_content.replace('\n', ' ')

        curr_cluster, curr_dimension = self.dictTree.getSelectedItem()

        if curr_dimension == None: # a cluster is selected
            if curr_cluster != new_name and self.dictionary.existingCategoryName(new_name) == True:
                WarningDialog("Warning", "{} is already used.".format(new_name))
                return False

        elif curr_cluster != None and curr_dimension != None: # a dim is selected
            if curr_dimension != new_name and self.dictionary.existingCategoryName(new_name) == True:
                WarningDialog("Warning", "{} is already used.".format(new_name))
                return False

        edited = False
        old_cluster, old_dimension = self.dictTree.categoryNameChanged(new_name)

        if old_cluster != None and old_dimension == None:
            edited = self.dictionary.updateClusterHelpContent(old_cluster, new_name, 
                                                                   new_label, new_content)
        elif old_dimension != None and old_cluster != None:
            edited = self.dictionary.updateDimensionHelpContent(old_cluster, old_dimension, new_name, 
                                                                     new_label, new_content)
        else: 
            edited = False
            
        if edited == True:
            self.dictChanged(18)

        return True 
        
    def clearHelpContent(self):
        if self.helpForm != None:
            self.helpForm.clearHelpContent()

    def generateHelpDoc(self):
        """
        generateHelpDoc() genreates a MS Word file that includes all the help strings at the cluster level.
        """
        dict_name = os.path.basename(self.dictionary.getDirectory())

        filepath, _ = QFileDialog.getSaveFileName(None, 'Enter File Name', './{}_HelpDoc.docx'.format(dict_name), "Text File (*.docx)")
        if filepath != "":

            docx = Document()
            styles = docx.styles

            style = styles['Title']   
            font = style.font
            font.color.rgb = RGBColor(0,0,0)

            style = styles['Heading 1']   
            font = style.font
            font.color.rgb = RGBColor(0,0,0)
            style.paragraph_format.keep_with_next = True

            style = styles['Heading 2']   
            font = style.font
            font.color.rgb = RGBColor(0,0,0)
            style.paragraph_format.keep_with_next = True

            style = styles['Normal']   
            font = style.font
            font.color.rgb = RGBColor(0,0,0)
            style.paragraph_format.keep_together = True
     
            p = docx.add_paragraph(dict_name)
            p.style = 'Title'

            for c in self.dictionary.getTones():
                if c['name'] == 'Other':
                    continue

                if len(c['dimensions']) == 0:
                    WarningDialog("Warning", "{} is not included becasue it has no dimensions.".format(c['name']))
                else:
                    p = docx.add_paragraph(c['label'])
                    p.style = 'Heading 1'
                    p = docx.add_paragraph(c['help'])
                    p.style = 'Normal'

                for d in c['dimensions']:
                    p = docx.add_paragraph(d['label'])
                    p.style = 'Heading 2'
                    p = docx.add_paragraph(d['help'])
                    p.style = 'Normal'

            docx.save(filepath)


    def invalidateTagger(self):
        self.tagger_instance = None
        if self.dictionary is not None:
            self.dictionary.invalidateTagger()

    def unloadDictionary(self):
        """
        Reset the dictionary.
        """
        self.dict_path  = ""
        self.dictionary = None

    def setupDictionary(self, folder):
        """
        setupDictionary() is used to create a DSDict object. It is used right after a new dictionary
        is created, or an existing dictionaryis opened.
        """
        self.dict_path = folder
        d = ds_dict.DSDict(self.controller, self.dict_path, compact=True)
        d.loadTones()
        d.loadWordClasses()
        d.loadHelp()
        d.loadInfo()
        self.setDictionary(d)

    def createDictionary(self, dict_path):
        """
        createDictionary() is used to create a new dictionary. The new
        dictionary includes a simplest possible dictionary structure by default.
        """
        if os.path.exists(dict_path):
            WarningDialog("Error", "{} already exists.".format(dict_path))
        else:
            os.makedirs(dict_path)
            tones_path       = os.path.join(dict_path, "_tones.txt")
            help_path        = os.path.join(dict_path, "_help.txt")
            lat_path         = os.path.join(dict_path, "LAT1.txt")
            wordclasses_path = os.path.join(dict_path, "_wordclasses.txt")

            with open(tones_path, 'w', encoding="utf-8") as fout:
                fout.write("CLUSTER: Cluster1\n")
                fout.write("DIMENSION: Dimension1\n")
                fout.write("LAT: LAT1\n\n")

            with open(help_path, 'w', encoding="utf-8") as fout:
                fout.write("CLUSTER: Cluster1\n")
                fout.write("Cluster 1\n")
                fout.write("n/a\n")
                fout.write("DIMENSION: Dimension1\n")
                fout.write("Dimension 1\n")
                fout.write("n/a\n\n")

            with open(lat_path, 'w', encoding="utf-8") as fout:
                fout.write("\n")

            with open(wordclasses_path, 'w', encoding="utf-8") as fout:
                fout.write("\n")

        self.unloadDictionary()
        self.setupDictionary(dict_path)

##################################################
#
# PopupMenu (used ad drop-down menus)
#
##################################################
class PopupMenu(QMenu):
    def __init__(self, button, parent=None):
        super(PopupMenu, self).__init__(parent)
        self.button = button

    def showEvent(self, event):
        p = self.pos()
        geo = self.button.geometry()
        self.move(p.x()+geo.width()-self.geometry().width(), p.y())

##################################################
#
# RepairDialog
#
##################################################
class RepairDialog(QDialog):
    def __init__(self, editor_win=None, parent=None):
        super(RepairDialog, self).__init__(parent)

        self.editor_win = editor_win

        self.repairCompleted = False
        self.collisions = []
        self.wordclasses = []
        self.current_index = 0

        width  = 500
        # height = 360
        if platform.system() == 'Windows':
            self.setMinimumWidth(1.5*width)
        else:
            self.setMinimumWidth(width)

        self.collisionButtons = list()

        vbox = QVBoxLayout()
        self.pattern = QLabel("Select the pattern to keep:")
        vbox.addWidget(self.pattern)

        vbox.addSpacing(5)

        self.buttonContainer = QVBoxLayout()

        vbox.addLayout(self.buttonContainer)
        vbox.addStretch()

        self.resolveButton = QPushButton("Resolve Collision(s)")
        self.resolveButton.clicked.connect(self.resolveCollisions)
        self.resolveButton.setEnabled(False)
        self.resolveButton.setAutoDefault(False)
        vbox.addWidget(self.resolveButton)

        self.setLayout(vbox)

        # Center the window
        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)

        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())  

    def patternSelected(self):
        self.resolveButton.setEnabled(True)        

    def resetButtons(self):
        self.resolveButton.setEnabled(False)
        for b in self.collisionButtons:
            b.setAutoExclusive(False)

        for b in self.collisionButtons:
            b.setChecked(False)

        for b in self.collisionButtons:
            b.setAutoExclusive(True)

    def findSelectedPatternAndRemove(self):
        checked_button_index = -1
        for i in range(len(self.collisionButtons)):
            b = self.collisionButtons[i]
            if b.isChecked():
                checked_button_index = i
                break

        keep_indexes = self.consolidatedCollisions[checked_button_index][1][1]

        collision = self.collisions[self.current_index]

        for i in range(len(collision)):
            if i not in keep_indexes:
                item = collision[i]
                p = ' '.join(item[1])
                self.editor_win.removePattern(p, item[0])

    def closeEvent(self, event):
        if self.repairCompleted:
            event.accept()
            return

        res  = YesNoDialog("Warning", "You haven't finished fixing all the collisions. Are you sure you want to close the repair dialog?")
        if res.retval == QMessageBox.Ok:
            event.ignore()
        elif res.retval == QMessageBox.No:
            event.accept()

    def setCollisions(self, collisions, wordclasses):
        self.repairCompleted = False
        self.collisions = collisions
        self.wordclasses = wordclasses
        self.current_index = 0
        self.num_collisions = len(collisions)
        self.setCurrentCollision()    

    def setCurrentCollision(self):
        self.resetButtons()
        patterns = self.collisions[self.current_index]

        self.setWindowTitle("Repair Collision ({} of {})".format(self.current_index+1, self.num_collisions))

        d = dict()
        count = 0
        for c in patterns:
            clist = list()
            if c[0] not in d:
                d[c[0]] = ["\"{}\"\n".format(' '.join(c[1])), [count]]
            else:
                s = d[c[0]][0] + "\"{}\"\n".format(' '.join(c[1]))
                clist = d[c[0]][1]
                clist.append(count)
                d[c[0]] = [s, clist]

            count+=1

        self.consolidatedCollisions = list(d.items())

        # Empyt the layout!
        self.collisionButtons.clear()
        for i in reversed(range(self.buttonContainer.count())):
           btn = self.buttonContainer.takeAt(i)
           btn.widget().setParent(None)

        # Create optional buttons.
        for item in self.consolidatedCollisions:
            label = "{} ({})".format(item[1][0], item[0])
            # label = "{} ({})".format(item[1], item[0])

            button = QRadioButton(label)
            button.clicked.connect(self.patternSelected)
            self.buttonContainer.addWidget(button)
            self.collisionButtons.append(button)    

    def resolveCollisions(self):
        """
        Here we'll revise the LAT:
        1. Find the LAT(s) where a pattern should be removed
        2. Remove it from the LAT
        3. Check the next set of collisionsa
        """

        self.findSelectedPatternAndRemove()

        #############################################

        # Let's ask the user to resolve the next collision.
        self.current_index+=1

        if self.current_index < len(self.collisions):
            self.setCurrentCollision()
        else:
            self.repairCompleted = True
            self.collisions = []
            self.close()
            self.editor_win.repairCompleted()

            if len(self.wordclasses) > 0:
                txt = ""
                for wc in self.wordclasses:
                    self.bFoundCollisions = True
                    txt += "{} ({})\n".format(wc[0], wc[1])
                WarningDialog("Missing word classes:", txt)
                self.wordclasses = []

##################################################
#
# EditableDictTrees
#
##################################################
class EditableDictTree(QTreeWidget):
    def __init__(self, editor_win=None, parent=None):
        super(EditableDictTree, self).__init__(parent)

        self.editor_win = editor_win
        self.curr_cluster = None
        self.curr_dimension = None
        self.curr_cluster_index = -1
        self.curr_dimension_index = -1
        self.selected_item = None
        self.help_form = None

        font_size = views.default_ui_font_size


        # Style
        # self.setStyleSheet("QTreeWidget {font-asize: " + str(font_size) +"pt;}")
        self.setStyleSheet("EditableDictTree { font-size: " + str(font_size) + " pt;" + \
                                          "background-color: #fff;" + \
                                          "selection-color: #fff;" + \
                                          "selection-background-color: #0080ff;};")

        self.setHeaderHidden(True)

        # Behavior
        self.itemClicked.connect(self.itemSelected)
        self.setFocusPolicy(Qt.NoFocus)

        self.setSelectionMode(QAbstractItemView.ContiguousSelection)

    def scrollToCurrentItem(self):
        if self.selected_item != None:
            self.scrollToItem(self.selected_item)
        
    def getSelectedItem(self):
        if self.selected_item != None:
            item = self.selected_item
            parent = item.parent()

            if parent == None:
                cluster_name   = item.text(0)
                dimension_name = None
            else:
                cluster_name   = parent.text(0)
                dimension_name = item.text(0)
            return cluster_name, dimension_name
        else:
            return None, None

    def getSelectedClusterIndex(self):
        if self.selected_item != None:
            item = self.selected_item
            parent = item.parent()

            if parent == None:
                index = self.indexOfTopLevelItem(item)
            else:
                index = self.indexOfTopLevelItem(parent)
        else:
            index = -1

        return index

    def saveLatestChanges(self):
        if self.selected_item != None:
            h = self.help_form.getHelpContent()
            return self.editor_win.categoryHelpContentChanged(h['name'],
                                                              h['label'],
                                                              h['help'])
    def mousePressEvent(self, event):
        self.editor_win.resetLATFilter()

        res = self.saveLatestChanges()
        if res == False:
            return

        if self.selected_item != None:
            cluster_name = self.editor_win.getClusterName()  # get a cluster name form the heelp form
            if cluster_name != None:
                self.selected_item.setText(0, cluster_name)

        # clear selection if the background of the tree view is clicked
        self.clearSelection()

        self.curr_cluster = None
        self.curr_dimension = None
        self.curr_cluster_index = -1
        self.curr_dimension_index = -1
        self.selected_item = None

        # if self.help_form:
        #     self.help_form.clearHelpContent()

        self.editor_win.validateSelectedLAT()

        self.editor_win.updateLATlist(None, None)
        self.editor_win.clearLATDetails()
        self.editor_win.disableGlobalReplaceDimNames()

        # send the event to its parent
        QTreeWidget.mousePressEvent(self, event)

        if self.selected_item == None and self.help_form:
            self.help_form.clearHelpContent()

    def setHelpForm(self, form):
        self.help_form = form

    def moveSelectedItem(self, direction):
        """
        Move up/dow the currently selected item.
        """

        # Moving a cluster item
        if self.curr_cluster_index < 0:
            return None

        if self.curr_dimension_index < 0:          
            num_clusters = self.topLevelItemCount()

            if self.curr_cluster_index == 0 and direction == UP:
                return None
            elif self.curr_cluster_index == (num_clusters-1) and direction == DOWN:
                return None
            else:
                child = self.topLevelItem(self.curr_cluster_index)
                bExpanded = child.isExpanded()
                child = self.takeTopLevelItem(self.curr_cluster_index)
                self.insertTopLevelItem(self.curr_cluster_index + direction, child) # move it up or down   
                self.setCurrentItem(child)
                child.setExpanded(bExpanded)
                self.curr_cluster_index += direction

            return child.text(0)

    def setHierarchy(self, hierarchy):
        """
        Set a new dictionary hierarchy. The old/existing
        hierarchy is cleared.
        """

        self.clear()

        self.isTreeInitialized = True

        self.curr_cluster = None
        self.curr_dimension = None
        self.curr_cluster_index = -1
        self.curr_dimension_index = -1
        self.selected_item = None

        for c in hierarchy:
            if c['name'] == 'Other':
                continue
            citem = QTreeWidgetItem(self, [c['name']])
            for d in c['dimensions']:
                ditem = QTreeWidgetItem(citem, [d['name']])

        self.selected_item == None
        self.isTreeInitialized = False

        first = self.topLevelItem(0)
        self.scrollToItem(first)

    def resetHierarchy(self, hierarchy, selected_clust, selected_dims=None):
        """
        Set a new dictionary hierarchy. The old/existing
        hierarchy is cleared.
        """

        # Let's find all the expanded clusteres
        expanded_clusts = [selected_clust]
        for i in range(self.topLevelItemCount()):
            item = self.topLevelItem(i)
            if item.isExpanded() == True:
                expanded_clusts.append(item.text(0))

        self.clear()

        self.isTreeInitialized = True

        self.curr_cluster = None
        self.curr_dimension = None
        self.curr_cluster_index = -1
        self.curr_dimension_index = -1
        self.selected_item = None

        for c in hierarchy:
            if c['name'] == 'Other':
                continue
            citem = QTreeWidgetItem(self, [c['name']])
            for d in c['dimensions']:
                ditem = QTreeWidgetItem(citem, [d['name']])

        self.selected_item == None
        self.isTreeInitialized = False

        for ci in range(self.topLevelItemCount()):
            clust_item = self.topLevelItem(ci)
            clust_name = clust_item.text(0)
            if clust_name in expanded_clusts:
                clust_item.setExpanded(True)

            if clust_name == selected_clust:
                if selected_dims != None:
                    for di in range(clust_item.childCount()):
                        dim_item = clust_item.child(di)
                        if dim_item.text(0) in selected_dims:
                            # self.setCurrentItem(dim_item)
                            dim_item.setSelected(True)
                            # break
                else:
                    self.setCurrentItem(item)

        first = self.topLevelItem(0)
        self.scrollToItem(first)

    def categoryNameChanged(self, new_label):

        # return if nothing is selected now
        if self.curr_cluster == None:
            return None, None

        if self.selected_item != None:
            # change the label of the selected item.n
            item = self.selected_item
            item.setText(0, new_label)

            if self.curr_dimension == None:
                # if a cluster is currently selected
                old_label = self.curr_cluster
                self.curr_cluster = new_label
                return old_label, None
            else:
                # if a dimension is currently selected
                old_label = self.curr_dimension
                self.curr_dimension = new_label
                parent = item.parent()
                if parent:
                    parent.sortChildren(0, Qt.AscendingOrder)
                return self.curr_cluster, old_label

        else:
            return None, None

    def selectADimension(self, dimension_name):
        items = self.findItems(dimension_name, Qt.MatchFixedString | Qt.MatchRecursive, 0)

        if len(items)>0:
            item = items[0]
            parent = item.parent()
            self.expandItem(parent)
            self.setCurrentItem(item)
            self.itemSelected(item, 0)
            return True
        else:
            return False

    def itemSelected(self, item, column):
        """
        This function is called if an item is clicked.
        """
        self.help_form.setIgnoreChanges(True)

        indexes = self.selectedIndexes() 

        if len(indexes) > 1:
            # if 2 or more items are selected, clear the LAT and the help form (i.e., nothing is selected)
            # self.clearSelection()     
            self.curr_cluster = None
            self.curr_dimension = None
            self.curr_cluster_index = -1
            self.curr_dimension_index = -1
            self.selected_item = None
            self.editor_win.validateSelectedLAT()
            self.editor_win.updateLATlist(None, None)
            self.editor_win.clearLATDetails()
            self.editor_win.enableGlobalReplaceDimNames()
            self.help_form.clearHelpContent()
            self.help_form.setIgnoreChanges(False)
            return


        self.selected_item = item
        parent = item.parent()

        if parent == None:  # No parent. 'item' is a cluster
            self.curr_cluster = item.text(0)
            self.curr_cluster_index = self.indexOfTopLevelItem(item)
            self.curr_dimension = None
            self.curr_dimension_index = -1
        else: # 'item' is a dimension
            self.curr_cluster   = parent.text(0)
            self.curr_cluster_index = self.indexOfTopLevelItem(parent)
            self.curr_dimension = item.text(0)
            self.curr_dimension_index = self.indexFromItem(item).row()

        self.editor_win.updateHelpForm(self.curr_cluster, self.curr_dimension)
        self.editor_win.updateLATlist(self.curr_cluster, self.curr_dimension)
        self.editor_win.disableGlobalReplaceDimNames()
        self.help_form.setIgnoreChanges(False)

    def deleteSelectedItem(self):
        """
        """
        if self.selected_item != None:
            dim_label = None
            # cluster
            if self.selected_item.parent() == None:  
                cluster_label = self.selected_item.text(0)
                index = self.indexOfTopLevelItem(self.selected_item)
                self.takeTopLevelItem(index)
                self.clearSelection()
                self.curr_dimension_index = -1
                self.curr_dimension = None
                self.curr_cluster_index = -1
                self.curr_cluster = None
                self.selected_item = None

            # dimension
            else:
                dim_label = self.selected_item.text(0)
                parent    = self.selected_item.parent()

                b = parent.foreground(0)
                if b.color() == Qt.gray: # The cluster is from the default dict.
                    WarningDialog("Warning", "{} cannot be deleted. It is part of an imported cluster.".format(dim_label))
                    return None, None

                index     = parent.indexOfChild(self.selected_item)
                parent.takeChild(index)
                num_children = parent.childCount()
                cluster_label = self.curr_cluster

                if num_children > 0:
                    next_item = parent.child(index) 
                    if next_item != None:
                        next_item.setSelected(True)
                        self.curr_dimension_index = index
                        self.curr_dimension = next_item.text(0)
                        self.selected_item  = next_item
                    else:
                        self.clearSelection()
                        self.curr_dimension_index = -1
                        self.curr_dimension = None
                        self.curr_cluster_index = -1
                        self.curr_cluster = None                        
                        self.selected_item = None
                else:
                    self.clearSelection()
                    self.curr_dimension_index = -1
                    self.curr_dimension = None
                    self.curr_cluster_index = -1
                    self.curr_cluster = None                        
                    self.selected_item = None

            return cluster_label, dim_label

        else:
            if self.topLevelItemCount() == 0:
                return None, None
            else:
                WarningDialog("Warning", "Nothing is selected.")
            return None, None

    def addNewCluster(self, index, cluster_name):
        new_item = QTreeWidgetItem([cluster_name])
        if index > 0:
            self.insertTopLevelItem(index, new_item)
        else:
            self.addTopLevelItem(new_item)

        self.setCurrentItem(new_item)
        self.itemSelected(new_item, 0)
        self.scrollToItem(new_item)

    def addNewDimension(self, index, dimension_name):
        parent = self.topLevelItem(index)
        new_item = QTreeWidgetItem([dimension_name])
        parent.addChild(new_item)
        parent.setExpanded(True)
        parent.sortChildren(0, Qt.AscendingOrder)

        self.setCurrentItem(new_item)
        self.itemSelected(new_item, 0)
        self.scrollToItem(new_item)

    def sortDimensions(self):
        for i in range(self.topLevelItemCount()):
            clust = self.topLevelItem(i)
            clust.sortChildren(0, Qt.AscendingOrder)

##################################################
#
# WordClass Dialog
#
##################################################
class WordClassDialog(QDialog):
    def __init__(self, editor_win=None, parent=None):

        super(WordClassDialog, self).__init__(parent)

        self.editor_win = editor_win;

        font_size = views.default_ui_font_size

        self.wcProxyModel = None

        width  = 640
        height = 640
        self.setGeometry(60,80,width,height)
        self.setMinimumWidth(width)
        self.setMinimumHeight(height)

        self.setStyleSheet("background-color: #eee;")

        vbox = QVBoxLayout()

        # The header aera.              
        hbox = QHBoxLayout()

        title = QLabel("Word Classes")
        title.setStyleSheet("font-size: {}pt; font-weight: bold;".format(font_size))
        hbox.addWidget(title)

        self.toolButton = QToolButton()
        self.toolButton.setIcon(QIcon(QPixmap(resource_path("data/icons/menu_icon.png"))))
        self.toolButton.setStyleSheet('QToolButton::menu-indicator { image: none; }')
        self.toolButton.setPopupMode(QToolButton.InstantPopup)

        toolsMenu = PopupMenu(self.toolButton)

        self.closeAction = QAction('Close', self)        
        self.closeAction.triggered.connect(self.close)
        toolsMenu.addAction(self.closeAction)

        self.toolButton.setMenu(toolsMenu)

        hbox.addWidget(self.toolButton)

        vbox.addLayout(hbox)

        line = QFrame()
        line.setFrameShape(QFrame.HLine);
        line.setFrameShadow(QFrame.Sunken);
        vbox.addWidget(line)

        # Control UI
        hbox = QHBoxLayout()
        add_button = QPushButton("Add")
        add_button.clicked.connect(self.addWC)
        add_button.setAutoDefault(False)
        hbox.addWidget(add_button)

        delete_button = QPushButton("Delete")
        delete_button.clicked.connect(self.deleteWC)
        delete_button.setAutoDefault(False)        
        hbox.addWidget(delete_button)

        self.wcFilterInput = QLineEdit()                                # filter input
        self.wcFilterInput.textChanged.connect(self.filterWCs)
        self.wcFilterInput.setStyleSheet("QLineEdit {background-color: #fff; font-size: " + str(font_size)+ "pt;}")
        hbox.addWidget(self.wcFilterInput)

        vbox.addLayout(hbox)

        # Main area
        hbox = QHBoxLayout()

        # List of Word Classes
        self.wcListView = QListView()
        self.wcListView.clicked.connect(self.wcSelected)
        self.wcListView.setStyleSheet("QListView { font-size: " + str(font_size) + " pt;" + \
                                            "background-color: #fff;" + \
                                        "selection-color: #fff;" + \
                                        "selection-background-color: #0080ff;};")

        hbox.addWidget(self.wcListView)

        # WC Name + list of words
        vbox_right = QVBoxLayout()

        hbox_ui = QHBoxLayout()
        header = QLabel("Properties")
        header.setStyleSheet("QLabel {font-weight: bold; font-size: " + str(font_size) + "pt;}")
        hbox_ui.addWidget(header)

        hbox_ui.addStretch()

        self.update_button = QPushButton("Update")                                # Wordclass
        self.update_button.clicked.connect(self.updateWCContent)
        self.update_button.setAutoDefault(False)        
        self.update_button.setEnabled(False)

        hbox_ui.addWidget(self.update_button)

        vbox_right.addLayout(hbox_ui)

        vbox_right.addWidget(QLabel("Name:"))

        self.wcName = QLineEdit()
        self.wcName.setStyleSheet("QLineEdit { background-color: #fff; font-size: " + \
                                   str(font_size) + "pt;}")
        vbox_right.addWidget(self.wcName)

        vbox_right.addWidget(QLabel("Words:"))
        self.editor = QPlainTextEdit()
        self.editor.setStyleSheet("QPlainTextEdit {background-color: #fff; font-size: " + \
                           str(font_size) + "pt;}" + \
                           "QMenu {selection-background-color: #0080ff;}")

        self.editor.modificationChanged.connect(self.wcChanged)
        doc = self.editor.document()
        doc.setDocumentMargin(2)
        vbox_right.addWidget(self.editor)

        hbox.addLayout(vbox_right)

        vbox.addLayout(hbox)

        self.setLayout(vbox)

        # Center the window
        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)

        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())   

        self.current_label = None
        self.just_selected_label = None
        self.bRevised = False
        self.untitled_count = 0

    def wcChanged(self, bModified):
        if bModified:
            self.update_button.setEnabled(True)
        else:
            self.update_button.setEnabled(False)

    def filterWCs(self):
        if self.wcProxyModel != None:
            self.cleanUpFields()
            filter_str = self.wcFilterInput.text()
            self.wcProxyModel.setFilterWildcard(filter_str)
            self.wcListView.clearSelection()

    def cleanUpFields(self):
        self.wcName.setText("")
        self.editor.setPlainText("")
#        self.wclist.setCurrentRow(0)
        self.current_label = None
        self.just_selected_label = None
        self.bRevised = False

    def closeEvent(self, event):
        self.updateCurrentSelection()        
        if self.bRevised:
            self.save()
            ConfirmationDialog("Some word classes have been changed.", "")
        else:
            ConfirmationDialog("No Changes are made in wore classes.", "")

        self.wcFilterInput.setText("")

    def addWC(self):
        self.cleanUpFields()
        self.wcFilterInput.setText("")

        self.untitled_count += 1
        name = '!UNTITLED{}'.format(self.untitled_count)

        self.wordclasses[name] = list()

        self.updateCurrentSelection()
        self.current_label = name        

        # if self.wcProxyModel != None:
        #     qModelIndex = self.wcProxyModel.match(self.wcProxyModel.index(0, 0), Qt.DisplayRole, name)[0]
        #     self.wcListView.setCurrentIndex(qModelIndex)
        #     self.wcListView.setFocus()
        #     self.wcListView.scrollTo(qModelIndex)

        self.updateEditor(name)
        self.updateList()
        self.bRevised = True

        # item = QListWidgetItem(name)
        # self.wclist.addItem(item)
        # self.wclist.sortItems()
        # self.wclist.setCurrentItem(item)
        # self.wclist.setFocus()


    def deleteWC(self):
        del self.wordclasses[self.current_label]
        self.current_label = None

        # self.wclist.clearSelection()
        self.bRevised = True

        self.updateList()

    def save(self):
        self.editor_win.updateWordClasses(self.wordclasses)
        self.cleanUpFields()
        self.bRevised = False

    def updateWCContent(self):
        if self.current_label != None:
            self.updateCurrentSelection()
            self.updateList()
            self.update_button.setEnabled(False)

    def updateCurrentSelection(self):

        if self.current_label != None:
            # let's save the current list of words first.
            new_label = self.wcName.text()
            new_label = new_label.strip().upper()
            if new_label != "" and new_label[0] != '!':
                new_label = '!' + new_label
            
            content = self.editor.toPlainText()
            words = content.lower().strip().splitlines()
            words = [w for w in words if w!='']
            words.sort()

            if self.current_label != new_label:
                # if the label has been edited, create a new key-value pair
                if self.wordclasses.get(self.current_label, None) != None:
                    del self.wordclasses[self.current_label]

                self.current_label = new_label
                self.wordclasses[self.current_label] = words

                self.bRevised = True

                # self.updateList()

            elif self.wordclasses[self.current_label] != words:
                # if the label hasn't been edited, but if the list of words
                # are different, update the list.
                self.wordclasses[self.current_label] = words                
                self.bRevised = True
        # else:
        #     self.updateList()

    def wcSelected(self, index):
        self.just_selected_label = index.data()

        # do nothing if the same item is clicked
        if self.current_label == self.just_selected_label:
            return

        self.updateCurrentSelection()
        self.updateEditor(self.just_selected_label)
        self.current_label = self.just_selected_label
        # self.updateList()
        self.just_selected_label = None

    def updateEditor(self, label):
        self.current_label = label
        self.wcName.setText(label)
        words = self.wordclasses.get(self.current_label, None)
        if words != None:
            self.editor.setPlainText('\n'.join(words))

    def updateList(self):
        words = sorted(self.wordclasses.keys())
        # for wc in  words:
        #     self.wclist.addItem(wc)

        self.wcModel = QStringListModel(self)

        if len(words)>0:
            self.wcModel.setStringList(words)

        self.wcProxyModel = QSortFilterProxyModel  ()
        self.wcProxyModel.setSourceModel(self.wcModel)
        self.wcProxyModel.setFilterCaseSensitivity(Qt.CaseInsensitive)
        self.wcListView.setModel(self.wcProxyModel)

        if self.current_label != None and self.wcProxyModel != None:
            qModelIndex = self.wcProxyModel.match(self.wcProxyModel.index(0, 0), 
                                        Qt.DisplayRole, self.current_label)[0]
            self.wcListView.setCurrentIndex(qModelIndex)
            self.wcListView.setFocus()
            self.wcListView.scrollTo(qModelIndex)

    def setContent(self, wordclasses):
        self.cleanUpFields()
        self.wordclasses = wordclasses.copy()

        words = sorted(wordclasses.keys())
        # for wc in words:
        #     self.wclist.addItem(wc)

        self.wcModel = QStringListModel(self)

        if len(words)>0:
            self.wcModel.setStringList(words)

        self.wcProxyModel = QSortFilterProxyModel  ()
        self.wcProxyModel.setSourceModel(self.wcModel)
        self.wcProxyModel.setFilterCaseSensitivity(Qt.CaseInsensitive)
        self.wcListView.setModel(self.wcProxyModel)


##################################################
#
# Dictionary Error Dialog
#
##################################################
class DictErrorDialog(QDialog):
    def __init__(self, html_str, editor_win=None, parent=None):

        super(DictErrorDialog, self).__init__(parent)

        self.editor_win = editor_win;

        font_size = views.default_ui_font_size

        width  = 900
        height = 640
        self.setGeometry(60,80,width,height)
        self.setMinimumWidth(width)
        self.setMinimumHeight(height)

        self.setStyleSheet("background-color: #eee;")

        vbox = QVBoxLayout()

        # Control UI
        hbox = QHBoxLayout()
        self.title = QLabel("Dictionary Errors")
        self.title.setStyleSheet("font-size: {}pt; font-weight: bold;".format(font_size))
        hbox.addWidget(self.title)

        vbox.addLayout(hbox)

        self.output_area = QTextEdit()
        self.output_area.setStyleSheet("QTextEdit { background-color: #fff; font-size: " + \
                                       str(font_size) + "pt;}")
        self.output_area.setHtml(html_str)

        vbox.addWidget(self.output_area)
        hbox = QHBoxLayout()
        hbox.addStretch()

        self.close_button = QPushButton("Close")
        self.close_button.clicked.connect(self.close)
        self.close_button.setAutoDefault(False)        
        hbox.addWidget(self.close_button)

        self.repair_button = QPushButton("Repair Collisions")
        self.repair_button.clicked.connect(self.repairCollisions)
        self.repair_button.setAutoDefault(False)        
        hbox.addWidget(self.repair_button)        

        vbox.addLayout(hbox)

        self.setLayout(vbox)

        # Center the window
        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)

        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())   

    def repairCollisions(self):
        self.editor_win.repairCollisions()

##################################################
#
# Help Form
#
##################################################
class HelpForm(QWidget):
    def __init__(self, editor_win=None, parent=None):
        super(HelpForm, self).__init__(parent)

        self.ignoreChanges = True

        self.editor_win = editor_win

        font_size = views.default_ui_font_size

        vbox = QVBoxLayout() # main layout

        #
        # Header
        #
        hbox_properties = QHBoxLayout()

        header = QLabel("Properties")
        header.setStyleSheet("QLabel {font-weight: bold; font-size: " + str(font_size) + "pt;}")
        hbox_properties.addWidget(header)

        hbox_properties.addStretch()

        self.update_button = QPushButton("Update")
        self.update_button.clicked.connect(self.categoryHelpContentChanged)
        self.update_button.setAutoDefault(False)
        self.update_button.setEnabled(False)
        hbox_properties.addWidget(self.update_button)

        vbox.addLayout(hbox_properties)

        # 
        # Help Form
        #
        fbox = QFormLayout()
        fbox.setContentsMargins(0,0,0,0)

        # help file 
        self.help_name = QLineEdit()
        self.help_name.setReadOnly(False)
        self.help_name.setStyleSheet("QLineEdit { background-color: #fff; font-size: " + str(font_size) + "pt;}")
        self.help_name.textEdited.connect(self.contentChanged)

        label = QLabel("Name: ")
        label.setStyleSheet("font-size: {}pt".format(font_size))
        fbox.addRow(label, self.help_name)

        self.help_label = QLineEdit()
        self.help_label.setReadOnly(False)
        self.help_label.setStyleSheet("QLineEdit { background-color: #fff; font-size: " + str(font_size) + "pt;}")
        self.help_label.textEdited.connect(self.contentChanged)
        label = QLabel("Label: ")
        label.setStyleSheet("font-size: {}pt".format(font_size))
        fbox.addRow(label, self.help_label)

        self.help_content = QPlainTextEdit()
        self.help_content.setMinimumHeight(200)
        self.help_content.setStyleSheet("QPlainTextEdit {background-color: #fff; font-size: " + str(font_size) + "pt;}")
        self.help_content.textChanged.connect(self.contentChanged)
        doc = self.help_content.document()
        doc.setDocumentMargin(2)
        label = QLabel("Content: ")
        label.setStyleSheet("font-size: {}pt".format(font_size))
        fbox.addRow(label,self.help_content)

        vbox.addLayout(fbox)

        self.setLayout(vbox) 

        self.ignoreChanges = False

    def contentChanged(self):
        if self.ignoreChanges == False:
            self.update_button.setEnabled(True)
            self.editor_win.helpContentChanged()

    def categoryHelpContentChanged(self):

        name    = self.help_name.text()
        name    = remove_punct_and_space(name)
        self.help_name.setText(name)

        label   = self.help_label.text()
        content = self.help_content.toPlainText()

        self.editor_win.categoryHelpContentChanged(name, label, content)

        self.update_button.setEnabled(False)


    def saveItem(self):
        pass

    def getHelpContent(self):
        res = dict()
        name = remove_punct_and_space(self.help_name.text())
        res['name']  = name
        res['label'] = self.help_label.text()
        # res['help']  = self.help_content.toPlainText()

        help_description = self.help_content.toPlainText()
        res['help']  = utils.remove_doublespaces_and_newlines(help_description)

        self.help_name.setText(name)

        return res

    def setIgnoreChanges(self, val):
        self.ignoreChanges = val

    def clearHelpContent(self):
        self.ignoreChanges = True
        self.help_name.setText("")
        self.help_label.setText("")
        self.help_content.setPlainText("")
        self.ignoreChanges = False

    def setHelpContent(self, category_d):

        if category_d != None:
            self.help_name.setText(category_d.get('name', "n/a"))
            self.help_label.setText(category_d.get('label', "n/a"))
            self.help_content.setPlainText(category_d.get('help', "n/a"))

            self.help_name.home(False)
            self.help_label.home(False)

            self.help_label.clearFocus()

        else:
            self.help_name.setText("")
            self.help_label.setText("")
            self.help_content.setPlainText("")

    def clearAll(self):
        self.help_name.setText("")
        self.help_label.setText("")
        self.help_content.setPlainText("")


##################################################
#
# DimensionPickerDialog
#
##################################################
class DimensionPickerDialog(QDialog):

    def __init__(self, tones, editor_win=None, parent=None):
        super(DimensionPickerDialog, self).__init__(parent)

        self.editor_win = editor_win

        self.setMinimumWidth(500)

        font_size = views.default_ui_font_size

        vbox = QVBoxLayout() # main layout

        label = QLabel("Select a Dimension")
        label.setStyleSheet("font-weight: bold; font-size: {}pt".format(font_size))
        vbox.addWidget(label)

        self.tree = DimPickerTree(tones, editor_win=editor_win)
        vbox.addWidget(self.tree)

        hbox = QHBoxLayout()
        hbox.addStretch()

        self.cancel_button = QPushButton("Cancel")
        self.cancel_button.setAutoDefault(False) 
        self.cancel_button.clicked.connect(self.reject)
        hbox.addWidget(self.cancel_button)

        self.ok_button = QPushButton("OK")
        self.ok_button.clicked.connect(self.accept)
        self.ok_button.setAutoDefault(False) 
        hbox.addWidget(self.ok_button)

        vbox.addLayout(hbox)

        self.setLayout(vbox)

    def getSelectedDimension(self):
        clust, dim = self.tree.getSelectedDimension()
        return clust, dim


class DimPickerTree(QTreeWidget):
    def __init__(self, tones, editor_win=None, parent=None):
        super(DimPickerTree, self).__init__(parent)

        self.editor_win = editor_win
        self.curr_cluster = None
        self.curr_dimension = None
        self.curr_cluster_index = -1
        self.curr_dimension_index = -1

        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        font_size = views.default_ui_font_size

        # wd = 400
        ht = int((2*720)/3)
        # self.setMaximumWidth(wd)
        self.setMinimumHeight(ht)

        self.setSizePolicy(QSizePolicy.MinimumExpanding, QSizePolicy.MinimumExpanding)

        # Style
        self.setStyleSheet("QTreeWidget {font-size: " + str(font_size) +"pt;}")

        # Behavior
        header= QTreeWidgetItem(["Categories"])
        self.setHeaderItem(header)  
        self.itemClicked.connect(self.itemSelected)

        self.setHierarchy(tones)

    def getSelectedDimension(self):
        return self.curr_cluster, self.curr_dimension

    def itemSelected(self, item, column):

#        self.selected_item = item
        parent = item.parent()

        if parent == None:  # No parent. 'item' is a cluster
            self.curr_cluster = item.text(0)
            self.curr_cluster_index = self.indexOfTopLevelItem(item)
            self.curr_dimension = None
            self.curr_dimension_index = -1
        else: # 'item' is a dimension or a lat
            grand_parent = parent.parent()
            if grand_parent == None:
                self.curr_cluster   = parent.text(0)
                self.curr_cluster_index = self.indexOfTopLevelItem(parent)
                self.curr_dimension = item.text(0)
                self.curr_dimension_index = self.indexFromItem(item).row()

    def setHierarchy(self, hierarchy):  
        self.clear()
        for c in hierarchy:
            citem = QTreeWidgetItem(self, [c['name'], '-'])
            for d in c['dimensions']:
                ditem = QTreeWidgetItem(citem, [d['name'], '-'])

        first = self.topLevelItem(0)
        self.scrollToItem(first)
        self.setCurrentItem(first)
        self.curr_cluster = first.text(0)
        self.curr_cluster_index = self.indexOfTopLevelItem(first)

##################################################
#
# ClusterPickerDialog
#
##################################################
class ClusterPickerDialog(QDialog):

    def __init__(self, tones, editor_win=None, parent=None):
        super(ClusterPickerDialog, self).__init__(parent)

        self.editor_win = editor_win
        self.setMinimumWidth(500)

        font_size = views.default_ui_font_size

        vbox = QVBoxLayout() # main layout

        label = QLabel("Select a Cluster")
        label.setStyleSheet("font-weight: bold; font-size: {}pt".format(font_size))
        vbox.addWidget(label)

        self.tree = ClusterPickerTree(tones, editor_win=editor_win)
        vbox.addWidget(self.tree)

        hbox = QHBoxLayout()
        hbox.addStretch()

        self.cancel_button = QPushButton("Cancel")
        self.cancel_button.setAutoDefault(False) 
        self.cancel_button.clicked.connect(self.reject)
        hbox.addWidget(self.cancel_button)

        self.ok_button = QPushButton("OK")
        self.ok_button.clicked.connect(self.accept)
        self.ok_button.setAutoDefault(False) 
        hbox.addWidget(self.ok_button)

        vbox.addLayout(hbox)

        self.setLayout(vbox)

    def getSelectedCluster(self):
        return self.tree.getSelectedCluster()

class ClusterPickerTree(QTreeWidget):
    def __init__(self, tones, editor_win=None, parent=None):
        super(ClusterPickerTree, self).__init__(parent)

        self.editor_win = editor_win
        self.curr_cluster = None
        self.curr_cluster_index = -1

        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        font_size = views.default_ui_font_size

        # wd = 400
        ht = int((2*720)/3)
        # self.setMaximumWidth(wd)
        self.setMinimumHeight(ht)

        self.setSizePolicy(QSizePolicy.MinimumExpanding, QSizePolicy.MinimumExpanding)

        # Style
        self.setStyleSheet("QTreeWidget {font-size: " + str(font_size) +"pt;}")

        # Behavior
        header= QTreeWidgetItem(["Clusters"])
        self.setHeaderItem(header)
        self.itemClicked.connect(self.itemSelected)

        self.setHierarchy(tones)

    def getSelectedCluster(self):
        return self.curr_cluster

    def itemSelected(self, item, column):
        self.curr_cluster = item.text(0)
        self.curr_cluster_index = self.indexOfTopLevelItem(item)

    def setHierarchy(self, hierarchy):  
        self.clear()
        for c in hierarchy:
            citem = QTreeWidgetItem(self, [c['name']])

        first = self.topLevelItem(0)
        self.scrollToItem(first)
        self.setCurrentItem(first)
        self.curr_cluster = first.text(0)
        self.curr_cluster_index = self.indexOfTopLevelItem(first)

# ----------------------------------------
# FindPatternDialog
# ----------------------------------------

class FindPatternDialog(QDialog):
    def __init__(self, editor_win=None, parent=None):

        super(FindPatternDialog, self).__init__(parent)

        self.editor_win = editor_win;

        width  = 400
        height = 80

        self.setGeometry(60,80,width,height)
        self.setMinimumWidth(width)

        self.setStyleSheet("QDialog {background-color: #eee;}")

        vbox = QVBoxLayout()

        # search field
        self.pattern_field = QLineEdit()
        vbox.addWidget(self.pattern_field)
        self.msg = QLabel("")
        vbox.addWidget(self.msg)

        vbox.addSpacing(10)

        hbox = QHBoxLayout()

        self.close_button = QPushButton("Close")
        self.close_button.clicked.connect(self.close)
        self.close_button.setAutoDefault(False) 
        hbox.addWidget(self.close_button)             

        hbox.addStretch()

        self.find_button = QPushButton("Find")
        self.find_button.clicked.connect(self.findNext)
        self.find_button.setAutoDefault(True) 
        hbox.addWidget(self.find_button)  

        self.find_from_top_button = QPushButton("Find from Top")
        self.find_from_top_button.clicked.connect(self.findFromTop)
        self.find_from_top_button.setAutoDefault(False) 
        hbox.addWidget(self.find_from_top_button)  

        vbox.addLayout(hbox)

        self.setLayout(vbox)

        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)
      
        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())   

        # self.exec_()   
        self.exec()   
        # self.show()
        # self.raise_()


    def find(self, from_the_top):
        pattern = self.pattern_field.text()
        res = self.editor_win.findPattern(pattern, reset=from_the_top)
        if res:
            self.msg.setText("")
        else:
            self.msg.setText("Not found.")

    def findNext(self):
        label = self.find_button.text()
        if label == "Find":
            self.find_button.setText("Find Next")
        self.find(False)

    def findFromTop(self):
        self.find(True)

# ----------------------------------------
# FindCategoryDialog
# ----------------------------------------

class FindCategoryDialog(QDialog):
    def __init__(self, dict, editor_win=None, parent=None):

        super(FindCategoryDialog, self).__init__(parent)

        self.dictionary = dict
        self.editor_win = editor_win;

        width  = 600
        height = 600

        self.setGeometry(60,80,width,height)
        self.setMinimumWidth(width)
        self.setMinimumHeight(height)

        self.setStyleSheet("background-color: #eee;")

        vbox = QVBoxLayout()

        # search field
        self.pattern_field = QLineEdit()
        vbox.addWidget(self.pattern_field)

        hbox = QHBoxLayout()
        self.find_dimensions_box = QCheckBox("Find Dimensions")
        self.find_dimensions_box.setChecked(True)
        hbox.addWidget(self.find_dimensions_box)

        self.find_lats_box = QCheckBox("Find LATs")
        self.find_lats_box.setChecked(True)
        hbox.addWidget(self.find_lats_box)

        hbox.addStretch()

        self.find_button = QPushButton("Find Categories")
        self.find_button.clicked.connect(self.findCategories)
        hbox.addWidget(self.find_button)
        vbox.addLayout(hbox)

        self.search_results = QListWidget()
        vbox.addWidget(self.search_results)

        hbox = QHBoxLayout()
        self.cancel_button = QPushButton("Cancel")
        self.cancel_button.clicked.connect(self.cancel)
        self.cancel_button.setAutoDefault(False) 
        hbox.addWidget(self.cancel_button)   
             
        self.select_button = QPushButton("Select")
        self.select_button.clicked.connect(self.selectACategory)
        self.select_button.setAutoDefault(False) 
        hbox.addWidget(self.select_button)        

        vbox.addLayout(hbox)

        self.setLayout(vbox)      

        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)
      
        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())   

        retval = self.exec_()   

    def cancel(self):
        self.hide()

    def selectACategory(self):
        item = self.search_results.currentItem()
        if item == None:
            WarningDialog("Select a category:", "No dimension or LAT is selected.")
        else:
            label = item.text()
            categories = label.split(' > ')

            if len(categories) >=2:
                self.editor_win.selectACategory(categories)

            self.close()

    def findCategories(self):

        self.search_results.clear()


        pattern = self.pattern_field.text()

        bLATs = self.find_lats_box.isChecked()
        bDims = self.find_dimensions_box.isChecked()

        if pattern == "":
            WarningDialog("Warning", "No keywords are entered.")
        else:
            pattern = pattern.replace('*', '[a-zA-Z]*')
            matched_dims, matched_lats = self.dictionary.findCategories(pattern, bDims, bLATs)

            for dim in matched_dims:
                self.search_results.addItem(' > '.join(dim))
            for lat in matched_lats:
                self.search_results.addItem(' > '.join(lat))

            self.find_button.setAutoDefault(False) 
            self.select_button.setDefault(True)

# ----------------------------------------
# ThesaurusDialog
# ----------------------------------------
THESAURUS_SEARCH_IN_LATNAMES = 1
THESAURUS_SEARCH_IN_PATTERNS = 2

def get_intersecting_lats(lst):
    lst = [set(lats) for lats in lst]
    return list(lst[0].intersection(*lst))

class ThesaurusDialog(QDialog):
    LATsFindCompleted = pyqtSignal()

    def __init__(self, dict, editor_win=None, parent=None):

        super(ThesaurusDialog, self).__init__(parent)

        self.dictionary = dict
        self.editor_win = editor_win;

        width  = 1000
        height = 800

        self.setGeometry(60,80,width,height)
        self.setMinimumWidth(width)
        self.setMinimumHeight(height)

        self.setStyleSheet("background-color: " + views.default_ui_background_color + ";")

        vbox = QVBoxLayout()

        header_hbox = QHBoxLayout()

        title = QLabel("Thesaurus Search")
        header_hbox.addWidget(title)

        header_hbox.addStretch()

        label = QLabel("Search In")
        header_hbox.addWidget(label)

        self.options_button_group = QButtonGroup()

        self.lat_dim_names_radio_button = QRadioButton("LAT and Dimension Names")
        self.lat_dim_names_radio_button.setChecked(False)
        header_hbox.addWidget(self.lat_dim_names_radio_button)
        self.options_button_group.addButton(self.lat_dim_names_radio_button, THESAURUS_SEARCH_IN_LATNAMES)

        self.latnames_radio_button = QRadioButton("LAT Names (Only)")
        self.latnames_radio_button.setChecked(False)
        header_hbox.addWidget(self.latnames_radio_button)
        self.options_button_group.addButton(self.latnames_radio_button, THESAURUS_SEARCH_IN_LATNAMES)

        self.patterns_radio_button = QRadioButton("LAT Patterns")
        self.patterns_radio_button.setChecked(True)
        header_hbox.addWidget(self.patterns_radio_button)
        self.options_button_group.addButton(self.patterns_radio_button, THESAURUS_SEARCH_IN_PATTERNS)

        vbox.addLayout(header_hbox)

        search_hbox = QHBoxLayout()

        self.pattern_field = QLineEdit()
        search_hbox.addWidget(self.pattern_field)

        self.find_button = QPushButton("Find")
        self.find_button.clicked.connect(self.findCategories)
        self.find_button.setAutoDefault(True) 
        search_hbox.addWidget(self.find_button)

        vbox.addLayout(search_hbox)

        main_hbox = QHBoxLayout()

        result_vbox = QVBoxLayout()
        self.search_results = QTreeWidget()
        self.search_results.itemClicked.connect(self.itemSelected)
        self.search_results.setHeaderHidden(True)
        result_vbox.addWidget(self.search_results)

        self.category_description = QPlainTextEdit()
        self.setMaximumHeight(100)
        result_vbox.addWidget(self.category_description)
        main_hbox.addLayout(result_vbox)

        self.lat_patterns = QPlainTextEdit()
        main_hbox.addWidget(self.lat_patterns)

        vbox.addLayout(main_hbox)

        self.setLayout(vbox)

        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)
      
        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())   

        self.LATsFindCompleted.connect(self.LATsFindCompletedAction)

        retval = self.exec_()   

    def itemSelected(self, item, column):
        data = item.data(0, Qt.UserRole)
        categories = data.split()
        num_items = len(categories)
        if num_items == 3:
            _, patterns = self.dictionary.getPatterns(categories[2])
            self.lat_patterns.setPlainText(patterns)
            self.category_description.setPlainText("")

        elif num_items == 2: # dim
            dim = self.dictionary.getDimension(categories[1])
            if dim is not None:
                dscr = dim.get('help', 'n/a')
                self.category_description.setPlainText(dscr)
                self.lat_patterns.setPlainText("")
        elif num_items == 1: # clust
            clust = self.dictionary.getCluster(categories[0])
            if clust is not None:
                dscr = clust.get('help', 'n/a')
                self.category_description.setPlainText(dscr)
                self.lat_patterns.setPlainText("")
                
    def cancel(self):
        self.hide()

    def findCategories(self):
        
        self.lat_patterns.setPlainText("")

        if self.lat_dim_names_radio_button.isChecked() or \
           self.latnames_radio_button.isChecked():
            self.findCategoriesByLATNames()
        elif self.patterns_radio_button.isChecked():
            self.findCategoriesByPatterns()

    def findCategoriesByLATNames(self):
        self.search_results.clear()        

        user_input = self.pattern_field.text().strip()
        patterns = user_input.split()
        if patterns == []:
            WarningDialog("Warning", "No search terms are entered.")
        else:
            all_matched_lats = []
            b_dim = self.lat_dim_names_radio_button.isChecked()
            for p in patterns:
                pattern = "({}(?=[A-Z])|{}$)".format(p.capitalize(), p.capitalize())
                matched_dims, matched_lats = self.dictionary.findCategories(pattern, b_dim, True)
                all_matched_lats.append(matched_lats)

            lats = get_intersecting_lats(all_matched_lats)

            self.tones = self.makeHierarchy(lats)

            if self.tones is not None:
                self.updateResults()

    def findCategoriesByPatterns(self):
        self.search_results.clear()

        user_input = self.pattern_field.text().strip()
        patterns = user_input.split()

        if patterns == []:
            WarningDialog("Warning", "No search terms are entered.")
        else:
            self.progdialog = QProgressDialog("Searching... It may take a few minutes.", 
                                              None, 0, 0, self)
            self.progdialog.setModal(True)
            self.progdialog.setMinimumWidth(400)      
            self.progdialog.show()

            threading.Timer(0.2, self.runFindLATs, [patterns]).start()

    def makeHierarchy(self, matched_lats):

        tones = dict()

        for lat_path in matched_lats:
            clust_name = lat_path[0]
            dim_name   = lat_path[1]
            lat_name   = lat_path[2]

            if clust_name not in tones:     # clust_name hasn't breen added.
                c = dict()                    
                c['name'] = clust_name
                c['dimensions'] = list()

                d = dict()
                d['name'] = dim_name
                d['lats'] = list()
                c['dimensions'].append(d)

                lat = dict()
                lat['name'] = lat_name
                d['lats'].append(lat)

                tones[clust_name] = c

            if clust_name in tones:       # clust_names has been added already
                c = tones[clust_name]

                # check if dim_name has arleady been added
                existing_dim = None
                for d in c['dimensions']:
                    if d['name'] == dim_name:
                        existing_dim = d
                        break

                # dim_name has been added
                if existing_dim is not None:

                    # check if lat_name has been added
                    existing_lat = None                        
                    for lat in existing_dim['lats']:
                        if lat['name'] == lat_name:
                            existing_lat = d
                            break

                    if existing_lat is None:    # add lat_name if it hasn't been added yet.
                        lat = dict()
                        lat['name'] = lat_name
                        existing_dim['lats'].append(lat)

                else: 
                    # dim_name hasn't been added yet
                    d = dict()
                    d['name'] = dim_name
                    d['lats'] = list()                        
                    c['dimensions'].append(d)

                    lat = dict()
                    lat['name'] = lat_name
                    d['lats'].append(lat)
        return tones

    def updateResults(self):
        self.search_results.clear()
        for c in self.tones.values():
            citem = QTreeWidgetItem(self.search_results, [c['name']]) 
            citem.setData(0, Qt.UserRole, c['name'])

            for d in c['dimensions']:
                ditem = QTreeWidgetItem(citem, [d['name']])
                ditem.setData(0, Qt.UserRole, "{} {}".format(c['name'], d['name']))

                for lat in d['lats']:
                    latitem = QTreeWidgetItem(ditem, [lat['name']])
                    latitem.setData(0, Qt.UserRole, "{} {} {}".format(c['name'], d['name'], lat['name']))

    @pyqtSlot()
    def LATsFindCompletedAction(self):
        if self.progdialog != None:
            self.progdialog.reset()

        if self.tones is not None:
            self.updateResults()

    def runFindLATs(self, search_patterns):
        all_matched_lats = []
        self.tones = None        
        for p in search_patterns:        
            pattern = p.lower().strip()
            self.matches = self.dictionary.findLATs(pattern, exact_match=False)
            if self.matches:
                matched_lats = [match[0] for match in self.matches]
                matched_lats = list(set(matched_lats))
                all_matched_lats.append(matched_lats)

        lats = get_intersecting_lats(all_matched_lats)
        self.tones = self.makeHierarchy(lats)

        self.LATsFindCompleted.emit()


##################################################
#
# List patterns Dialog
#
##################################################
class ListPatternsDialog(QDialog):

    def __init__(self, title, patterns, editor_win=None, parent=None):
        super(ListPatternsDialog, self).__init__(parent)

        self.editor_win = editor_win

        font_size = views.default_ui_font_size

        width  = 300
        height = 400
        self.setGeometry(60,80,width,height)
        self.setMinimumWidth(width)
        self.setMinimumHeight(height)

        vbox = QVBoxLayout() # main layout

        label = QLabel(title)
        label.setStyleSheet("font-weight: bold; font-size: {}pt".format(font_size))
        vbox.addWidget(label)

        self.patterns = QPlainTextEdit()
        self.patterns.setStyleSheet("background-color: #fff; font-size: {}pt;".format(font_size))
        self.patterns.setReadOnly(True)
        self.patterns.setPlainText('\n'.join(patterns))
        vbox.addWidget(self.patterns)

        self.close_button = QPushButton("Close")
        self.close_button.setAutoDefault(False) 
        self.close_button.clicked.connect(self.close)
        vbox.addWidget(self.close_button)

        self.setLayout(vbox)

        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)
      
        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())   

        self.exec_()

class NewDictDialog(QDialog):
    def __init__(self, editor_win=None, parent=None):

        super(NewDictDialog, self).__init__(parent)

        self.editor_win = editor_win;

        width  = 800
        height = 120

        self.setGeometry(60,80,width,height)
        self.setMinimumWidth(width)
        self.setMinimumHeight(height)

        self.setStyleSheet("background-color: #eee;")

        fbox = QFormLayout()
        fbox.setVerticalSpacing(0)

        hbox = QHBoxLayout()
        self.default_dict_dir_field = QLineEdit()
        self.select_button = QPushButton("Select")
        self.select_button.clicked.connect(self.selectDefaultDictDir)
        hbox.addWidget(self.default_dict_dir_field)
        hbox.addWidget(self.select_button)
        fbox.addRow(QLabel("Parent Folder: "), hbox)
        fbox.addRow(QLabel(""), QLabel(""))
        hbox = QHBoxLayout()

        self.dict_name_field = QLineEdit()
        fbox.addRow(QLabel("Dictionary Name: "), self.dict_name_field)

        fbox.addRow(QLabel(""), QLabel(""))

        hbox = QHBoxLayout()
        self.cancel_button = QPushButton("Cancel")
        self.cancel_button.clicked.connect(self.cancel)
        self.create_button = QPushButton("Create")
        self.create_button.clicked.connect(self.create)
        self.create_button.setDefault(True)
        hbox.addStretch()
        hbox.addWidget(self.cancel_button)
        hbox.addWidget(self.create_button)
        fbox.addRow(QLabel(""), hbox)

        self.setLayout(fbox)      

        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)
      
        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())   

        retval = self.exec_()   

    def cancel(self):
        self.hide()

    def selectDefaultDictDir(self):
        folder = QFileDialog.getExistingDirectory(None, 'Select a parent folder:', '', QFileDialog.ShowDirsOnly)
        if folder == "":
            return
        else:
            self.default_dict_dir_field.setText(folder)
            self.create_button.setDefault(True)

    def create(self):
        folder    = self.default_dict_dir_field.text()
        dict_name = self.dict_name_field.text()

        if folder == "" or dict_name == "":
            WarningDialog("Warning", "The dictionary name or the default dictionary folder is missing.")
        else:
            self.close()
            self.editor_win.createDictionary(os.path.join(folder, dict_name))

class FindLATsDialog(QDialog):

    LATsFindCompleted = pyqtSignal()

    def __init__(self, dictionary, editor_win=None, parent=None):

        super(FindLATsDialog, self).__init__(parent)

        self.editor_win = editor_win
        self.dictionary = dictionary

        width  = 600
        height = 480

        self.setGeometry(60,80,width,height)
        self.setMinimumWidth(width)
        self.setMinimumHeight(height)

        vbox = QVBoxLayout()

        # search field
        self.pattern_field = QLineEdit()
        vbox.addWidget(self.pattern_field)

        hbox = QHBoxLayout()
        self.msg_field = QLabel("")
        hbox.addWidget(self.msg_field)

        hbox.addStretch()
        self.find_button = QPushButton("Find")
        self.find_button.clicked.connect(self.findLATs)
        hbox.addWidget(self.find_button)
        vbox.addLayout(hbox)

        self.search_results = QListWidget()
        vbox.addWidget(self.search_results)

        hbox = QHBoxLayout()
        hbox.addStretch()

        self.cancel_button = QPushButton("Close")
        self.cancel_button.clicked.connect(self.cancel)
        self.cancel_button.setAutoDefault(False) 
        hbox.addWidget(self.cancel_button)   
             
        self.export_button = QPushButton("Export")
        self.export_button.clicked.connect(self.exportResults)
        self.export_button.setAutoDefault(False) 
        hbox.addWidget(self.export_button)        

        if self.editor_win is not None:
            self.openLAT_button = QPushButton("Open LAT")
            self.openLAT_button.clicked.connect(self.openLAT)
            self.openLAT_button.setAutoDefault(False)
            hbox.addWidget(self.openLAT_button)

        vbox.addLayout(hbox)

        self.setLayout(vbox)      

        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)
      
        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())   

        self.LATsFindCompleted.connect(self.LATsFindCompletedAction)

        retval = self.exec_()   


    def cancel(self):
        self.hide()

    def updateFindLATsProgressBar(self):
        pass
        # print("updateFindLATsProgressBar()")

    def openLAT(self):
        item = self.search_results.currentItem()
        if item == None:
            WarningDialog("Select a category:", "No dimension or LAT is selected.")
        else:
            label      = item.text()
            pattern    = label[:label.find("(")].strip()
            pattern    = pattern.replace('\"', '')
            label      = label[label.find("(")+1:-1]
            categories = label.split(' > ')

            if len(categories) >=2:
                self.editor_win.selectACategory(categories)
                self.editor_win.findPattern(pattern, reset=True)

            self.close()

    def findLATs(self):
        self.search_results.clear()
        pattern = self.pattern_field.text()
        self.search_pattern = pattern.lower().strip()

        if self.search_pattern == "":
            WarningDialog("Warning", "Enter a pattern first.")
        else:
            # Create a progress dialog
            self.progdialog = QProgressDialog("Searching... It may take a few minutes.", 
                                              None, 0, 0, self)
            self.progdialog.setModal(True)
            self.progdialog.setMinimumWidth(400)      
            self.progdialog.show()

            threading.Timer(0.2, self.runFindLATs, []).start()
            threading.Timer(0.5, self.updateFindLATsProgressBar, []).start()

    @pyqtSlot()
    def LATsFindCompletedAction(self):
        if self.progdialog != None:
            self.progdialog.reset()
        self.find_button.setAutoDefault(False) 

    def runFindLATs(self):
        self.msg_field.setText("Searching...")
        self.matches = self.dictionary.findLATs(self.search_pattern)

        if self.matches:
            self.msg_field.setText("{} patterns.".format(len(self.matches)))
            for match in self.matches:
                item = "\"{}\" ({})".format(match[1], ' > '.join(match[0]))
                self.search_results.addItem(item)
        else:
            self.msg_field.setText("No matches")

        self.LATsFindCompleted.emit()

    def exportResults(self):
        if len(self.matches)==0:
            WarningDialog("No matches", "There are no matches to export.")
            return

        dialog = QFileDialog(None, 'Enter File Name', './', "Text File (*.csv)")
        dialog.setFileMode(QFileDialog.AnyFile)
        dialog.setAcceptMode(QFileDialog.AcceptSave)
        dialog.selectFile("{}.csv".format(self.search_pattern));
        ret = dialog.exec();

        if (ret == QDialog.Accepted):
            lst = dialog.selectedFiles();

            if type(lst) == list:
                filepath = lst[0]
            else:
                return

            if filepath != "":
                with open(filepath, 'w', encoding="utf-8") as fout:
                    fout.write("Search pattern:, {}\n".format(self.search_pattern))
                    if len(self.matches)==1:
                        fout.write("Resuls:, 1 pattern\n")
                    else:
                        fout.write("Results:, {} patterns\n".format(len(self.matches)))

                    dict_path = self.dictionary.getDirectory()
                    if dict_path == None:
                        dict_path = self.dictionary.getDATPath()
                    fout.write("Dictionary:, {}\n".format(dict_path))

                    fout.write("\n")
                    for match in self.matches:
                        fout.write("\"{}\", {}\n".format(match[1], ' > '.join(match[0])))




class GlobalReplaceLATNamesDialog(QDialog):
    def __init__(self, dim_name):
        super(GlobalReplaceLATNamesDialog, self).__init__()

        self.setStyleSheet("GlobalReplaceLATNamesDialog {background-color: " + views.default_ui_background_color + ";}" + \
                                                  "QLabel {color: " + views.default_ui_text_color + ";}" + \
                                  "QPushButton {background-color: " + views.default_ui_button_color + ";" + \
                                                          "color: " + views.default_ui_text_color + ";}" + \
                                    "QLineEdit {background-color: " + views.default_ui_input_background_color + ";" + \
                                                          "color: " + views.default_ui_text_color + ";}")
        self.setModal(True)
        self.retval = 0

        self.setFixedWidth(500)

        main_vbox = QVBoxLayout()

        # title
        title = QLabel("Global Find & Replace")
        title.setStyleSheet("font-weight: bold;")
        main_vbox.addWidget(title)

        fbox = QFormLayout()
        fbox.setContentsMargins(0,0,0,0)

        self.find_field = QLineEdit()
        fbox.addRow(QLabel("Find: "), self.find_field)

        self.replace_field = QLineEdit()
        fbox.addRow(QLabel("Replace: "), self.replace_field)

        main_vbox.addLayout(fbox)

        msg = QTextEdit()
        html_str  = "<p>Are you sure you want to rename <b>all the LATs</b> under the dimension <b>{}</b>?</p>".format(dim_name)
        html_str += "<p>The changes will be saved automatically and cannot be undone.</p>"
        msg.setHtml(html_str)
        msg.setReadOnly(True)
        msg.setStyleSheet("QTextEdit {border: 0; background-color: " + views.default_ui_background_color + "}")
        msg.setFixedHeight(80)
        msg.setFocusPolicy(Qt.NoFocus)
        main_vbox.addWidget(msg)

        # buttons
        hbox = QHBoxLayout()
        hbox.addStretch()

        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        cancel_button.setAutoDefault(False)        
        hbox.addWidget(cancel_button)

        ok_button = QPushButton("Replace All")
        ok_button.clicked.connect(self.accept)
        ok_button.setAutoDefault(True)
        hbox.addWidget(ok_button)

        main_vbox.addLayout(hbox)

        self.setLayout(main_vbox)

        self.find_field.setFocus()

        self.retval = self.exec_() 

    def reject(self):
        self.done(QMessageBox.Cancel)

    def accept(self):
        self.done(QMessageBox.Apply)

    def getFindAndReplace(self):
        find    = self.find_field.text()
        replace = self.replace_field.text()
        return find, replace


class GlobalReplaceDimNamesDialog(QDialog):
    def __init__(self):
        super(GlobalReplaceDimNamesDialog, self).__init__()

        self.setStyleSheet("GlobalReplaceDimNamesDialog {background-color: " + views.default_ui_background_color + ";}" + \
                                                  "QLabel {color: " + views.default_ui_text_color + ";}" + \
                                  "QPushButton {background-color: " + views.default_ui_button_color + ";" + \
                                                          "color: " + views.default_ui_text_color + ";}" + \
                                    "QLineEdit {background-color: " + views.default_ui_input_background_color + ";" + \
                                                          "color: " + views.default_ui_text_color + ";}")
        self.setModal(True)
        self.retval = 0

        self.setFixedWidth(500)

        main_vbox = QVBoxLayout()

        # title
        title = QLabel("Global Find & Replace")
        title.setStyleSheet("font-weight: bold;")
        main_vbox.addWidget(title)

        fbox = QFormLayout()
        fbox.setContentsMargins(0,0,0,0)

        self.find_field = QLineEdit()
        fbox.addRow(QLabel("Find: "), self.find_field)

        self.replace_field = QLineEdit()
        fbox.addRow(QLabel("Replace: "), self.replace_field)

        hbox = QHBoxLayout()
        self.name_button = QCheckBox("Name")
        self.name_button.setChecked(True)
        hbox.addWidget(self.name_button) 
        hbox.addSpacing(10)   

        self.label_button = QCheckBox("Label")
        self.label_button.setChecked(True)
        hbox.addWidget(self.label_button)  
        hbox.addStretch()
        fbox.addRow(QLabel(""), hbox)

        main_vbox.addLayout(fbox)

        msg = QTextEdit()
        html_str  = "<p>Are you sure you want to <b>rename all the selected dimensions</b></p>"
        html_str += "<p>The changes will be saved automatically and cannot be undone.</p>"
        msg.setHtml(html_str)
        msg.setReadOnly(True)
        msg.setStyleSheet("QTextEdit {border: 0; background-color: " + views.default_ui_background_color + "}")
        msg.setFixedHeight(80)
        msg.setFocusPolicy(Qt.NoFocus)
        main_vbox.addWidget(msg)

        # buttons
        hbox = QHBoxLayout()
        hbox.addStretch()

        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        cancel_button.setAutoDefault(False)        
        hbox.addWidget(cancel_button)

        ok_button = QPushButton("Replace All")
        ok_button.clicked.connect(self.accept)
        ok_button.setAutoDefault(True)
        hbox.addWidget(ok_button)

        main_vbox.addLayout(hbox)

        self.setLayout(main_vbox)

        self.find_field.setFocus()

        self.retval = self.exec_() 

    def reject(self):
        self.done(QMessageBox.Cancel)

    def accept(self):
        self.done(QMessageBox.Apply)


    def getFindAndReplace(self):
        find    = self.find_field.text()
        replace = self.replace_field.text()

        b_name  = self.name_button.isChecked()
        b_label = self.label_button.isChecked()

        return find, replace, b_name, b_label





