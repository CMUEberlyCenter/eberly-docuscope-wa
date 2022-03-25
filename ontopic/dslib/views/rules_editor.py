#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
    rules_editor.py

"""

import os, sys
import platform

from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *

import string
import re
import copy
import shutil
import traceback
import json
from datetime import datetime, date

import shlex, subprocess
import time
import threading

import bs4
from bs4 import BeautifulSoup as bs  

from operator import itemgetter

from docx import Document                        # python_docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt

import dslib.views.dialogs      as dialogs
import dslib.models.rule        as rule
import dslib.utils              as utils
import dslib.views              as views
import dslib.models.dict        as ds_dict

import pprint     
pp = pprint.PrettyPrinter(indent=4)

UP   = -1
DOWN = 1

ACTIVE           = 1
QUIET            = 2
OPTIONAL         = 3
BOUNDED_OPTIONAL = 4
GROUP            = 5

TOPIC_CLUSTER    = 1
LITERAL_TOPIC    = 2
IMPRESSION_CLUSTER = 3
DOCUSCOPE_CATEGORY = 4

IMPRESSION_NEUTRAL  = 0
IMPRESSION_COMMON   = 1
IMPRESSION_UNCOMMON = 2

topic_blue_brush   = QBrush(QColor("#023da1"))
exper_yellow_brush = QBrush(QColor("#ffc505"))
transparent_brush  = QBrush(Qt.transparent)
text_black_brush   = QBrush(Qt.black)

def blueBoldSelection(editor):
    cursor = editor.textCursor()        
    fmt = cursor.charFormat()
    font = fmt.font()

    if font.weight() == QFont.Bold:
        fmt.setFontWeight(QFont.Normal)
        fmt.setForeground(text_black_brush)
    else:
        fmt.setFontWeight(QFont.Bold)
        fmt.setForeground(topic_blue_brush)

    cursor.setCharFormat(fmt)    

def highlightSelection(editor):
    cursor = editor.textCursor()
    fmt = cursor.charFormat()
    font = fmt.font()

    if fmt.background() == exper_yellow_brush:
        fmt.setBackground(transparent_brush)
    else:
        fmt.setBackground(exper_yellow_brush)

    cursor.setCharFormat(fmt)

def boldSelection(editor):
    cursor = editor.textCursor()        
    curr_fmt = cursor.charFormat()
    font = curr_fmt.font()

    fmt = QTextCharFormat()
    if font.weight() == QFont.Bold:
        fmt.setFontWeight(QFont.Normal)
    else:
        fmt.setFontWeight(QFont.Bold)

    if font.italic():
        fmt.setFontItalic(True)

    cursor.setCharFormat(fmt)    

def italicizeSelection(editor):
    cursor = editor.textCursor()
    curr_fmt = cursor.charFormat()
    font = curr_fmt.font()

    fmt = QTextCharFormat()
    if font.italic():
        fmt.setFontItalic(False)
    else:
        fmt.setFontItalic(True)

    if font.weight() == QFont.Bold:
        fmt.setFontWeight(QFont.Bold)

    cursor.setCharFormat(fmt)

def bulletSelection(editor):
    cursor = editor.textCursor()

    curr_list = cursor.currentList()
    if curr_list is not None:
        block = cursor.block()
        curr_list.remove(block)
        block_fmt = cursor.blockFormat()
        block_fmt.setIndent(0)
        cursor.setBlockFormat(block_fmt)
    else:
        cursor.createList(QTextListFormat.ListDisc)

class RulesEditor(QFrame):
    # def __init__(self, custom_dict, default_dict, parent=None)
    def __init__(self, app_win=None, controller=None, parent=None):

        super(RulesEditor, self).__init__(parent)

        self.directory = None
        self.app_win = app_win
        self.controller = controller

        self.rules_not_saved = False
        self.rules_changed   = False

        self.just_selected_rule = None
        self.current_rule = None
        self.b_revised = False
        self.b_rs_revised = False
        self.untitled_count = 0

        self.current_cv_panel = None
        self.cv_panel_descr_changed = False
        self.ignore_cv_panel_updates = True
        self.cluster_selectors = None
        self.ruleset = None

        # Center the dialog window
        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)
      
        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft()) 

        self.setCursor(Qt.ArrowCursor)

        self.showMaximized()
        self.show()

    def setContentPath(self, path):
        self.path.setText(path)        

    def loadContent(self, directory):

        self.directory = directory

        self.ruleset = rule.DSRuleSet(directory)

        self.topic_clusters = self.ruleset.getTopicClusters()

        self.initUI()
        self.setContentPath(self.directory)

        ruleset_name = self.ruleset.getName()
        self.rs_name.setText(ruleset_name)

        if self.ruleset.getRules():
            self.untitled_count = self.rs_list_view.setRules(self.ruleset)
            self.rs_list_view.setFocus()

            self.current_rule = self.ruleset.getRule(0)
            self.just_selected_rule = self.current_rule
            self.updateRuleFields(self.current_rule)


        self.setRuleEdited(False)
        self.setRuleSetEdited(False)

        if ruleset_name == '' or ruleset_name is None:
            self.app_win.setWindowTitle('Untitled')
        else:
            self.app_win.setWindowTitle(ruleset_name)


    # ----------        
    # 'initUI' initializes all the UI elements for the app.
    # ----------        
    def initUI(self):               

        button_ht = 26

        #
        # Main Window
        #

        self.wcProxyModel = None

        main_vbox = QVBoxLayout() # main vertical layout
        main_vbox.setContentsMargins(6,6,6,0)
        # The header aera.              
        header_hbox = QHBoxLayout()
        header_hbox.setContentsMargins(0,0,0,0)

        rs_vbox = QVBoxLayout()
        self.rs_name = QLineEdit()
        self.rs_name.setMaximumWidth(200)
        header_hbox.addWidget(QLabel("Genre/Assignment:"))
        self.rs_name.textEdited.connect(self.ruleSetNameEdited)
        header_hbox.addWidget(self.rs_name)

        header_hbox.addSpacing(10)

        self.path = QLineEdit()
        self.path.setReadOnly(True)
        self.path.setStyleSheet("background-color: " + views.default_ui_read_only_color + ";")
        header_hbox.addWidget(QLabel("Content Folder:"))
        header_hbox.addWidget(self.path)

        main_vbox.addLayout(header_hbox)  # add the rule set info area to the man layout

        ##### End of the header area #####

        bold_icon_path       = utils.resource_path('data/icons/bold_icon.png')
        italic_icon_path     = utils.resource_path('data/icons/italic_icon.png')
        blue_bold_icon_path  = utils.resource_path('data/icons/blue_bold_icon.png')
        yellow_bg_icon_path  = utils.resource_path('data/icons/yellow_bg_icon.png')
        bullet_icon_path     = utils.resource_path('data/icons/bullet_icon.png')
        clear_icon_path      = utils.resource_path('data/icons/clear_icon.png')

        # Draw a horizontal line
        line = QFrame()
        line.setFrameShape(QFrame.HLine);
        line.setFrameShadow(QFrame.Sunken);
        main_vbox.addWidget(line)

        ##############################
        #
        # Left 
        #
        ##############################

        left_vbox = QVBoxLayout()
        left_vbox.setContentsMargins(6,0,6,6)
        left_container = QFrame()
        left_container.setLayout(left_vbox)
        left_container.setSizePolicy(QSizePolicy.MinimumExpanding, QSizePolicy.Preferred)

        header = QLabel("Composing Patterns")
        header.setStyleSheet("QLabel {font-weight: bold;}")
        left_vbox.addWidget(header)

        # Control UI (Add/Delet Buttons)
        ctrl_hbox = QHBoxLayout()
        self.add_group_button = QPushButton()
        self.add_group_button.setIcon(QIcon(QPixmap(utils.resource_path("data/icons/add_folder_icon.png"))))
        self.add_group_button.clicked.connect(self.addNewGroup)
        ctrl_hbox.addWidget(self.add_group_button)              

        self.new_button = QPushButton()
        self.new_button.setIcon(QIcon(QPixmap(utils.resource_path("data/icons/new_icon.png"))))
        self.new_button.clicked.connect(self.addNewRule)
        ctrl_hbox.addWidget(self.new_button)              

        self.delete_button = QPushButton()
        self.delete_button.setIcon(QIcon(QPixmap(utils.resource_path("data/icons/delete_icon.png"))))
        self.delete_button.clicked.connect(self.deleteSelectedRule)
        ctrl_hbox.addWidget(self.delete_button)

        self.move_up_button = QPushButton()
        self.move_up_button.setIcon(QIcon(QPixmap(utils.resource_path("data/icons/up_icon.png"))))
        self.move_up_button.clicked.connect(self.moveSelectedRuleUp)
        ctrl_hbox.addWidget(self.move_up_button)

        self.move_down_button = QPushButton()
        self.move_down_button.setIcon(QIcon(QPixmap(utils.resource_path("data/icons/down_icon.png"))))
        self.move_down_button.clicked.connect(self.moveSelectedRuleDown)
        ctrl_hbox.addWidget(self.move_down_button)

        ctrl_hbox.addStretch()

        left_vbox.addLayout(ctrl_hbox)

        #
        # List of Rules
        #
        self.rs_list_view = RuleTreeWidget(editor=self, ruleset=self.ruleset)
        self.rs_list_view.setStyleSheet("RuleTreeWidget {background-color: #fff;" + \
                                        "selection-color: #fff;" + \
                                        "selection-background-color: #0080ff;};")
        left_vbox.addWidget(self.rs_list_view)

        ###
        left_vbox.addSpacing(5)
        header = QLabel("<b>Comm. Values Associated with Tools</b> (Tabs)")
        left_vbox.addWidget(header)

        cv_title_fbox = QFormLayout()
        self.panel_menu = QComboBox()
        self.panel_menu.addItems(['Expectations', 'Coherence', 'Clarity', 'Impressions'])
        self.panel_menu.currentIndexChanged.connect(self.cvPanelSelected)
        cv_title_fbox.addRow(QLabel("Tab: "), self.panel_menu)

        self.cv_panel_title = QLabel("")
        self.cv_panel_title.setStyleSheet("font-weight: bold;")
        cv_title_fbox.addRow(QLabel("Title: "), self.cv_panel_title)
        left_vbox.addLayout(cv_title_fbox)

        hbox = QHBoxLayout()  # formatting gicons

        self.panel_bold_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(bold_icon_path), QIcon.Normal, QIcon.On)
        self.panel_bold_button.setIcon(icon)
        self.panel_bold_button.setAutoDefault(False)
        self.panel_bold_button.setFixedSize(19,20)
        self.panel_bold_button.clicked.connect(self.boldPanelDescription)

        self.panel_italic_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(italic_icon_path), QIcon.Normal, QIcon.On)
        self.panel_italic_button.setIcon(icon)
        self.panel_italic_button.setAutoDefault(False)
        self.panel_italic_button.setFixedSize(19,20)
        self.panel_italic_button.clicked.connect(self.italicizePanelDescription)

        self.panel_bullet_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(bullet_icon_path), QIcon.Normal, QIcon.On)
        self.panel_bullet_button.setIcon(icon)
        self.panel_bullet_button.setAutoDefault(False)        
        self.panel_bullet_button.setFixedSize(19,20)
        self.panel_bullet_button.clicked.connect(self.bulletPanelDescription)

        self.panel_clear_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(clear_icon_path), QIcon.Normal, QIcon.On)
        self.panel_clear_button.setIcon(icon)
        self.panel_clear_button.setAutoDefault(False)        
        self.panel_clear_button.setFixedSize(19,20)
        self.panel_clear_button.clicked.connect(self.clearPanelDescrFormat)

        hbox.addWidget(self.panel_bold_button)
        hbox.addWidget(self.panel_italic_button)
        hbox.addWidget(self.panel_bullet_button)
        hbox.addStretch()
        hbox.addWidget(self.panel_clear_button)

        left_vbox.addLayout(hbox)

        self.cv_panel_description = QTextEdit()
        self.cv_panel_description.textChanged.connect(self.cvPanelDescriptionChanged)
        left_vbox.addWidget(self.cv_panel_description)

        d = self.cv_panel_description.document()
        d.setDocumentMargin(8)
        s =  "p   {font-size: " + str(views.default_ui_font_size) + "pt;"
        s += "     margin-bottom: 0.4em;}"        
        s += "li  {font-size: " + str(views.default_ui_font_size) + "pt;"
        s += "     margin-bottom: 0.4em;}"

        d.setDefaultStyleSheet(s)

        html_dscr = self.ruleset.getPanelDescriptions(views.EXPECTATIONS)
        self.ignore_cv_panel_updates = True
        self.cv_panel_description.setHtml(html_dscr)
        self.ignore_cv_panel_updates = False
        self.current_cv_panel = views.EXPECTATIONS
        self.cv_panel_title.setText(views.getSectionHeading(views.EXPECTATIONS))

        ##############################
        # Definitions
        ##############################

        center_vbox = QVBoxLayout()
        center_vbox.setContentsMargins(6,0,6,6)        
        center_container = QFrame()
        center_container.setLayout(center_vbox)
        center_container.setSizePolicy(QSizePolicy.MinimumExpanding, QSizePolicy.Preferred)

        # Header Area
        hbox_ui = QHBoxLayout()
        header = QLabel("Definitions")
        header.setStyleSheet("QLabel {font-weight: bold;}")
        hbox_ui.addWidget(header)

        hbox_ui.addStretch()

        center_vbox.addLayout(hbox_ui)

        ##### End of the Header Area

        fbox = QFormLayout()
        fbox.setContentsMargins(0,0,0,0)
        center_vbox.addLayout(fbox)

        # Name
        hbox = QHBoxLayout()
        hbox.setContentsMargins(0,2,0,5)        
        self.rule_name = QLineEdit()
        self.rule_name.setMinimumWidth(50)          
        self.rule_name_header = QLabel("Question:")
        hbox.addWidget(self.rule_name)
        fbox.addRow(self.rule_name_header, hbox)
        self.rule_name.textEdited.connect(self.ruleEdited)

        # TOPIC 1
        topic_vbox = QVBoxLayout()
        topic_vbox.setContentsMargins(0,5,0,5)

        topic_hbox = QHBoxLayout()
        self.topic_button_group = QButtonGroup()
        self.topic_button_group.idClicked.connect(self.topicClusterToggled)
        self.topic_cluster_radio_button = QRadioButton("Topic Cluster")
        self.topic_cluster_radio_button.setChecked(True)
        topic_hbox.addWidget(self.topic_cluster_radio_button)
        self.topic_button_group.addButton(self.topic_cluster_radio_button, TOPIC_CLUSTER)

        self.literal_topic_radio_button = QRadioButton("Literal Topic")
        self.literal_topic_radio_button.setChecked(False)
        topic_hbox.addWidget(self.literal_topic_radio_button)
        self.topic_button_group.addButton(self.literal_topic_radio_button, LITERAL_TOPIC)

        topic_vbox.addLayout(topic_hbox)

        #
        # Topic Settings
        #
        topic_settings_vbox = QVBoxLayout()
        topic_settings_vbox.setContentsMargins(0,3,0,5)

        self.literal_topic_field = QLineEdit()
        topic_settings_vbox.addWidget(self.literal_topic_field)
        self.literal_topic_field.hide()

        tc_vbox = QVBoxLayout()
        tc_vbox.setContentsMargins(0,3,0,5)        
        # tc_vbox.setSpacing(0)        
        self.tc_settings_container = QWidget()
        self.tc_settings_container.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self.tc_settings_container.setLayout(tc_vbox)
        topic_settings_vbox.addWidget(self.tc_settings_container)

        hbox = QHBoxLayout()
        options = ['']
        options.extend(list(self.topic_clusters.keys()))

        self.topic1_menu = QComboBox()
        self.topic1_menu.addItems(options)
        self.topic1_menu.setStyleSheet("selection-background-color: #0080ee;")
        self.topic1_menu.currentIndexChanged.connect(self.topicClusterChanged)
        hbox.addWidget(self.topic1_menu)

        self.topic1_add_button = QPushButton()
        self.topic1_add_button.setFixedSize(20,20)
        self.topic1_add_button.setToolTip("Add a new topic cluster")
        pic = QPixmap(utils.resource_path('data/icons/new_icon.png'))
        pic = pic.scaled(11, 11, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.topic1_add_button.setIcon(QIcon(pic))
        self.topic1_add_button.clicked.connect(self.addTopicCluster)
        self.topic1_add_button.setAutoDefault(False)   
        hbox.addWidget(self.topic1_add_button)

        self.topic1_edit_button = QPushButton()
        self.topic1_edit_button.setFixedSize(20,20)
        self.topic1_edit_button.setToolTip("Edit the selected topic cluster")
        pic = QPixmap(utils.resource_path('data/icons/edit_icon.png'))
        pic = pic.scaled(11, 11, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.topic1_edit_button.setIcon(QIcon(pic))
        self.topic1_edit_button.clicked.connect(self.editTopicCluster)
        self.topic1_edit_button.setAutoDefault(False)   
        hbox.addWidget(self.topic1_edit_button)

        self.topic1_delete_button = QPushButton()
        self.topic1_delete_button.setFixedSize(20,20)
        self.topic1_delete_button.setToolTip("Delete the selected topic cluster")
        pic = QPixmap(utils.resource_path('data/icons/delete_icon.png'))
        pic = pic.scaled(11, 11, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.topic1_delete_button.setIcon(QIcon(pic))
        self.topic1_delete_button.clicked.connect(self.deleteTopicCluster)
        self.topic1_delete_button.setAutoDefault(False)   
        hbox.addWidget(self.topic1_delete_button)

        tc_vbox.addLayout(hbox)

        self.pre_defined_topics_label = QLabel("Pre-defined Topics")
        tc_vbox.addWidget(self.pre_defined_topics_label)
        self.pre_defined_topics_label.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)

        self.pre_defined_topics_field = QPlainTextEdit()
        self.pre_defined_topics_field.setMinimumWidth(50)        
        self.pre_defined_topics_field.setMaximumHeight(200)
        self.pre_defined_topics_field.setEnabled(True)
        self.pre_defined_topics_field.setReadOnly(True)
        self.pre_defined_topics_field.viewport().setCursor(Qt.ArrowCursor)
        self.pre_defined_topics_field.setStyleSheet("background-color: " + views.default_ui_read_only_color + ";")
        self.pre_defined_topics_field.textChanged.connect(self.ruleEdited)
        self.pre_defined_topics_field.setToolTip("Click the edit button above to revise the prompt.")
        tc_vbox.addWidget(self.pre_defined_topics_field)

        topic_vbox.addLayout(topic_settings_vbox)

        fbox.addRow(QLabel("Topic:"), topic_vbox)

        #
        # Descriptions
        #
        dscr_vbox = QVBoxLayout()
        dscr_vbox.setContentsMargins(0,5,0,5)
        hbox = QHBoxLayout()

        self.rule_bold_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(bold_icon_path), QIcon.Normal, QIcon.On)
        self.rule_bold_button.setIcon(icon)
        self.rule_bold_button.setAutoDefault(False)
        self.rule_bold_button.setFixedSize(19,20)
        self.rule_bold_button.clicked.connect(self.boldRuleDescription)

        self.rule_italic_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(italic_icon_path), QIcon.Normal, QIcon.On)
        self.rule_italic_button.setIcon(icon)
        self.rule_italic_button.setAutoDefault(False)
        self.rule_italic_button.setFixedSize(19,20)
        self.rule_italic_button.clicked.connect(self.italicizeRuleDescription)

        self.rule_bullet_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(bullet_icon_path), QIcon.Normal, QIcon.On)
        self.rule_bullet_button.setIcon(icon)
        self.rule_bullet_button.setAutoDefault(False)        
        self.rule_bullet_button.setFixedSize(19,20)
        self.rule_bullet_button.clicked.connect(self.bulletRuleDescription)

        self.rule_clear_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(clear_icon_path), QIcon.Normal, QIcon.On)
        self.rule_clear_button.setIcon(icon)
        self.rule_clear_button.setAutoDefault(False)        
        self.rule_clear_button.setFixedSize(19,20)
        self.rule_clear_button.clicked.connect(self.clearRuleDescrFormat)

        hbox.addWidget(self.rule_bold_button)
        hbox.addWidget(self.rule_italic_button)
        hbox.addWidget(self.rule_bullet_button)
        hbox.addStretch()
        hbox.addWidget(self.rule_clear_button)
        dscr_vbox.addLayout(hbox)

        # Description
        self.rule_description = QTextEdit()
        self.rule_description.textChanged.connect(self.ruleEdited)
        self.rule_description.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.MinimumExpanding)        
        self.rule_description.setMinimumWidth(50)
        dscr_vbox.addWidget(self.rule_description)
        fbox.addRow(QLabel("Description:    "), dscr_vbox)

        ##############################
        #
        # Right Side Panel
        #
        ##############################

        right_vbox = QVBoxLayout()
        right_vbox.setContentsMargins(6,0,6,6)        
        right_container = QFrame()
        right_container.setLayout(right_vbox)
        right_container.setSizePolicy(QSizePolicy.MinimumExpanding, QSizePolicy.Preferred)

        fbox = QFormLayout()

        sents_vbox = QVBoxLayout()
        sents_vbox.addSpacing(24)

        hbox = QHBoxLayout()  # formatting gicons

        self.sent_blue_bold_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(blue_bold_icon_path), QIcon.Normal, QIcon.On)
        self.sent_blue_bold_button.setIcon(icon)
        self.sent_blue_bold_button.setAutoDefault(False)
        self.sent_blue_bold_button.setFixedSize(19,20)
        self.sent_blue_bold_button.clicked.connect(self.blueBoldExamples)

        self.sent_clear_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(clear_icon_path), QIcon.Normal, QIcon.On)
        self.sent_clear_button.setIcon(icon)
        self.sent_clear_button.setAutoDefault(False)        
        self.sent_clear_button.setFixedSize(19,20)
        self.sent_clear_button.clicked.connect(self.clearExamplesFormat)

        hbox.addWidget(self.sent_blue_bold_button)
        hbox.addStretch()
        hbox.addWidget(self.sent_clear_button)

        sents_vbox.addLayout(hbox)

        self.rule_examples = QTextEdit()
        self.rule_examples.setMinimumWidth(50)        
        self.rule_examples.textChanged.connect(self.ruleEdited)
        sents_vbox.addWidget(self.rule_examples)

        d = self.rule_examples.document()
        d.setDocumentMargin(8)
        style  = "span.topic {font-weight: bold; color: #023da1;} "
        style += "span.experience {background-color: #ffc505;}"
        d.setDefaultStyleSheet(style)

        label = QLabel("Sample\nsentences: ")
        label.setStyleSheet("margin-top: 11px")
        fbox.addRow(label, sents_vbox)


        #
        # Communication Values
        #
        values_vbox = QVBoxLayout()
        values_vbox.setContentsMargins(0,3,0,0)
        self.cv_compelling_checkbox = QCheckBox("Compelling")
        self.cv_compelling_checkbox.setChecked(False)
        self.cv_credible_checkbox = QCheckBox("Credible")
        self.cv_credible_checkbox.setChecked(False)
        self.cv_considerate_checkbox = QCheckBox("Considerate")
        self.cv_considerate_checkbox.setChecked(False)
        self.cv_ethical_checkbox = QCheckBox("Ethical")
        self.cv_ethical_checkbox.setChecked(False)
        hbox = QHBoxLayout()
        hbox.addWidget(self.cv_compelling_checkbox)
        hbox.addWidget(self.cv_credible_checkbox)
        hbox.addWidget(self.cv_considerate_checkbox)
        hbox.addWidget(self.cv_ethical_checkbox)
        hbox.addStretch()
        values_vbox.addLayout(hbox)

        # CV Descriptions
        hbox = QHBoxLayout()
        self.cv_bold_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(bold_icon_path), QIcon.Normal, QIcon.On)
        self.cv_bold_button.setIcon(icon)
        self.cv_bold_button.setAutoDefault(False)
        self.cv_bold_button.setFixedSize(19,20)
        self.cv_bold_button.clicked.connect(self.boldCVDescription)

        self.cv_italic_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(italic_icon_path), QIcon.Normal, QIcon.On)
        self.cv_italic_button.setIcon(icon)
        self.cv_italic_button.setAutoDefault(False)
        self.cv_italic_button.setFixedSize(19,20)
        self.cv_italic_button.clicked.connect(self.italicizeCVDescription)

        self.cv_bullet_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(bullet_icon_path), QIcon.Normal, QIcon.On)
        self.cv_bullet_button.setIcon(icon)
        self.cv_bullet_button.setAutoDefault(False)        
        self.cv_bullet_button.setFixedSize(19,20)
        self.cv_bullet_button.clicked.connect(self.bulletCVDescription)

        self.cv_clear_button = QPushButton()
        icon = QIcon()
        icon.addPixmap(QPixmap(clear_icon_path), QIcon.Normal, QIcon.On)
        self.cv_clear_button.setIcon(icon)
        self.cv_clear_button.setAutoDefault(False)        
        self.cv_clear_button.setFixedSize(19,20)
        self.cv_clear_button.clicked.connect(self.clearCVDescrFormat)

        hbox.addWidget(self.cv_bold_button)
        hbox.addWidget(self.cv_italic_button)
        hbox.addWidget(self.cv_bullet_button)
        hbox.addStretch()
        hbox.addWidget(self.cv_clear_button)
        values_vbox.addLayout(hbox)

        self.values_description = QTextEdit()
        self.values_description.setMinimumWidth(50)
        self.values_description.textChanged.connect(self.ruleEdited)
        self.values_description.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.MinimumExpanding)
        values_vbox.addWidget(self.values_description)

        fbox.addRow(QLabel("Values:"), values_vbox)

        right_vbox.addLayout(fbox)

        ##############################
        #
        # Far Right Panel
        #
        ##############################

        far_right_vbox = QVBoxLayout()
        far_right_container = QFrame()
        far_right_container.setLayout(far_right_vbox)
        far_right_container.setSizePolicy(QSizePolicy.MinimumExpanding, QSizePolicy.Preferred)

        header = QLabel("Impressions")
        header.setStyleSheet("QLabel {font-weight: bold;}")
        far_right_vbox.addWidget(header)

        hbox = QHBoxLayout()
        label = QLabel("Dictionary")
        hbox.addWidget(label)
        hbox.addStretch()
        dict_button = QPushButton("Select")
        dict_button.clicked.connect(self.selectDictionary)
        hbox.addWidget(dict_button)
        far_right_vbox.addLayout(hbox)

        self.dict_path_field = QPlainTextEdit()
        self.dict_path_field.setMaximumHeight(120)
        far_right_vbox.addWidget(self.dict_path_field)

        far_right_vbox.addSpacing(5)        

        clust_header = QLabel("Clusters")
        clust_header.setStyleSheet("QLabel {font-weight: bold;}")
        far_right_vbox.addWidget(clust_header)

        self.clust_container = QFrame()

        self.clust_layout    = QVBoxLayout()
        self.clust_container.setLayout(self.clust_layout)

        self.scroll_area  = QScrollArea()        
        self.scroll_area.setWidgetResizable(True)   
        self.scroll_area.setWidget(self.clust_container)

        far_right_vbox.addWidget(self.scroll_area)

        self.main_hsplit_panel = QSplitter(Qt.Horizontal)        
        self.main_hsplit_panel.addWidget(left_container)
        self.main_hsplit_panel.addWidget(center_container) 
        self.main_hsplit_panel.addWidget(right_container)
        self.main_hsplit_panel.addWidget(far_right_container)
        self.main_hsplit_panel.setStyleSheet("QSplitter::handle {background-color: " + views.splitter_handle_color + ";}")

        self.main_hsplit_panel.setStretchFactor(0, 2)
        self.main_hsplit_panel.setStretchFactor(1, 4)
        self.main_hsplit_panel.setStretchFactor(2, 1)
        self.main_hsplit_panel.setStretchFactor(3, 2)
        # self.main_hsplit_panel.setCollapsible(0, False)
        # self.main_hsplit_panel.setCollapsible(1, False)
        # self.main_hsplit_panel.setCollapsible(2, False)
        # self.main_hsplit_panel.setCollapsible(3, False)

        main_vbox.addWidget(self.main_hsplit_panel)
        self.setLayout(main_vbox)

        # Center the window
        frameGm = self.frameGeometry()
        screen = QApplication.desktop().screenNumber(QApplication.desktop().cursor().pos())
        screen_rect = QApplication.desktop().screenGeometry(screen)

        centerPoint = screen_rect.center()   
        frameGm.moveCenter(centerPoint)
        self.move(frameGm.topLeft())

    #################################################################
    # Rule Description
    #################################################################

    def boldCVDescription(self):
        boldSelection(self.values_description)
        
    def italicizeCVDescription(self):
        italicizeSelection(self.values_description)

    def bulletCVDescription(self):
        bulletSelection(self.values_description)

    def clearCVDescrFormat(self):

        html = self.values_description.toHtml()
        soup = bs(html, "html.parser")

        new_html_str = ""
        for p in soup.find_all('p'):
            new_html_str += "<p>" + p.text + "</p>\n"

        style = "p {font-family: sans-serif; font-style: normal; font-size: " + str(views.editor_normal_font_size) + "pt;}" + \
                "span {font-family: sans-serif; font-style: normal; font-size: " + str(views.editor_normal_font_size) + "pt;}"

        doc = self.values_description.document()
        doc.setDefaultStyleSheet(style)

        self.values_description.setHtml(new_html_str)


    #################################################################
    # Communication Values (CV) Description
    #################################################################

    def boldRuleDescription(self):
        boldSelection(self.rule_description)
        
    def italicizeRuleDescription(self):
        italicizeSelection(self.rule_description)

    def bulletRuleDescription(self):
        bulletSelection(self.rule_description)

    def clearRuleDescrFormat(self):

        html = self.rule_description.toHtml()
        soup = bs(html, "html.parser")

        new_html_str = ""
        for p in soup.find_all('p'):
            new_html_str += "<p>" + p.text + "</p>\n"

        style = "p {font-family: sans-serif; font-style: normal; font-size: " + str(views.editor_normal_font_size) + ";}" + \
                "span {font-family: sans-serif; font-style: normal; font-size: " + str(views.editor_normal_font_size) + ";}"

        doc = self.rule_description.document()
        doc.setDefaultStyleSheet(style)

        self.rule_description.setHtml(new_html_str)

    #################################################################
    # Examples
    #################################################################

    def blueBoldExamples(self):
        blueBoldSelection(self.rule_examples)
        
    def highlightExamples(self):
        highlightSelection(self.rule_examples)

    # def bulletExamples(self):
        # bulletSelection(self.rule_examples)

    def clearExamplesFormat(self):

        html = self.rule_examples.toHtml()
        soup = bs(html, "html.parser")

        new_html_str = ""
        for p in soup.find_all('p'):
            new_html_str += "<p>" + p.text + "</p>\n"

#        style = "p {font-family: sans-serif; font-style: normal; font-size: " + str(normal_font_size) + ";}" + \
#                "span {font-family: sans-serif; font-style: normal; font-size: " + str(normal_font_size) + ";}"

        doc = self.rule_examples.document()
        # doc.setDefaultStyleSheet(style)

        self.rule_examples.setHtml(new_html_str) 


    #################################################################
    # Comm Values Panel Description
    #################################################################

    def boldPanelDescription(self):
        boldSelection(self.cv_panel_description)
        
    def italicizePanelDescription(self):
        italicizeSelection(self.cv_panel_description)

    def bulletPanelDescription(self):
        bulletSelection(self.cv_panel_description)

    def clearPanelDescrFormat(self):

        html = self.cv_panel_description.toHtml()
        soup = bs(html, "html.parser")

        new_html_str = ""
        for p in soup.find_all('p'):
            new_html_str += "<p>" + p.text + "</p>\n"

        style = "p {font-family: sans-serif; font-style: normal; font-size: " + str(views.editor_normal_font_size) + ";}" + \
                "span {font-family: sans-serif; font-style: normal; font-size: " + str(views.editor_normal_font_size) + ";}"

        doc = self.cv_panel_description.document()
        doc.setDefaultStyleSheet(style)

        self.cv_panel_description.setHtml(new_html_str)

    def cvPanelDescriptionChanged(self):
        if self.ignore_cv_panel_updates == False:
            self.cv_panel_descr_changed = True
            self.ruleSetEdited()

    def cvPanelSelected(self, val):

        selected_cv_panel = views.panelID2String(val)

        if self.cv_panel_descr_changed:
            # Save the text if the content has been edited. 
            # This will convert inch mark pairs to curly quotes.
            prev_html_str = rule.simplifyHtml(self.cv_panel_description.toHtml(), 'description')
            self.ruleset.setValuesDescription(self.current_cv_panel, prev_html_str)

        self.ignore_cv_panel_updates = True
        # update the html text in the text box
        self.cv_panel_description.setHtml(self.ruleset.getValuesDescription(selected_cv_panel))
        self.cv_panel_title.setText(views.getSectionHeading(selected_cv_panel))
        self.current_cv_panel = selected_cv_panel
        self.cv_panel_descr_changed = False
        self.ignore_cv_panel_updates = False

    #################################################################

    #################################################################
    def noLexOverlapCheckboxChanged(self):
        self.ruleEdited()

    def userDefinedTopicChanged(self, arg):

        if arg == 1:
            if self.topic_cluster_radio_button.isChecked():
                self.topic1_menu.setEnabled(True)
                self.topic1_add_button.setEnabled(True)
                self.topic1_delete_button.setEnabled(True)
                self.topic1_edit_button.setEnabled(True)

                self.pre_defined_topics_field.setEnabled(True)
                self.pre_defined_topics_field.show()          
                self.pre_defined_topics_label.show()
            else:
                self.topic1_menu.setEnabled(False)
                self.topic1_add_button.setEnabled(False)
                self.topic1_delete_button.setEnabled(False)
                self.topic1_edit_button.setEnabled(False)

                self.pre_defined_topics_field.clear()
                self.pre_defined_topics_field.setEnabled(False)
                self.pre_defined_topics_field.hide()             
                self.pre_defined_topics_label.hide()

        self.ruleEdited()

    # def rsChanged(self, bModified):
        # if bModified:
            # self.update_button.setEnabled(True)
        # else:
            # self.update_button.setEnabled(False)

    def cleanUpRuleFields(self):
        rule_revised = self.isRuleEdited()
        rs_revised   = self.isRuleSetEdited()

        self.rule_name.setText("")
        self.rule_description.setHtml("")
        self.rule_examples.setHtml("")
        self.values_description.setHtml("")

        self.cv_credible_checkbox.setChecked(False)
        self.cv_ethical_checkbox.setChecked(False)
        self.cv_considerate_checkbox.setChecked(False)
        self.cv_compelling_checkbox.setChecked(False)

        self.topic1_menu.setCurrentIndex(0)
        self.pre_defined_topics_field.setPlainText('')

        self.setRuleEdited(rule_revised)
        self.setRuleSetEdited(rs_revised)

    def editTopicCluster(self):
        selected_lemma  = self.topic1_menu.currentText()
        if selected_lemma == "-":
            dialogs.WarningDialog("Select a Topic Cluster", "Select a Topic Cluster from the menu, or press the add button to create a new Topic Cluster.")
            return

        d = TopicClusterDialog(lemma=selected_lemma, pre_defined_topics=self.topic_clusters.get(selected_lemma, []))
        if d.retval == QMessageBox.Cancel:
            d.close()
            return

        if d.retval == QMessageBox.Save:
            new_lemma = d.getLemma()
            phrases   = d.getWordsAndPhrases()

            if new_lemma != selected_lemma:
                del self.topic_clusters[selected_lemma]               # update the data
                self.topic_clusters[new_lemma] = phrases
                index = self.topic1_menu.findText(selected_lemma)     # update the item text
                self.topic1_menu.setItemText(index, new_lemma)
            else:
                self.topic_clusters[selected_lemma] = phrases

            self.pre_defined_topics_field.setPlainText('\n'.join(phrases))

            topic  = self.current_rule.getTopic(0)
            print(topic)
            topic['pre_defined_topics'] = phrases


    def addTopicCluster(self):
        d = TopicClusterDialog()

        if d.retval == QMessageBox.Cancel:
            d.close()
            return

        if d.retval == QMessageBox.Save:
            name = d.getLemma()
            phrases = d.getWordsAndPhrases()

            self.topic_clusters[name] = phrases

            self.topic1_menu.addItem(name) 
            self.topic1_menu.setCurrentText(name)

            self.pre_defined_topics_field.setPlainText('\n'.join(phrases))

            t = dict()
            t['lemma'] = name
            t['pre_defined_topics'] = phrases
            t['no_lexical_overlap'] = False
            t['user_defined']       = True
            self.current_rule.setTopics([t])

    def deleteTopicCluster(self, lemma):
        lemma  = self.topic1_menu.currentText()
        if lemma == "-":
            dialogs.WarningDialog("Select a Topic Cluster", "Select a Topic Cluster from the menu first.")
            return

        if lemma in self.topic_clusters:
            d = dialogs.YesNoDialog("Delete a Topic Cluster: {}".format(lemma), 
                                "{} will be removed from all the composigin patterns.\n\n".format(lemma) + \
                                "Are you sure you want to delete {} permanently?".format(lemma))

            if d.retval == QMessageBox.Cancel:
                return

            elif d.retval == QMessageBox.Ok:

                del self.topic_clusters[lemma]

                index = self.topic1_menu.findText(lemma)
                self.topic1_menu.removeItem(index)
                self.topic1_menu.setCurrentIndex(0)
                self.ruleset.removeTopicCluster(lemma)    

    def topicClusterChanged(self):
        lemma = self.topic1_menu.currentText()
        topics = self.topic_clusters.get(lemma, [])
        topics.sort()
        prompt = '\n'.join(topics)
        self.pre_defined_topics_field.setPlainText(prompt)

    def topicClusterToggled(self, val):
        if val == TOPIC_CLUSTER:
            self.topic1_menu.setEnabled(True)
            self.topic1_add_button.setEnabled(True)
            self.topic1_delete_button.setEnabled(True)
            self.topic1_edit_button.setEnabled(True) 
            self.pre_defined_topics_label.setStyleSheet("color: " + views.default_ui_text_color + ";")
            self.pre_defined_topics_field.setEnabled(True)            
            self.tc_settings_container.show()
            self.literal_topic_field.hide()
            self.literal_topic_field.setText("")
        else:
            self.topic1_menu.setCurrentIndex(0)
            self.topic1_menu.setEnabled(False)
            self.topic1_add_button.setEnabled(False)
            self.topic1_delete_button.setEnabled(False)
            self.topic1_edit_button.setEnabled(False) 
            self.pre_defined_topics_label.setStyleSheet("color: " + views.default_ui_text_inactive_color + ";")
            self.pre_defined_topics_field.setPlainText('')
            self.pre_defined_topics_field.setEnabled(False)
            self.tc_settings_container.hide()
            self.literal_topic_field.show()
            self.literal_topic_field.setText("")

        self.ruleEdited()

    def ruleSetNameEdited(self):
        name = self.rs_name.text()
        if name.find(":") > -1:
            dialogs.WarningDialog("Invalid Character", "You cannot use colon signs (\":\") in the genre/assignment name.")
            self.rs_name.setText(name.replace(":", ""))

        self.ruleSetEdited()

    def isRuleEdited(self):
        return self.b_revised

    def isRuleSetEdited(self):
        return self.b_revised or self.b_rs_revised

    def ruleEdited(self, val=None):
        self.setRuleEdited(True)

    def ruleSetEdited(self):
        self.setRuleSetEdited(True)

    def setRuleEdited(self, val):
        self.b_revised = val
        self.setRuleSetEdited(True)    

    def setRuleSetEdited(self, val):
        self.b_rs_revised = val
        if val:
            curr_title = self.app_win.windowTitle()
            if curr_title and curr_title.endswith("*") == False:
                edited_title = curr_title + "*"
                self.app_win.setWindowTitle(edited_title)
        else:
            curr_title = self.app_win.windowTitle()
            saved_title = curr_title.strip("*")
            self.app_win.setWindowTitle(saved_title)

    def closeEditor(self):
        self.close()

    def closeEvent(self, event):
        if self.isRuleSetEdited():

            wd = dialogs.SaveWarningDialog("Composing Patterns Not Saved", "Do you want to save the changes?")

            if wd.retval == QMessageBox.Discard:
                event.accept()
                return

            elif wd.retval == QMessageBox.Cancel:
                event.ignore()
                # return
            elif wd.retval == QMessageBox.Save:
                # The user says we should save the changes.
                if self.isRuleEdited(): # Let's update the current selection (if it has been eidted).
                    self.updateCurrentSelection()

                # save the rule
                res = self.save()
                if res == False:
                    return

                event.accept()
        else:
            event.accept()

    def addNewGroup(self):
        if self.current_rule is not None:
            if self.current_rule.isGroup() == False and self.isRuleEmpty():
                dialogs.WarningDialog("Warning", "The rule \"{}\" is not defined.".format(self.current_rule.getName()))
                return False

        self.updateCurrentSelection()
        self.cleanUpRuleFields()

        self.untitled_count += 1
        name = 'Group {}'.format(self.untitled_count)

        # Create a rule object
        cp_rule = rule.DSRule(group=True)
        cp_rule.setName(name)

        index = self.rs_list_view.addNewRule(cp_rule)

        self.ruleset.insertRule(index, cp_rule)

        self.current_rule = cp_rule
        self.updateRuleFields(cp_rule)

    def addNewRule(self):

        success = self.updateRule()

        if success == False:
            return

        self.cleanUpRuleFields()

        self.untitled_count += 1
        name = 'Rule {}'.format(self.untitled_count)

        # Create a rule object
        cp_rule = rule.DSRule()
        cp_rule.setName(name)
        index = self.rs_list_view.addNewRule(cp_rule)

        parent_rule = None
        if self.current_rule is not None:

            parent_rule = self.current_rule.getParent()

            if parent_rule is not None:
                parent_rule.insertChild(index, cp_rule)

            elif self.current_rule.isGroup():
                self.current_rule.insertChild(index, cp_rule)
            else:
                self.ruleset.insertRule(index, cp_rule)
        else:

            self.ruleset.insertRule(index, cp_rule)

        self.current_rule = cp_rule
        self.updateRuleFields(cp_rule)

    def deleteRule(self, deleted_rule):
        self.ruleset.deleteRule(deleted_rule)

    def insertRule(self, index, rule):
        self.ruleset.insertRule(index, rule)

    def deleteSelectedRule(self):
        deleted_rule = self.rs_list_view.deleteSelectedRule()

        if self.current_rule is not None:
            # delete the rule
            self.ruleset.deleteRule(deleted_rule)
            self.cleanUpRuleFields()
            self.just_selected_rule = None
            self.current_rule = None
            self.setRuleSetEdited(True)

    def moveSelectedRuleUp(self):
        self.rs_list_view.moveSelectedItem(UP)
        self.setRuleSetEdited(True)

    def moveSelectedRuleDown(self):
        self.rs_list_view.moveSelectedItem(DOWN)
        self.setRuleSetEdited(True)

    def save(self, directory=None, info=None):

        if self.directory is None and directory is None:
            return

        if directory is not None:
            self.directory = directory

        if self.directory is not None and os.path.exists(self.directory):
            shutil.rmtree(self.directory)

        os.makedirs(self.directory)
    
        self.ruleset.setDirectory(self.directory)        

        if self.isRuleSetEdited():
            name = self.rs_name.text()
            name = name.strip()
            self.ruleset.setName(name)

            res = self.updateRule()

            if res == False:
                return False

        common_clusters, uncommon_clusters = self.getClusterTypes()
        self.ruleset.setClusterTypes(common_clusters, uncommon_clusters)

        if self.cv_panel_descr_changed:
            prev_html_str = rule.simplifyHtml(self.cv_panel_description.toHtml(), 'description')
            self.ruleset.setValuesDescription(self.current_cv_panel, prev_html_str)

        self.ruleset.save(self.directory)

        if info is not None:
            timestamp = datetime.now().strftime("%Y-%m-%d-%H%M%S").replace(":", "-")            
            info['saved'] = timestamp
            info_path = os.path.join(self.directory, "info.json")
            with open(info_path, 'w') as fout:
                json.dump(info, fout, indent=4)

        self.path.setText(self.directory)

        self.setRuleEdited(False)
        self.setRuleSetEdited(False)

    def moveRule(self, moved_rule, dest_rule, index):
        self.ruleset.moveRule(moved_rule, dest_rule, index)

    def isRuleEmpty(self):
        # lemma = self.topic1.text()
        lemma = self.topic1_menu.currentText()        
        # expr  = self.expr1_menu.currentText()

        if lemma == '-' and expr == '-':
            return True
        else:
            return False

    def getDirectory(self):
        return self.directory

    def getRuleset(self):
        return self.ruleset

    def getRulesetName(self):
        rs_name_str = self.rs_name.text()            
        return rs_name_str.strip()

    def updateCurrentSelection(self):

        # def getRuleType():
            # type_id = self.type_button_group.checkedId()

            # if type_id == ACTIVE:
            #     type_name = 'active'
            # elif type_id == QUIET:
            #     type_name = 'quiet'
            # elif type_id == OPTIONAL:
            #     type_name = 'optional'
            # elif type_id == BOUNDED_OPTIONAL:
            #     type_name = 'bounded_optional'
            # elif type_id == GROUP:
            #     type_name = 'group'

            # return type_name

        def getValues():
            result = list()
            if self.cv_credible_checkbox.isChecked():
                result.append('credible')

            if self.cv_considerate_checkbox.isChecked():
                result.append('considerate')

            if self.cv_compelling_checkbox.isChecked():
                result.append('compelling')

            if self.cv_ethical_checkbox.isChecked():
                result.append('ethical')

            return result

        if self.current_rule != None and self.isRuleEdited():
            # let's update the current list of words first.
            new_name = self.rule_name.text()
            new_name = new_name.strip()

            self.current_rule.setName(new_name)
            self.rs_list_view.updateCurrentSelection(new_name)

            self.current_rule.setDescription(rule.simplifyHtml(self.rule_description.toHtml(), 'description'))
            self.current_rule.setExamples(rule.simplifyHtml(self.rule_examples.toHtml(), 'example'))
            # self.current_rule.setType(getRuleType())
            self.current_rule.setCVDescription(rule.simplifyHtml(self.values_description.toHtml(), 'description'))
            self.current_rule.setValues(getValues())


            if self.topic_cluster_radio_button.isChecked():
                lemma = self.topic1_menu.currentText()
            else:
                lemma = self.literal_topic_field.text()
                lemma = lemma.strip()

            topics = list()
            if lemma:
                topic1 = dict()
                topic1['lemma'] = lemma
                topic1['user_defined'] = self.topic_cluster_radio_button.isChecked()

                if topic1['user_defined']:
                    phrases = self.pre_defined_topics_field.toPlainText()                    
                    topic1['pre_defined_topics'] = [w for w in phrases.splitlines() if w]
                else:
                    topic1['pre_defined_topics'] = []

                # topic1['no_lexical_overlap'] = self.no_lex_overlap_checkbox1.isChecked()
                topic1['no_lexical_overlap'] = False
                topics.append(topic1)

            self.current_rule.setTopics(topics)


    def updateRule(self):
        if self.current_rule is not None:

            if self.current_rule.isGroup():
                pass
            elif self.isRuleEmpty():
                dialogs.WarningDialog("Warning", "The rule \"{}\" is not defined.".format(self.current_rule.getName()))
                return False

            self.updateCurrentSelection()

            return True
        else:
            return True

    def ruleChanged(self, current, previous):
        self.previous_rule_item = previous     

    def ruleSelected(self, cp_rule):

        self.just_selected_rule = cp_rule

        # do nothing if the same item is clicked
        if self.current_rule == self.just_selected_rule:
            return True

        res = self.updateRule()
        if res:
            self.cleanUpRuleFields()
            self.current_rule = self.just_selected_rule
            self.updateRuleFields(self.current_rule)

            return True
        else:
            self.just_selected_rule = self.current_rule
            return False

    def updateRuleFields(self, cp_rule):

        if cp_rule is None:
            self.cleanUpRuleFields()

            self.rule_description.setDisabled(False)
            self.rule_examples.setDisabled(False)
            # self.active_radio_button.setDisabled(False)
            # self.quiet_radio_button.setDisabled(False)
            # self.optional_radio_button.setDisabled(False)
            # self.bounded_optional_radio_button.setDisabled(False)
            # self.group_radio_button.setDisabled(False)
            self.topic_cluster_radio_button.setDisabled(False)
            self.literal_topic_radio_button.setDisabled(False)
            self.literal_topic_field.setDisabled(False)

            self.topic1_menu.setDisabled(False)
            self.topic1_add_button.setDisabled(False)
            self.topic1_delete_button.setDisabled(False)
            self.topic1_edit_button.setDisabled(False)     

            self.pre_defined_topics_label.setDisabled(False)
            self.pre_defined_topics_field.setDisabled(False)
            # self.no_lex_overlap_checkbox1.setDisabled(False)
            # self.expr1_menu.setDisabled(False)
            self.values_description.setDisabled(False)
            self.cv_credible_checkbox.setDisabled(False)
            self.cv_compelling_checkbox.setDisabled(False)
            self.cv_ethical_checkbox.setDisabled(False)
            self.cv_considerate_checkbox.setDisabled(False)
            self.rule_bold_button.setDisabled(False)
            self.rule_clear_button.setDisabled(False)
            self.rule_italic_button.setDisabled(False)
            self.rule_bullet_button.setDisabled(False)
            self.cv_bold_button.setDisabled(False)
            self.cv_clear_button.setDisabled(False)
            self.cv_italic_button.setDisabled(False)            
            self.cv_bullet_button.setDisabled(False)
            self.sent_clear_button.setDisabled(False)
            self.sent_blue_bold_button.setDisabled(False)
            # self.sent_yellow_button.setDisabled(False)

            return

        rule_revised = self.isRuleEdited()
        rs_revised   = self.isRuleSetEdited()

        if cp_rule.isGroup():
            name = cp_rule.getName()
            self.rule_name.setText(name)
            self.rule_name_header.setText("Group Name")

            self.rule_description.setDisabled(False)
            self.rule_examples.setDisabled(True)

            # self.active_radio_button.setDisabled(True)
            # self.quiet_radio_button.setDisabled(True)
            # self.optional_radio_button.setDisabled(True)
            # self.bounded_optional_radio_button.setDisabled(True)
            # self.group_radio_button.setDisabled(False)
            # self.group_radio_button.setChecked(True)

            self.topic_cluster_radio_button.setDisabled(True)
            self.literal_topic_radio_button.setDisabled(True)
            self.literal_topic_field.setDisabled(True)

            # self.type_button_group.setExclusive(False)
            self.topic_cluster_radio_button.setChecked(False)
            self.literal_topic_radio_button.setChecked(False)
            # self.type_button_group.setExclusive(True)

            self.topic1_menu.setDisabled(True)
            self.topic1_menu.setCurrentIndex(0)
            self.topic1_add_button.setDisabled(True)
            self.topic1_delete_button.setDisabled(True)
            self.topic1_edit_button.setDisabled(True) 
            self.pre_defined_topics_field.setDisabled(True)
            self.pre_defined_topics_field.setDisabled(True)
            # self.no_lex_overlap_checkbox1.setDisabled(True)

            # self.impression1_menu.setDisabled(True)
            # self.impression1_add_button.setDisabled(True)
            # self.impression1_delete_button.setDisabled(True)
            # self.impression1_edit_button.setDisabled(True)      
            # self.user_defined_impression_instruction.setDisabled(True)            
            # self.user_defined_impression_prompt1.setDisabled(True)       
            # self.expr1_menu.setDisabled(True)

            self.values_description.setDisabled(True)
            self.cv_credible_checkbox.setDisabled(True)
            self.cv_compelling_checkbox.setDisabled(True)
            self.cv_ethical_checkbox.setDisabled(True)
            self.cv_considerate_checkbox.setDisabled(True)

            self.rule_bold_button.setDisabled(False)
            self.rule_clear_button.setDisabled(False)
            self.rule_italic_button.setDisabled(False)
            self.rule_bullet_button.setDisabled(False)

            self.cv_bold_button.setDisabled(True)
            self.cv_clear_button.setDisabled(True)
            self.cv_italic_button.setDisabled(True)            
            self.cv_bullet_button.setDisabled(True)

            self.sent_clear_button.setDisabled(True)
            self.sent_blue_bold_button.setDisabled(True)
            # self.sent_yellow_button.setDisabled(True)

            description = cp_rule.getDescription()
            self.rule_description.setHtml(description)

        else:            
            self.rule_description.setDisabled(False)
            self.rule_examples.setDisabled(False)
            # self.active_radio_button.setDisabled(False)
            # self.quiet_radio_button.setDisabled(False)
            # self.optional_radio_button.setDisabled(False)
            # self.bounded_optional_radio_button.setDisabled(False)
            # self.group_radio_button.setDisabled(True)
            self.topic_cluster_radio_button.setDisabled(False)
            self.literal_topic_radio_button.setDisabled(False)            
            self.literal_topic_field.setDisabled(False)

            self.topic1_menu.setDisabled(False)
            self.topic1_add_button.setDisabled(False)
            self.topic1_delete_button.setDisabled(False)
            self.topic1_edit_button.setDisabled(False)      
            self.pre_defined_topics_field.setDisabled(False)            
            self.pre_defined_topics_field.setDisabled(False)                                          

            self.values_description.setDisabled(False)
            self.cv_credible_checkbox.setDisabled(False)
            self.cv_compelling_checkbox.setDisabled(False)
            self.cv_ethical_checkbox.setDisabled(False)
            self.cv_considerate_checkbox.setDisabled(False)
            self.rule_bold_button.setDisabled(False)
            self.rule_clear_button.setDisabled(False)
            self.rule_italic_button.setDisabled(False)
            self.rule_bullet_button.setDisabled(False)
            self.cv_bold_button.setDisabled(False)
            self.cv_clear_button.setDisabled(False)
            self.cv_italic_button.setDisabled(False)            
            self.cv_bullet_button.setDisabled(False)
            self.sent_clear_button.setDisabled(False)
            self.sent_blue_bold_button.setDisabled(False)

            self.rule_name_header.setText("Question")

            name = cp_rule.getName()

            rule_description = cp_rule.getDescription()
            examples    = cp_rule.getExamples()
            cv_description = cp_rule.getCVDescription()

            self.rule_name.setText(name)
            self.rule_description.setHtml(rule_description)
            self.values_description.setHtml(cv_description)
            self.rule_examples.setHtml(examples)

#            type_name = cp_rule.getType()

            # if type_name == 'active':
            #     self.active_radio_button.setChecked(True)
            # elif type_name == 'quiet':
            #     self.quiet_radio_button.setChecked(True)
            # elif type_name == 'optional':
            #     self.optional_radio_button.setChecked(True)
            # elif type_name == 'bounded_optional':
            #     self.bounded_optional_radio_button.setChecked(True)
            # elif type_name == 'group':
            #     self.group_radio_button.setChecked(True)

            values = cp_rule.getValues()
            if 'credible' in values:
                self.cv_credible_checkbox.setChecked(True)
            else:
                self.cv_credible_checkbox.setChecked(False)

            if 'compelling' in values:
                self.cv_compelling_checkbox.setChecked(True)
            else:
                self.cv_compelling_checkbox.setChecked(False)                

            if 'ethical' in values:
                self.cv_ethical_checkbox.setChecked(True)
            else:
                self.cv_ethical_checkbox.setChecked(False)

            if 'considerate' in values:
                self.cv_considerate_checkbox.setChecked(True)
            else:
                self.cv_considerate_checkbox.setChecked(False)

            topics      = cp_rule.getTopics()

            num_topics = len(topics)
            if num_topics > 0:  # 1 or 2
                topic1 = topics[0]

                i = self.topic1_menu.findText(topic1['lemma'])
                self.topic1_menu.setCurrentIndex(i)

                if topic1['user_defined']:
                    self.topic_cluster_radio_button.setChecked(True)
                    self.literal_topic_radio_button.setChecked(False)
                    self.topic1_menu.setEnabled(True)
                    self.topic1_add_button.setEnabled(True)
                    self.topic1_delete_button.setEnabled(True)
                    self.topic1_edit_button.setEnabled(True)

                    self.pre_defined_topics_field.show()
                    self.pre_defined_topics_field.setEnabled(True)

                    lst_of_phrases = topic1.get('pre_defined_topics', [])
                    if type(lst_of_phrases) == str:
                        lst_of_phrases = []
                    lst_of_phrases.sort()
                    self.pre_defined_topics_field.setPlainText('\n'.join(lst_of_phrases))

                    self.pre_defined_topics_label.show()

                    self.tc_settings_container.show()
                    self.literal_topic_field.hide()
                else:
                    self.topic_cluster_radio_button.setChecked(False)
                    self.literal_topic_radio_button.setChecked(True)

                    self.topic1_menu.setEnabled(False)
                    self.topic1_add_button.setEnabled(False)
                    self.topic1_delete_button.setEnabled(False)
                    self.topic1_edit_button.setEnabled(False)

                    self.pre_defined_topics_field.hide()
                    self.pre_defined_topics_field.setPlainText('')
                    self.pre_defined_topics_label.hide()

                    self.tc_settings_container.hide()
                    self.literal_topic_field.show()
                    self.literal_topic_field.setText(topic1['lemma'])

        self.setRuleEdited(rule_revised)
        self.setRuleSetEdited(rs_revised)

    def selectDictionary(self):
        dict_path = QFileDialog.getExistingDirectory(None, 'Select a DocuScope dictionary:', '', QFileDialog.ShowDirsOnly)
        if dict_path:
            self.dict_path = dict_path
            self.dict_path_field.setPlainText(dict_path)
            self.updateClusters()

    def updateClusters(self):
        self.dict_path
        self.dictionary = ds_dict.DSDict(self.controller, self.dict_path)
        self.dictionary.loadTones() 
        tones = self.dictionary.getTones()

        self.cluster_selectors = list()

        for i in reversed(range(self.clust_layout.count())):
            item = self.clust_layout.itemAt(i)
            w = item.widget()
            if w is not None:
                self.clust_layout.removeWidget(w)
                w.setParent(None)
                w.hide()
            del item

        header_hbox = QHBoxLayout()
        header_hbox.addStretch()

        self.common_icon = QLabel()
        pic = QPixmap(utils.resource_path('data/icons/common_icon.png'))
        self.common_icon.setPixmap(pic.scaled(14, 14, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        header_hbox.addWidget(self.common_icon)

        self.uncommon_icon = QLabel()
        pic = QPixmap(utils.resource_path('data/icons/uncommon_icon.png'))
        self.uncommon_icon.setPixmap(pic.scaled(14, 14, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        header_hbox.addWidget(self.uncommon_icon)

        self.neutral_icon = QLabel()
        pic = QPixmap(utils.resource_path('data/icons/neutral_icon.png'))
        self.neutral_icon.setPixmap(pic.scaled(14, 14, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        header_hbox.addWidget(self.neutral_icon)

        self.clust_layout.addLayout(header_hbox)

        for clust in tones:
            s = ImpressionTypeSelector(clust['name'])
            self.clust_layout.addWidget(s)
            self.cluster_selectors.append(s)
        self.clust_layout.addStretch()

    def getClusterTypes(self):
        if self.cluster_selectors is None:
            return [], []
            
        common_clusts   = list()
        uncommon_clusts = list()
        for selector in self.cluster_selectors:
            if selector.getType() == IMPRESSION_COMMON:
                common_clusts.append(selector.getClustName())
            elif selector.getType() == IMPRESSION_UNCOMMON:
                uncommon_clusts.append(selector.getClustName())

        return common_clusts, uncommon_clusts

    def saveAsDocx(self, savepath):

        def parse_html(tag, docx_elem):
            if tag.name == 'p':
                   

                if len(list(tag.children))>0:
                    p = docx_elem.add_paragraph('')
                    p.style = 'Normal' 
                    for c in tag.children:
                        parse_html(c, p)
                else:
                    s = tag.text
                    if s.replace('\n', '').strip() != '':
                        p = docx_elem.add_paragraph('')
                        p.style = 'Normal'                         
                        r = p.add_run(tag.text)

            elif tag.name == 'li':

                if len(list(tag.children))>0:
                    p = docx_elem.add_paragraph('')
                    p.style = 'List Bullet'   
                    for c in tag.children:
                        parse_html(c, p)
                else:
                    s = tag.text
                    if s.replace('\n', '').strip() != '':
                        p = docx_elem.add_paragraph('')
                        p.style = 'List Bullet'                           
                        r = p.add_run(tag.text)

            elif type(tag) == bs4.element.NavigableString:
                if type(docx_elem).__name__ == 'Document':
                    # p = docx_elem.add_paragraph(str(tag))
                    pass
                else:
                    s = str(tag)                    
                    if s.replace('\n', '').strip() != '':
                        s = s.replace('\n', ' ')
                        r = docx_elem.add_run(s)

            elif tag.name == 'span':
                if type(docx_elem).__name__ == 'Document':
                    p = docx_elem.add_paragraph(str(tag.text))
                else:                
                    s = str(tag.text)
                    s = s.replace('\n', ' ')     
                    r = docx_elem.add_run(s)
                    if tag.get('style') == "font-weight:600;":
                        r.bold = True
                    
            else:
                for c in tag.children:
                    parse_html(c, docx_elem)   

        docx = Document()
        styles = docx.styles

        style = styles['Title']   
        font = style.font
        font.color.rgb = RGBColor(0,112,192)

        style = styles['Heading 1']   
        font = style.font
        font.color.rgb = RGBColor(0,112,192)
        style.paragraph_format.keep_with_next = True

        style = styles['Heading 2']
        font = style.font
        font.color.rgb = RGBColor(25,25,25)
        style.paragraph_format.keep_with_next = True

        style = styles['Normal']   
        font = style.font
        font.color.rgb = RGBColor(25,25,25)
        style.paragraph_format.keep_together = True
 
        p = docx.add_paragraph(self.ruleset.getName())
        p.style = 'Title'

        for rule in self.ruleset.getRules():
            if rule.isGroup():

                name        = rule.getName()
                description = rule.getDescription()
                p = docx.add_paragraph(name)
                p.style = 'Heading 1'

                soup = bs(description, "html.parser")
                parse_html(soup, docx)

                for sub_rule in rule.getChildren():
                    name        = sub_rule.getName()
                    description = sub_rule.getDescription()
                    examples    = sub_rule.getExamples()
                    p = docx.add_paragraph(name)
                    p.style = 'Heading 2'

                    soup = bs(description, "html.parser")
                    parse_html(soup, docx)

            else:
                name        = sub_rule.getName()
                description = sub_rule.getDescription()
                examples    = sub_rule.getExamples()
                p = docx.add_paragraph(name)
                p.style = 'Heading 1'

                soup = bs(description, "html.parser")
                parse_html(soup, docx)                

        docx.save(savepath)

class ImpressionTypeSelector(QFrame):

    def __init__(self, clust_name, parent=None):
        super(ImpressionTypeSelector, self).__init__(parent)
        hbox = QHBoxLayout()
        hbox.setContentsMargins(0,0,0,0)
        self.setLayout(hbox)

        self.clust_name = clust_name
        label     = QLabel(clust_name)
        hbox.addWidget(label)

        hbox.addStretch()
            
        self.group = QButtonGroup()
        common_rb = QRadioButton()
        self.group.addButton(common_rb, IMPRESSION_COMMON)
        hbox.addWidget(common_rb)

        uncommon_rb   =  QRadioButton()
        self.group.addButton(uncommon_rb, IMPRESSION_UNCOMMON)
        hbox.addWidget(uncommon_rb)

        neutral_rb   =  QRadioButton()
        neutral_rb.setChecked(True)
        self.group.addButton(neutral_rb, IMPRESSION_NEUTRAL)
        hbox.addWidget(neutral_rb)

    def getType(self):
       return self.group.checkedId()

    def getClustName(self):
        return self.clust_name

class DescrTextEditor(QTextEdit):
    def __init__(self, parent=None):
        super(DescrTextEditor, self).__init__(parent)
        self.doc = self.document()

    def insertFromMimeData(self, src):
        current_html = self.toHtml()

        new_html_str = ""
        soup = bs(current_html, "html.parser")
        for p in soup.find_all('p'):
            new_html_str += "<p>" + p.text + "</p>\n"

        if src.hasHtml():

            html_str = src.html()
            soup = bs(html_str, "html.parser")

            body = soup.find('body')

            if soup.find_all('span'):
                for p in soup.find_all('span'):
                    new_html_str += "<p>" + p.text + "</p>\n"
            else:
                new_html_str = "<p>" + body.text + "</p>\n"

            style = "p {font-family: \'Arial\',\'sans-serif\'; font-style: normal; font-size: " + str(views.editor_normal_font_size) + ";}" + \
                    "span {font-family: \'Arial\',\'sans-serif\'; font-style: normal; font-size: " + str(views.editor_normal_font_size) + ";}"

            self.doc.setDefaultStyleSheet(style)
            self.setHtml(new_html_str)


class RuleTreeWidget(QTreeWidget):
    def __init__(self, editor=None, ruleset=None, parent=None):
        super(RuleTreeWidget, self).__init__(parent)

        self.editor  = editor
        self.ruleset = ruleset

        self.selected_item = None

        self.setHeaderHidden(True)

        # Behavior
        self.itemClicked.connect(self.itemSelected)
        self.setFocusPolicy(Qt.NoFocus)
        self.setDragDropMode(QAbstractItemView.InternalMove)

    def setSelectedItemExpand(self, b):
        if self.selected_item != None:
            self.selected_item.setExpanded(b)
                                                               
    def setSelectedItemText(self, label):
        if self.selected_item != None:
            self.selected_item.setText(0, label)
            self.selected_item.setData(0, Qt.UserRole, label)  # NEW!!

            parent = self.selected_item.parent()
            if parent is not None:
                parent.sortChildren(0, Qt.AscendingOrder)                

    def getRuleHierarchy(self):
        res = dict()
        for index in range(self.topLevelItemCount()):
            rule_item = self.topLevelItem(index)
            children = list()            

            for i in range(rule_item.childCount()):
               sub_rule_item = rule_item.child(i)
               children.append(sub_rule_item.text(0))

            res[rule_item.text(0)] = children

    def dropEvent(self, event):
        super().dropEvent(event)

        if self.pressed_item is not None:
            rule_name = self.pressed_item.data(0, Qt.UserRole)
            moved_rule = self.ruleset.getRuleByName(rule_name)

            parent = self.pressed_item.parent()

            if parent is not None:                                # dropped under a group
                index = parent.indexOfChild(self.pressed_item)
                rule_name = parent.data(0, Qt.UserRole)
                parent_rule = self.ruleset.getRuleByName(rule_name)

                self.editor.moveRule(moved_rule, parent_rule, index)
            else:
                index = self.indexOfTopLevelItem(self.pressed_item)
                self.editor.moveRule(moved_rule, None, index)

            self.setCurrentItem(self.pressed_item)
            self.pressed_item = None
            self.editor.setRuleSetEdited(True)

    def mousePressEvent(self, event):
        pos = event.pos()
        self.pressed_item = self.itemAt(pos)
        super().mousePressEvent(event)

    def moveSelectedItem(self, direction):
        """
        Move up/dow the currently selected item.
        """
        if self.selected_item is None:
            return

        item = self.selected_item
        parent = item.parent()

        # A top level category is selected.
        if parent is None:
            num_rules = self.topLevelItemCount()

            index = self.indexOfTopLevelItem(self.selected_item)

            expanded_items = list()

            if index == 0 and direction == UP:
                return None
            elif index == (num_rules-1) and direction == DOWN:
                return None
            else:
                rule_item = self.topLevelItem(index)
                bExpanded = rule_item.isExpanded()

                if bExpanded:
                    # record all the expanded items in the hierarchy
                    for i in range(rule_item.childCount()):
                        sub_rule_item = rule_item.child(i)
                        if sub_rule_item.isExpanded():
                            expanded_items.append(sub_rule_item)

                rule_item = self.takeTopLevelItem(index)
                self.insertTopLevelItem(index + direction, rule_item) # move it up or down   
                self.setCurrentItem(rule_item)
                rule_item.setExpanded(bExpanded)

                for item in expanded_items:
                    item.setExpanded(True)

            rule_name = rule_item.data(0, Qt.UserRole)
            moved_rule = self.ruleset.getRuleByName(rule_name)
            self.editor.moveRule(moved_rule, None, index + direction)
        else:
            num_sub_rules = parent.childCount()   # num children/sub-rules
            index = parent.indexOfChild(self.selected_item)
            if index == 0 and direction == UP:
                return None
            elif index == (num_sub_rules-1) and direction == DOWN:
                return None
            else:
                sub_rule_item = parent.takeChild(index)
                parent.insertChild(index + direction, sub_rule_item) # move it up or down   
                self.setCurrentItem(sub_rule_item)

                rule_name = parent.data(0, Qt.UserRole)
                parent_rule = self.ruleset.getRuleByName(rule_name)

                rule_name = sub_rule_item.data(0, Qt.UserRole)
                moved_rule = self.ruleset.getRuleByName(rule_name)

                self.editor.moveRule(moved_rule, parent_rule, index + direction)

    def setRules(self, ruleset):
        """
        """
        self.clear()
        untitled_count = 0

        for cp_rule in ruleset.getRules():
            rule_name = cp_rule.getName()

            if cp_rule.isGroup():
                if re.match('Group\s[0-9]+', rule_name):          # increment 'untitled count'.
                    untitled_count += 1

                rule_item = QTreeWidgetItem(self, [rule_name])
                # rule_item.setData(0, Qt.UserRole, cp_rule)
                rule_item.setData(0, Qt.UserRole, rule_name)

                for sub_rule in cp_rule.getChildren():
                    sub_rule_name = sub_rule.getName()
                    if re.match('Rule\s[0-9]+', sub_rule_name):   # increment 'untitled count'.
                        untitled_count += 1

                    sub_rule_item = QTreeWidgetItem(rule_item, [sub_rule_name])
                    sub_rule_item.setFlags(sub_rule_item.flags() ^ Qt.ItemIsDropEnabled)
                    # sub_rule_item.setData(0, Qt.UserRole, sub_rule)
                    sub_rule_item.setData(0, Qt.UserRole, sub_rule_name)
            else:
                if re.match('Rule\s[0-9]+', rule_name):           # increment 'untitled count'.
                    untitled_count += 1

                item = QTreeWidgetItem(self, [rule_name])
                item.setFlags(item.flags() ^ Qt.ItemIsDropEnabled)
                # item.setData(0, Qt.UserRole, cp_rule)
                item.setData(0, Qt.UserRole, rule_name)

        self.selected_item == None

        first = self.topLevelItem(0)
        self.scrollToItem(first)

        self.setFocusPolicy(Qt.NoFocus)

        self.expandAll()

        return untitled_count

    def itemSelected(self, item, column):

        self.pressed_item = None

        parent = item.parent()

        rule_name = item.data(0, Qt.UserRole)
        cp_rule = self.ruleset.getRuleByName(rule_name)

        success = self.editor.ruleSelected(cp_rule)

        if success == False:
            self.setCurrentItem(self.selected_item)
        else:
            self.selected_item = item            

    def deleteSelectedRule(self):
        if self.selected_item != None:

            rule_label     = None
            sub_rule_label = None
            removed_sub_rules = None

            rule_name = self.selected_item.data(0, Qt.UserRole)
            deleted_rule = self.ruleset.getRuleByName(rule_name)

            parent = self.selected_item.parent()

            if parent is None:  # rule
                rule_label = self.selected_item.text(0)
                index = self.indexOfTopLevelItem(self.selected_item)
                self.takeTopLevelItem(index) # remove it

            elif parent is not None: # sub-rule
                sub_rule_label = self.selected_item.text(0)
                rule_label     = parent.text(0)

                index     = parent.indexOfChild(self.selected_item)
                parent.takeChild(index)

            self.clearSelection()
            self.selected_item = None

            return deleted_rule
        else:
            return None

    def addNewRule(self, new_rule):
        if self.selected_item is not None:
            parent    = self.selected_item.parent()
            curr_rule_name = self.selected_item.data(0, Qt.UserRole)
            curr_rule = self.ruleset.getRuleByName(curr_rule_name)
        else:
            parent = None
            curr_rule = None

        name = new_rule.getName()
        new_item = QTreeWidgetItem([name])
        new_item.setData(0, Qt.UserRole, name)

        self.editor.setRuleEdited(True)

        index = 0
        if parent is not None:         # the currently selected item is a sub-rule.
            if new_rule.isGroup():     # if the new_rule is a group, we can't add it to the parent,
                index = self.indexOfTopLevelItem(parent)   # so, we'll add it next to the parent.
                if index >= 0:
                    self.insertTopLevelItem(index+1, new_item)
                    index = index + 1
                else:
                    self.addTopLevelItem(new_item)
                    index = self.topLevelItemCount() - 1

                if new_rule.isGroup() == False:
                    new_item.setFlags(new_item.flags() ^ Qt.ItemIsDropEnabled)                    
            else:
                # the new rule is not a group, so we'll just add the new regular rule to the parent group.
                index = parent.indexOfChild(self.selected_item)
                if index >= 0:
                    parent.insertChild(index+1, new_item)
                else:
                    parent.addChild(new_item)                
                index = index + 1

                new_item.setFlags(new_item.flags() ^ Qt.ItemIsDropEnabled)

        elif curr_rule is not None and curr_rule.isGroup() and new_rule.isGroup() == False:
            self.selected_item.addChild(new_item)
            index = self.selected_item.childCount() - 1
            new_item.setFlags(new_item.flags() ^ Qt.ItemIsDropEnabled)

        else: # top level
            index = self.indexOfTopLevelItem(self.selected_item)
            if index >= 0:
                self.insertTopLevelItem(index+1, new_item)
                index = index + 1
            else:
                self.addTopLevelItem(new_item)
                index = self.topLevelItemCount() - 1

            if new_rule.isGroup() == False:
                new_item.setFlags(new_item.flags() ^ Qt.ItemIsDropEnabled)

        self.setCurrentItem(new_item)
        self.selected_item = new_item

        return index

    def updateCurrentSelection(self, new_name):
        if self.selected_item is not None:
            self.selected_item.setText(0, new_name)
            self.selected_item.setData(0, Qt.UserRole, new_name)

class TopicClusterDialog(QDialog):
    def __init__(self, lemma=None, pre_defined_topics=None):
        super(TopicClusterDialog, self).__init__()

        self.setStyleSheet("TopicClusterDialog {background-color: " + views.default_ui_background_color + ";}" + \
                                                  "QLabel {color: " + views.default_ui_text_color + ";}" + \
                                  "QPushButton {background-color: " + views.default_ui_button_color + ";" + \
                                                          "color: " + views.default_ui_text_color + ";}" + \
                                    "QLineEdit {background-color: " + views.default_ui_input_background_color + ";" + \
                                                          "color: " + views.default_ui_text_color + ";}")
        self.setModal(True)
        self.retval = 0

        self.setFixedWidth(500)

        main_vbox = QVBoxLayout()

        # title
        msg = QLabel("Topic Cluster")
        msg.setStyleSheet("font-weight: bold;")
        main_vbox.addWidget(msg)

        fbox = QFormLayout()
        fbox.setContentsMargins(0,0,0,0)

        # # topic cluster lemma
        if lemma is not None:
            self.lemma_field = QLineEdit(lemma)
        else:
            self.lemma_field = QLineEdit()
        fbox.addRow(QLabel("Topic Cluster Name:  "), self.lemma_field)

        if pre_defined_topics is not None:        
            self.pre_defined_topics_field = QPlainTextEdit('\n'.join(pre_defined_topics))
        else:
            self.pre_defined_topics_field = QPlainTextEdit()

        # self.pre_defined_topics_field.setFixedHeight(80)
        fbox.addRow(QLabel("Pre-Defined Words/Phrases"), self.pre_defined_topics_field)

        main_vbox.addLayout(fbox)

        # buttons
        hbox = QHBoxLayout()
        hbox.addStretch()

        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        cancel_button.setAutoDefault(False)        
        hbox.addWidget(cancel_button)

        ok_button = QPushButton("Save")
        ok_button.clicked.connect(self.accept)
        ok_button.setAutoDefault(True)
        hbox.addWidget(ok_button)

        main_vbox.addLayout(hbox)

        self.setLayout(main_vbox)

        self.setFixedWidth(400)
        self.lemma_field.setFocus()

        # self.show()
        self.retval = self.exec_() 

    def reject(self):
        self.done(QMessageBox.Cancel)

    def accept(self):
        self.done(QMessageBox.Save)

    def getLemma(self):
        text = self.lemma_field.text()
        text = text.strip()
        return text

    def getWordsAndPhrases(self):
        text = self.pre_defined_topics_field.toPlainText()
        text = text.strip()
        lst = text.splitlines()
        lst.sort()
        return lst






